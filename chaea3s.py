#!/user/bin/env python
#===================================================
#===================================================
# CHAEA3S:
# CHAEA AUTOMATIC ANALYSIS OF LEARNING STYLES
#===================================================
#===================================================
# VERSION 7.i
#
# This program implements a principal component 
# analysis of the learning styles of a group of
# students as prescribed by CHAEA:
# activist, theorist, pragmatist, and reflector.
# The program is based on the program 
# learning_styles_v7e.py, by Fabio Revuelta.
#
# Further information can be found in the reference:
# J. Ablanque, V. Gabaldon, P. Almendros, J. C. Losada,
# R. M. Benito,  and F. Revuelta
# “CHAEA3S: A software package for comprehensive analysis
# of learning styles and academic performance”
# PLOS ONE (2026).
#
# The software provided here is distributed on an “as-is” basis,
# without any warranties or guarantees of any kind.
# While we have made every effort to ensure its accuracy
# and reliability, we cannot be held responsible for any
# unintended consequences, errors, or issues that may arise from its use.
# Users are encouraged to thoroughly test the software,
# review the source code, and exercise due diligence before deploying it 
# in any critical environment.
#
# Version 7.i also performs two aditional optional calculations,
# if the corresponding input data is provided.
# On the one hand, if the learning styles of the
# teacher(s)/instructor(s) are given as input(s),
# the learning styles of the students are plotted as
# a function of them on the plain of the two main principal
# components using a certain color.
# On the other hand, if the marks of the students are provided,
# a study of their relationship with the learning styles is also
# conducted as a function of the original learning-styles classification,
# and as a function of the principal components.
# Furthermore, a clustering analysis based on K-means algorithm
# is also performed.
#
# In version 7.i, the default web browser is opened after execution
# in case the used wants to contact the software authors.
#
# By using this software, you agree to hold the developers harmless 
# from any liability, damages, or losses resulting from its use.
# Additionally, we kindly request that you cite the reference mentioned
# above when sharing or distributing it, and 
# acknowledge Fabio Revuelta as the author of this software.
# 
# Feel free to adapt this disclaimer further to match your specific context.
# And remember to give credit to us :)
#
# Copyright 2026 Fabio Revuelta
# 
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# 
#     http://www.apache.org/licenses/LICENSE-2.0
# 
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
# Author: Fabio Revuelta
#         Grupo de Sistemas Complejos
#         Universidad Politecnica de Madrid
#         Madrid, July 2026
#
#===================================================
# PRELIMINARIES
#===================================================
from mpl_toolkits.mplot3d import Axes3D
from mpl_toolkits.mplot3d import proj3d
from matplotlib.patches import FancyArrowPatch
from matplotlib import cm
#
import pandas as pd 
import matplotlib
import matplotlib.pyplot as plt
import numpy as np
import os, fnmatch
from matplotlib.ticker import  MultipleLocator, FormatStrFormatter
from decimal import Decimal
from numpy.linalg import eig
#
import statistics
import scipy.stats as st
from mpl_toolkits.mplot3d.art3d import Poly3DCollection
import random
#
from random import seed
from random import gauss
from scipy.stats import linregress
from scipy.optimize import curve_fit
#
from sklearn.cluster import KMeans
#
import itertools
import datetime
#
# Library to contact via e-mail
import webbrowser
#
# Necessary libraries for the summary report
import subprocess
#
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
#
#
# Necessary for the log file
import os
import re
import sys
import datetime
#
## For Windows
#import codecs
##sys.stdout = codecs.getwriter('utf-8')(sys.stdout.detach())   
#
#
#===================================================
# CHAEA3S.LOG FILE
#===================================================
# Current folder
current_folder = os.getcwd()
print('Current folder : ' + current_folder)
print(' ')
#
# Output folders
output_gen = current_folder + '/output'
os.makedirs(output_gen, exist_ok=True) 
print('Output general folder         : ' + output_gen)
print(' ')
#
# Define the file name for logging
log_file = "output/chaea3s.log"
#
if os.path.exists(log_file):
  os.remove(log_file)
  print(f"File '{log_file}' removed successfully.")
else:
  print(f"File '{log_file}' did not previously exist.")
#
print(f"File '{log_file}' created.")
print()
#
# Create a new log_file
with open(log_file, "w") as f:
  f.write('==============================\n')
  f.write('CHAEA3S.LOG\n')
  f.write('==============================\n')
  f.write('This file contains all the information on the execution of CHAEA3S package\n')
  f.write('\n')
  pass
#
# Define a custom print function
def printt(*args):
  print(*args, end=' ')
#   
  with open(log_file, "a", encoding='utf-8') as f:
    f.write(str(*args)+'\n')
#
# Redirect stdout to log file
#sys.stdout = open(log_file, "a")
#
# Get the current date and time
current_datetime = datetime.datetime.now()
#
# Format the date and time as a string
formatted_datetime = current_datetime.strftime("%Y-%m-%d %H:%M:%S")
printt('Program starts on: ' + formatted_datetime)
printt(' ')
#
# Check operating system
# (necessary to transform the docx report in pdf)
if(sys.platform == 'linux' or sys.platform == 'linux2'):
  operating_system = 'linux'
  printt('The operating system is Linu.x')
elif(sys.platform == 'darwin'):
  operating_system = 'mac'
  printt('The operating system is macOS (Darwin).')
elif(sys.platform == 'win32' or sys.platform == 'win64'):
  operating_system = 'win'
  printt('The operating system is Windows.')
#
printt(' ')
#
#===================================================
# PLOT PARAMETERS
#===================================================
#----------------------------------------------------
def write_number_with_decimals(x, n):
  return "{:.{}f}".format(x, n)
#----------------------------------------------------
def find_number_position(string, number):
  try:
    position = string.index(str(number))
    return position
  except ValueError:
    return -1  # Number not found in the string
#----------------------------------------------------
def mean_uncert(x, dx):
# This subroutine returns a string with the average mean
# and the corresponding uncertainty with the correct
# number of decimals
# Convert numbers to strings if necessary
  if isinstance(x, str):
    x_str  = x
    xin    = float(x)
  else:
    x_str  = str(x)
    xin    = x

  if isinstance(dx, str):
    dx_str = dx
    dxin   = float(dx)
  else:
    dx_str = str(dx)

# Check if dx==nan
  if ( dx_str == 'nan' or dx_str == '0.0'):
     x_out = write_number_with_decimals(xin, 1) + '(nan)'
  else:
#   Get the three most significant digits of dx
#   (in case the point is included)
    for digit in dx_str:
      if(digit != '0' and digit != '.'):
        dx0 = int(digit)
        break

    pos0 = find_number_position(dx_str, dx0)
    dec  = find_number_position(dx_str, '.')
    if(dec == -1):
      dx_str = dx_str + '.0'
      dec  = find_number_position(dx_str, '.')

    if(dx0 != 1):
#     dx0 > 1
      if( pos0 < dec ):
        dx_out = dx_str[pos0]
        for j in range(pos0, dec-1):
          dx_out = dx_out + '0'
      else:
        dx_out = dx_str[0:pos0+1]
    else:
#     dx0 = 1
      if( pos0 > dec ):
        if(pos0+1 == len(dx_str)):
          dx_out = dx_str + '0'
        else:
          dx_out = dx_str[0:pos0+2]
      else:
        if(pos0+1 == dec):
          if(pos0+2 == len(dx_str)):
            dx_out = dx_str + '0'
          else:
            dx_out = dx_str[0:pos0+3]
        else:
          dx_out = dx_str[pos0] + dx_str[pos0+1]
          for j in range(pos0+1, dec-1):
            dx_out = dx_out + '0'

#   Count the number of decimals
    num_decimals = len(dx_out.split('.')[1]) if '.' in dx_out else 0

#   Get the same number of decimals from x
    x_out = write_number_with_decimals(xin, num_decimals) + '(' + dx_out + ')'

  return x_out
#----------------------------------------------------
def round_until_two_non_9(xin):

    if isinstance(xin, str):
      x = float(xin)
    else:
      x = xin

    if (-1 < x < 1):
      s = f"{x:.20f}"  # Convert to string with many decimals
      sign = "-" if x < 0 else ""
      decimals = s.split(".")[1]

      result = "0."
      non9_count = 0
      for d in decimals:
        result += d
        if d != '9':
          non9_count += 1
        if non9_count == 2:
           break
      output = str(sign + result)
    else:
      if( x == -1 ):
        output = '-1.0'
      else:
        if( x == 1 ):
          output = '1.0'
        else:
          if(np.isnan(x)):
            output = 'nan'
          else:
            raise ValueError("x must be between -1 and 1")

    return output
#----------------------------------------------------
#----------------------------------------------------
printt('===========================================')
printt('PLOTTING PARAMETERS...')
printt('===========================================')
printt(' ')
#
# Labels for the learning styles (LSs)
List_LS        = ['activist', 'reflector', 'theorist', 'pragmatist']
Label_LS       = ['Activist', 'Reflector', 'Theorist', 'Pragmatist']
Label_LS_print = ['Activist  ', 'Reflector ', 'Theorist  ', 'Pragmatist']
#
# Labels for the principal components (PCs)
Label_PC = [0, 1, 2, 3]
#
Label_PC   = list(itertools.chain(range(0, len(Label_LS))))
Label_PCPC = list(itertools.chain(range(0, len(Label_LS))))

for i in range(0, len(Label_LS)):
  Label_PC[i] = str(Label_PC[i])
  Label_PCPC[i] = 'PC' + Label_PC[i]
#
# Tendency labels
Label_tendencies       = ['Very low', 'Low', 'Moderate', 'High', 'Very high']
Label_tendencies_print = ['Very low ', 'Low      ', 'Moderate ', 'High     ', 'Very high']
#
#---------------------------------------------------
# Size, Label_LS and ticks of the figures
w_fig = 28
h_fig = w_fig / 1.61
#
# Original: labels = 30, ticks = 0.9*labels
labelsize = 50
ticksize  = 0.90 * labelsize
#
# Quartiles
labelsize_quartile = 50
ticksize_quartile  = 0.9 * labelsize_quartile
#
# W(PR)
labelsize_PR = 50
ticksize_PR  = 0.9 * labelsize_PR
#
# W(PC)
labelsize_PC = 50
ticksize_PC  = 0.9 * labelsize_PC
#
# 3D plots
labelsize_3D = 50
ticksize_3D  = 0.9 * labelsize_3D
labelpad_3D  = labelsize_3D
#
width_PC_vectors_3D = [9, 7, 5, 4]
width_PC_vectors_2D = [9, 7, 5, 5]
#
colors_proj     = 'tab:blue'
edgecolors_proj = 'tab:blue'
#
# Colors of the projections on the Cartesian planes
# if each PC corresponds to a different LS
colors_proj_PC     = 'tab:blue'
edgecolors_proj_PC = 'tab:blue'
#
# Colors of the projections on the Cartesian planes
# if more than one PC is associated with the same LS
colors_proj_PC2     = 'tab:blue'
edgecolors_proj_PC2 = 'tab:blue'
#
# Mean average potting parameters
size_mean_ref = 30000
mean_symbol   = '*'
mean_color    = 'tab:blue'
mean_color2   = 'gray'
#
# Uncertainty of the average mean
mean_uncert_color = 'pink'
mean_alpha        = 0.5
#
# 2D and 3D plots for the PCs
twenty_symbol = ['s', 's', 's', 's', 's', 'D']
twenty_color  = ['r', 'green', 'purple', 'orange', 'k', 'brown']
twenty_line   = ['--', '--', '--', '--', '--', 'dotted']
twenty_alpha  = 0.5
#
twenty_width_3D = [2, 2, 1, 1, 1, 1, 2]
twenty_width_2D = [2, 2, 1, 1, 2, 2, 2]
twenty_width_W  = [6, 6, 6, 6, 3, 4, 2]
#
tendency_color = ['pink', 'brown', 'b']
#
origin_symbol = '^'
origin_color  = 'k'
origin_alpha  = 0.5


# Clustering parameters
cluster_symbol = ['o', 's', '^', 'v', '*', 'D']
cluster_color  = ['dodgerblue', 'darkred', 'forestgreen', 'navy', 'goldenrod', 'saddlebrown' ]
cluster_alpha  = 0.5
#---------------------------------------------------
# Number of plots
Nplots = 1
printt('  Nplots : ' + str(Nplots))
# If Nplots = 1, then only one 2D and 3D
# representions as a function of the LSs are
# performed. Otherwise, all (2D and 3D)
# possible combinations are consided
#
#---------------------------------------------------
# Plotting limits for the plots as a function 
# of the LSs
xmin =  0
xmax = 20.4
#
ymin = xmin
ymax = xmax
#
zmin = xmin
zmax = xmax
#
xmin3D = xmin
xmax3D = xmax
#
ymin3D = xmin3D
ymax3D = xmax3D
#
zmin3D = xmin3D
zmax3D = xmax3D
#
printt( ' xmin : ' + str(xmin))
printt( ' xmax : ' + str(xmax))
printt(' ')
printt( ' ymin : ' + str(ymin))
printt( ' ymax : ' + str(ymax))
printt(' ')
printt( ' zmin : ' + str(zmin))
printt( ' zmax : ' + str(zmax))
printt(' ')
#---------------------------------------------------
# Parameters for the LS statistical analysis
LS_mean_line  = 'dotted'
LS_mean_width = 2.5
LS_mean_alpha = 0.05
#
LS_color = ['r', 'green', 'purple', 'orange', 'pink']
LS_line  = ['-',     '--', 'dashdot',    'dotted']
LS_areas = ['b', 'orange', 'green', 'pink']
LS_alpha = 0.25
LS_alpha2 = 0.60
LS_alpha3 = 1.0 #0.60
#
affinity_alpha = 1.0 #0.60
#
LS_hist_alpha  = 0.25
LS_hist_alpha2 = 0.40
#
LS_W_width      = 4
#
# Parameters of the participation ratios (PRs)
Label_PR = ['PR (LS)', 'PR (PC)']
#
PR_color = ['r', 'b']
PR_line  = ['-',     '--', 'dashdot',    'dotted']
PR_areas = ['brown', 'orange', 'green', 'pink']
PR_alpha = 0.40
#
PR_mean_line  = 'dotted'
PR_mean_width = 2.5
PR_mean_alpha = 0.05
#
PR_W_width      = 2.5
PC_W_width      = 4
#
# Parameters of the quartiles of the PRs
quartile_width = 2
quartile_color = 'k'
quartile_line  = 'dotted'
#
PR_hist_quartile_color = ['r', 'b', 'green', 'orange', 'purple', 'pink']
PR_hist_quartile_alpha = 1.0
#---------------------------------------------------
# Parameters of the eigenvectors and the projections
# on them (principal components, PCs)
PC_color = ['r', 'green', 'purple', 'orange', 'pink', 'brown']
PC_line  = ['-', '--', 'dashdot', 'dotted']
PC_line_bis = ['-', '-', '-', '-']
PC_areas = ['b', 'orange', 'green', 'pink']
PC_width = 2.5
#
PC_mean_line  = 'dotted'
PC_mean_width = 2.5
PC_mean_alpha = 0.05
#
PC_alpha = 0.25
PC_alpha2 = 0.60
#
PC_hist_alpha  = 0.25
PC_hist_alpha2 = 0.40
#
Weibull_width = 2
#
max_iter_Weib = 5000
#---------------------------------------------------
# Parameters the PCs
colors_proj_lines = ['r', 'purple']
style_proj_lines  = ['-', '--']
colors_proj_areas = ['b', 'orange', 'green', 'pink']
#---------------------------------------------------
# Possible combinations of the learning styles in pairs
learning_pairs = [[0,1], [0,2], [0,3], [1,2], [1,3], [2, 3]]
#
# Possible combinations of the learning styles in trios
learning_trios = [[0,1,2], [0,1,3], [0,2,3], [1,2,3]]
#----------------------------------------------------------
printt('-------------------------------------------')
printt('  Definition of tendency/plotting functions...')
printt('-------------------------------------------')
printt(' ')
# Tendency depending on the reference LS
# (activist -a-, theorist-t-, pragmatist -p-, reflector -r-).
#
def scatter_tendency(ls, value):
  # Activist
  if(ls == 'Activist'):
    if value<=6:
      tendency='vl'
    else:
      if value<=8:
        tendency='l'
      else:
        if value<=12:
          tendency='m'
        else:
          if value<=14:
            tendency='h'
          else:
            if value<=20:
              tendency='vh'
            else:
              printt('Wrong input for the ls :' + ls)
  else:
    # Reflector
    if(ls == 'Reflector'):
      if value<=10:
        tendency='vl'
      else:
        if value<=13:
          tendency='l'
        else:
          if value<=17:
            tendency='m'
          else:
            if value<=19:
              tendency='h'
            else:
              if value<=20:
                tendency='vh'
              else:
                printt('Wrong input for the ls :' + ls)
    else:
      # Theorist
      if(ls == 'Theorist'):
        if value<=6:
          tendency='vl'
        else:
          if value<=9:
            tendency='l'
          else:
            if value<=13:
              tendency='m'
            else:
              if value<=15:
                tendency='h'
              else:
                if value<=20:
                  tendency='vh'
                else:
                  printt('Wrong input for the ls :', ls)
      else:
        # Pragmatist
        if(ls == 'Pragmatist'):
          if value<=8:
            tendency='vl'
          else:
            if value<=10:
              tendency='l'
            else:
              if value<=13:
                tendency='m'
              else:
                if value<=15:
                  tendency='h'
                else:
                  if value<=20:
                    tendency='vh'
                  else:
                    printt('Wrong input for the ls :', ls)
  return tendency
#===================================================
def tendency_long_name(tendency):
  if(tendency == 'vl'):
    return 'Very low'
  else:
    if(tendency == 'vll'):
      return 'Very low/ Low'
    else:
      if(tendency == 'l'):
        return 'Low'
      else:
        if(tendency == 'lm'):
          return 'Low/ Moderate'
        else:
          if(tendency == 'm'):
            return 'Moderate'
          else:
            if(tendency == 'mh'):
              return 'Moderate/ High'
            else:
              if(tendency == 'h'):
                return 'High'
              else:
                if(tendency == 'hvh'):
                  return 'High/ Very high'
                else:
                  if(tendency == 'vh'):
                    return 'Very high'
                  else:
                    return 'Error'
#===================================================
def tendency_intermediate(ls, value):
  # Here, we also take into account the transition values
  # for the group of students, i.e., 6.5 points for active
  # is not very low (it is larger than 6) but smaller than 
  # low (it is smaller than 7).
  # Activist
  if(ls == 'Activist'):
    if value<=6:
      tendency='vl'
    else:
      if value<7:
        tendency='vll'
      else:
        if value<=8:
          tendency='l'
        else:
          if value<9:
            tendency='lm'
          else:
            if value<=12:
              tendency='m'
            else:
              if value<13:
                tendency='mh'
              else:
                if value<=14:
                  tendency='h'
                else:
                  if value<15:
                    tendency='hvh'
                  else:
                    if value<=20:
                      tendency='vh'
                    else:
                      printt('Wrong input for the ls :', ls)
  else:
    # Reflector
    if(ls == 'Reflector'):
      if value<=10:
        tendency='vl'
      else:
        if value<11:
          tendency='vll'
        else:
          if value<=13:
            tendency='l'
          else:
            if value<14:
              tendency='lm'
            else:
              if value<=17:
                tendency='m'
              else:
                if value<18:
                  tendency='mh'
                else:
                  if value<=19:
                    tendency='h'
                  else:
                    if value<20:
                      tendency='hvh'
                    else:
                      if value<=20:
                        tendency='vh'
                      else:
                        printt('Wrong input for the ls :', ls)
    else:
      # Theorist
      if(ls == 'Theorist'):
        if value<=6:
          tendency='vl'
        else:
          if value<7:
            tendency='vll'
          else:
            if value<=9:
              tendency='l'
            else:
              if value<10:
                tendency='lm'
              else:
                if value<=13:
                  tendency='m'
                else:
                  if value<14:
                    tendency='mh'
                  else:
                    if value<=15:
                      tendency='h'
                    else:
                      if value<16:
                        tendency='hvh'
                      else:
                        if value<=20:
                          tendency='vh'
                        else:
                          printt('Wrong input for the ls :', ls)
      else:
      # Pragmatist
        if(ls == 'Pragmatist'):
          if value<=8:
            tendency='vl'
          else:
            if value<9:
              tendency='vll'
            else:
              if value<=10:
                tendency='l'
              else:
                if value<11:
                  tendency='lm'
                else:
                  if value<=13:
                    tendency='m'
                  else:
                    if value<14:
                      tendency='mh'
                    else:
                      if value<=15:
                        tendency='h'
                      else:
                        if value<16:
                          tendency='hvh'
                        else:
                          if value<=20:
                            tendency='vh'
                          else:
                            printt('Wrong input for the ls :', ls)
  return tendency
#
#===================================================
def scatter_properties(tendency, value):
  # It returns a color and a number, depending on whether
  # the tendency towards a certain learning style is
  # very low/low (blueish), moderate (greenish) or high/very high (redish).
  #
  # Scatter color and symbol
  if(tendency=='vl' or tendency=='l'):
    scatter_color='b'
    scatter_symbol='o'
  else:
    if(tendency=='m'):
      scatter_color='b'
      scatter_symbol='o'
    else:
      if(tendency=='h' or tendency=='vh'):
        scatter_color='b'
        scatter_symbol='o'
  #
  # Scatter filling
  alphamin=0.1
  alphamax=0.8
  scatter_alpha=alphamin+(alphamax-alphamin)*value/20
  #
  # Scatter size
  scattersizemax = 150*100/L # The more points, the smaller the scatter size
  scattersizemin =  20*100/L
  #
  scatter_size=scattersizemin+(scattersizemax-scattersizemin)*value/20
  if ( scatter_size < 1):
    scatter_size = 1
  else:
    if ( scatter_size > 500):
      scatter_size = 500
  #               
  return [tendency, scatter_color, scatter_size, scatter_alpha, scatter_symbol]
printt('-------------------------------------------')
printt('  Definition of tendency/plotting functions done!')
printt('-------------------------------------------')
#
#
#---------------------------------------------------
# Fitting functions
#---------------------------------------------------
#
printt('-------------------------------------------')
printt('    Fitting functions...')
printt('-------------------------------------------')
printt(' ')
# We fit the projections of the LSs of each student and the
# corresponding probabilities using Weibull distributions.
# The fitting is performed considering the cumulative 
# distribution functions 
# (regular Weibull distribution for the probabilities
# and the translated Weibull distribution for the
# projections)
#
# Cumulative function for Weibull distribution
def Wweibull(x, alpha_Weibull, k_Weibull):
  return 1 - np.exp(-(x/alpha_Weibull)**k_Weibull)
#
# Probability density function for Weibull distribution
def Pweibull(x, alpha_Weibull, k_Weibull):
  return (k_Weibull/alpha_Weibull) * (x/alpha_Weibull)**(k_Weibull-1) * np.exp(-(x/alpha_Weibull)**k_Weibull)
#  
# Cumulative function for the translated Weibull distribution
# Translated Weibull distribution for the projections on the PCs
def Wweibull_translated(x, alpha_Weibull, k_Weibull, theta_Weibull):
  return 1 - np.exp(-((x-theta_Weibull)/alpha_Weibull)**k_Weibull)
# 
# Probability density function for the translated Weibull distribution
def Pweibull_translated(x, alpha_Weibull, k_Weibull, theta_Weibull):
  return (k_Weibull/alpha_Weibull) * ((x-theta_Weibull)/alpha_Weibull)**(k_Weibull-1) * np.exp(-((x-theta_Weibull)/alpha_Weibull)**k_Weibull)
printt(' ')
#
printt('-------------------------------------------')
printt('    Fitting functions done!')
printt('-------------------------------------------')
printt(' ')
#
#===================================================
#===================================================
#
#===================================================
printt('===========================================')
printt('PLOTTING PARAMETERS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# FOLDERS
#===================================================
printt('===========================================')
printt('FOLDERS...')
printt('===========================================')
printt(' ')
#
# Current folder
current_folder = os.getcwd()
printt('Current folder : ' + current_folder)
printt(' ')
#
# Input folder
input_folder = current_folder + '/input/'
printt('Input folder   : ' + input_folder)
printt(' ')
#
# Output folders
output_gen = current_folder + '/output'
os.makedirs(output_gen, exist_ok=True) 
printt('Output general folder         : ' + output_gen)
printt(' ')
#
output = output_gen + '/figs/'
os.makedirs(output, exist_ok=True) 
printt('Output figures folder         : ' + output)
printt(' ')
#       
output_participants = output + '/participants'
os.makedirs(output_participants, exist_ok=True) 
printt('Output-participants folder : ' + output_participants)
printt(' ')
#
output_statistics = output + '/statistics'
os.makedirs(output_statistics, exist_ok=True) 
printt('Output-statistics folder : ' + output_statistics)
printt(' ')
#
output_statistics_pr = output_statistics + '/pr'
os.makedirs(output_statistics_pr, exist_ok=True) 
printt('Output-PR folder : ' + output_statistics_pr)
printt(' ')
#
output_statistics_ls = output_statistics + '/ls'
os.makedirs(output_statistics_ls, exist_ok=True) 
printt('Output-statistics-LS folder : ' + output_statistics_ls)
printt(' ')
#
# Marks folders
output_participants_ls_marks = output_participants + '/marks'
os.makedirs(output_participants_ls_marks, exist_ok=True)
printt('Output-LS-marks folder : ' + output_participants_ls_marks)
#
output_participants_pc_marks = output_participants + '/marks'
printt('Output-PC-marks folder : ' + output_participants_pc_marks)
printt('  ')
printt('===========================================')
printt('FOLDERS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#
#
#
#===================================================
# INPUT READING
#===================================================
printt('===========================================')
printt('READING INPUT FILES...')
printt('===========================================')
printt(' ')
nan = 0
#
studentsin = [] # List with the names of all the students
students   = [] # List with the names of the students that have correct input data
data       = [] # Input data

data_marks = ['Excel', 'Sheet', 'Student'] # List with the names of the students that have correct input data and their marks or grades (if given)
datain_error_students = []
#
# Iteration over all input files (.xls and .xlsx)
for files in os.listdir(input_folder):
# 
# xls and xlsx files
  if ((files.endswith('.xls')) or (files.endswith('.xlsx'))):
    studentsin.append(files)
#
studentsin = sorted(studentsin)
#
for filei in studentsin:
#
    data_in = []
    opt = 0
    try:
      data_in = pd.read_excel(input_folder + filei, header=None, sheet_name='CÁLCULO')
    except FileNotFoundError:
      print(f"File '{filei}' not found.")
    except ValueError:
      print("Worksheet 'CÁLCULO' not found. Trying alternative sheet names...")
      try:
        data_in = pd.read_excel(input_folder + filei, header=None, sheet_name='CALCULO')
      except ValueError:
        try:
          data_in = pd.read_excel(input_folder + filei, header=None, sheet_name='CALCULATION')
          opt = 1
        except ValueError:
          print("Neither 'CALCULATIONS', 'CALCULO' nor 'CÁLCULO' sheet found.")

    if not data_in.empty:
      data_LS = data_in.loc[29].values          # Extract the LS data (I)
      if( opt != 1):
        line    = np.delete(data_LS, [0,1,6,7,8]) # Extract the LS data (II)
      else:
        line    = np.delete(data_LS, [0,1,2,7,8]) # Extract the LS data (II)
#
    file_line = filei+'      '+str(line[0])+' '+str(line[1])+' '+str(line[2])+' '+str(line[3])
#
    try:   
      n0 = float(line[0])
      n1 = float(line[1])
      n2 = float(line[2])
      n3 = float(line[3])
      error = 'False'      
    except ValueError:
      error = 'True'
#
    if ( (error == 'True') or
       ( np.isnan(n0) or np.isnan(n1) or np.isnan(n2) or np.isnan(n3) 
         or (n0 < 0) or (n0 > 20) or (n1 < 0) or (n1 > 20)
         or (n2 < 0) or (n2 > 20) or (n3 < 0) or (n3 > 20) ) ):

      printt('Invalid input file')
      printt(file_line)
      if(len(datain_error_students) == 0):
        datain_error_students.append(filei)
      else:
        datain_error_students.append(', ' + filei)

    else:
      printt(file_line)
      students.append(filei)                        # Students' names added (I)
      data_marks.append(os.path.splitext(filei)[0]) # Students' names added (II)
      data.append([n0, n1, n2, n3])                 # LS values added to data matrix

datain_error_students.append('.')
data = np.vstack(data) # Stack the list
#
shape = np.shape(data)
#
L = shape[0]
K = shape[1]
#
xa = data[:, 0]
xr = data[:, 1]
xt = data[:, 2]
xp = data[:, 3]
#
printt(' xa : ' + str(xa))
printt(' ')
printt(' xr : ' + str(xr))
printt(' xt : ' + str(xt))
printt(' ')
printt(' xp : ' + str(xp))
printt(' ')
#
printt(' ')
#
#-------------------------------------------------------
# OPTIONAL INPUT FILE(S) WITH EDUCATORS' LEARNING STYLES
#-------------------------------------------------------
# The corresponding Excel files must be
# all contained in a folder placed in the input folder
# named educator, professor, teacher, or instructor.
# If more than one folder is included,
# the program only considers the data included first
# in the folder educator. If professor folder does
# not exist, then the data in professor is considered.
# If none of them appear, then the data in the
# teacher folder is used. If none of the previous
# three folders are included, then the data in the
# instructor folder is analyzed.
# Finally, if no folder named educator, professor,
# teacher, or instructor is inluded in the input
# folder, no data for the educator is used.

# Input file with the LSs of the educator(s)/professor(s)/teacher(s)/instuctor(s)
vprofessor = ['educator', 'professor', 'teacher', 'instructor']
vProfessor = ['Educator', 'Professor', 'Teacher', 'Instructor']

nfiles_prof = 0
Lprof = 0

prof_ref = vprofessor[0]
Prof_ref = vProfessor[0]
# Check if one or more of the folders exists
for professor in vprofessor:
  input_folder_prof = input_folder + professor + '/'
  if os.path.exists(input_folder_prof) and os.path.isdir(input_folder_prof):
    printt(f"The folder '{input_folder_prof}' exists.")
    prof_ref = professor
    Prof_ref = vProfessor[nfiles_prof]
    nfiles_prof += 1
    input_folder_prof_ref = input_folder_prof
  else:
    printt(f"The folder '{input_folder_prof}' does not exist.")

printt('nfiles_prof :' + str(nfiles_prof))

if ( nfiles_prof == 0):
  printt(' The learning styles of the educator(s) have not been provided')
else:
  datain_error_profs = []
  if ( nfiles_prof >= 2):
    printt(' Error: Conflict detected. Multiple folders associated with the learning styles of the educator(s). Please, change so that all the input files are contained in a single folder with one of these names:')
    for i in range(0, len(vprofessor)):
      printt(vprofessor[i])
  else:
    if ( nfiles_prof == 1):
      printt(' The learning styles of the educator(s) have been provided')

      input_folder_prof = input_folder_prof_ref

      profs       = []
      data_prof   = []
      profs_files = []

      for files in os.listdir(input_folder_prof):
#
#       xls and xlsx files
        if ((files.endswith('.xls')) or (files.endswith('.xlsx'))):
          profs_files.append(files)
#
      profs_files = sorted(profs_files)
#
      for filei in profs_files:
#
          profi = filei.rsplit(".", 1)[0]
          data_prof_in = []
          opt = 0
          try:
            data_prof_in = pd.read_excel(input_folder_prof + filei, header=None, sheet_name='CÁLCULO')
          except FileNotFoundError:
            printt(f"File '{filei}' not found.")
          except ValueError:
            printt("Worksheet 'CÁLCULO' not found. Trying alternative sheet names...")
            try:
              data_prof_in = pd.read_excel(input_folder_prof + filei, header=None, sheet_name='CALCULO')
            except ValueError:
              try:
                data_prof_in = pd.read_excel(input_folder_prof + filei, header=None, sheet_name='CALCULATION')
                opt = 1
              except ValueError:
                printt("Neither 'CALCULATIONS' nor 'CALCULO' sheet found.")

          if not data_in.empty:
            data_prof_LS = data_prof_in.loc[29].values          # Extract the LS data (I)
            if( opt != 1):
              line    = np.delete(data_prof_LS, [0,1,6,7,8]) # Extract the LS data (II)
            else:
              line    = np.delete(data_prof_LS, [0,1,2,7,8]) # Extract the LS data (II)
#
          file_line = filei+'      '+str(line[0])+' '+str(line[1])+' '+str(line[2])+' '+str(line[3])
#
          try:
            n0 = float(line[0])
            n1 = float(line[1])
            n2 = float(line[2])
            n3 = float(line[3])
            error = 'False'
          except ValueError:
            error = 'True'
#
          if ( (error == 'True') or
             ( np.isnan(n0) or np.isnan(n1) or np.isnan(n2) or np.isnan(n3)
               or (n0 < 0) or (n0 > 20) or (n1 < 0) or (n1 > 20)
               or (n2 < 0) or (n2 > 20) or (n3 < 0) or (n3 > 20) ) ):
#
            printt('error in input file : ' + str(file_line))
            if(len(datain_error_profs) == 0):
              datain_error_profs.append(filei)
            else:
              datain_error_profs.append(', ' + filei)
#
          else:
            printt(file_line)
            profs.append(profi)                  # Educator/Professor/Teacher/Instructor added
            data_prof.append([n0, n1, n2, n3])   # LS values added to data matrix
#
      datain_error_profs.append('.')

      data_prof = np.vstack(data_prof) # Stack the list data_prof
      Lprof = len(profs)
      printt(' ')
      printt('Educators : ')
#
#     Folder to save the output figures with the educators
      output_average_profile_prof = output_statistics_ls + '/' + prof_ref
      os.makedirs(output_average_profile_prof, exist_ok=True)
      printt('Output-average profile-prof folder : ' + output_average_profile_prof)

      output_participants_pc_prof = output_participants + '/' + prof_ref
      os.makedirs(output_participants_pc_prof, exist_ok=True)
      printt('Output-participants-PC_prof folder : ' + output_participants_pc_prof)
      printt('  ')

      for i in range(0, Lprof):
        printt(profs[i])
        printt(str(data_prof[i]))
#===================================================
#===================================================
#
#
#
#------------------------------------------------
# OPTIONAL INPUT FILE(S) WITH THE STUDENTS' MARKS
#------------------------------------------------
# The corresponding Excel files must be
# all contained in a folder placed in the input
# folder named marks, grades or qualitications.
# The corresponding data are named as a function of
# the names of the (i) Excel document, (ii) Excel sheet,
# and (iii) column in the Excel sheet.
# Each column is related to the marks obtained by the
# students in a different activity (laboratory, exams, etc.)
# and/or academic year.
# Thus, the results for different academic years and/or
# subjects can be saved all of them in a single Excel sheet,
# or in different documents.
# If the input folder contains folders named marks, grades and
# qualitications, only the data in the folder marks is considered;
# If the input folder contains folders named grades and
# qualitications, only the data in the folder grades is considered;
# Finally, if no folder named marks, grades or qualitications
# is inluded in the input folder,
# no calculation with the students' marks is conducted.

# Input file(s) with the students' marks/grades/qualifications
# (all of them must be in the same folder)

K_value = 0
vmarks  = ['marks', 'grades', 'qualifications']

pattern_K = re.compile(r"^marks_K(\d+)$")  # Pattern to match "marks_K<number>"

nfiles_marks = 0

# Check if one or more of the folders exists

# Check for marks_K* folders inside input_folder
for marks in vmarks:
#     printt('marks  :  ' + marks)
#     Check if the folder without the number of clusters (K) exists
      input_folder_mark = input_folder + marks + '/'
      if os.path.exists(input_folder_mark) and os.path.isdir(input_folder_mark):
        printt(f"The folder '{input_folder_mark}' exists.")
        nfiles_marks += 1
        input_folder_marks_ref = marks
        input_folder_marks     = input_folder_mark
#       The default K value is 2
        K_value = 2
      else:
        printt(f"The folder '{input_folder_mark}' does not exist.")

#     Check if the folder with the number of clusters exists, e.g., marks_K3 for 3 clusters
      for folder in os.listdir(input_folder):
       folder_path = os.path.join(input_folder, folder)
#      printt('folder  :  ' + folder)
       match = re.match(rf"^{marks}_K(\d+)$", folder)
       if match:
        K_value = int(match.group(1))
        print(f"Found file '{folder}' with K value: {K_value}")
        nfiles_marks += 1
        input_folder_marks_ref = folder
        input_folder_marks     = folder_path + '/'

printt('nfiles_marks :' + str(nfiles_marks))
printt( )

if ( nfiles_marks == 0):
  printt(' No file with the marks is provided, and then no clustering analysis is performed')
  K_value = 0
else:
  if ( nfiles_marks >= 2):
    printt(' Error: Conflict detected. Multiple folders associated with the marks of the educator(s). Please, change so that all the input files are contained in a single folder with one of these names:')
    for i in range(0, len(vmarks)):
      printt(vmarks[i])
    printt(' (If desired, change the number of the files, e.g., to marks_K3 to perform a K-means clustering with K=3 )')
    K_value = 0
  else:
    if ( nfiles_marks == 1):
      printt(' The marks of the students have been provided in the file')
      printt(' ' + input_folder_marks_ref)
      printt(' have been provided')
      printt(' ' + input_folder_marks)
#
# Marks folders (if K_value > 0)
if ( K_value == 0 ):
#
# K = 0
# No marks are provided, and the no clustering analysis is conducted.
#
# Check if the mark folders exist and, if not empty, remove them
  if os.path.exists(output_participants_ls_marks) and os.path.isdir(output_participants_ls_marks):
    if not os.listdir(output_participants_ls_marks):
      os.rmdir(output_participants_ls_marks)

  if os.path.exists(output_participants_pc_marks) and os.path.isdir(output_participants_pc_marks):
    if not os.listdir(output_participants_pc_marks):
      os.rmdir(output_participants_pc_marks)

else:
#
# K > 0
# Marks are provided.
# For K = 1 the whole group is studied; for K > 1, a K-means clustering is performed.

  if os.path.exists(output_participants_ls_marks) and os.path.isdir(output_participants_ls_marks):
    if not os.listdir(output_participants_ls_marks):
      os.rmdir(output_participants_ls_marks)

  output_participants_ls_marks = output_participants + '/' + input_folder_marks_ref
  os.makedirs(output_participants_ls_marks, exist_ok=True)
  printt('Output-LS-marks folder : ' + output_participants_ls_marks)
  printt('  ')

  if os.path.exists(output_participants_pc_marks) and os.path.isdir(output_participants_pc_marks):
    if not os.listdir(output_participants_pc_marks):
      os.rmdir(output_participants_pc_marks)

  output_participants_pc_marks = output_participants + '/' + input_folder_marks_ref
  os.makedirs(output_participants_pc_marks, exist_ok=True)
  printt('Output-PC-marks folder : ' + output_participants_pc_marks)
  printt('  ')

  excel_marks = []
  jref        = 0

# Run over all the Excel documents (and their corresponding sheets)
# to read the marks of the students
  for files in os.listdir(input_folder_marks):
#
#   xls and xlsx files
    if ((files.endswith('.xls')) or (files.endswith('.xlsx'))):
      excel_marks.append(files)
#
  excel_marks = sorted(excel_marks)
  printt("Excel files with the marks of the students : " + str(excel_marks))
  printt("\n")
#
  for filei in excel_marks:
    print()
    print('===============================')
    print('Excel document : ', filei)
    print('===============================')

    try:
      all_sheets = pd.read_excel(input_folder_marks + filei, sheet_name=None)  # Read all sheets into a dictionary
      for sheet_name, data_marks_in in all_sheets.items():
        print()
        print('-------------------------------')
        print(f"Sheet name     : {sheet_name}")
        print('-------------------------------')
        print("\n")
        print(data_marks_in)
        print('size : ', data_marks_in.shape)
        (Lstudents_tot, Lmarks) = data_marks_in.shape
        print('size : ', Lstudents_tot, Lmarks)
        print("\n")

        add_marks = 0

        if( data_marks_in.columns[0].strip().lower() == 'student' or data_marks_in.columns[0].strip().lower() == 'students'):
          print("The first column is named 'Student(s)' (case-insensitive).")
        else:
          print(f"The first column is named '{data_marks_in.columns[0]}', not 'Student'.")
          break

#       Save only the marks of those students whose LS questionnaires are provided
#       'student(s)' column (first column)
#
#       Check if any of the students whose LS marks are provided is also
#       contained in the list of students whose LS have been previouly given
        for j in range(0, L):
          studentj = os.path.splitext(students[j])[0]
          print('studentj : ', studentj)

          for i in range(Lstudents_tot):
            studenti = data_marks_in.iloc[i, 0]
            #print('  studenti :' + str(studenti))

            if(studenti == studentj):
              print('  studenti = studentj = ' + str(studenti))  # Print all elements in the column

#             Add new rows for the marks (if required)
              if(add_marks == 0):
                print('Enlarge data_marks')
                print('np.shape(data_marks) : ', np.shape(data_marks))
                print('data_marks : ', data_marks)
                data_marks = np.vstack([data_marks, np.full((Lmarks - 1, L + 3), np.nan)])
                print('np.shape(data_marks) : ', np.shape(data_marks))
                add_marks = 1

              for k in range(1, Lmarks):
                print('jref, k : ', jref, k)
#               Add to the first line of data_marks the name of the Excel document
                data_marks[k+jref, 0] = filei
#               Add to the second line of data_marks the name of the Excel sheet
                data_marks[k+jref, 1] = sheet_name
#               Add to the third line of data_marks the name of the column (mark)
                data_marks[k+jref, 2] = data_marks_in.columns[k]

              for k in range(1, Lmarks):
#               Add the marks in the corresponding positions
                markij = data_marks_in.iloc[i, k]
                print('    markij              : ', markij)
                if isinstance(markij, str):  # Check if it's a string
                  print(f"data_marks_in[{i}, {j}] is an (empty?) string.")
                  markij =  np.nan
                elif pd.isna(markij):  # Check for NaN or None (emptiness)
                  print(f"data_marks_in[{i}, {j}] is empty (NaN or None).")
                  markij =  np.nan

                data_marks[k + jref, j + 3] = markij
              break
        print(f"Sheet name     : {sheet_name}")
        if ( add_marks == 1 ):
          jref += Lmarks - 1
#        jref = jref + Lmarks - 1

    except ValueError:
        print('')

  for i in range(0, len(data_marks)):
    print('    student ' + str(data_marks[i]))
  print('end students')

  data_marks = np.vstack(data_marks)
  shape = np.shape(data_marks)
  printt('data_marks : ' + str(data_marks))
  printt('shape : ' + str(shape))

  if ( len(shape) > 0 and shape[1] > 1 ):
    Lmarks_activities = shape[0]
    Lmarks_students   = shape[1]
  else:
    Lmarks_activities = 0
    Lmarks_students   = 0
    Lmarks = 0
  printt(' Lmarks_activities : ' + str(Lmarks_activities))
  printt(' Lmarks_students   : ' + str(Lmarks_students))
  printt(' ')
#
#
#===================================================
#===================================================
#
#
#
#
#===================================================
# AVERAGE MEAN AND UNCERTAINTIES
#===================================================
printt('===========================================')
printt('AVERAGE MEAN AND UNCERTAINTIES...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
# Average mean
printt('-------------------------------------------')
printt(' Average mean...')
printt('-------------------------------------------')
shape = np.shape(data)
#
L = shape[0]
K = shape[1]
printt('  Input data size')
printt('    L :' + str(L))
printt('    K :' + str(K))
printt(' ')
#
# Size of the average value 0n the plots
size_mean = size_mean_ref / L
if ( size_mean < 1):
  size_mean = 4
else:
  if ( size_mean > 500):
    size_mean = 800
#
xamean = np.mean(xa)
xrmean = np.mean(xr)
xtmean = np.mean(xt)
xpmean = np.mean(xp)
#
xall     = np.concatenate( data )
xmeanall = np.mean(xall)
#
xmean = [xamean, xrmean, xtmean, xpmean]
#
printt('  Average mean')
printt('    mean(xa),    # : ' + str(xamean))
printt('    mean(xr),    # : ' + str(xrmean))
printt('    mean(xt),    # : ' + str(xtmean))
printt('    mean(xp),    # : ' + str(xpmean))
printt(' ')
printt('    mean(all),   # : ' + str(xmeanall))
printt('------------------------------------------')
printt(' Average mean done!')
printt('------------------------------------------')
printt(' ')
#
# Length of the eigenvectors when plotted
# (in 2D and 3D) as a function of the LSs
distmeanmax = 30. - np.min( xmean )
#
#---------------------------------------------------
# Uncertainties
printt('------------------------------------------')
printt('  Uncertainties...')
printt('------------------------------------------')
# Create 95% confidence interval for population mean weight.
#
confidence_interval_ACT=st.t.interval(confidence=0.95, df=len(data)-1, loc=np.mean(xa), scale=st.sem(xa))
uncert_abs_ACT=0.5*(confidence_interval_ACT[1]-confidence_interval_ACT[0])
#
confidence_interval_REF=st.t.interval(confidence=0.95, df=len(data)-1, loc=np.mean(xr), scale=st.sem(xr)) 
uncert_abs_REF=0.5*(confidence_interval_REF[1]-confidence_interval_REF[0])
#
confidence_interval_THEO=st.t.interval(confidence=0.95, df=len(data)-1, loc=np.mean(xt), scale=st.sem(xt)) 
uncert_abs_THEO=0.5*(confidence_interval_THEO[1]-confidence_interval_THEO[0])
#
confidence_interval_PRA=st.t.interval(confidence=0.95, df=len(data)-1, loc=np.mean(xp), scale=st.sem(xp)) 
uncert_abs_PRA=0.5*(confidence_interval_PRA[1]-confidence_interval_PRA[0])
#
dxmean = [uncert_abs_ACT, uncert_abs_REF, uncert_abs_THEO, uncert_abs_PRA]
#
confidence_interval_ALL=st.t.interval(confidence=0.95, df=len(xall)-1, loc=np.mean(xall), scale=st.sem(xall)) 
uncert_abs_ALL=0.5*(confidence_interval_ALL[1]-confidence_interval_ALL[0])
#
printt('  Uncertainties')
printt('    Uncert(xa),  # : ' + str(uncert_abs_ACT))
printt('    Uncert(xr),  # : ' + str(uncert_abs_REF))
printt('    Uncert(xt),  # : ' + str(uncert_abs_THEO))
printt('    Uncert(xp),  # : ' + str(uncert_abs_PRA))
printt(' ')
printt('    Uncert(all), # : ' + str(uncert_abs_ALL))
printt('------------------------------------------')
printt(' Uncertainties done!')
printt('------------------------------------------')
printt(' ')
#---------------------------------------------------
# Average profile
printt('------------------------------------------')
printt(' Average profile...')
printt('------------------------------------------')
# Length of axes and spacing between ticks.
xmin, xmax, ymin, ymax = -20, 20, -20, 20
ticks_frequency = 5
#
# Plot
fig, ax = plt.subplots(figsize = ( w_fig, h_fig )) 
#
# Coordenates (x, y) for the average mean
xs = [xrmean, 0, -xpmean,0]
ys = [0, xamean, 0, -xtmean]
ax.scatter(xs, ys)
#
# Connect the previous points
vertices_x=xs.copy()
vertices_x.append(xrmean)
vertices_y=ys.copy()
vertices_y.append(0)
#
# Average values
printt(' vertices_x : ' + str(vertices_x))
printt(' vertices_y : ' + str(vertices_y))
plt.plot(vertices_x, vertices_y, ls='--')
#
# Plot error barrs
printt(str(-xpmean-uncert_abs_PRA))
printt(str(xamean+uncert_abs_ACT))
printt(str(xpmean+uncert_abs_PRA))
#
#---------------------------------------------------
# Function that returns the value at x
# of a line that crosses points (x1, y1) and (x2, y2)
printt('-------------------------------------------')
printt('    Interpolating functions...')
printt('-------------------------------------------')
def line(x1, y1, x2, y2, x):
  m = (y2-y1)/(x2-x1)
  return y1 + m * ( x - x1 )  
#
ysupneg = line(-xpmean-uncert_abs_PRA, 0,
               0, xamean+uncert_abs_ACT,
               -xpmean+uncert_abs_PRA)
#
yinfneg = line(-xpmean-uncert_abs_PRA, 0,
               0, -xtmean-uncert_abs_THEO,
               -xpmean+uncert_abs_PRA)
#
ysuppos = line(xrmean+uncert_abs_REF, 0,
               0, xamean+uncert_abs_ACT,
               xrmean-uncert_abs_REF)
#
yinfpos = line(xrmean+uncert_abs_REF, 0,
               0, -xtmean-uncert_abs_THEO,
               xrmean-uncert_abs_REF)
#
printt('-------------------------------------------')
printt('    Interpolating functions done!')
printt('-------------------------------------------')
#---------------------------------------------------
printt(' ')
#
plt.fill_between([-xpmean-uncert_abs_PRA, -xpmean+uncert_abs_PRA, 0,
                   xrmean-uncert_abs_REF,  xrmean+uncert_abs_REF],
                 [0,  0, xamean-uncert_abs_ACT,  0, 0],
                 [0,  ysupneg, xamean+uncert_abs_ACT, ysuppos, 0],
                 facecolor='pink', alpha=0.5)
#
plt.fill_between([-xpmean-uncert_abs_PRA, -xpmean+uncert_abs_PRA, 0,
                   xrmean-uncert_abs_REF,  xrmean+uncert_abs_REF],
                 [0, yinfneg, -xtmean-uncert_abs_THEO, yinfpos, 0],
                 [0,  0, -xtmean+uncert_abs_THEO, 0, 0],
                 facecolor='pink', alpha=0.5)
#
x_error=[uncert_abs_REF,0,uncert_abs_PRA,0]
y_error=[0,uncert_abs_ACT,0,uncert_abs_THEO]
#
plt.errorbar(xs, ys, xerr = x_error, yerr=y_error, fmt='o', ecolor = 'red',color='blue',elinewidth = 5, capsize=10)
#
# Same scale for all axes
ax.set(xlim=(xmin-1, xmax+1), ylim=(ymin-1, ymax+1), aspect='equal')
#
# Set bottom and left spines as x and y axes of coordinate system
ax.spines['bottom'].set_position('zero')
ax.spines['left'].set_position('zero')
#
# Remove top and right spines
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
#
# Create 'x' and 'y' labels placed at the end of the axes
plt.text(0,  22, Label_LS[0], fontsize=labelsize, horizontalalignment='center')
plt.text( 22, 0, Label_LS[1], fontsize=labelsize, verticalalignment  ='center')
plt.text(0, -24, Label_LS[2], fontsize=labelsize, horizontalalignment='center')
plt.text(-34, 0, Label_LS[3], fontsize=labelsize, verticalalignment  ='center')
#
# Create custom major ticks to determine position of tick labels
x_ticks = np.arange(xmin, xmax+1, ticks_frequency)
y_ticks = np.arange(ymin, ymax+1, ticks_frequency)
ax.set_xticks(x_ticks[x_ticks != 0])
ax.set_yticks(y_ticks[y_ticks != 0])
#
# Rename negative parts of the axes
labels = [item.get_text() for item in ax.get_xticklabels()]
printt('labels : ' + str(labels))
labels[0] = 20
labels[1] = 15
labels[2] = 10
labels[3] = 5
ax.set_xticklabels(labels)
#
labels = [item.get_text() for item in ax.get_yticklabels()]
printt('labels : ' + str(labels))
labels[0] = 20
labels[1] = 15
labels[2] = 10
labels[3] = 5
ax.set_yticklabels(labels)
#
ax.set_xticks(np.arange(xmin, xmax+1), minor=True)
ax.set_yticks(np.arange(ymin, ymax+1), minor=True)
#
ax.tick_params(axis='both', which='both', labelsize = ticksize)
#
#plt.show()
filename_average_profile=output_statistics_ls + '/Fig_averageprofile.png'
printt('Saving ' + filename_average_profile)
plt.savefig(filename_average_profile)
printt(' ')
printt('------------------------------------------')
printt(' Average profile done!')
printt('------------------------------------------')
printt(' ')

# Average profile
printt('------------------------------------------')
printt(' Average profile '+ professor + '...')
printt('------------------------------------------')
#
filename_average_profile_prof = [ ]
for i in range(0, Lprof):
# Plot
  fig, ax = plt.subplots(figsize = ( w_fig, h_fig ))
#
# Coordenates (x, y) for the average mean
  ax.scatter(xs, ys)
  plt.plot(vertices_x, vertices_y, ls='--')

# Coordinates (x, y) for the professors
  xs_prof = [data_prof[i, 1], 0, -data_prof[i, 2],0]
  ys_prof = [0, data_prof[i,0], 0, -data_prof[i, 3]]
  ax.scatter(xs_prof, ys_prof, c='g', s=2)
#
# Connect the previous points
  vertices_x_prof=xs_prof.copy()
  vertices_x_prof.append(data_prof[i, 1])
  vertices_y_prof=ys_prof.copy()
  vertices_y_prof.append(0)
#
# Average values
  printt(' vertices_x_prof : ' + str(vertices_x_prof))
  printt(' vertices_y_prof : ' + str(vertices_y_prof))
  plt.plot(vertices_x_prof, vertices_y_prof, ls='-', c = 'g', lw=2)
#
# Plot error barrs
  printt(str(-xpmean-uncert_abs_PRA))
  printt(str(xamean+uncert_abs_ACT))
  printt(str(xpmean+uncert_abs_PRA))
#
#---------------------------------------------------
  printt(' ')
#
  plt.fill_between([-xpmean-uncert_abs_PRA, -xpmean+uncert_abs_PRA, 0,
                   xrmean-uncert_abs_REF,  xrmean+uncert_abs_REF],
                   [0,  0, xamean-uncert_abs_ACT,  0, 0],
                   [0,  ysupneg, xamean+uncert_abs_ACT, ysuppos, 0],
                   facecolor='pink', alpha=0.5)
#
  plt.fill_between([-xpmean-uncert_abs_PRA, -xpmean+uncert_abs_PRA, 0,
                     xrmean-uncert_abs_REF,  xrmean+uncert_abs_REF],
                   [0, yinfneg, -xtmean-uncert_abs_THEO, yinfpos, 0],
                   [0,  0, -xtmean+uncert_abs_THEO, 0, 0],
                   facecolor='pink', alpha=0.5)
#
  x_error=[uncert_abs_REF,0,uncert_abs_PRA,0]
  y_error=[0,uncert_abs_ACT,0,uncert_abs_THEO]
#
  plt.errorbar(xs, ys, xerr = x_error, yerr=y_error, fmt='o', ecolor = 'red',color='blue',elinewidth = 5, capsize=10)
#
# Same scale for all axes
  ax.set(xlim=(xmin-1, xmax+1), ylim=(ymin-1, ymax+1), aspect='equal')
#
# Set bottom and left spines as x and y axes of coordinate system
  ax.spines['bottom'].set_position('zero')
  ax.spines['left'].set_position('zero')
#
# Remove top and right spines
  ax.spines['top'].set_visible(False)
  ax.spines['right'].set_visible(False)
#
# Create 'x' and 'y' labels placed at the end of the axes
  plt.text(0,  22, Label_LS[0], fontsize=labelsize, horizontalalignment='center')
  plt.text( 22, 0, Label_LS[1], fontsize=labelsize, verticalalignment  ='center')
  plt.text(0, -24, Label_LS[2], fontsize=labelsize, horizontalalignment='center')
  plt.text(-34, 0, Label_LS[3], fontsize=labelsize, verticalalignment  ='center')
#
# Create custom major ticks to determine position of tick labels
  x_ticks = np.arange(xmin, xmax+1, ticks_frequency)
  y_ticks = np.arange(ymin, ymax+1, ticks_frequency)
  ax.set_xticks(x_ticks[x_ticks != 0])
  ax.set_yticks(y_ticks[y_ticks != 0])
#
# Rename negative parts of the axes
  labels = [item.get_text() for item in ax.get_xticklabels()]
  printt('labels : ' + str(labels))
  labels[0] = 20
  labels[1] = 15
  labels[2] = 10
  labels[3] = 5
  ax.set_xticklabels(labels)
#
  labels = [item.get_text() for item in ax.get_yticklabels()]
  printt('labels : ' + str(labels))
  labels[0] = 20
  labels[1] = 15
  labels[2] = 10
  labels[3] = 5
  ax.set_yticklabels(labels)
#
  ax.set_xticks(np.arange(xmin, xmax+1), minor=True)
  ax.set_yticks(np.arange(ymin, ymax+1), minor=True)
#
  ax.tick_params(axis='both', which='both', labelsize = ticksize)
#
#plt.show()
  filename_average_profile_prof.append( output_average_profile_prof + '/Fig_averageprofile_prof_' + profs[i] + '.png' )
  printt('Saving ' + filename_average_profile_prof[i])
  plt.savefig(filename_average_profile_prof[i])
  printt(' ')
  printt('------------------------------------------')
  printt(' Average profile with '+ professor + profs[i] + ' done!')
  printt('------------------------------------------')
printt(' ')
printt('===========================================')
printt('AVERAGE MEAN AND UNCERTAINTIES DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# TENDENCIES TO THE DIFFERENT LSs
#===================================================
printt('===========================================')
printt('TENDENCIES...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
# Average mean
printt('-------------------------------------------')
printt(' Tendencies in %...')
printt('-------------------------------------------')
# Matrix with the tendencies (very low, low,
# moderate, high and very high) towards the
# different LSs.
# Each line is associated with a different LS.
# Each column corresponds to a different
# tendency.
tendency_matrix     = np.zeros((K, 5))
tendency_matrix_all = [['' for _ in range(K)] for _ in range(L)]
tendency_vector  = ['', '', '', '', '']
#
for i in range(0,L):
  printt(str(i))
  tendency_vector[0] = scatter_tendency('Activist',   xa[i])
  tendency_vector[1] = scatter_tendency('Reflector',  xr[i])
  tendency_vector[2] = scatter_tendency('Theorist',   xt[i])
  tendency_vector[3] = scatter_tendency('Pragmatist', xp[i])
#
  tendency_matrix_all[i][0] = tendency_vector[0]
  tendency_matrix_all[i][1] = tendency_vector[1]
  tendency_matrix_all[i][2] = tendency_vector[2]
  tendency_matrix_all[i][3] = tendency_vector[3]
#
  for j in range(0,K):
    if ( tendency_vector[j] == 'vl' ) :
      tendency_matrix[j,0] = tendency_matrix[j,0] + 1
    else:
      if ( tendency_vector[j] == 'l' ) :
        tendency_matrix[j,1] = tendency_matrix[j,1] + 1
      else:
        if ( tendency_vector[j] == 'm' ) :
          tendency_matrix[j,2] = tendency_matrix[j,2] + 1
        else:
          if ( tendency_vector[j] == 'h' ) :
            tendency_matrix[j,3] = tendency_matrix[j,3] + 1
          else:
            if ( tendency_vector[j] == 'vh' ) :
              tendency_matrix[j,4] = tendency_matrix[j,4] + 1
#
tendency_matrix_percentage = tendency_matrix * 100 / L
#
for i in range(0,K):
  printt('  ' + Label_LS[i] + ' learning style')
  printt('    Tendency      No.     %')
  for j in range(0,len(Label_tendencies)):
    printt('   ' + Label_tendencies_print[j] + '   ' + str(int(tendency_matrix[i, j])) + '   ' + str(round(tendency_matrix_percentage[i, j], 1)))
  printt(' ')
#
printt('Learning Style ' + Label_tendencies_print[0] + Label_tendencies_print[1] +
                         Label_tendencies_print[2]   +
                         Label_tendencies_print[3]   + Label_tendencies_print[4])
printt('                No.  %    No.  %    No.  %    No.  %    No.  %')
for i in range(0,K):
  printt('  ' + Label_LS_print[i] + ' ' + str(int(tendency_matrix[i, 0])) + str(round(tendency_matrix_percentage[i, 0], 1))  + '   ' + str(int(tendency_matrix[i, 1])) + str(round(tendency_matrix_percentage[i, 1], 1)) + '   ' + str(int(tendency_matrix[i, 2])) + str(round(tendency_matrix_percentage[i, 2], 1)) + '   ' + str(int(tendency_matrix[i, 3])) + str(round(tendency_matrix_percentage[i, 3], 1)) + '   ' + str(int(tendency_matrix[i, 4])) + str(round(tendency_matrix_percentage[i, 4], 1)) )
printt(' ')
#
Ntot_tendency    = np.zeros(len(Label_tendencies))
percent_tendency = np.zeros(len(Label_tendencies))
#
for j in range(0,len(Label_tendencies)):
  for i in range(0,K):
    Ntot_tendency[j]    = Ntot_tendency[j] + tendency_matrix[i, j]
    percent_tendency[j] = percent_tendency[j] + tendency_matrix[i, j] * tendency_matrix_percentage[i, j]  
#
  percent_tendency[j] = percent_tendency[j] / Ntot_tendency[j]
#
printt('   Average    ' + str(int(Ntot_tendency[0])) + str(round(percent_tendency[0], 1)) + '   ' + str(int(Ntot_tendency[1])) + str(round(percent_tendency[1], 1)) + '   ' + str(int(Ntot_tendency[2])) + str(round(percent_tendency[2], 1)) + '   ' + str(int(Ntot_tendency[3])) + str(round(percent_tendency[3], 1)) + '   ' + str(int(Ntot_tendency[4])) + str(round(percent_tendency[4], 1)))
#
fig, ax = plt.subplots(figsize = ( w_fig, h_fig )) #, layout='constrained')
#
ymax = np.max(tendency_matrix_percentage)
#
# Fill areas below the curve
for j in range(0,K):
  for i in range(0,len(Label_tendencies)):
#
    percentage = tendency_matrix_percentage[j,i]
    xmin_percentage = -2+i+(j-2)*0.2
    xmax_percentage = xmin_percentage + 0.2
    plt.fill_between([xmin_percentage, xmax_percentage], [percentage, percentage], facecolor = LS_color[j], alpha = LS_alpha3)  
#
# Tick parameters
  ax.set_xticks([-2, -1, 0, 1, 2], Label_tendencies)
  ax.tick_params(axis='both', which='major', labelsize = ticksize_quartile)
#
# Labels
  ax.set_xlabel('Tendency',   fontsize = labelsize_quartile)
  ax.set_ylabel('% Students', fontsize = labelsize_quartile)
#
ymax = ymax*1.05
ax.set_xlim( [ -2.5, 2.5] )
ax.set_ylim( [ 0.05, ymax] )
#
filename_tendencies=output_statistics_ls + '/Fig_tendencies.png'
printt('Saving ' + filename_tendencies)
plt.savefig(filename_tendencies)
printt('------------------------------------------')
printt(' Tendencies done!')
printt('------------------------------------------')
printt(' ')
#
printt('===========================================')
printt('TENDENCIES DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# AFINITY
#===================================================
printt('===========================================')
printt('AFFINITY...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Affinity in %...')
printt('-------------------------------------------')
#
affinity  = np.zeros(K)
#
for i in range(0,K):
  for j in range(2,len(Label_tendencies)):
    affinity[i] = affinity[i] + tendency_matrix_percentage[i,j] 
#
fig, ax = plt.subplots(figsize = ( w_fig, h_fig ))
#
ymax = np.max(affinity)
#
# Fill areas below the curve
printt( '   Learning style     %')
for j in range(0,K):
  printt( '  ' + Label_LS_print[j] + '    ' + str(round(affinity[j], 1)))
  plt.fill_between([j-0.4, j+0.4], [affinity[j], affinity[j]], facecolor = LS_color[j], alpha = affinity_alpha)  
#
# Tick parameters
  x = np.arange(len(Label_LS))
  ax.set_xticks(x, Label_LS)
  ax.tick_params(axis='both', which='major', labelsize = ticksize_quartile)
#
# Labels
  ax.set_xlabel('LS',         fontsize = labelsize_quartile)
  ax.set_ylabel('% Students', fontsize = labelsize_quartile)
#
printt( '   Average        ' + str(round(percent_tendency[2]+percent_tendency[3]+percent_tendency[4], 1)))
#
ymax = ymax*1.05
ax.set_xlim( [ -0.5, 3.5] )
ax.set_ylim( [ 0.05, ymax] )
#
printt(' ')
filename_affinity=output_statistics_ls + '/Fig_affinity.png'
printt('Saving ' + filename_affinity)
plt.savefig(filename_affinity)
printt('------------------------------------------')
printt(' Tendencies done!')
printt('------------------------------------------')
printt(' ') 
#
printt('===========================================')
printt('AFFINITY DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# PROBABILITIES IN THE ORIGINAL LS BASIS SET
#===================================================
#
printt('-------------------------------------------')
printt(' Probability of each student on the LSs...')
printt('-------------------------------------------')
printt(' ')
# Calculate the probability to have a
# certain LS for each student
probLS = np.zeros((L,K))
#
for i in range(0,L):
  normLSj2 = 0
  for j in range(0,K):
    normLSj2 = normLSj2 + data[i,j]*data[i,j]
#
# Compute the probability
# (it is set to zero if the vector is zero)
  for j in range(0,K):
    if ( normLSj2 > 0):
      probLS[i, j] = 100 * data[i,j] * data[i,j] / normLSj2
#
printt('   data   : ' + str(data))
printt(' ')
printt('   probLS : ' + str(probLS))
printt(' ')
#
# Mean values and uncertainties of the probabilities
probLSmean = np.array([np.mean(probLS[:,0]), np.mean(probLS[:,1]), np.mean(probLS[:,2]), np.mean(probLS[:,3])])
#
uncert_abs_probLS = [0, 0, 0, 0]
for j in range(0,K):
  probj = probLS[:,j]
  confidence_interval_probLS=st.t.interval(confidence=0.95, df=len(probj)-1, loc=probLSmean[j], scale=st.sem(probj)) 
  uncert_abs_probLS[j]=0.5*(confidence_interval_probLS[1]-confidence_interval_probLS[0])
#
printt('  The mean values of the corresponding probabilities')
printt('  do not nullify (as all of them are positive or zero)')
printt('    probLSmean     : ' + str(probLSmean))
printt('    Uncert(probLS) : ' + str(uncert_abs_probLS))
printt(' ')
#
printt(' Probability of each student on the LSs done!')
printt('-------------------------------------------')
printt(' ')
printt(' ')
printt('===========================================')
printt('PROBABILITIES IN ORIGINAL BASIS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# STATISTICAL ANALYSIS IN THE LS ORIGINAL BASIS
#===================================================
printt('===========================================')
printt('STATISTICAL ANALYSIS LS ORIGINAL BASIS...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Analysis of the cumulative probability')
printt(' distribution (W(prob)) of the probabilities')
printt(' of the different LSs (plot and fitting')
printt(' using a Weibull distribution)...')
printt('-------------------------------------------')
printt(' ')
printt('-------------------------------------------')
printt(' W(LS)...')  
printt('-------------------------------------------')
xWeib   = np.arange(0,20.1, 0.001)
#
ymin  = 0
ymax  = 1.01
ymaxx = 10000000000
#
parameters_Weibull_LS20 = np.zeros((K,2))
#
for ils in range(0, K):
  printt('   LS 0-20 :' + Label_LS[ils])
  x = data[:, ils]
  x = np.sort(x)
#
  y = np.zeros(L)
  y[0] = 0
  dy = 1 / float(L-1)
  for i in range(1,L):
    y[i] = y[i-1] + dy
#
# Fitting using Weibull distribution
  param = curve_fit(Wweibull, x, y, [xmean[ils], 1], maxfev = max_iter_Weib)
  printt('    ' + str(param))
  [alpha_Weibull, k_Weibull] = param[0]
  parameters_Weibull_LS20[ils] = [alpha_Weibull, k_Weibull]
  printt('    alpha_Weibull : ' + str(alpha_Weibull))
  printt('    k_Weibull     : ' + str(k_Weibull))  
  printt(' ')
#
# Create a figure with 4 panels,
# each one related to a different LS
fig, axs = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), sharex = True, sharey = True, layout='constrained')  
#
panels = [[0, 0], [0, 1], [1, 0], [1, 1]]
#
for ils in range(0, K):
  printt('   LS 0-20 :' + Label_LS[ils])
  x = data[:, ils]
  x = np.sort(x)
#
  y = np.zeros(L)
  y[0] = 0
  dy = 1 / float(L-1)
  for i in range(1,L):
    y[i] = y[i-1] + dy
#
# Staircase
  x2 = np.zeros(2*L+2)
  y2 = np.zeros(2*L+2)
  x2[0]  =   0
  y2[0]  =   0
  x2[2*L+1] = x[L-1]
  y2[2*L+1] = 1
#
# Fitting using Weibull distribution
  [alpha_Weibull, k_Weibull] = parameters_Weibull_LS20[ils]
  printt('    alpha_Weibull : ' + str(alpha_Weibull))
  printt('    k_Weibull     : ' + str(k_Weibull))  
  printt(' ')
#
  for i in range(0,L):
    x2[2*i+1] = x[i]
    x2[2*i+2] = x[i]

    y2[2*i+1] = y[i]
    y2[2*i+2] = y[i] + dy
  y2[2*L] = 1    
#
# Find the points close to 5, 10, and 15
# in order to correcly color the
# LS that lie in the ranges [0, 5), 
# (10, 15), and (15,20]
  ic   = [-1,  -1, -1]
  xc   = [ 0,  5, 10]
  xref = [ 5, 10, 15]
#
  for i in range(0,len(ic)):
    for j in range(0, L):
      if( x[j] > xc[i] and x[j] < xref[i] ):
        ic[i] = j
        xc[i] = x[j]
#
  axs[panels[ils][0], panels[ils][1]].set_xlim( [ 0, 20] )
  axs[panels[ils][0], panels[ils][1]].set_ylim( [ ymin, ymax] )
#
  axs[panels[ils][0], panels[ils][1]].tick_params(axis='both', which='major', labelsize = ticksize_PC)
#
  axs[panels[ils][0], panels[ils][1]].plot(x2, y2,                                 lw = PC_width,   color = PC_color[ils], ls = PC_line[ils])
  axs[panels[ils][0], panels[ils][1]].plot(xWeib,Wweibull(xWeib, alpha_Weibull, k_Weibull) , lw = PC_W_width, color = PC_color[ils], ls = PC_line_bis[ils])  
#
# Plot the average value as a vertical line
# we do not plot the uncertainties as they can be very large
  axs[panels[ils][0], panels[ils][1]].plot([xmean[ils] , xmean[ils] ], [ymin, ymaxx], lw = LS_mean_width, color = LS_color[ils], ls = LS_mean_line)
  axs[panels[ils][0], panels[ils][1]].fill_between([xmean[ils]-dxmean[ils], xmean[ils]+dxmean[ils]], [ymin, ymaxx], color = LS_color[ils], alpha = LS_mean_alpha )
#      
# Text label
  axs[panels[ils][0], panels[ils][1]].text(20*0.05, 0.9*ymax, Label_LS[ils],
        fontsize = labelsize,
        horizontalalignment='left',
        verticalalignment='top')
#
axs[1, 0].tick_params(axis='both', which='major', labelsize=ticksize)
#
axs[1, 0].set_xlabel('Contribution', fontsize=labelsize)
axs[1, 0].set_ylabel('W [Contribution]', fontsize=labelsize)
#
filename_statistics_ls_w = output_statistics_ls + '/Fig_LS20_W2.png'
printt('Saving ' + filename_statistics_ls_w)
plt.savefig(filename_statistics_ls_w)
#
printt('-------------------------------------------')
printt(' W(LS) done!')  
printt('-------------------------------------------')
printt(' ')
#
printt('-------------------------------------------')
printt(' W(probLS)...')  
printt('-------------------------------------------')
xWeib   = np.arange(0, 1.1, 0.001)
parameters_Weibull_LS = np.zeros((K,2))
#
#
for ils in range(0, K):
  printt('   LS :' + Label_LS[ils])
  x = probLS[:, ils] / 100
  x = np.sort(x)
#
  y = np.zeros(L)
  y[0] = 0
  dy = 1 / float(L-1)
  for i in range(1,L):
    y[i] = y[i-1] + dy
#
# Fitting using Weibull distribution
  param = curve_fit(Wweibull, x, y, [probLSmean[ils], 1], maxfev = max_iter_Weib)
  printt('    ' + str(param))
  [alpha_Weibull, k_Weibull] = param[0]
  parameters_Weibull_LS[ils] = [alpha_Weibull, k_Weibull]
  printt('    alpha_Weibull : ' + str(alpha_Weibull))
  printt('    k_Weibull : ' + str(k_Weibull))
  printt(' ')
#
printt('-------------------------------------------')
printt(' W(probLS) done!')  
printt('-------------------------------------------')
printt(' ')
#
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Probabilities of the LSs...')
printt('-------------------------------------------')
#
xWeib   = np.arange(0, 120, 0.005)
ymaxtot = 0
ymin  = 0
ymax  = 1.01
ymaxx   = 10000000000
#
# Create a figure with 4 panels,
# each one related to a different LS
fig, axs = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), sharex = True, sharey = True, layout='constrained')  
#
panels = [[0, 0], [0, 1], [1, 0], [1, 1]]
#
for ils in range(0, K):
  printt('   LS 0-20 :' + Label_LS[ils])
#
  x = data[:, ils]
  x = np.sort(x)
#
  [alpha_Weibull, k_Weibull] = parameters_Weibull_LS20[ils]
  printt('    alpha_Weibull : ' + str(alpha_Weibull))
  printt('    k_Weibull : ' + str(k_Weibull))
  printt(' ')
#
# Plot the average value as a vertical line
# we do not plot the uncertainties as they can be very large
  axs[panels[ils][0], panels[ils][1]].plot([xmean[ils], xmean[ils]], [ymin, ymaxx], lw = LS_mean_width, color = LS_color[ils], ls = LS_mean_line)
  axs[panels[ils][0], panels[ils][1]].fill_between([xmean[ils]-dxmean[ils], xmean[ils]+dxmean[ils]],  [ymin, ymaxx], color = LS_color[ils], alpha = LS_mean_alpha )
#
  y, x, _= axs[panels[ils][0], panels[ils][1]].hist(x, bins=20, color = LS_color[ils], alpha = LS_hist_alpha)
  ymax = y.max()
#
  ymax = y.max()
  if ymax > ymaxtot:
    ymaxtot = ymax
#
# Probability distribution for Weibull function (scaled)
  factor = ymax / Pweibull(alpha_Weibull, alpha_Weibull, k_Weibull)
  axs[panels[ils][0], panels[ils][1]].plot(xWeib,factor*Pweibull(xWeib, alpha_Weibull, k_Weibull) , lw = LS_W_width,   color = LS_color[ils], ls = LS_line[0])  
#
  axs[panels[ils][0], panels[ils][1]].tick_params(axis='both', which='major', labelsize=ticksize)
#
  axs[panels[ils][0], panels[ils][1]].set_xlim( [ 0, 20] )
  axs[panels[ils][0], panels[ils][1]].set_ylim( [ 0, ymaxtot] )   
#
# Text label
  axs[panels[ils][0], panels[ils][1]].text(20*0.05, 0.9*ymaxtot, Label_LS[ils],
        fontsize = labelsize,
        horizontalalignment='left',
        verticalalignment='top')
#
axs[1, 0].tick_params(axis='both', which='major', labelsize=ticksize)
#
axs[1, 0].set_xlabel('Contribution', fontsize=labelsize)
axs[1, 0].set_ylabel('P [Contribution]', fontsize=labelsize)
#
filename_statistics_ls = output_statistics_ls + '/Fig_LS20_P2.png'
printt('Saving ' + filename_statistics_ls)
plt.savefig(filename_statistics_ls)
printt('-------------------------------------------')
printt(' Probabilities of the LSs done!')
printt('-------------------------------------------')
printt(' ')
printt('===========================================')
printt('STATISTICAL ANALYSIS LS ORIGINAL BASIS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# COVARIANCE MATRIX
#===================================================
printt('===========================================')
printt('COVARIANCE MATRIX...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Differences between input data and mean...')
printt('-------------------------------------------')
printt(' ')
# Difference between the input data and the mean
ddata = np.zeros((L,K))
for i in range(0,L):
  for j in range(0,K):
    ddata[i,j] = data[i,j] - xmean[j]
#
printt('------------------------------------------')
printt(' Differences done!')
printt('------------------------------------------')
printt(' ')
#---------------------------------------------------
# Covariance matrix
printt('------------------------------------------')
printt(' Construction of the covariance matrix...')
printt('------------------------------------------')
#
covX = np.cov([xa, xr, xt, xp])
#
printt('  covX = ' + str(covX))
printt(' ')
#
trace_covX = 0
for i in range(0,K):
  trace_covX = trace_covX + covX[i,i]
#
printt('  tr(covX) = ' + str(trace_covX))
printt(' ')
printt('------------------------------------------')
printt(' Covariance matrix constructed!')
printt('------------------------------------------')
printt(' ')
#---------------------------------------------------
printt('------------------------------------------')
printt(' Eigensystem...')
printt('------------------------------------------')
eigenValues, eigenVectors = eig(covX)
printt('------------------------------------------')
printt(' done!')
printt('------------------------------------------')
printt(' ')
printt('------------------------------------------')
printt(' Sorting the eigenvalues...')
printt('------------------------------------------')
idx = eigenValues.argsort()[::-1]  
printt('  Eigenvalues ordering')
printt('   idx    : ' + str(idx))
printt(' ')
eigenValues = eigenValues[idx]
eigenVectors = eigenVectors[:,idx]
#
# The eigenvectores are saved in the columns 
# of the matrix of the eigenVectors
printt('  EigenValues  : ' + str(eigenValues))
printt(' ')
disp0 = 100/trace_covX *  eigenValues[0]
disp1 = 100/trace_covX *  eigenValues[1]
disp2 = 100/trace_covX *  eigenValues[2]
disp3 = 100/trace_covX *  eigenValues[3]
printt('  % Dispersion (PC0) : ' + str(disp0))
printt('  % Dispersion (PC1) : ' + str(disp1) + ' % Dispersion (PC0+PC1)         : ' + str(disp0 + disp1))
printt('  % Dispersion (PC2) : ' + str(disp2) + ' % Dispersion (PC0+PC1+PC2)     : ' + str(disp0 + disp1 + disp2))
printt('  % Dispersion (PC3) : ' + str(disp3) + ' % Dispersion (PC0+PC1+PC2+PC3) : ' + str(disp0 + disp1 + disp2 + disp3))
printt(' ')
printt('  EigenVectors : ' + str(eigenVectors))
printt(' ')
printt('------------------------------------------')
printt(' Sorting the eigenvalues done!')
printt('------------------------------------------')
printt(' ')
printt('------------------------------------------')
printt(' EigenVectors normalization (norm, max)...')
printt('------------------------------------------')
eigenVectors_max1       = np.zeros((K,K))
eigenVectors_percentage = np.zeros((K,K))
#
for j in range(0,K):
  norm=0.0
  vj = eigenVectors[:,j]
#
  vjmax = np.max(np.abs(vj))
  vjmin = np.min(vj)  
  if vjmax == -vjmin:
    vjmax = vjmin
  eigenVectors_max1[:,j] = vj/vjmax
#
  for l in range(0,K):
    norm = norm + vj[l] * vj[l]
  norm = np.sqrt(norm)    
#
  eigenVectors[:,j] = vj/norm
  eigenVectors_percentage[:,j] = eigenVectors[:,j]*eigenVectors[:,j]*100
#
printt('  (Renormalized) Sorted eigenVectors (norm=1): ' + str(eigenVectors))
printt(' ')
printt('  Sorted eigenVectors (maximum = 1): ' + str(eigenVectors_max1))
printt(' ')
printt('  Sorted eigenVectors (%)          : ' + str(eigenVectors_percentage))
printt(' ')
printt('------------------------------------------')
printt(' EigenVectors normalization done!')
printt('------------------------------------------')

printt(' ')
printt('------------------------------------------')
printt(' Correspondence between eigenvectors-LS...')
printt('------------------------------------------')
# We only plot the two/three principal directions.
# Thus, we redifine the learning_pairs and
# learning_trios vectors so that their first
# component gives the two LSs of interest
#
lsref = [0, 0, 0]
for j in range(0,3):
  Vector = np.abs(eigenVectors[:,j])
  Vectormax = Vector[0]
  for i in range(1,K):
    if (Vector[i] > Vectormax):
      lsref[j]  = i
      Vectormax = Vector[i]
#
# If two of the terms in lsref are equal
# we simply take lsref equal to the
# first element in learning_trios
if( lsref[0] == lsref[1] or lsref[0] == lsref[2] or lsref[1] == lsref[2] ):
  lsref = learning_trios[0]
# The points in the plots are shown as diamonds instead as circles
#  
printt(' lsref : ' + str(lsref))
printt(' ')
#
for i in range(0, len(learning_pairs)):
  if(([lsref[0], lsref[1]] == learning_pairs[i]) or 
      [lsref[1], lsref[0]] == learning_pairs[i]):
    ipairs = i
#    
for i in range(0, len(learning_trios)):
  if(([lsref[0], lsref[1], lsref[2]] == learning_trios[i]) or 
     ([lsref[0], lsref[2], lsref[1]] == learning_trios[i]) or
     ([lsref[1], lsref[0], lsref[2]] == learning_trios[i]) or 
     ([lsref[1], lsref[2], lsref[0]] == learning_trios[i]) or
     ([lsref[2], lsref[0], lsref[1]] == learning_trios[i]) or 
     ([lsref[2], lsref[1], lsref[0]] == learning_trios[i])):
    itrios = i
#
printt('  Pairs of learning styles')
printt('    ipairs                 : ' + str(ipairs))
printt('    learning_pairs[ipairs] : ' + str(learning_pairs[ipairs]))
printt(' ')
printt('  Trios of learning styles')
printt('    itrios                 : ' + str(itrios))
printt('    learning_trios[itrios] : ' + str(learning_trios[itrios]))
printt(' ')
printt('------------------------------------------')
printt(' Correspondence eigenvectors-LS done!')
printt('------------------------------------------')
printt(' ')
printt('===========================================')
printt('COVARIANCE MATRIX DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# PRINCIPAL COMPONENT ANALYSIS (PCA).
# PROJECTIONS ON THE EIGENVECTORS OF THE 
# COVARIANCE MATRIX, AND STATISTICAL ANALYSIS
#===================================================
printt('===========================================')
printt('PRINCIPAL COMPONENT ANALYSIS...')
printt('===========================================')
printt(' ')
printt('===========================================')
printt('PROJECTION ON THE EIGENVECTORS...')
printt('===========================================')
#---------------------------------------------------
# Projections on the principal directions of the 
# maximal "pure" LSs given by vectors with one 
# component equal to 20, and the rest to 0
# (original LS in lines 0 to 3 of the
# matrix LS; the projections are given in lines 0 to 3
# of the matrix projectLS), the origin (line 4) and the
# maximal possible state (20, 20, 20, 20) (line 5).
# The j-th of the matrix projectLS gives the projection 
# on the j-th eigenvector
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Projections on the principal directions')
printt(' of some reference states...')
printt('-------------------------------------------')
#
LS        = np.zeros((K+2,K))
projectLS = np.zeros((K+2,K))
#
# Maximal "pure" LS
for i in range(0,K):
  LS[i,i] = 20
#
# Maximal possible state (20, 20, 20, 20)
for j in range(0,K):
  LS[5,j] = 20
#
printt('  LS associated with the pure states,')
printt('    the origin and the maximal state')
printt('                    (20, 20, 20, 20)')
printt(' ')
printt('    LS        : '  + str(LS))
printt(' ')
#
# Difference with the mean value
for i in range(0,K+2):
  for j in range(0,K):
    LS[i,j] = LS[i,j] - xmean[j]
#
for i  in range(0,K+2):  
  for j  in range(0,K):
    vj = eigenVectors[:,j]
#
    for l  in range(0,K):
      projectLS[i,j]    = projectLS[i,j] + vj[l]*LS[i,l]  
printt('    LS-xmean  : ' + str(LS))
printt(' ')
printt('    projectLS : ' + str(projectLS))
printt(' ')
printt('-------------------------------------------')
printt(' Projections of some reference states done!')
printt('-------------------------------------------')
printt(' ')
#---------------------------------------------------
# Projection of the (ith) data on the jth eigenvector
printt('-------------------------------------------')
printt(' Projections of data on eigenvectors...')
printt('-------------------------------------------')
printt(' ')
proj = np.zeros((L,K))
prob = np.zeros((L,K))
#
for i  in range(0,L):
  norm_ddata = 0.0
  for j  in range(0,K):
    norm_ddata = norm_ddata + ddata[i,j]*ddata[i,j]
  norm_ddata = np.sqrt( norm_ddata )
#
  for j  in range(0,K):
    vj = eigenVectors[:,j]
#
    for l  in range(0,K):
      proj[i,j] = proj[i,j] + vj[l]*ddata[i,l]  
#
    cos = proj[i,j]/norm_ddata
    prob[i,j]  = cos * cos * 100
printt('-------------------------------------------')
printt(' Projections of data on eigenvectors done!')
printt('-------------------------------------------')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Mean values for the projections')
printt(' and probabilities...')
printt('-------------------------------------------')
printt(' ')
# Mean values of the projections (they must zero)
projmean = np.array([np.mean(proj[:,0]), np.mean(proj[:,1]), np.mean(proj[:,2]), np.mean(proj[:,3])])
printt('  The mean values of the projections must nullify')
printt('    projmean : ' + str(projmean))
printt(' ')
printt('  Uncertainty associated with the mean values of the projections')
uncert_abs_proj = [0, 0, 0, 0]
for j in range(0,K):
  projj = proj[:,j]
  confidence_interval_proj=st.t.interval(confidence=0.95, df=len(projj)-1, loc=projmean[j], scale=st.sem(projj))  
  uncert_abs_proj[j]=0.5*(confidence_interval_proj[1]-confidence_interval_proj[0])
#
printt('    Uncert(proj) : ' + str(uncert_abs_proj))
printt(' ')
#---------------------------------------------------
# Mean values and uncertainties of the probabilities
probmean = np.array([np.mean(prob[:,0]), np.mean(prob[:,1]), np.mean(prob[:,2]), np.mean(prob[:,3])])
#
uncert_abs_prob = [0, 0, 0, 0]
for j in range(0,K):
  probj = prob[:,j]
  confidence_interval_prob=st.t.interval(confidence=0.95, df=len(probj)-1, loc=probmean[j], scale=st.sem(probj)) 
  uncert_abs_prob[j]=0.5*(confidence_interval_prob[1]-confidence_interval_prob[0])
#
printt('  The mean values of the corresponding probabilities')
printt('  do not nullify (as all of them are positive or zero)')
printt('    probmean     : ' + str(probmean))
printt('    Uncert(prob) : ' + str(uncert_abs_prob))
printt(' ')
#---------------------------------------------------
# Cumulative values of the probabilities obtained by adding all of them
# from 0 to 0, 0 to 1, 0 to 2, and 0 to 3)
probmeantot = np.array([np.mean(prob[:,0]), np.mean(prob[:,0])+np.mean(prob[:,1]), np.mean(prob[:,0])+np.mean(prob[:,1])+np.mean(prob[:,2]), np.mean(prob[:,0])+np.mean(prob[:,1])+np.mean(prob[:,2])+np.mean(prob[:,3])])
printt('  Cumulative values of the probabilities obtained by adding all of them')
printt('    probmeantot     : ' + str(probmeantot))
printt(' ')
printt('-------------------------------------------')
printt(' Project. and prob. mean values done!')
printt('-------------------------------------------')
printt(' ')
printt('===========================================')
printt('PROJECTION ON THE EIGENVECTORS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#
#

#===================================================
# PARTICIPATION RATIOS OF THE DATA IN THE
# LEARNING-STYLES'S ORIGINAL BASIS SET
# (ACTIVIST, REFLECTOR, THEORIST, AND PRAGMATIST)
# AND IN THE BASIS SET FORMED BY THE EIGENVECTORS
# OF THE COVARIANCE MATRIX
#===================================================
printt('===========================================')
printt('PARTICIPATION RATIOS (PRs)...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Computation of the PRs...')
printt('-------------------------------------------')
printt(' ')
# Participation ratios in the original LS basis set
# (column 0) and in the basis set formed by the
# covariance eigenvectors (column 1)
pr   = np.zeros((L,2))
#
# Notice that the PRs for a state with coefficientes C_i
# are defined as
#
# PR = (\sum Ci^2)^2/\sum C_i^4
#
# Instead of as
#
# PR = \sum Ci^2 /\sum C_i^4
#
# In this way, in a basis set formed by N elements
#
#  1 <= PR <= N.
#
# Notice that PR may equal 0 when all coefficients
# nullify. Then, we impose a value equal to 1.
#
for i  in range(0,L):
# LSs basis set
# (activist, reflector, theorist, pragmatist)
  sum2 = 0.0
  sum4 = 0.0
  for j  in range(0,K):
    p2   = data[i, j] * data[i, j]
    sum2 = sum2 + p2
    sum4 = sum4 + p2 * p2
#
  if( p2 == 0):
#   We impose that the PR equals 1
    pr[i,0] = 1 # 0
  else:
    pr[i,0] = sum2 * sum2 / sum4
#
# Basis set formed by the eigenfunctions
  sum2 = 0.0
  sum4 = 0.0
  for j  in range(0,K):
    p2   = proj[i, j] * proj[i, j]
    sum2 = sum2 + p2
    sum4 = sum4 + p2 * p2
#
  if( p2 == 0):
#   We impose that the PR equals 1
    pr[i,1] = 1
  else:
    pr[i,1] = sum2 * sum2 / sum4
#
printt('-------------------------------------------')
printt(' Computation of the PRs done!')
printt('-------------------------------------------')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Mean value and uncertainties of the PRs...')
printt('-------------------------------------------')
printt(' ')
prmean = [ np.mean(pr[:,0]), np.mean(pr[:,1])]
#
confidence_interval_pr0=st.t.interval(confidence=0.95, df=len(pr[:,0])-1, loc=np.mean(pr[:,0]), scale=st.sem(pr[:,0])) 
uncert_abs_pr0    =0.5*(confidence_interval_pr0[1]-confidence_interval_pr0[0])
#
confidence_interval_pr1=st.t.interval(confidence=0.95, df=len(pr[:,1])-1, loc=np.mean(pr[:,1]), scale=st.sem(pr[:,1])) 
uncert_abs_pr1    =0.5*(confidence_interval_pr1[1]-confidence_interval_pr1[0])
#
uncert_abs_pr = [ uncert_abs_pr0, uncert_abs_pr1 ]
printt('  prmean     : ' + str(prmean))
printt('  Uncert(PR) : ' + str(uncert_abs_pr))
printt(' ')
printt('-------------------------------------------')
printt(' Mean value and uncertainties of PRs done!')
printt('-------------------------------------------')
printt(' ')
#
#---------------------------------------------------
# Plot of the participation ratios...
printt('-------------------------------------------')
printt(' Analysis of the cumulative probability')
printt(' distribution (W(PR)) of the PRs')
printt(' (plot and fitting using a Weibull)')
printt(' distribution)...')
printt('-------------------------------------------')
#
fig, ax = plt.subplots(figsize = ( w_fig, h_fig )) #, layout='constrained')
#
ymaxx   = 10000000000
#
parameters_Weibull_PR = np.zeros((K,2))
xWeib   = np.arange(1, 4.1, 0.001)
#
for ipr in range(0, 2):
  if(ipr == 0):
    printt('  W(PR) in the LS original basis set')
  else:
    printt('  W(PR) in the eigenvectors basis set')
  printt(' ')
#
  x = pr[:, ipr]
  x = np.sort(x)
#
  y = np.zeros(L)
  y[0] = 0
  dy = 1 / float(L-1)
  for i in range(1,L):
    y[i] = y[i-1] + dy
#
# Staircase
  x2 = np.zeros(2*L+2)
  y2 = np.zeros(2*L+2)
  x2[0]  =   0
  y2[0]  =   0
  x2[2*L+1] = x[L-1]
  y2[2*L+1] = 1
#
# Find the points close to 1, 2, and 3
# in order to correcly color the
# PR that lie in the ranges [0, 1), 
# (1, 2), (2,3), and (3,4]
  ic   = [-1, -1]
  xc   = [ 1,  2]
  xref = [ 2,  3]
#
  for i in range(0,len(ic)):
    for j in range(0, L):
      if( x[j] > xc[i] and x[j] < xref[i] ):
        ic[i] = j
        xc[i] = x[j]
#
# Fitting using Weibull distribution
  param = curve_fit(Wweibull, x, y, [prmean[ipr], 1], maxfev = max_iter_Weib)

  [alpha_Weibull, k_Weibull] = param[0]
  pcov = param[1]
  parameters_Weibull_PR[ipr] = [alpha_Weibull, k_Weibull]
  printt('  alpha_Weibull : ' + str(alpha_Weibull))
  printt('  k_Weibull : ' + str(k_Weibull))
  printt(' ')
  printt(' ')
#
  for i in range(0,L):
    x2[2*i+1] = x[i]
    x2[2*i+2] = x[i]

    y2[2*i+1] = y[i]
    y2[2*i+2] = y[i] + dy
  y2[2*L] = 1    
#
# Plotting limits
  ymin_W = 0.00
  ymax_W = 1.01 
  ax.set_xlim( [ 1, 4] )
  ax.set_ylim( [ ymin_W, ymax_W] )
#
# Tick parameters
  ax.tick_params(axis='both', which='major', labelsize = ticksize_PR)
#
# Labels
  ax.set_xlabel('PR', fontsize = labelsize_PR)
  ax.set_ylabel('W [PR]',      fontsize = labelsize_PR)
#
# Plot cumulative distribution and Weibull fitting function
  ax.plot(x2, y2, lw = PR_W_width, color = PR_color[ipr], ls = PR_line[ipr])
  ax.plot(xWeib,Wweibull(xWeib, alpha_Weibull, k_Weibull) , lw = PR_W_width, color = PR_color[ipr], ls = PR_line[ipr])  
#
# Plot the average value as a vertical line
  ax.plot([prmean[ipr], prmean[ipr]], [ymin, ymax], lw = PR_mean_width, color = PR_color[ipr], ls = PR_mean_line)
  plt.fill_between([prmean[ipr]-uncert_abs_pr[ipr], prmean[ipr]+uncert_abs_pr[ipr]], [ymin, ymaxx], color = PR_color[ipr], alpha = PR_mean_alpha )  
#
filename_PR_W = output_statistics_pr + '/Fig_PR_W.png'
printt('Saving ' + filename_PR_W)
plt.savefig(filename_PR_W)
printt(' ')
printt('-------------------------------------------')
printt(' W(PR) done!')
printt('-------------------------------------------')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Probability distribution (P(PR)) of the PRs')
printt(' (plot and fitting using Weibull function...)')
printt('-------------------------------------------')
#
fig, ax = plt.subplots(figsize = ( w_fig, h_fig )) #, layout='constrained')
#
xWeib   = np.arange(0.8, 4.2, 0.001)
xmaxtot = 0
ymaxtot = 0
ymaxx   = 10000000000
#
for ipr in range(0, 2):
  if(ipr == 0):
    printt('  W(PR) in the LS original basis set')
  else:
    printt('  W(PR) in the eigenvectors basis set')
  printt(' ')
  x = pr[:, ipr]
#
# Tick parameters
  ax.tick_params(axis='both', which='major', labelsize = ticksize_PR)
#
# Labels
  ax.set_xlabel('PR', fontsize = labelsize_PR)
  ax.set_ylabel('P [PR]',      fontsize = labelsize_PR)
#
  y, x, _= ax.hist(x, bins=15, color = PR_color[ipr], alpha = PR_alpha) # We create a histogram with 15 blocks instead of just bins=10
  xmax = x.max()
  if xmax > xmaxtot:
    xmaxtot = xmax
#
  ymax = y.max()
  if ymax > ymaxtot:
    ymaxtot = ymax
#
# Probability distribution for Weibull function, scaled
# in such a way that its maximum coincides with that of the histogram
  [alpha_Weibull, k_Weibull] = parameters_Weibull_PR[ipr]
  factor = ymax / Pweibull(alpha_Weibull, alpha_Weibull, k_Weibull)
  ax.plot(xWeib,factor*Pweibull(xWeib, alpha_Weibull, k_Weibull) , lw = PR_W_width,   color = PR_color[ipr], ls = PR_line[ipr])  
#
# Plot the average value as a vertical line
  plt.fill_between([prmean[ipr]-uncert_abs_pr[ipr], prmean[ipr]+uncert_abs_pr[ipr]], [ymin, ymaxx], color = PR_color[ipr], alpha = PR_mean_alpha )
  ax.plot([prmean[ipr], prmean[ipr]], [ymin, ymaxx], lw = PR_mean_width, color = PR_color[ipr], ls = PR_mean_line)
#
# Plotting limits 
ymax = ymaxtot * 1.05
ax.set_xlim( [ 0.99, 4.01] )
ax.set_xlim( [ 1, xmaxtot] )
ax.set_ylim( [ 0, ymax] )

filename_PR_P = output_statistics_pr + '/Fig_PR_P.png'
printt('Saving ' + filename_PR_P)
plt.savefig(filename_PR_P)
printt(' ')
printt('-------------------------------------------')
printt(' P(PR) done!')
printt('-------------------------------------------')
printt(' ')
printt(' ')
printt('===========================================')
printt('PARTICIPATION RATIOS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#
#
#===================================================
# 4D PLOT OF THE LS
# (THE FOURTH LS IS INCLUDED IN THE CHARACTERISTICS
# OF THE SYMBOLS)
#===================================================
printt('===========================================')
printt('4D PLOT OF THE LSs...')
printt('===========================================')
printt(' ')
#----------------------------------------------------------
Llt = len(learning_trios)
printt('-------------------------------------------')
printt(' Llt : ' + str(Llt))
printt('-------------------------------------------')
printt(' ')

for ilt in range(itrios,itrios+1):

  fig = plt.subplots(figsize = ( w_fig, h_fig )) #, layout='constrained')
  ax  = plt.axes(projection ="3d")

  ax.set_xlim([ xmin3D, xmax3D] )
  ax.set_ylim([ ymin3D, ymax3D] )
  ax.set_zlim([ zmin3D, zmax3D] )

  ax.set_xticks([5, 10, 15])
  ax.set_yticks([5, 10, 15])
  ax.set_zticks([5, 10, 15])
  ax.tick_params(axis='both', which='major', labelsize = ticksize_3D)
  
#----------------------------------------------------------
# Scatter as a function of PCi, PCj and PCk
# with i, j, k = 0, 1, 2, or 3.
  printt('-------------------------------------------')
  printt('   ilt : ' + str(ilt))

  index0 = learning_trios[ilt][0]
  index1 = learning_trios[ilt][1]
  index2 = learning_trios[ilt][2]
  
  for i in range(0,K):
    if( i != index0 and i!= index1  and i!= index2 ):
      index3 = i

# Label_LS
  Label_LS0 = Label_LS[index0]
  Label_LS1 = Label_LS[index1]
  Label_LS2 = Label_LS[index2]

  printt('Labels LS   :     ' + str(Label_LS0) + ' vs ' + str(Label_LS1) + ' vs ' + str(Label_LS2))
  printt(' ')        
  printt('     ' + str(Label_LS0) + ' vs ' + str(Label_LS1) + ' vs ' + str(Label_LS2))
  printt('      Mean :' + str(xmean[index0]) + str(xmean[index1]) + str(xmean[index2]))
  printt(' ')
  printt('      index0 : ' + str(index0))
  printt('      index1 : ' + str(index1))     
  printt('      index2 : ' + str(index2))
  printt('     (index3 : ' + str(index3) + ')')
  
  ax.set_xlabel(Label_LS0, fontsize = labelsize_3D, labelpad = labelpad_3D)
  ax.set_ylabel(Label_LS1, fontsize = labelsize_3D, labelpad = labelpad_3D)
  ax.set_zlabel(Label_LS2, fontsize = labelsize_3D, labelpad = labelpad_3D)
                   
  vx0 = data[:,index0]
  vx1 = data[:,index1]
  vx2 = data[:,index2]
  vx3 = data[:,index3]

# Eigenvectors
  for j in range(0,K):
    vj = eigenVectors_max1[:,j]*distmeanmax

    v0 = xmean[index0]+vj[index0]
    v1 = xmean[index1]+vj[index1]
    v2 = xmean[index2]+vj[index2]              
    v3 = xmean[index3]+vj[index3]

    ax.plot([xmean[index0], v0], [xmean[index1], v1], [xmean[index2], v2], lw=width_PC_vectors_3D[j], color=PC_color[j], ls=PC_line_bis[j])

# Shaded are in the PC1-PC2 plane
  vj0 = eigenVectors[:,0]*distmeanmax
  v00 = xmean[index0]+vj0[index0]
  v01 = xmean[index1]+vj0[index1]
  v02 = xmean[index2]+vj0[index2]        
  v03 = xmean[index3]+vj0[index3]
    
  vj1 = eigenVectors[:,1]*distmeanmax
  v10 = xmean[index0]+vj1[index0]
  v11 = xmean[index1]+vj1[index1]
  v12 = xmean[index2]+vj1[index2]        
  v13 = xmean[index3]+vj1[index3]
  
  xx = [xmean[index0], v00, v10]
  yy = [xmean[index1], v01, v11]
  zz = [xmean[index2], v02, v12]

# Points
  for i in range(0,L):
    x0 = vx0[i]
    x1 = vx1[i]
    x2 = vx2[i]
    x3 = vx3[i]

    tendency3 = scatter_tendency(Label_LS[index3], x3)
    [tendency3, scatter_color3, scatter_size3, scatter_alpha3, scatter_symbol3]=scatter_properties(tendency3, x3)

    ax.scatter3D(x0, x1, x2, s=scatter_size3, marker=scatter_symbol3, color=scatter_color3,      alpha=scatter_alpha3)
    
#   Projections on the Cartesian planes
    colorsx     = colors_proj
    edgecolorsx = edgecolors_proj
      
    colorsy     = colors_proj
    edgecolorsy = edgecolors_proj

    colorsz     = colors_proj
    edgecolorsz = edgecolors_proj
    
    ax.scatter(x1, x2, s=scatter_size3, marker=scatter_symbol3, color=colorsx,  alpha=scatter_alpha3, zdir='x')
    ax.scatter(x0, x2, s=scatter_size3, marker=scatter_symbol3, color=colorsy,  alpha=scatter_alpha3, zdir='y')
    ax.scatter(x0, x1, s=scatter_size3, marker=scatter_symbol3, color=colorsz,  alpha=scatter_alpha3, zdir='z')
            
# Average mean
  ax.scatter3D(xmean[index0], xmean[index1], xmean[index2], s=size_mean, marker=mean_symbol, color='k')
  
# Projections of the average mean on the Cartesian planes
  ax.scatter(xmean[index1], xmean[index2], s=size_mean, marker=mean_symbol, color=colorsx, zdir='x')
  
  ax.scatter(xmean[index0], xmean[index2], s=size_mean, marker=mean_symbol, color=colorsy, zdir='y')
    
  ax.scatter(xmean[index0], xmean[index1], s=size_mean, marker=mean_symbol, color=colorsz, zdir='z')
  
# Fill uncertainty areas for the average mean on the projections
  xx = [xmin3D, xmin3D, xmin3D, xmin3D]
  yy = [xmean[index1]-dxmean[index1],
        xmean[index1]+dxmean[index1],
        xmean[index1]+dxmean[index1],
        xmean[index1]-dxmean[index1]]
  zz = [xmean[index2]-dxmean[index2],
        xmean[index2]-dxmean[index2],
        xmean[index2]+dxmean[index2],
        xmean[index2]+dxmean[index2]]
  verts = [(xx[i],yy[i],zz[i]) for i in range(len(xx))]
  ax.add_collection3d(Poly3DCollection([verts],facecolor = mean_uncert_color, alpha = mean_alpha))

  xx = [xmean[index0]-dxmean[index0],
        xmean[index0]+dxmean[index0],
        xmean[index0]+dxmean[index0],
        xmean[index0]-dxmean[index0]]
  zz = [xmean[index2]-dxmean[index2],
        xmean[index2]-dxmean[index2],
        xmean[index2]+dxmean[index2],
        xmean[index2]+dxmean[index2]]
  yy = [ymin3D, ymin3D, ymin3D, ymin3D]
  verts = [(xx[i],yy[i],zz[i]) for i in range(len(xx))]
  ax.add_collection3d(Poly3DCollection([verts],facecolor = mean_uncert_color, alpha = mean_alpha))
  
  xx = [xmean[index0]-dxmean[index0],
        xmean[index0]+dxmean[index0],
        xmean[index0]+dxmean[index0],
        xmean[index0]-dxmean[index0]]
  yy = [xmean[index1]-dxmean[index1],
        xmean[index1]-dxmean[index1],
        xmean[index1]+dxmean[index1],
        xmean[index1]+dxmean[index1]]
  zz = [zmin3D, zmin3D, zmin3D, zmin3D]
  verts = [(xx[i],yy[i],zz[i]) for i in range(len(xx))]
  ax.add_collection3d(Poly3DCollection([verts],facecolor = mean_uncert_color, alpha = mean_alpha))

# Origin
  orig0 = LS[4, index0] + xmean[index0]
  orig1 = LS[4, index1] + xmean[index1]
  orig2 = LS[4, index2] + xmean[index2]
  ax.scatter3D(orig0, orig1, orig2, s = size_mean, marker = origin_symbol, color = origin_color, alpha = origin_alpha)  

# Pure maximal learning styles
# (only those whith the LS that appear in the axes)
  for i in [index0, index1, index2]:  
    x0 = LS[i, index0] + xmean[index0]
    x1 = LS[i, index1] + xmean[index1]
    x2 = LS[i, index2] + xmean[index2]
    ax.scatter3D(x0, x1, x2,  s = size_mean, marker = twenty_symbol[i], color = PC_color[i], alpha = twenty_alpha)  
    ax.plot([orig0, x0], [orig1, x1], [orig2, x2], color = twenty_color[i], ls = twenty_line[i], lw = twenty_width_3D[i])

# Maximal state (20, 20, 20, 20)
  ax.plot(xx,yy,zz,label="line plot")

  ax.view_init(25,40)

  filename_ls_3d = output_participants + '/Fig_'+Label_LS0+'_'+Label_LS1+'_'+Label_LS2+'.png'
  printt('Saving ' + filename_ls_3d)
  plt.savefig(filename_ls_3d)

  printt(' Plots for the LSs done!')
  printt('-------------------------------------------')
  printt(' ')
printt('===========================================')
printt('4D PLOT OF THE LSs DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#
#

printt('===========================================')
printt(' PROJECTIONS ON PCs AND STATISTICS...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' W(s) for the projections on the PCs...')
printt('-------------------------------------------')
parameters_Weibull_PC20 = np.zeros((K,3))
  
xmin = proj.min()
xmax = proj.max()
  
xmin_proj_tot = xmin
xmax_proj_tot = xmax

xWeib_proj = np.arange(xmin*2.5, xmax*2.5, (xmax-xmin)/100000)

for ipc in range(0, K):
  printt('   PC -20 to 20 : ' + str(ipc))
  x = proj[:, ipc]
  x = np.sort(x)

  xmin_proj = x.min()
 
  y = np.zeros(L)
  y[0] = 0
  dy = 1 / float(L-1)
  for i in range(1,L):
    y[i] = y[i-1] + dy

# Fitting using Weibull distribution
  param = curve_fit(Wweibull, x-xmin_proj, y, [projmean[ipc]-xmin_proj, 1], maxfev = max_iter_Weib)

  [alpha_Weibull, k_Weibull] = param[0]
  theta_Weibull = xmin_proj
  parameters_Weibull_PC20[ipc] = [alpha_Weibull, k_Weibull, theta_Weibull]
  printt('   alpha_Weibull     : ' + str(alpha_Weibull))
  printt('   k_Weibull         : ' + str(k_Weibull))
  printt('   theta_Weibull     : ' + str(theta_Weibull))
  printt(' ')
  
for ilp in range(0, 1):

  fig, ((ax_hor, ax_no), (ax, ax_ver)) = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), layout='constrained')

# Limits
  ax_no.axis('off')
  ax.set_xlim([ xmin, xmax] )
  ax.set_ylim([ ymin, ymax] )

# Ticks
  ax.tick_params(axis='both', which='major', labelsize = ticksize)

#----------------------------------------------------------
# Scatter as a function of PC0, PC1, PC2, and PC3
  printt('-------------------------------------------')
  printt(' Plotting scatter 2D on the principal axes')
  printt(' and probabilities on the PCs...')

  printt(' PC : ' + str(ilp))
  printt(' ')
  index0 = learning_pairs[ilp][0]
  index1 = learning_pairs[ilp][1]
  index2 = -1
 
  for i in range(0,K):
    if( i != index0 and i!= index1 ):
      if ( index2 == -1):
        index2 = i
      else:
        index3 = i
#
# MAIN PANEL 
# Label PCs
  Label_PC0 = Label_PC[index0]
  Label_PC1 = Label_PC[index1]
  
  Label_PCPC0 = Label_PCPC[index0]
  Label_PCPC1 = Label_PCPC[index1]

  printt(' ( ' + str(Label_PCPC0) + ' vs ' + str(Label_PCPC1) + ' )')
  printt(' ')

  ax.set_xlabel(Label_PCPC0, fontsize=labelsize)
  ax.set_ylabel(Label_PCPC1, fontsize=labelsize)
    
# LS values
  vx0 = data[:,index0]
  vx1 = data[:,index1]
  vx2 = data[:,index2]
  vx3 = data[:,index3]

# Plot the projection of the coefficients PC0 and PC1
  vy0 = proj[:,index0]
  vy1 = proj[:,index1]
  vy2 = proj[:,index2]
  vy3 = proj[:,index3]

# Plot limits
  xmin = np.min( [np.min(vy0), np.min(projectLS[:,index0])])
  xmax = np.max( [np.max(vy0), np.max(projectLS[:,index0])])
  
  ymin = np.min( [np.min(vy1), np.min(projectLS[:,index1])])
  ymax = np.max( [np.max(vy1), np.max(projectLS[:,index1])])
  
  if ( xmin < 0 ):
   xmin = xmin * 1.05
  else:
   xmin = xmin * 0.95
   
  if ( xmax < 0 ):
   xmax = xmax * 0.95
  else:
   xmax = xmax * 1.05
   
  if ( ymin < 0 ):
   ymin = ymin * 1.05
  else:
   ymin = ymin * 0.95
   
  if ( ymax < 0 ):
   ymax = ymax * 0.95
  else:
   ymax = ymax * 1.05
   
  printt('   xmin : ' + str(xmin))
  printt('   xmax : ' + str(xmax))
  printt('   ymin : ' + str(ymin))
  printt('   ymax : ' + str(ymax))
  printt(' ')
  
# Points
  for i in range(0,L):
#  printt(i)
    x0 = vx0[i]
    x1 = vx1[i]
    x2 = vx2[i]
    x3 = vx3[i]
    
    y0 = vy0[i]
    y1 = vy1[i]
    y2 = vy2[i]
    y3 = vy3[i]

    tendency2 = scatter_tendency(Label_LS[index2], x2)
    tendency3 = scatter_tendency(Label_LS[index3], x3)
    [tendency2, scatter_color2, scatter_size2, scatter_alpha2, scatter_symbol2]=scatter_properties(tendency2, x2)
    [tendency3, scatter_color3, scatter_size3, scatter_alpha3, scatter_symbol3]=scatter_properties(tendency3, x3)

#   Color depends on whether we are making a plot with PC' vs PC1 or not
#   (see projections in the 3D plots as a funciton of the original LSs)
    if (ilp == 0):
      ax.scatter(y0, y1, s=scatter_size2, marker=scatter_symbol2,      color = colors_proj_PC, alpha=scatter_alpha3)
      ax.scatter(y0, y1, s=scatter_size2, marker=scatter_symbol2, edgecolors = colors_proj_PC, facecolors='none')
    else:
      ax.scatter(y0, y1, s=scatter_size2, marker=scatter_symbol2,      color = colors_proj, alpha=scatter_alpha3)
      ax.scatter(y0, y1, s=scatter_size2, marker=scatter_symbol2, edgecolors = colors_proj, facecolors='none')
      
# Origin
  orig0 = projectLS[4, index0]
  orig1 = projectLS[4, index1]
  ax.scatter(orig0, orig1, s = size_mean, marker = origin_symbol, color = origin_color, alpha = origin_alpha)  

# Pure maximal learning styles
  for i in range(0,K):  
    y0 = projectLS[i, index0]
    y1 = projectLS[i, index1]
    ax.scatter(y0, y1,   s = size_mean, marker = twenty_symbol[i], color = PC_color[i], alpha = twenty_alpha)  
    ax.plot([orig0, y0], [orig1, y1], color = twenty_color[i], ls = twenty_line[i], lw = twenty_width_2D[i])

  i = K + 1
  y0 = projectLS[i, index0]
  y1 = projectLS[i, index1]
  ax.scatter(y0, y1,   s = size_mean, marker = twenty_symbol[i], color = twenty_color[i], alpha = twenty_alpha)  
  ax.plot([orig0, y0], [orig1, y1], color    = twenty_color[i], ls = twenty_line[i], lw = twenty_width_2D[i])
  
# Average mean
  ax.scatter(0, 0, s = size_mean, marker=mean_symbol, color = mean_color2)
  
  ax.set_xlim([ xmin, xmax] )
  ax.set_ylim([ ymin, ymax] )
#
#**********************************************************************
# SECONDARY PANELS (PROBABILITY DENSITY FUNCTIONS)
#
  xWeib   = np.arange(xmin, xmax, (xmax-xmin)/10000)
  yWeib   = np.arange(ymin, ymax, (ymax-ymin)/10000)
#
#-------------------------------------------------------
# Projection on the PC associated with the horizontal axis
  ax_hor.set_ylabel('PDF', fontsize=labelsize)
  ax_hor.set_xlim([ xmin, xmax] )
  
  ipc = index0
  x = proj[:, ipc]
  x = np.sort(x)
     
  [alpha_Weibull, k_Weibull, theta_Weibull] = parameters_Weibull_PC20[ipc]
  printt('   alpha_Weibull : ' + str(alpha_Weibull))
  printt('   k_Weibull     : ' + str(k_Weibull))
  printt('   theta_Weibull : ' + str(theta_Weibull))
  printt(' ')
  
  ax_hor.tick_params(axis='both', which='major', labelsize=ticksize, labelbottom = False)
 
  y, x, _= ax_hor.hist(x, bins=20, color = PC_color[ipc], alpha=0.5)
  ymaxx = y.max()

# Probability distribution for Weibull function, scaled
  factor = ymaxx / Pweibull(alpha_Weibull, alpha_Weibull,  k_Weibull)
  ax_hor.plot(xWeib_proj,factor*Pweibull_translated(xWeib_proj, alpha_Weibull, k_Weibull, theta_Weibull), lw = PC_W_width,   color = PC_color[ipc], ls = PC_line[0]) #, ls = PC_line[ipc])
  
  ymaxx = ymaxx * 1.205
  ax_hor.set_ylim( [ 0, ymaxx] )   
 
# Plot as a vertical line the origin
  ax_hor.plot([projectLS[4,ipc], projectLS[4,ipc]], [ymin, ymaxx], color = twenty_color[4], ls = twenty_line[4], lw = twenty_width_W[6])
  
# Mark with lines the origin and the maximal (pure) learning styles  
  for i in range(0,K):  
    ax_hor.plot([projectLS[i, ipc], projectLS[i, ipc]], [ymin, ymaxx], color = twenty_color[i], ls = twenty_line[i], lw = twenty_width_W[6])

# Mark as lines the maximal point (20, 20, 20, 20)
  ax_hor.plot([projectLS[5,ipc], projectLS[5,ipc]], [ymin, ymaxx], color=twenty_color[5], ls = twenty_line[5], lw = twenty_width_W[6])
#
#-------------------------------------------------------
# Projection on the PC associated with the vertical axis
  ax_ver.set_xlabel('PDF', fontsize=labelsize)
  ax_ver.set_ylim([ ymin, ymax] )
  
  ipc = index1
  x = proj[:, ipc]
  x = np.sort(x)
     
  [alpha_Weibull, k_Weibull, theta_Weibull] = parameters_Weibull_PC20[ipc]
  printt('   alpha_Weibull : ' + str(alpha_Weibull))
  printt('   k_Weibull     : ' + str(k_Weibull))
  printt('   theta_Weibull : ' + str(theta_Weibull))
  printt(' ')
  
  ax_ver.tick_params(axis='both', which='major', labelsize=ticksize, labelleft = False)
 
  y, x, _= ax_ver.hist(x, bins=20, color = PC_color[ipc], alpha=0.5, orientation='horizontal')
  ymaxx = y.max()

# Probability distribution for Weibull function, scaled
  factor = ymaxx / Pweibull(alpha_Weibull, alpha_Weibull,  k_Weibull)
  ax_ver.plot(factor*Pweibull_translated(yWeib, alpha_Weibull, k_Weibull, theta_Weibull), yWeib, lw = PC_W_width,   color = PC_color[ipc], ls = PC_line[0]) #, ls = PC_line[ipc])
  
  ymaxx = ymaxx * 1.20
  ax_ver.set_xlim( [ 0, ymaxx ] )    # Plot as a vertical line the origin
  ax_ver.plot([xmin, ymaxx], [projectLS[4,ipc], projectLS[4,ipc]], color = twenty_color[4], ls = twenty_line[4], lw = twenty_width_W[6])
  
# Mark with lines the origin and the maximal (pure) learning styles  
  for i in range(0,K):  
    ax_ver.plot([xmin, ymaxx], [projectLS[i, ipc], projectLS[i, ipc]], color = twenty_color[i], ls = twenty_line[i], lw = twenty_width_W[6])

# Mark as lines the maximal point (20, 20, 20, 20)
  ax_ver.plot([xmin, ymaxx], [projectLS[5,ipc], projectLS[5,ipc]], color=twenty_color[5], ls = twenty_line[5], lw = twenty_width_W[6])

  filename_PC0PC1 = output + '/Fig_'+Label_PCPC0+'_'+Label_PCPC1+'_P.png'
  plt.savefig(filename_PC0PC1)
  printt('   Saving ' + filename_PC0PC1)
  plt.savefig(filename_PC0PC1)
  printt(' ')
  printt(' Plots for the PCs done!')
  printt('-------------------------------------------')
  printt(' ')
printt('===========================================')
printt(' PROJECTIONS ON PCs AND STATISTICS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#********************************************

printt('===========================================')
printt(' PROJECTIONS ON PCs AND PROFESSORS LS...')
printt('===========================================')
printt(' ')

printt('-------------------------------------------')
printt(' Projections of professors data on eigenvectors...')
printt('-------------------------------------------')
printt(' ')
# Difference between the input data and the mean
if(Lprof!=0):
 ddata_prof = np.zeros((Lprof,K))
 for i in range(0,Lprof):
  for j in range(0,K):
    ddata[i,j] = data_prof[i,j] - xmean[j]

 proj_prof = np.zeros((L,K))
#
 for i  in range(0,Lprof):
  norm_ddata = 0.0
  for j  in range(0,K):
    norm_ddata = norm_ddata + ddata_prof[i,j]*ddata_prof[i,j]
  norm_ddata = np.sqrt( norm_ddata )
#
  for j  in range(0,K):
    vj = eigenVectors[:,j]
#
    for l  in range(0,K):
      proj_prof[i,j] = proj_prof[i,j] + vj[l]*ddata[i,l]
#
 printt('-------------------------------------------')
 printt(' Projections of professors data on eigenvectors done!')
 printt('-------------------------------------------')
 printt(' ')

pr_prof   = np.zeros((Lprof,2))

for i  in range(0,Lprof):
  if (i == 0):
    printt('-------------------------------------------')
    printt(' Participation ratios of the professors... ')
    printt('-------------------------------------------')

# LSs basis set
# (activist, reflector, theorist, pragmatist)
  sum2 = 0.0
  sum4 = 0.0
  for j  in range(0,K):
    p2   = data_prof[i, j] * data_prof[i, j]
    sum2 = sum2 + p2
    sum4 = sum4 + p2 * p2
#
  if( p2 == 0):
#   We impose that the PR equals 1
    pr_prof[i,0] = 1 # 0
  else:
    pr_prof[i,0] = sum2 * sum2 / sum4
#
# Basis set formed by the eigenfunctions
  sum2 = 0.0
  sum4 = 0.0
  for j  in range(0,K):
    p2   = proj_prof[i, j] * proj_prof[i, j]
    sum2 = sum2 + p2
    sum4 = sum4 + p2 * p2
#
  if( p2 == 0):
#   We impose that the PR equals 1
    pr_prof[i,1] = 1 # 0
  else:
    pr_prof[i,1] = sum2 * sum2 / sum4

  printt('-------------------------------------------')
  printt(' Participation ratios of the professors done!')
  printt('-------------------------------------------')
  printt(' ')



for i_prof in range(0, Lprof):
  profi = profs[i_prof]
  printt('Professor : ' + profi )
  fig, axs = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), sharex = True, sharey = True, layout='constrained')  
  
  for ils in range(0,K):

    axs[panels[ils][0],panels[ils][1]].tick_params(axis='both', which='major', labelsize=ticksize)  
#----------------------------------------------------------
#   Scatter as a function of PC0, PC1, PC2, and PC3
    printt('-------------------------------------------')
    printt(' Plotting scatter 2D on PCs (professor clusters)...')

    printt(' ls : ' + str(ilp))
    printt('  ')
  
    tendencyref = scatter_tendency(Label_LS[ils], data_prof[i_prof, ils])
    [tendencyprof, scatter_colorref, scatter_sizeref, scatter_alpharef, scatter_symbolprof]=scatter_properties(tendencyref,  data_prof[i_prof, ils])
    
    index0 = learning_pairs[ilp][0]
    index1 = learning_pairs[ilp][1]
    index2 = -1
 
    for i in range(0,K):
      if( i != index0 and i!= index1 ):
        if ( index2 == -1):
          index2 = i
        else:
          index3 = i
    
#   Label PCs
    Label_PC0 = Label_PC[index0]
    Label_PC1 = Label_PC[index1]
  
    Label_PCPC0 = Label_PCPC[index0]
    Label_PCPC1 = Label_PCPC[index1]

    printt(' ( ' + Label_PCPC0 + ' vs ' + Label_PCPC1 + ' )')
    printt('  ')

    ax.set_xlabel(Label_PCPC0, fontsize=labelsize)
    ax.set_ylabel(Label_PCPC1, fontsize=labelsize)
    
#   LS values
    vx0 = data[:,index0]
    vx1 = data[:,index1]
    vx2 = data[:,index2]
    vx3 = data[:,index3]

#   Considered LS for clustering
#   vxref = data[:,ils]
  
#   Plot the projection of the coefficients PC0 and PC1
    vy0 = proj[:,index0]
    vy1 = proj[:,index1]
    vy2 = proj[:,index2]
    vy3 = proj[:,index3]

#   Plot limits
    xmin = np.min( [np.min(vy0), np.min(projectLS[:,index0])])
    xmax = np.max( [np.max(vy0), np.max(projectLS[:,index0])])
  
    ymin = np.min( [np.min(vy1), np.min(projectLS[:,index1])])
    ymax = np.max( [np.max(vy1), np.max(projectLS[:,index1])])
  
    if ( xmin < 0 ):
      xmin = xmin * 1.05
    else:
     xmin = xmin * 0.95
   
    if ( xmax < 0 ):
      xmax = xmax * 0.95
    else:
      xmax = xmax * 1.05
   
    if ( ymin < 0 ):
      ymin = ymin * 1.05
    else:
      ymin = ymin * 0.95
   
    if ( ymax < 0 ):
      ymax = ymax * 0.95
    else:
      ymax = ymax * 1.05
   
    printt( ' xmin : ' + str(xmin))
    printt( ' xmax : ' + str(xmax))
    printt( ' ymin : ' + str(ymin))
    printt( ' ymax : ' + str(ymax))
    printt('  ')
  
#   Points
    for i in range(0,L):
      x0 = vx0[i]
      x1 = vx1[i]

      vxrefi = data[i,ils]
        
      y0 = vy0[i]
      y1 = vy1[i]

      tendencyref = scatter_tendency(Label_LS[ils], vxrefi)
      [tendencyref, scatter_colorref, scatter_sizeref, scatter_alpharef, scatter_symbolref]=scatter_properties(tendencyref, vxrefi)

      if(tendencyref == tendencyprof):
        scatter_sizeref   = 2*scatter_sizeref
        scatter_symbolref = 'X'

#     Color depends on the tendency
      if( tendencyref == 'vl' or tendencyref == 'l' ):
        colorref = tendency_color[0]
      else:
        if( tendencyref == 'm' ):
          colorref = tendency_color[1]
        else:
          colorref = tendency_color[2]
    
      axs[panels[ils][0],panels[ils][1]].scatter(y0, y1, s=scatter_sizeref, marker=scatter_symbolref,      color = colorref, alpha=scatter_alpharef)
      axs[panels[ils][0],panels[ils][1]].scatter(y0, y1, s=scatter_sizeref, marker=scatter_symbolref, edgecolors = colorref, facecolors='none')

#   Points professors
    for i in range(0, Lprof):

#     Plot the projection of the coefficients PC0 and PC1
      tendencyref = scatter_tendency(Label_LS[ils], data_prof[i, ils])
      [tendencyref, scatter_colorref, scatter_sizeref, scatter_alpharef, scatter_symbolref]=scatter_properties(tendencyref, data_prof[i, ils])

      if( tendencyref == 'vl' or tendencyref == 'l' ):
        colorref = tendency_color[0]
      else:
        if( tendencyref == 'm' ):
          colorref = tendency_color[1]
        else:
          colorref = tendency_color[2]

      y0 = proj_prof[i, index0]
      y1 = proj_prof[i, index1]
      if(i == i_prof ):
        axs[panels[ils][0],panels[ils][1]].scatter(y0, y1, s=2*scatter_sizeref, marker='P', color = colorref, alpha=scatter_alpharef)
        axs[panels[ils][0],panels[ils][1]].scatter(y0, y1, s=2*scatter_sizeref, marker='P', edgecolors = colorref, facecolors='none')
      else:
        axs[panels[ils][0],panels[ils][1]].scatter(y0, y1, s=scatter_sizeref, marker='P', edgecolors = colorref, facecolors='none')

#   Origin
    orig0 = projectLS[4, index0]
    orig1 = projectLS[4, index1]
    axs[panels[ils][0],panels[ils][1]].scatter(orig0, orig1, s = size_mean, marker = origin_symbol, color = origin_color, alpha = origin_alpha)  

#   Pure maximal learning styles
    for i in range(0,K):  
      y0 = projectLS[i, index0]
      y1 = projectLS[i, index1]
      axs[panels[ils][0],panels[ils][1]].scatter(y0, y1,   s = size_mean, marker = twenty_symbol[i], color = PC_color[i], alpha = twenty_alpha)  
      axs[panels[ils][0],panels[ils][1]].plot([orig0, y0], [orig1, y1], color = twenty_color[i], ls = twenty_line[i], lw = twenty_width_2D[i])

    i = K + 1
    y0 = projectLS[i, index0]
    y1 = projectLS[i, index1]
    axs[panels[ils][0],panels[ils][1]].scatter(y0, y1,   s = size_mean, marker = twenty_symbol[i], color = twenty_color[i], alpha = twenty_alpha)  
    axs[panels[ils][0],panels[ils][1]].plot([orig0, y0], [orig1, y1], color    = twenty_color[i], ls = twenty_line[i], lw = twenty_width_2D[i])

#   Average mean
    axs[panels[ils][0],panels[ils][1]].scatter(0, 0, s = size_mean, marker=mean_symbol, color = mean_color2)

#   Educators
    axs[panels[ils][0],panels[ils][1]].scatter(proj_prof[i_prof, 0], proj_prof[i_prof, 1], s = size_mean, marker = 'P', color = 'k', alpha = origin_alpha)

# Text label
    axs[panels[ils][0], panels[ils][1]].text(xmin + 0.1*(xmax-xmin), ymax - 0.1*(ymax-ymin), Label_LS[ils],
        fontsize = labelsize,
        horizontalalignment='left',
        verticalalignment='top')
  
  axs[panels[ils][0],panels[ils][1]].set_xlim([ xmin, xmax] )
  axs[panels[ils][0],panels[ils][1]].set_ylim([ ymin, ymax] )

  axs[1, 0].set_xlabel(Label_PCPC0, fontsize=labelsize)
  axs[1, 0].set_ylabel(Label_PCPC1, fontsize=labelsize)

  filename = output_participants_pc_prof + '/Fig_'+Label_PCPC0+'_'+Label_PCPC1+'_prof_' + profi + '.png'
  plt.savefig(filename)
  printt('   Saving ' + filename)
  plt.savefig(filename)

printt('  ')
printt(' Plots for the professor clusters done!')
printt('-------------------------------------------')
printt('  ')

printt('===========================================')
printt('2D PLOT OF THE PCs WITH THE SCATTER')
printt('SIZE AND COLORS A FUNCTION OF PROFESSORS LS')
printt('DONE!')
printt('===========================================')
printt('  ')
printt('  ')
printt('  ')
printt('  ')
#
#===================================================
#===================================================
# OPTIONAL ANALYSIS OF STUDENTS' MARKS
# (peformed only if their marks are provided)
#===================================================
#===================================================
print('K_value : ', K_value)

if ( K_value > 0 and Lmarks > 0 ):
 printt('===========================================')
 printt('ANALYSIS OF THE MARKS OF THE STUDENTS...')
 printt('===========================================')

 printt('-------------------------------------------')
 printt(' Plots of students marks vs original LS...')
 printt('-------------------------------------------')

#Clustering figure for the LS clustering
 fig_clustering, ax_clustering = plt.subplots(figsize=(w_fig, h_fig), layout='constrained')

 panels = [[0, 0], [0, 1], [1, 0], [1, 1]]

 printt('np.shape(data_marks) : ' + str(np.shape(data_marks)))

 printt('data_marks    : ' + str(data_marks))
 marks_vs_LS = []

 printt(' K_value          : ' + str(K_value) )

 if(K_value > len(cluster_color)):
   printt('Please, include more colors in the cluster_color array.')
   sys.exit()

 for i in range(1, Lmarks_activities):

   printt(' Evaluating activity : ' + data_marks[i, 2] + ' contained in the sheet ' + data_marks[i, 1] + ' of ' + data_marks[i, 0])
   printt('  ')

   fig, axs = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), sharex = True, sharey = True, layout='constrained')

   students = data_marks[0, 3:]
   marks    = np.array(data_marks[i, 3:], dtype=float)

#  Create two vectors with the marks and the LS
#  and a string with the names of the students
#  (if marks are correctly defined)
   vls       = []
   vstudents = []
   vmarks    = []
   for j in range(0, L):
    if(~np.isnan(marks[j])):
      vstudents.append(students[j])
      vls.append([data[j,0], data[j,1], data[j,2], data[j,3]])
      vmarks.append(marks[j])

   vstudents = np.stack(vstudents)
   vls    = np.stack(vls)
   vmarks = np.array(vmarks)
   vmarks = vmarks.astype(float)

#  Plot limits
   xmin = np.min( np.min(data) )
   xmax = np.max( np.max(data) )

   ymin = -0.5/1.05
   ymax = 10
   if( np.max(vmarks) > ymax):
    ymax = np.max(vmarks)

   if ( xmin < 0 ):
    xmin = xmin * 1.05
   else:
    xmin = xmin * 0.95

   if ( xmax < 0 ):
    xmax = xmax * 0.95
   else:
    xmax = xmax * 1.05

   if ( ymin < 0 ):
    ymin = ymin * 1.05
   else:
    ymin = ymin * 0.95

   if ( ymax < 0 ):
    ymax = ymax * 0.95
   else:
    ymax = ymax * 1.05

   printt( ' xmin : ' + str(xmin))
   printt( ' xmax : ' + str(xmax))
   printt( ' ymin : ' + str(ymin))
   printt( ' ymax : ' + str(ymax))
   printt('  ')

   if( K_value <= len(vmarks) ):
    for ils in range(0,K):
     print('panels : ', panels[ils][0], panels[ils][1])
     axs[panels[ils][0], panels[ils][1]].tick_params(axis='both', which='major', labelsize=ticksize)
#----------------------------------------------------------
#    Scatter of the marks as a function of the LS  points
     printt(' LS : ' + str(ils))
     printt('  ')

     Label_LS0 = Label_LS[ils]
     printt(' ( marks vs ' + Label_LS0 + ' )')
     printt('  ')

     vx0 = vls[:,ils]
     vx1 = vmarks

     printt('   Clustering for ' + Label_LS0 + '...')

     array=vx1.reshape(-1, 1)

#    Create and fit the K-Means model with the optimal number of clusters
     kmeans = KMeans(n_clusters=K_value, init='k-means++', random_state=42)
     kmeans.fit(array)

#    Get the cluster labels for each data point
     k_labels = kmeans.labels_

#    Average means and uncertainties for the clusters
     clusters_means  = [ ]
     clusters_uncert = [ ]
     clusters_index  = [ ]
     list_students = ''
     for k in range(0, K_value):
      values = []

      for kk in range(0,len(k_labels)):
        if(k == k_labels[kk]):
          if(list_students == ''):
            list_students = vstudents[kk]
          else:
            list_students = list_students + ', ' + vstudents[kk]
          values.append( [ vx0[kk], vx1[kk] ] )
      shape = np.shape(values)
      clusters_index.append(shape[0])
      values = np.stack(values)

#     Average means
      x0_mean = np.mean( values[:,0] )
      x1_mean = np.mean( values[:,1] )

#     Uncertainties
      confidence_interval=st.t.interval(confidence=0.95, df=len(values[:,0])-1, loc=x0_mean, scale=st.sem(values[:,0]))
      dx0 = 0.5*(confidence_interval[1]-confidence_interval[0])
      confidence_interval=st.t.interval(confidence=0.95, df=len(values[:,1])-1, loc=x0_mean, scale=st.sem(values[:,1]))
      dx1 = 0.5*(confidence_interval[1]-confidence_interval[0])
      clusters_means.append([ x0_mean, x1_mean ])
      clusters_uncert.append( [dx0, dx1] )

     clusters_means  = np.stack(clusters_means)
     clusters_uncert = np.stack(clusters_uncert)
     printt('   Clustering for ' + Label_LS0 + ' done!')

#    Linear regression
     if np.all(vx0 == vx0[0]):
       print("Error: All vx0 values are identical. Linear regression cannot be performed.")
       m = b = delta_m = delta_b = rvalue = pvalue = np.nan
     else:
       m, b, rvalue, pvalue, delta_m = linregress(vx0, vx1)
#      Compute uncertainty in intercept (y0)
       delta_b  = delta_m * np.sqrt( np.sum(vx0 * vx0) / len(vx0) )

#    Average means
     x0_mean = np.mean( vx0 )
     x1_mean = np.mean( vx1 )
#    Uncertainties
     confidence_interval=st.t.interval(confidence=0.95, df=len(vx0)-1, loc=x0_mean, scale=st.sem(vx0))
     dx0 = 0.5*(confidence_interval[1]-confidence_interval[0])
     confidence_interval=st.t.interval(confidence=0.95, df=len(vx1)-1, loc=x1_mean, scale=st.sem(vx1))
     dx1 = 0.5*(confidence_interval[1]-confidence_interval[0])

     marks_vs_LS.append([ data_marks[i, 0], data_marks[i, 1], data_marks[i, 2], len(vx0), list_students, -1, Label_LS0, x0_mean, dx0, x1_mean, dx1, m,  delta_m, b, delta_b, rvalue ] )

     def lin_regr2(m, b, x):
      return m*x+b

#    The confidence interval is well defined within the lines
#    ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
#    since x is always positive
     x_linregr  = [ xmin, xmax ]
     y_linregr  = [ lin_regr2(m, b, x_linregr[0]), lin_regr2(m, b, x_linregr[1]) ]
     y_linregr1 = [ lin_regr2(m-delta_m, b-delta_b, x_linregr[0]), lin_regr2(m-delta_m, b-delta_b, x_linregr[1]) ]
     y_linregr2 = [ lin_regr2(m+delta_m, b+delta_b, x_linregr[0]), lin_regr2(m+delta_m, b+delta_b, x_linregr[1]) ]

     axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr,  l, c='k')
     axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr1, l, ls='--', c='gray') # LS_color[ils])
     axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr2, l, ls='--', c='gray') # LS_color[ils])
     axs[panels[ils][0], panels[ils][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1[0], y_linregr1[1]], [y_linregr2[0], y_linregr2[1]], facecolor = 'gray', alpha = LS_alpha*0.25) # LS_color[ils]

#    Scatter points
     for k in range(0, len(vx0)):
      axs[panels[ils][0], panels[ils][1]].scatter(vx0[k], vx1[k], s = scatter_size2, marker=cluster_symbol[k_labels[k]], color = cluster_color[k_labels[k]], alpha = cluster_alpha)

     for k in range(0,K_value):
      vx0_K = []
      vx1_K = []
      list_students = ''
      for kk in range(0, len(vx0)):
        if(k == k_labels[kk]):
          if( list_students == ''):
            list_students = vstudents[kk]
          else:
            list_students = list_students + ', ' + vstudents[kk]
          vx0_K.append(vx0[kk])
          vx1_K.append(vx1[kk])

      vx0_K =  np.array( vx0_K )
      vx1_K =  np.array( vx1_K )
      printt('k          : ' + str(k))
      printt('list_students : ' + str(list_students))
      printt('vx0_K      : ' + str(vx0_K))
      printt('vx1_K      : ' + str(vx1_K))

#     Check if all vx0_K values are identical
      if np.all(vx0_K == vx0_K[0]):
        print("Error: All vx0_K values are identical. Linear regression cannot be performed.")
        m_K = b_K = delta_m_K = delta_b_K = rvalue_K = np.nan
      else:
        m_K, b_K, rvalue_K, pvalue_K, delta_m_K = linregress(vx0_K, vx1_K)
        delta_b_K  = delta_m_K * np.sqrt( np.sum(vx0_K * vx0_K) / len(vx0_K) )

      printt('    m_K            : ' + str(m_K))
      printt('    delta_m_K      : ' + str(delta_m_K))
      printt('    b_K            : ' + str(b_K))
      printt('    delta_b_K      : ' + str(delta_b_K))
      printt('    rvalue_K            : ' + str(rvalue_K))
      printt('              ')

      marks_vs_LS.append([ data_marks[i, 0], data_marks[i, 1], data_marks[i, 2], clusters_index[k], list_students, k, Label_LS0, clusters_means[k, 0], clusters_uncert[k, 0], clusters_means[k, 1], clusters_uncert[k, 1], m_K, delta_m_K, b_K, delta_b_K, rvalue_K ] )

#       The confidence interval is well defined within the lines
#       ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
#        since x is always positive
      y_linregr_K  = [ lin_regr2(m_K, b_K, x_linregr[0]), lin_regr2(m_K, b_K, x_linregr[1]) ]
      y_linregr1_K = [ lin_regr2(m_K-delta_m_K, b_K-delta_b_K, x_linregr[0]), lin_regr2(m_K-delta_m_K, b_K-delta_b_K,  x_linregr[1]) ]
      y_linregr2_K = [ lin_regr2(m_K+delta_m_K, b_K+delta_b_K, x_linregr[0]), lin_regr2(m_K+delta_m_K, b_K+delta_b_K,  x_linregr[1]) ]

      axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr_K,  l, c=cluster_color[k])
      axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr1_K, l, ls='--', c=cluster_color[k])
      axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr2_K, l, ls='--', c=cluster_color[k])
      axs[panels[ils][0], panels[ils][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1_K[0], y_linregr1_K[1]],  [y_linregr2_K[0], y_linregr2_K[1]], facecolor = cluster_color[k], alpha = LS_alpha*0.25) # LS_color[ils]

#    Average mean
     axs[panels[ils][0], panels[ils][1]].scatter(x0_mean, x1_mean, s = size_mean, marker = mean_symbol, color = mean_color2)

     for k in range(0,K_value):
#     Average means for the clusters
      axs[panels[ils][0], panels[ils][1]].scatter(clusters_means[k, 0], clusters_means[k, 1], s = size_mean, marker=mean_symbol, color = cluster_color[k])

#    Text label
     axs[panels[ils][0], panels[ils][1]].text(xmin + 0.1*(xmax-xmin), ymax - 0.1*(ymax-ymin), Label_LS0,
        fontsize = labelsize,
        horizontalalignment='left',
        verticalalignment='top')

    axs[1, 0].set_xlim([ xmin, xmax] )
    axs[1, 0].set_ylim([ ymin, ymax] )

    axs[1, 0].set_xlabel('Points', fontsize=labelsize)
    axs[1, 0].set_ylabel('Marks', fontsize=labelsize)

    filename = output_participants_ls_marks + '/Fig_'+ data_marks[i, 0] + '_' + data_marks[i, 1] + '_' + data_marks[i, 2] + '_LS_K' +str(K_value) + '.png'
    plt.savefig(filename)
    printt('   Saving ' + filename)
    plt.savefig(filename)

 printt('   LS clustering done!')
 marks_vs_LS = np.stack(marks_vs_LS)
 print('marks_vs_LS : ', marks_vs_LS)

 printt('  ')
 printt('-------------------------------------------')
 printt(' Plots of students marks vs original LS done!')
 printt('-------------------------------------------')


 printt('-------------------------------------------')
 printt(' Plots of students marks vs PCs...')
 printt('-------------------------------------------')

 marks_vs_PC = []

 for i in range(1, Lmarks_activities):

   printt(' Evaluating activity : ' + data_marks[i, 2] + ' contained in the sheet ' + data_marks[i, 1] + ' of ' + data_marks[i, 0])
   printt('  ')

   fig, axs = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), sharex = True, sharey = True, layout='constrained')

   students = data_marks[0, 3:]
   marks    = np.array(data_marks[i, 3:], dtype=float)

#  Create two vectors with the marks and the LS
#  (if marks are correctly defined)
   vpc   = []
   vmarks = []
   for j in range(0, L):
     if(~np.isnan(marks[j])):
       vpc.append([proj[j,0], proj[j,1], proj[j,2], proj[j,3]])
       vmarks.append(marks[j])

   vpc    = np.stack(vpc)
   vmarks = np.array(vmarks)
   vmarks = vmarks.astype(float)

#  Plot limits
   xmin = np.min( np.min(proj) )
   xmax = np.max( np.max(proj) )

   ymin = -0.5/1.05
   ymax = 10
   if( np.max(vmarks) > ymax):
    ymax = np.max(vmarks)

   if ( xmin < 0 ):
    xmin = xmin * 1.05
   else:
    xmin = xmin * 0.95

   if ( xmax < 0 ):
    xmax = xmax * 0.95
   else:
    xmax = xmax * 1.05

   if ( ymin < 0 ):
    ymin = ymin * 1.05
   else:
    ymin = ymin * 0.95

   if ( ymax < 0 ):
    ymax = ymax * 0.95
   else:
    ymax = ymax * 1.05

   printt( ' xmin : ' + str(xmin))
   printt( ' xmax : ' + str(xmax))
   printt( ' ymin : ' + str(ymin))
   printt( ' ymax : ' + str(ymax))
   printt('  ')

   if( K_value <= len(vmarks) ):
    for ipc in range(0,K):
     print('panels : ', panels[ipc][0], panels[ipc][1])
     axs[panels[ipc][0], panels[ipc][1]].tick_params(axis='both', which='major', labelsize=ticksize)
#----------------------------------------------------------
#    Scatter of the marks as a function of the LS  points
     printt(' PC : ' + str(ipc))
     printt('  ')

     Label_PC0 = Label_PCPC[ipc]
     printt(' ( marks vs ' + Label_PC0 + ' )')
     printt('  ')

     vx0 = vpc[:,ipc]
     vx1 = vmarks

     printt('   Clustering for ' + Label_PC0 + '...')

     array=vx1.reshape(-1, 1)

#    Create and fit the K-Means model with the optimal number of clusters
     kmeans = KMeans(n_clusters=K_value, init='k-means++', random_state=42)
     kmeans.fit(array)

#    Get the cluster labels for each data point
     k_labels = kmeans.labels_

#    Average means and uncertainties for the clusters
     clusters_means  = [ ]
     clusters_uncert = [ ]
     clusters_index  = [ ]
     for k in range(0, K_value):
      values = []

      for kk in range(0,len(k_labels)):
        if(k == k_labels[kk]):
          values.append( [ vx0[kk], vx1[kk] ] )
      shape = np.shape(values)
      clusters_index.append(shape[0])
      values = np.stack(values)

#     Average means
      x0_mean = np.mean( values[:,0] )
      x1_mean = np.mean( values[:,1] )

#     Uncertainties
      confidence_interval=st.t.interval(confidence=0.95, df=len(values[:,0])-1, loc=x0_mean, scale=st.sem(values[:,0]))
      dx0 = 0.5*(confidence_interval[1]-confidence_interval[0])
      confidence_interval=st.t.interval(confidence=0.95, df=len(values[:,1])-1, loc=x0_mean, scale=st.sem(values[:,1]))
      dx1 = 0.5*(confidence_interval[1]-confidence_interval[0])
      clusters_means.append([ x0_mean, x1_mean ])
      clusters_uncert.append( [dx0, dx1] )

     clusters_means  = np.stack(clusters_means)
     clusters_uncert = np.stack(clusters_uncert)

#    Linear regression
     if np.all(vx0 == vx0[0]):
      print("Error: All vx0 values are identical. Linear regression cannot be performed.")
      m = b = delta_m = delta_b = rvalue = pvalue = np.nan
     else:
      m, b, rvalselected_kue, pvalue, delta_m = linregress(vx0, vx1)
#     Compute uncertainty in intercept (y0)
      delta_b  = delta_m * np.sqrt( np.sum(vx0 * vx0) / len(vx0) )

#    Average means
     x0_mean = np.mean( vx0 )
     x1_mean = np.mean( vx1 )

#    Uncertainties
     confidence_interval=st.t.interval(confidence=0.95, df=len(vx0)-1, loc=x0_mean, scale=st.sem(vx0))
     dx0 = 0.5*(confidence_interval[1]-confidence_interval[0])
     confidence_interval=st.t.interval(confidence=0.95, df=len(vx1)-1, loc=x1_mean, scale=st.sem(vx1))
     dx1 = 0.5*(confidence_interval[1]-confidence_interval[0])

     marks_vs_PC.append([ data_marks[i, 0], data_marks[i, 1], data_marks[i, 2], len(vx0), '', -1, Label_PC0,  x0_mean, dx0, x1_mean, dx1, m, delta_m, b, delta_b, rvalue ] )

     def lin_regr2(m, b, x):
      return m*x+b

#   The confidence interval for x > 0
#   ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
     x_linregr  = [ 0, xmax ]
     y_linregr  = [ lin_regr2(m, b, x_linregr[0]), lin_regr2(m, b, x_linregr[1]) ]
     y_linregr1 = [ lin_regr2(m-delta_m, b-delta_b, x_linregr[0]), lin_regr2(m-delta_m, b-delta_b, x_linregr[1]) ]
     y_linregr2 = [ lin_regr2(m+delta_m, b+delta_b, x_linregr[0]), lin_regr2(m+delta_m, b+delta_b, x_linregr[1]) ]

     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr,  l, c='k')
     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr1, l, ls='--', c='gray')
     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr2, l, ls='--', c='gray')
     axs[panels[ipc][0], panels[ipc][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1[0], y_linregr1[1]], [y_linregr2[0], y_linregr2[1]], facecolor = 'gray', alpha = LS_alpha*0.25)

#   The confidence interval for x < 0
#   ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
     x_linregr  = [ xmin, 0 ]
     y_linregr  = [ lin_regr2(m, b, x_linregr[0]), lin_regr2(m, b, x_linregr[1]) ]
     y_linregr1 = [ lin_regr2(m-delta_m, b+delta_b, x_linregr[0]), lin_regr2(m-delta_m, b+delta_b, x_linregr[1]) ]
     y_linregr2 = [ lin_regr2(m+delta_m, b-delta_b, x_linregr[0]), lin_regr2(m+delta_m, b-delta_b, x_linregr[1]) ]

     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr,  l, c='k')
     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr1, l, ls='--', c='gray')
     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr2, l, ls='--', c='gray')
     axs[panels[ipc][0], panels[ipc][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1[0], y_linregr1[1]], [y_linregr2[0], y_linregr2[1]], facecolor = 'gray', alpha = LS_alpha*0.25)

#    Scatter points
     for k in range(0, len(vx0)):
      axs[panels[ipc][0], panels[ipc][1]].scatter(vx0[k], vx1[k], s = scatter_size2, marker=cluster_symbol[k_labels[k]], color = cluster_color[k_labels[k]], alpha = cluster_alpha)

     for k in range(0,K_value):
      vx0_K = []
      vx1_K = []
      for kk in range(0, len(vx0)):
        if(k == k_labels[kk]):
          vx0_K.append(vx0[kk])
          vx1_K.append(vx1[kk])

      vx0_K =  np.array( vx0_K )
      vx1_K =  np.array( vx1_K )

      printt('k     : ' + str(k))
      printt('vx0_K : ' + str(vx0_K))
      printt('vx1_K : ' + str(vx1_K))

      if np.all(vx0_K == vx0_K[0]):
        print("Error: All vx0_K values are identical. Linear regression cannot be performed.")
        m_L = b_K = delta_m_K = delta_b_K = delta_b_K = rvalue_K = np.nan
      else:
        m_K, b_K, rvalue_K, pvalue_K, delta_m_K = linregress(vx0_K, vx1_K)
        delta_b_K  = delta_m_K * np.sqrt( np.sum(vx0_K * vx0_K) / len(vx0_K) )

      printt('    m_K            : ' + str(m_K))
      printt('    delta_m_K      : ' + str(delta_m_K))
      printt('    b_K            : ' + str(b_K))
      printt('    delta_b_K      : ' + str(delta_b_K))
      printt('    rvalue_K            : ' + str(rvalue_K))
      printt('              ')

      marks_vs_PC.append([ data_marks[i, 0], data_marks[i, 1], data_marks[i, 2], clusters_index[k], '', k, Label_PC0, clusters_means[k, 0], clusters_uncert[k, 0], clusters_means[k, 1], clusters_uncert[k, 1], m_K, delta_m_K, b_K, delta_b_K, rvalue_K ] )

#     The confidence interval for x > 0
#     ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
      x_linregr  = [ 0, xmax ]
      y_linregr_K  = [ lin_regr2(m_K, b_K, x_linregr[0]), lin_regr2(m_K, b_K, x_linregr[1]) ]
      y_linregr1_K = [ lin_regr2(m_K-delta_m_K, b_K-delta_b_K, x_linregr[0]), lin_regr2(m_K-delta_m_K, b_K-delta_b_K, x_linregr[1]) ]
      y_linregr2_K = [ lin_regr2(m_K+delta_m_K, b_K+delta_b_K, x_linregr[0]), lin_regr2(m_K+delta_m_K, b_K+delta_b_K, x_linregr[1]) ]

      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr_K,  l, c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr1_K, l, ls='--', c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr2_K, l, ls='--', c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1_K[0], y_linregr1_K[1]], [y_linregr2_K[0], y_linregr2_K[1]], facecolor = cluster_color[k], alpha = LS_alpha*0.25)

#     The confidence interval for x < 0
#     ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
      x_linregr  = [ xmin, 0 ]
      y_linregr_K  = [ lin_regr2(m_K, b_K, x_linregr[0]), lin_regr2(m_K, b_K, x_linregr[1]) ]
      y_linregr1_K = [ lin_regr2(m_K-delta_m_K, b_K-delta_b_K, x_linregr[0]), lin_regr2(m_K-delta_m_K, b_K+delta_b_K, x_linregr[1]) ]
      y_linregr2_K = [ lin_regr2(m_K+delta_m_K, b_K+delta_b_K, x_linregr[0]), lin_regr2(m_K+delta_m_K, b_K-delta_b_K, x_linregr[1]) ]

      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr_K,  l, c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr1_K, l, ls='--', c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr2_K, l, ls='--', c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1_K[0], y_linregr1_K[1]], [y_linregr2_K[0], y_linregr2_K[1]], facecolor = cluster_color[k], alpha = LS_alpha*0.25)

#    Average mean
     axs[panels[ipc][0], panels[ipc][1]].scatter(x0_mean, x1_mean, s = size_mean, marker = mean_symbol, color = mean_color2)

     for k in range(0,K_value):
#     Average means for the clusters
      axs[panels[ipc][0], panels[ipc][1]].scatter(clusters_means[k, 0], clusters_means[k, 1], s = size_mean, marker=mean_symbol, color = cluster_color[k])

#    Text label
     axs[panels[ipc][0], panels[ipc][1]].text(xmin + 0.1*(xmax-xmin), ymax - 0.1*(ymax-ymin), Label_PC0,
        fontsize = labelsize,
        horizontalalignment='left',
        verticalalignment='top')

    axs[1, 0].set_xlim([ xmin, xmax] )
    axs[1, 0].set_ylim([ ymin, ymax] )

    axs[1, 0].set_xlabel('Points', fontsize=labelsize)
    axs[1, 0].set_ylabel('Marks', fontsize=labelsize)

    filename = output_participants_pc_marks + '/Fig_'+ data_marks[i, 0] + '_' + data_marks[i, 1] + '_' + data_marks[i, 2] + '_PC_K' +str(K_value) + '.png'
    plt.savefig(filename)
    printt('   Saving ' + filename)
    plt.savefig(filename)
    printt('  ')

 marks_vs_PC = np.stack(marks_vs_PC)

 printt('-------------------------------------------')
 printt(' Plots of students marks vs PCs done!')
 printt('-------------------------------------------')
 printt('  ')
 printt('===========================================')
 printt('ANALYSIS OF THE MARKS OF THE STUDENTS DONE!')
 printt('===========================================')
 printt('  ')
 printt('  ')
 printt('  ')
 printt('  ')
#
#
# Get the current date and time
current_datetime = datetime.datetime.now()
#
# Format the date and time as a string
formatted_datetime = current_datetime.strftime("%Y-%m-%d %H:%M:%S")
printt('Program continues on: ' + formatted_datetime)
printt(' ')
#
#===================================================
#===================================================
#===================================================
#===================================================
# SUMMARY REPORT
#===================================================
#===================================================
# Most of the results of the analysis are
# included in the summary report, which is saved in
# docx and pdf formats.
# Further information can be found in the
# log file (output/chaea3s.log).
#
# If the number of students is larger than
# Lmax = 500, then only the first 150 are saved.
#
if ( L > 150):
  Lsave = 150
else:
  Lsave = L

#===================================================
# File location and title
printt('===========================================')
printt('SUMMARY REPORT...')
printt('===========================================')
printt('output : ' + output)
folder_name = os.path.basename(output)
document    = Document()
#
#===================================================
# TITLE AND INTRO
#===================================================
# Title
title = 'Summary report of CHAEA learning styles by CHAEA'
title += u'\u00B3'  # Adding superscript 3
title += 'S package'

document.add_heading(title, level=0)

document.add_paragraph('This report contains the most important results of the analysis that is conducted to unveil the learning styles that are present in the group of students under study. The analysis is based on the learning styles considered by CHAEA: activist, reflector, theorist, and pragmatist. Unless otherwise stated, the uncertainties throughout the document (in parenthesis) have been obtained using a t-Student distribution with a confidence interval of 95%.')

if ( K_value == 0 ):
# No marks provided, i.e., no clustering conducted
  document.add_paragraph('The report is structured as follows. First, Section 1 is devoted to the individual and global statistical analysis. Here, the importance of the different learning styles for each individual student can be found (both quantitatively as well as qualitatively). Then, average means and confidence intervals are presented, along with the affinities and the Probability Density Functions. Second, Section 2 discusses the principal component analysis. The eigenvalues and eigenvectors of the covariance matrix are first introduced. Subsequently, the learning styles of the students in the principal components basis set is presented. Next, a reduced dimensional representation of the data is conducted. Third, the participation ratios are presented in Section 3. The values for the original (activist, reflector, theorist, and pragmatist) and in the principal components basis sets are listed. Here, a statistical analysis of the distribution of the participation ratios is performed. The report concludes in Section 4 with a brief summary of the main characteristics of the students, along with some recommendations.')
else:
  if ( K_value == 1 ):
# Marks provided, but no clustering conducted as K=1
    document.add_paragraph('The report is structured as follows. First, Section 1 is devoted to the individual and global statistical analysis. Here, the importance of the different learning styles for each individual student can be found (both quantitatively as well as qualitatively). Then, average means and confidence intervals are presented, along with the affinities and the Probability Density Functions. Second, Section 2 discusses the principal component analysis. The eigenvalues and eigenvectors of the covariance matrix are first introduced. Subsequently, the learning styles of the students in the principal components basis set is presented. Next, a reduced dimensional representation of the data is conducted. Third, the participation ratios are presented in Section 3. The values for the original (activist, reflector, theorist, and pragmatist) and in the principal components basis sets are listed. Here, a statistical analysis of the distribution of the participation ratios is performed. Section 4 is devoted to the analysis of marks of the whole group of students, and their relationship to the CHAEA learning styles and the principal components. The report concludes in Section 5 with a brief summary of the main characteristics of the students, along with some recommendations.')
  else:
#   Marks provided, and no clustering conducted as K>1
    document.add_paragraph('The report is structured as follows. First, Section 1 is devoted to the individual and global statistical analysis. Here, the importance of the different learning styles for each individual student can be found (both quantitatively as well as qualitatively). Then, average means and confidence intervals are presented, along with the affinities and the Probability Density Functions. Second, Section 2 discusses the principal component analysis. The eigenvalues and eigenvectors of the covariance matrix are first introduced. Subsequently, the learning styles of the students in the principal components basis set is presented. Next, a reduced dimensional representation of the data is conducted. Third, the participation ratios are presented in Section 3. The values for the original (activist, reflector, theorist, and pragmatist) and in the principal components basis sets are listed. Here, a statistical analysis of the distribution of the participation ratios is performed. Section 4 presents the clustering analysis of the data based on K-means algorithm with K=' + str(K_value)+ ' clusters. In this section, the relationship between the students marks, and the CHAEA learning styles and the principal components The report concludes in Section 5 with a brief summary of the main characteristics of the students, along with some recommendations.')

# Further information in our reference
our_reference = document.add_paragraph()

our_reference.add_run("Further information at: ").italic = True

document.add_paragraph(' ')
our_reference.add_run("J. Ablanque, V. Gabaldon, P. Almendros, J. C. Losada, R. M. Benito, and F. Revuelta. ").italic = True
document.add_paragraph(' ')
#
#
our_reference.add_run("CHAEA3S: A software package for comprehensive analysis of learning styles and academic performance. ").italic = True
document.add_paragraph()

our_reference.add_run("PLOS ONE (2026).").italic = True

#our_reference.add_run("If you have any comments or suggestions, feel free to reach out by sending an email with your feedback to fabio.revuelta@upm.es.").italic = True

document.add_page_break()
#===================================================
# SECTION 1
#===================================================
nsec    = 1
nsubsec = 1
#==================================================
document.add_heading(str(nsec) + '. Individual and global statistical analysis', level=1)
printt('******************************************')
printt(str(nsec) + '. Individual and global statistical analysis')
printt('******************************************')
#==================================================
document.add_paragraph('In this section we present the individual and global statistical analysis of the learning styles as originally defined in CHAEA (activist, reflector, theorist, and pragmatist). This section is divided in two parts. First, the (quantitative and qualitative) importance of the different learning styles for each individual student can be found. Second, a global analysis is performed, where average means, confidence intervals, affinities, and the probability density functions can be found.')
#
#---------------------------------------------------
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Individual analysis', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Individual analysis')
ntab = 1
ntab = 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the number of points...')
printt('-------------------------------------------')
document.add_heading('Quantitative description of the learning styles for each individual student', level=3)

if ( L <= Lsave ):
  document.add_paragraph('Table ' + str(ntab) + ' shows the number of points (from 0 to 20) that each student gets in CHAEA for the different learning styles.')
else:
  document.add_paragraph('Table ' + str(ntab) + ' shows the number of points (from 0 to 20) that each student gets in CHAEA for the different learning styles (only the first ' + str(Lsave) + ' students are included).')

# Table 1 title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Points related the learning styles for each of the students.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

# Table 1
t = document.add_table(Lsave+1, K+1)

# Table header
t.cell(0,0).text = 'Student'
for j in range(0,K):
  t.cell(0,j+1).text = Label_LS_print[j]

for i in range(0,Lsave):
  t.cell(i+1,0).text = students[i]
  for j in range(0,K):
    t.cell(i+1,j+1).text = str(data[i,j])

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the number of points done!')
printt('-------------------------------------------')
printt(' ')

#Lerror = len(datain_error_students)
if ( datain_error_students != ['.'] ):
  document.add_paragraph()
  document.add_paragraph('The following students were not considered in the analysis because of errors in the CHAEA files:')
  document.add_paragraph(datain_error_students)

document.add_page_break()

ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the individual tendencies...')
printt('-------------------------------------------')
document.add_heading('Qualitative description of the learning styles for each individual student', level=3)

if ( L <= Lsave ):
  document.add_paragraph('Table ' + str(ntab) + ' shows the individual qualitative tendency (very low, low, moderate, high, or very high) that each student has towards each of the learning styles.')
else:
  document.add_paragraph('Table ' + str(ntab) + ' shows the individual qualitative tendency (very low, low, moderate, high, or very high) that each student has towards each of the learning styles (only the first ' + str(Lsave) + ' students are included).')



# Table 2 title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Tendency towards the learning styles for each of the students.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(Lsave+1, K+1)

# Table header
t.cell(0,0).text = 'Student'
for j in range(0,K):
  t.cell(0,j+1).text = Label_LS_print[j]

for i in range(0,Lsave):
  t.cell(i+1,0).text = students[i]
  for j in range(0,K):
    t.cell(i+1,j+1).text = tendency_long_name(tendency_matrix_all[i][j])

document.add_page_break()
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the individual tendencies done!')
printt('-------------------------------------------')
printt(' ')







nsubsec += 1
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Global analysis', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Global analysis')
ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the tendencies...')
printt('-------------------------------------------')
document.add_heading('Global tendencies towards the learning styles of the students (I)', level=3)

document.add_paragraph('Table ' + str(ntab) + ' presents the number and percentage of students with the same tendency (very low, low, moderate, high, or very high) towards each of the learning styles. The bottom line shows the average tendency. These results are also shown as a barr graphic in Fig. 1.')

# Table 4 title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Number and percentage of students with the same tendency towards each of the learning styles.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+3, 6)

t.cell(0,0).text = 'Tendency'
for j in range(0,5):
  t.cell(0,j+1).text = Label_tendencies_print[j]
  t.cell(1,j+1).text = 'No.  %'

for i in range(0,K):
  t.cell(i+2,0).text = Label_LS_print[i]
  for j in range(0,5):
    t.cell(i+2,j+1).text = str(int(tendency_matrix[i, j])) + '   ' + str( round(tendency_matrix_percentage[i, j], 1))

t.cell(K+2,0).text = 'Total'
for j in range(0,5):
  t.cell(K+2,j+1).text = str(int(Ntot_tendency[j])) + '   ' + str(round(percent_tendency[j], 1))

document.add_page_break()
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the tendencies done!')
printt('-------------------------------------------')
printt(' ')

nfig = 1
# Figure 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the tendencies...')
printt('-------------------------------------------')
document.add_picture(filename_tendencies, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Bar graphic showing the percentage of students with a very low (leftmost), low (left), moderate (middle), high (right) and very high (rightmost) tendency towards the activist (red), reflector (green), theorist (purple) and pragmatist (orange) learning styles. The shown results correspond to the values of Table 3.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

document.add_page_break()

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the tendencies done!')
printt('-------------------------------------------')



ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the global tendencies...')
printt('-------------------------------------------')
document.add_heading('Global tendencies towards the learning styles of the students (II)', level=3)

document.add_paragraph('Table ' + str(ntab) + ' shows the average mean and the uncertainties for the different learning styles, along with the corresponding qualitative tendency, and the corresponding affinity. For the shake of clarity, the average profile of the students given by the average means and the corresponding uncertainties is shown in Fig. 2, while the affinities are represented as a barr graphic in Fig. 3.')
#
#The bottom line shows the average tendency.
#
# Table 4 title
table_title = document.add_paragraph('Table 4. Numerical average values of each learning style with the uncertainties (in parenthesis) along with their qualitative tendency, and the corresponding affinity.')
# The last line provides the average results for the whole data set (no tendency is provided in this case).')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+1, 4)

# Table header
t.cell(0,0).text = 'Learning style'
t.cell(0,1).text = 'Average mean (Uncertainty)' #'μ (Δμ)'
t.cell(0,2).text = 'Tendency'
t.cell(0,3).text = 'Affinity (%)'

for j in range(0,K):
  t.cell(j+1,0).text = Label_LS_print[j]
  xdx = mean_uncert(xmean[j], dxmean[j])
  t.cell(j+1,1).text = xdx
  tendency = tendency_long_name(tendency_intermediate(Label_LS[j], xmean[j]))
  t.cell(j+1,2).text = tendency
  t.cell(j+1,3).text = str(round(affinity[j], 1))

document.add_page_break()
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the global tendencies done!')
printt('-------------------------------------------')
printt(' ')

# Figure 2
nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the average profile...')
printt('-------------------------------------------')
document.add_picture(filename_average_profile, width=Inches(5))
x
# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Average profile of the learning styles of the students.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the average profile done!')
printt('-------------------------------------------')


# Figure 3
nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the affinities...')
printt('-------------------------------------------')

document.add_picture(filename_affinity, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure 3. Affinity of the learning styles. The bar graphic gives the percentage of students that have a noticeable tendency towards the activist (red), reflector (green), theorist (purple) and pragmatist (orange) learning styles.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the affinities done!')
printt('-------------------------------------------')





nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the PDFs of the LSs...')
printt('-------------------------------------------')

document.add_paragraph('Figure ' + str(nfig) + ' shows histograms with the probability distributions of the results of Table 1. The continuous lines show fittings with Weibull distributions with the parameters shown in Table 5. These fittings have been performed on the corresponding cumulative distributions given by the staircases shown in Figure 5.')

document.add_picture(filename_statistics_ls, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure 4. Histograms showing the probability distributions for the activist (red), reflector (green), theorist (purple), and pragmatist (yellow) learning styles. The continuous lines show fittings provided by the Weibull distributions with the parameters contained in Table 5. The vertical dotted lines mark the average values, and the shaded areas around them the corresponding confidence intervals.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the PDFs of the LSs done!')
printt('-------------------------------------------')

nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with W(s) of the LSs...')
printt('-------------------------------------------')

document.add_picture(filename_statistics_ls_w, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure 5. Same as Fig. 4 for the cumulative distributions (staircases).')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

document.add_page_break()
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with W(s) of the LSs done!')
printt('-------------------------------------------')


ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the Weibull parameters for the PDFs of the LSs...')
printt('-------------------------------------------')

# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Parameters of the Weibull distributions that fit the probability distributions (histograms) shown in Figs. 4 and 5. The location parameter is set equal to θ = 0.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+1, 3)

# Table header
t.cell(0,0).text = 'Learning style'
t.cell(0,1).text = 'α'
t.cell(0,2).text = 'k'

for j in range(0,K):
  t.cell(j+1,0).text = Label_LS_print[j]
  [alpha_Weibull, k_Weibull] = parameters_Weibull_LS20[j]
  t.cell(j+1,1).text = str(round(alpha_Weibull,2))
  t.cell(j+1,2).text = str(round(k_Weibull,2))

document.add_page_break()
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the Weibull parameters for the PDFs of the LSs done!')
printt('-------------------------------------------')





#---------------------------------------------------
# OPTIONAL SUBSECTION WITH THE PROFESSOR'S
# LEARNING STYLES

if(Lprof == 0):
  printt('No results for the ' + prof_ref + 's are provided.')
else:
  nsubsec += 1
  ntab += 1
  if(Lprof == 1):
    nsubsec += 1
    document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Summary of the learning styles of the ' + prof_ref, level=2)
    printt(str(nsec) + '.' + str(nsubsec) + ' Summary of the learning styles of the ' + prof_ref)
    document.add_paragraph('In this section we present the learning styles of the ' + prof_ref + ' as originally defined in CHAEA (activist, reflector, theorist, and pragmatist), and compare them to those of the students.')
  else:
    document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Summary of the learning styles of the ' + prof_ref + 's', level=2)
    printt(str(nsec) + '.' + str(nsubsec) + ' Summary of the learning styles of the ' + prof_ref + 's')
    document.add_paragraph('In this section we present the learning styles of the ' + prof_ref + 's as originally defined in CHAEA (activist, reflector, theorist, and pragmatist), and compare them to those of the students.')

  printt('-------------------------------------------')
  printt(' Table ' + str(ntab) + ' with the number of points of the ' + prof_ref + '(s)...')
  printt('-------------------------------------------')
  if(Lprof == 1):
    document.add_heading('Quantitative description of the learning styles for the ' + prof_ref, level=3)
    document.add_paragraph('Table ' + str(ntab) + ' shows the number of points (from 0 to 20) that the ' + prof_ref + ' gets in CHAEA for the different learning styles.')
  else:
    document.add_heading('Quantitative description of the learning styles for each individual ' + prof_ref, level=3)
    document.add_paragraph('Table ' + str(ntab) + ' shows the number of points (from 0 to 20) that each ' + prof_ref + ' gets in CHAEA for the different learning styles.')

  # Table 6 title
  if(Lprof == 1):
    table_title = document.add_paragraph('Table ' + str(ntab) + '. Points related the learning styles for the ' + prof_ref + '.')
  else:
    table_title = document.add_paragraph('Table ' + str(ntab) + '. Points related the learning styles for each of the ' + prof_ref + 's.')
  table_title.alignment = 1  # Center alignment
  title_run = table_title.runs[0]
  title_run.bold = True

  # Table 6
  t = document.add_table(Lprof+1, K+1)

  # Table header
  t.cell(0,0).text = Prof_ref
  for j in range(0,K):
    t.cell(0,j+1).text = Label_LS_print[j]

  for i in range(0,Lprof):
    t.cell(i+1,0).text = profs[i]
    for j in range(0,K):
      t.cell(i+1,j+1).text = str(data_prof[i,j])

  if ( datain_error_profs != ['.'] ):
    document.add_paragraph()
    document.add_paragraph('The following ' + prof_ref + '(s) were not considered in the analysis because of errors in the CHAEA files:')
    document.add_paragraph(datain_error_profs)
  printt('-------------------------------------------')
  printt(' Table ' + str(ntab) + ' with the number of points of the ' + prof_ref + '(s) done!')
  printt('-------------------------------------------')
  printt(' ')

  ntab += 1
  printt('-------------------------------------------')
  printt(' Table ' + str(ntab) + ' with the individual tendencies of the ' + prof_ref + '(s)...')
  printt('-------------------------------------------')

  if(Lprof == 1):
    document.add_heading('Qualitative description of the learning styles for the ' + prof_ref, level=3)
    document.add_paragraph('Table ' + str(ntab) + ' shows the qualitative tendency (very low, low, moderate, high, or very high) that the ' + prof_ref + ' has towards each of the learning styles.')
  else:
    document.add_heading('Qualitative description of the learning styles for each individual ' + prof_ref, level=3)
    document.add_paragraph('Table ' + str(ntab) + ' shows the individual qualitative tendency (very low, low, moderate, high, or very high) that each ' + professor + ' has towards each of the learning styles.')

  # Table 7 title
  table_title = document.add_paragraph('Table ' + str(ntab) + '. Tendency towards the learning styles for each of the ' + prof_ref + 's.')
  table_title.alignment = 1  # Center alignment
  title_run = table_title.runs[0]
  title_run.bold = True

  t = document.add_table(Lprof+1, K+1)
  tendency_prof = ['', '', '', '', '']

  # Table header
  t.cell(0,0).text = Prof_ref
  for j in range(0,K):
    t.cell(0,j+1).text = Label_LS_print[j]

  for i in range(0,Lprof):
    t.cell(i+1,0).text = profs[i]

    tendency_prof[0] = scatter_tendency('Activist',   data_prof[i, 0])
    tendency_prof[1] = scatter_tendency('Reflector',  data_prof[i, 1])
    tendency_prof[2] = scatter_tendency('Theorist',   data_prof[i, 2])
    tendency_prof[3] = scatter_tendency('Pragmatist', data_prof[i, 3])

    for j in range(0,K):
      t.cell(i+1,j+1).text = tendency_long_name(tendency_prof[j])

  document.add_page_break()
  printt('-------------------------------------------')
  printt(' Table ' + str(ntab) + ' with the individual tendencies of the ' + prof_ref + '(s) done!')
  printt('-------------------------------------------')
  printt(' ')


  printt('-------------------------------------------')
  printt(' Figures with the average profile of the students + the profile of the ' + prof_ref + 's ...')
  printt('-------------------------------------------')
  nfig_prof0 = nfig + 1
  if(Lprof == 1):
    document.add_heading('Comparison of the learning-styles profile of the students and of the ' + prof_ref, level=3)
    document.add_paragraph('Figure ' + str(nfig+1) + ' shows the learning style profiles of the ' + prof_ref + ' superimposed with the average profile of the students shown in Fig. 2.')
  else:
    document.add_heading('Comparison of the learning-styles profile of the students and of each individual ' + prof_ref, level=3)
    document.add_paragraph('Figures ' + str(nfig+1) + ' to ' + str(nfig + Lprof) +' show the learning style profiles of the ' + prof_ref + 's superimposed with the average profile of the students shown in Fig. 2.')

  for i_prof in range(0,Lprof):
    profi = profs[i_prof]
    nfig  = nfig + 1
    filename_average_profile_prof=output_average_profile_prof + '/Fig_averageprofile_prof_' + profi +'.png'
    document.add_picture(filename_average_profile_prof, width=Inches(5))

#   Figure title
    if(i_prof == 0):
      figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Profile of the learning styles of the ' + prof_ref + ' "' + profi + '" (green continuous line) superimposed with the average profile of the students shown in Fig. 2.')
    else:
      figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Same as Fig. ' + str(nfig_prof0) + ' for the ' + prof_ref + ' of ' + profi + '.')
    figure_title.alignment = 1  # Center alignment
    figure_run = figure_title.runs[0]
    figure_run.bold = True

  printt('-------------------------------------------')
  printt(' Figure with the average profile of the students + the profile of the ' + professor + 's done!')
  printt('-------------------------------------------')










document.add_page_break()
#===================================================
# SECTION 2
#==================================================
nsec   += 1
nsubsec = 1
document.add_heading(str(nsec) + '. Principal component analysis', level=1)
printt('******************************************')
printt(str(nsec) + '. Principal component analysis')
printt('******************************************')
#==================================================

document.add_paragraph('This section presents the principal component analysis of the learning styles. It is organized as follows. First, the eigenvalues and eigenvectors of the covariance matrix are presented. Second, the description of the learning styles of the students in the principal components basis set is discussed. Finally, a reduced dimensional representation of the data is conducted.')

document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Structure of the principal components', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Structure of the principal components')

document.add_heading('Eigenvalues and dispersion of the covariance matrix', level=3)

ntab += 1
document.add_paragraph('In this section, we discuss the essentials of the principal components. For this purpose, Table ' + str(ntab) + ' shows the four eigenvalues of the covariance matrix. Here, not only their values are listed but also their contribution to the total dispersion (given by the trace of the covariance matrix tr(K)='+str(round(trace_covX,2))+', which equals the sum of all the eigenvalues) as a percentage. Likewise, the dispersion accounted by solely the principal component with the largest eigenvalue Σ_0(%)=λ_0*100/tr(K), and by combining the principal components with the two, three, and four largest eigenvalues Σ_1(%)=(λ_0+λ_1)*100/tr(K), Σ_2(%)=(λ_0+λ_1+λ_2)*100/tr(K), and by four Σ_3(%)=(λ_0+λ_1+λ_2+λ_3)*100/tr(K)=100%, respectively, are given.')

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the eigenvalues of the PCs...')
printt('-------------------------------------------')

# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Eigenvalues of the covariance matrix given by their corresponding values λ_i and as a percentage of the total dispersion, and percentage of total dispersion Σ_i accounted by combination of the principal components with the eigenvalues λ_j, being j≤i.')

table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+1, 4)

# Table header
t.cell(0,0).text = 'Principal component (i)'
t.cell(0,1).text = 'λ_i'
t.cell(0,2).text = 'λ_i(%)'
t.cell(0,3).text = 'Σ_i(%)'

sigma_percent = 0.0
for j in range(0,K):
  t.cell(j+1,0).text = Label_PC[j]
  lj            = eigenValues[j]
  lj_percent    = 100/trace_covX * lj
  sigma_percent = sigma_percent + lj_percent
  t.cell(j+1,1).text = str(round(lj,2))
  t.cell(j+1,2).text = str(round(lj_percent,1))
  t.cell(j+1,3).text = str(round(sigma_percent,1))
  if(j==2):
    sigma01 = str(round(sigma_percent,1))

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the eigenvalues of the PCs done!')
printt('-------------------------------------------')

document.add_page_break()
document.add_heading('Eigenvectors of the covariance matrix', level=3)
ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the structure of the eigenvectors...')
printt('-------------------------------------------')

document.add_paragraph('Table ' + str(ntab) + ' shows the structure of the eigenvectors of the covariance matrix in the basis set of CHAEA learning styles. Each cell contains the percentage of the eigenvector of the principal component 0, 1, 2, and 3 that is projected on the corresponding learning style (activist, reflector, theorist, and pragmatist).')

# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Structure (as percentages) of the eigenvectors of the covariance matrix in the basis set of CHAEA learning styles.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+1, 5)

# Table header
t.cell(0,0).text = 'Principal component'

for j in range(0,K):
  t.cell(0,j+1).text = Label_LS[j]

sigma_percent = 0.0
for i in range(0,K):
  t.cell(i+1,0).text = Label_PC[i]
  for j in range(0,K):
    t.cell(i+1,j+1).text = str(round(eigenVectors_percentage[j,i],1))

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the structure of the eigenvectors done!')
printt('-------------------------------------------')







document.add_page_break()
nsubsec += 1
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Individual analysis of the principal components', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Individual analysis of the principal components')
ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the individual analysis of PCs...')
printt('-------------------------------------------')
document.add_heading('Quantitative description of the students in the basis set of the principal components', level=3)

document.add_paragraph('Table ' + str(ntab) + ' shows (as percentages) the structure of the learning styles in the basis set formed by the principal components. The table also includes the percentage of the learning styles of each student that is described by combining the two (sum of the percentages for the principal components 0 and 1) or three (sum of the percentages for the principal components 0, 1, and 2) principal components with the largest eigenvalues. Recall that when the four principal components are considered, 100% of the learning style of the student is reproduced.')

# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Structure of the learning styles of each of the students in the basis set of principal components.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

# Table 1
t = document.add_table(Lsave+1, 7)

# Table header
t.cell(0,0).text = 'Student'
for j in range(0,K):
  t.cell(0,j+1).text = Label_PC[j]

#Σ
t.cell(0,5).text = '0+1'
t.cell(0,6).text = '0+1+2'

for i in range(0,Lsave):
  t.cell(i+1,0).text = students[i]
  for j in range(0,K):
    t.cell(i+1,j+1).text = str(round(prob[i,j], 2))

  t.cell(i+1,5).text = str(round(prob[i,0]+prob[i,1], 2))
  t.cell(i+1,6).text = str(round(prob[i,0]+prob[i,1]+prob[i,2], 2))

document.add_page_break()
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the individual analysis of PCs done!')
printt('-------------------------------------------')
printt(' ')


































nsubsec += 1
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Global analysis of the principal components', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Global analysis of the principal components')
nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the 4D plot of the LSs...')
printt('-------------------------------------------')
nfig_4d = nfig

document.add_paragraph('Figure ' + str(nfig) + ' shows a four-dimensional representation of the learning styles of the students as a function of the contributions to three of the CHAEA learning styles. The tendency towards the remaining learning style is implicitly shown in the size and color shades of the points (bigger and darker colors imply a larger tendency). The principal directions shown as continuous lines (PC0 , red; PC1, green; PC2, purple; PC3, orange) emerge from centroid, which is given by the average mean of the data set and is shown as a black star. The projection of the points and the average mean is shown in gray. The projection confidence interval of the average mean is also shown. The origin, which is defined as (0, 0, 0, 0) in the 4D space of the learning styles is shown as a black triangle. The colored points correspond to the maximal pure learning styles (e.g., (20, 0, 0, 0) for activist (red), or (0, 20, 0, 0) for reflector (green) are also marked.')
document.add_picture(filename_ls_3d, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Four-dimensional representation of the learning styles of the students as a function of the CHAEA learning styles.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

document.add_page_break()
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the 4D plot of the LSs done!')
printt('-------------------------------------------')



nfig += 1
nfig_PC_statistics = nfig
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the projection on PC0-PC1...')
printt('-------------------------------------------')
document.add_paragraph('The central panel in Fig. ' + str(nfig) + ' shows the projection of the learning styles of the student represented in Fig. ' + str(nfig_4d) + ' on the two main principal components (PC0 and PC1), along with their histograms. The data surround the average mean (gray star). The points (20, 0, 0, 0), (0, 20, 0, 0), (0, 0, 20, 0), and (0, 0, 0, 20), which correspond, respectively, to the maximal pure activist (red), reflector (green), theorist (purple), and pragmatist (orange) learning styles in the original basis set are also marked, and joined with a dashed line to the corresponding origin (0, 0, 0, 0) (black triangle). The point (20, 20, 20, 20) is also shown (brown diamond).')

document.add_paragraph('The reduced dimensional representation of the data given by Fig. ' + str(nfig_4d) + ' is usually meaningful if Σ_1(%)≥70.0%; in the case under study we have that Σ_1(%)='+sigma01+'. The top and left panels of Fig. ' + str(nfig) + ' show the histrograms of the data as a function of PC0 (red) and PC1 (green), which are fitted using Weibull distributions with the parameters listed in Table ' + str(ntab + 1) + '.')
document.add_picture(filename_PC0PC1, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. The projection of the learning styles of the students represented in Fig. ' + str(nfig_4d) + ' on the two main principal components (PC0 and PC1), along with their histograms as a function of PC0 (red) and PC1 (green).')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

document.add_page_break()
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the projection on PC0-PC1 done!')
printt('-------------------------------------------')

ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the Weibull parameters for the PDFs of the PCs...')
printt('-------------------------------------------')
# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Parameters of the Weibull distributions that fit the probability distributions (histograms) of the projection of the learning styles on the principal-components basis set (see histograms in Fig. ' + str(nfig) + '). The location parameter θ is set equal to the smallest projection for each component.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+1, 4)

# Table header
t.cell(0,0).text = 'Principal component'
t.cell(0,1).text = 'α'
t.cell(0,2).text = 'k'
t.cell(0,3).text = 'θ'

for j in range(0,K):
  t.cell(j+1,0).text = Label_PC[j]
  [alpha_Weibull, k_Weibull, theta_Weibull] = parameters_Weibull_PC20[j]
  t.cell(j+1,1).text = str(round(alpha_Weibull,2))
  t.cell(j+1,2).text = str(round(k_Weibull,2))
  t.cell(j+1,3).text = str(round(theta_Weibull,2))

document.add_page_break()

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the Weibull parameters for the PDFs of the PCs done!')
printt('-------------------------------------------')












if ( Lprof > 0 ):
 nsubsec += 1
 nfig_prof0 = nfig + 1
 if ( Lprof == 1 ):
  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Relation between the learning styles of the students and of the '+ prof_ref +' in the principal-component space', level=2)
  printt(str(nsec) + '.' + str(nsubsec) + ' Relation between the learning styles of the students and of the '+ prof_ref +' in the principal-component space')
  document.add_paragraph('Figure ' + str(nfig_prof0) + ' shows the projection of the learning styles of the students on the two main principal components (PC0 and PC1), as represented in the central panel of Fig. ' + str(nfig_PC_statistics) + ', along with the results for the ' + prof_ref + ' (gray cross). The students with the same qualitative tendency as the ' + prof_ref + ' have been represented with colored crosses.')
 else:
  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Relation between the learning styles of the students and of the '+ prof_ref +'s in the principal-component space', level=2)
  printt(str(nsec) + '.' + str(nsubsec) + ' Relation between the learning styles of the students and of the '+ prof_ref +'s in the principal-component space')
  document.add_paragraph('Figures ' + str(nfig_prof0) + ' to '  + str(nfig_prof0 + Lprof) + ' show the projection of the learning styles of the students on the two main principal components (PC0 and PC1), as represented in the central panel of Fig. ' + str(nfig_PC_statistics) + ', along with the results for the ' + prof_ref + 's (gray crosses). The students with the same qualitative tendency as the ' + prof_ref + 's have been represented with colored crosses.')

 printt('-------------------------------------------')
 printt(' Figure(s) with the LSs of the students and of the '+ prof_ref +'(s) on PC0-PC1...')
 printt('-------------------------------------------')
 for i_prof in range(0, Lprof):
  profi = profs[i_prof]
  nfig  = nfig + 1

  printt('output_participants_pc_prof:' + output_participants_pc_prof)
  filename = output_participants_pc_prof + '/Fig_'+Label_PCPC0+'_'+Label_PCPC1+'_prof_' + profi + '.png'
  document.add_picture(filename, width=Inches(5))

# Figure title
  printt('i_prof : ' + str(i_prof))
  printt('prof_ref : ' + profi)
  if(i_prof == 0):
    figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Projection of the learning styles of the students on the two main principal components (PC0 and PC1), as represented in the central panel of Fig. ' + str(nfig_PC_statistics) + ' along with the results of the ' + prof_ref + ' of ' + profi + ' (gray cross). The colored crosses mark the students with the same qualitative tendency towards the corresponding learning style (light shade for very low and low tendencies, medium shade for moderate tendency, and dark shade for high and very high tendencies).')
  else:
    figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Same as Fig. ' + str(nfig_prof0) + ' for the ' + prof_ref + ' of ' + profi + '.')

  figure_title.alignment = 1  # Center alignment
  figure_run = figure_title.runs[0]
  figure_run.bold = True

  printt('-------------------------------------------')
  printt(' Figure ' + str(nfig) + ' with the projection of the LS of the students and professor ' + profi +' on PC0-PC1 done!')
  printt('-------------------------------------------')


document.add_page_break()
#===================================================
# SECTION 3
#===================================================
#==================================================
nsec   += 1
nsubsec = 1
document.add_heading(str(nsec) + '. Participation ratios', level=1)
printt('******************************************')
printt(str(nsec) + '. Participation ratios')
printt('******************************************')
#==================================================

document.add_paragraph('Third, the participation ratios are presented in this section. The values for the CHAEA (activist, reflector, theorist, and pragmatist), and for the principal-components basis sets are listed. A statistical analysis of the distribution of the participation ratios is also included.')

nsubsec = 1
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Individual analysis of the participation ratios', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Individual analysis of the participation ratios')

ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the PRs...')
printt('-------------------------------------------')

document.add_paragraph('Table ' + str(ntab) + ' shows the participation ratios of each individual student for CHAEA (activist, reflector, theorist, and pragmatist) and for the principal-components basis sets. For the case under study, this parameter lies between 1 and 4. The smaller the participation ratio, the better.')

# Table 1 title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Participation ratios of each individual student for CHAEA learning styles (LS) and for the principal-components (PC) basis sets.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

# Table 1
t = document.add_table(Lsave+1, 3)

# Table header
t.cell(0,0).text = 'Student'
t.cell(0,1).text = 'LS'
t.cell(0,2).text = 'PC'

for i in range(0,Lsave):
  t.cell(i+1,0).text = students[i]
  t.cell(i+1,1).text = str(round(pr[i,0],2))
  t.cell(i+1,2).text = str(round(pr[i,1],2))

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the PRs done!')
printt('-------------------------------------------')
printt(' ')


document.add_page_break()
nsubsec += 1
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Global analysis of the participation ratios', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Global analysis of the participation ratios')

nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the PDF of the PRs...')
printt('-------------------------------------------')
ntab += 1
document.add_paragraph('Figure ' + str(nfig) + ' shows the probability distribution functions of the participation ratios for the learning-styles basis set (red) and for the principal components (blue). The vertical dashed lines mark the average means along with their corresponding confidence intervals (shaded areas), whose values can be found in Table ' + str(ntab) + '. The corresponding cumulative distributions used in the fitting are shown in Fig. ' + str(nfig+1) + '.')

document.add_picture(filename_PR_P, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Probability distribution functions of the participation ratios associated with the learning-styles basis set (red) and with the principal components (blue). The vertical dashed lines mark the average means along with their corresponding confidence intervals (shaded areas), whose values can be found in Table ' + str(ntab) + '. The continuous lines show fitting functions given by Weibull distributions with parameters shown in Table ' + str(ntab) + '.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the PDF of the PRs done!')
printt('-------------------------------------------')


nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with W(s) of the PRs...')
printt('-------------------------------------------')

document.add_paragraph('')
document.add_picture(filename_PR_W, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Cumulative distributions for the results of Fig. ' + str(nfig-1) + '.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with W(s) of the PRs done!')
printt('-------------------------------------------')




printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the average parameters for the PRs...')
printt('-------------------------------------------')

# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Average mean and corresponding uncertainty (in parenthesis) of the participation ratios for the learning-styles basis set and for the principal components. α and k are, respectively, the shape and scale parameters of the Weibull distributions the fit the probability distributions (histograms) of Fig. ' + str(nfig) + '. The location parameter is set equal to θ = 0.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(3, 4)

# Table header
t.cell(0,0).text = 'Basis set'
t.cell(0,1).text = 'Average mean (Uncertainty)'
t.cell(0,2).text = 'α'
t.cell(0,3).text = 'k'

t.cell(1,0).text = 'Learning styles'
t.cell(2,0).text = 'Principal components'

for j in range(0,2):
  xdx = mean_uncert(prmean[j], uncert_abs_pr[j])
  t.cell(j+1,1).text = xdx
  [alpha_Weibull, k_Weibull] = parameters_Weibull_PR[j]
  t.cell(j+1,2).text = str(round(alpha_Weibull,2))
  t.cell(j+1,3).text = str(round(k_Weibull,2))

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the average parameters for the PRs done!')
printt('-------------------------------------------')



print('ntab : ', ntab)
ntab0_mark = ntab + 1
document.add_page_break()
#===================================================
# SECTION 4 (IF MARKS ARE PROVIDED)
#===================================================
if ( K_value > 0 ):
 nsec   += 1
 nsubsec = 1
 document.add_heading(str(nsec) + '. Analysis of the students marks', level=1)
 printt('******************************************')
 printt(str(nsec) + '. Analysis of the students marks')
 printt('******************************************')
 if ( Lmarks == 0 ):
  document.add_paragraph('The marks provided in the input folder ' + input_folder_marks_ref + ' have not been analyzed as the students included are not the same ones as those that fullfilled CHAEA.')
 else:
  document.add_paragraph('This section is devoted to the analysis of the students marks contained in the input folder ' + input_folder_marks_ref + '.')

  document.add_paragraph('The tables present the averages and uncertainties (in parenthesis) of the learning styles And the principal components (x) and the marks (y) of the whole group, along with the parameters of the linear regressions that fit the marks as a function of the number of points in the corresponding learning style or principal components, being m the slope, b the intercept, and r the regression parameter. Recall that the closest the regression parameter to -1 or 1, the better the fitting. Conversely, a regression parameter close to zero indicates the lack of correlations between the marks and the corresponding variable.')

  document.add_paragraph('The results marked as nan indicate failures in the linear regressions because, e.g., having too few points.')

  df = pd.DataFrame(marks_vs_LS)
  if df.isna().values.any():
    document.add_paragraph('The cells with nan indicate the failure of the linear regression.')

  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Marks as a function of CHAEA learning styles', level=2)
  printt(str(nsec) + '.' + str(nsubsec) + ' Marks as a function of CHAEA learning styles')
  document.add_paragraph('In this section, we report on the relationship between the marks and CHAEA learning styles.')

  column0_strings = [str(row[0]) for row in data_marks]
  Excel_documents_marks = column0_strings
  del Excel_documents_marks[0]
  Excel_documents_marks = list(dict.fromkeys(Excel_documents_marks))
  print('Excel_documents_marks : ', Excel_documents_marks)

  Lmarks_vs_LS = np.shape(marks_vs_LS)[0]
  Lmarks_vs_PC = np.shape(marks_vs_PC)[0]

  ntab_mean = ntab + 1
  nfig_mean = nfig + 1

  for Excel_doc in Excel_documents_marks:
    print('Excel_doc  : ', Excel_doc)

    document.add_heading('Marks of the document ' + Excel_doc, level=3)
    printt(str(nsec) + '.' + str(nsubsec) + ' Marks of the document ' + Excel_doc)

#   Subsubsections (Excel sheets)
    Excel_sheets = [str(row[1]) for row in data_marks if row[0] == Excel_doc]
    print('Excel sheets : ', Excel_sheets)
    Lsheets = len(Excel_sheets)

    for nExcel_sheet in range(0,Lsheets):
#      if(Lsheets == 1):marks_vs_LS
#        Excel_sheet = Excel_sheets
#      else:
      Excel_sheet = Excel_sheets[nExcel_sheet]
      print('Excel sheet : ', Excel_sheet)

#     Tables (each Excel column corresponds to the marks for an evaluating activity
      Excel_columns = [str(row[2]) for row in data_marks if row[0] == Excel_doc and row[1] == Excel_sheet ]
      print('Excel_columns : ', Excel_columns)
      Lcolumns = len(Excel_columns)

      for nExcel_column in range(0,Lcolumns):

        Excel_column = Excel_columns[nExcel_column]
        print('Excel_column : ', Excel_column)
        print('Marks of the column ' + Excel_column + ' of the Excel sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc) # + ' of the Excel document ' + Excel_doc)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Marks of the column {Excel_column} of the sheet {Excel_sheet}")
        run.bold = True

        nfig += 1
        # Figure
        printt('-------------------------------------------')
        printt(' Figure ' + str(nfig) + ' with the marks vs LS...')
        printt('-------------------------------------------')

        filename = output_participants_ls_marks + '/Fig_'+ Excel_doc + '_' + str(Excel_sheet) + '_' + Excel_column + '_K' + str(K_value) + '.png'
        if not os.path.exists(filename):
         nfig -= 1
         print(f"The file '{filename}' does NOT exist.")
         document.add_paragraph('No clustering analysis based on K-means algorithm is performed due to the low number of students for the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '. Consequently, no figures or tables relating the marks and the CHAEA punctuations for the different learning styles are included here.')
        else:
         print(f"The file '{filename}' exists.")
         if( K_value == 1):
          document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the CHAEA punctuations for the different learning styles. The results are presented in Table ' + str(ntab+1) + ', and correspond to the whole group.')
         else:
          if( K_value == 2):
            document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the CHAEA punctuations for the different learning styles. The results in gray are presented in Table ' + str(ntab+1) + ', and correspond to the whole group. The colored results are listed in the Tables ' + str(ntab + 2) + ' and ' +str(ntab + 1 + K_value) + ', and are associated with the different clusters formed by the K-mean algorithm.')
          else:
            document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the CHAEA punctuations for the different learning styles. The results in gray are presented in Table ' + str(ntab+1) + ', and correspond to the whole group. The colored results are listed in the Tables ' + str(ntab + 2) + ' to ' +str(ntab + 1 + K_value) + ', and are associated with the different clusters formed by the K-mean algorithm.')

         document.add_picture(filename, width=Inches(5))

#        Figure title
         if(nfig == nfig_mean):
          if(K_value == 1):
           figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the CHAEA punctuations for the different learning styles. Students are depicted with circles, while cluster centroids are marked with stars. Linear fits are shown as continuous lines, and their confidence intervals are illustrated in shaded areas.')
          else:
           figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the CHAEA punctuations for the different learning styles. Each color represents a distinct cluster determined using the K-means algorithm with K=' + str(K_value)+'. Students are depicted with circles, while cluster centroids are marked with stars. Linear fits are shown as continuous lines, and their confidence intervals are illustrated in shaded areas. The results for the whole group are indicated in gray.')
         else:
           figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Same as Fig. ' + str(nfig_mean) + ' for the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '.')

         figure_title.alignment = 1  # Center alignment
         figure_run = figure_title.runs[0]
         figure_run.bold = True

         printt('-------------------------------------------')
         printt(' Figure ' + str(nfig) + ' with the marks vs LS done!')
         printt('-------------------------------------------')

         ntab += 1
         printt('-------------------------------------------')
         printt(' Table ' + str(ntab) + ' with the marks and the LS for the whole group...')
         printt('-------------------------------------------')

         j = 0
         title = 0
         for i in range(0, Lmarks_vs_LS):
          if ( (marks_vs_LS[i, 0] == Excel_doc) and (marks_vs_LS[i, 1] == Excel_sheet) and (marks_vs_LS[i, 2] == Excel_column) and (int(marks_vs_LS[i, 5]) == -1) ):
            if(title == 0):
              title = 1
              iref  = i
#             Table title
              if(ntab == ntab_mean):
                if(float(marks_vs_LS[i, 3]) == 1):
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Results for the whole group, which is formed by only ' + str(marks_vs_LS[i, 3]) + ' student. The table presents the average values and uncertainties (in parenthesis) for the learning styles (x), for the marks (y) contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ', along with the parameters of the linear regression, which cannot fit the mark as a function of the number of points in the corresponding learning style because only one point is provided (m: slope, b: intercept, r: regression parameter).')
                else:
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Results for the whole group, which is formed by ' + str(marks_vs_LS[i, 3]) + ' students. The table presents the average values and uncertainties (in parenthesis) for the learning styles (x), for the marks (y) contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ', along with the parameters of the linear regressions that fit the marks as a function of the number of points in the corresponding learning style (m: slope, b: intercept, r: regression parameter).')
              else:
                if(float(marks_vs_LS[i, 3]) == 1):
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean) + ' for the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '. The linear regression fails because only student is included.')
                else:
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean) + ' for the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '.')

              table_title.alignment = 1  # Center alignment
              title_run = table_title.runs[0]
              title_run.bold = True
#
#             Average results for the whole group
              t = document.add_table(5, 6)
#             Table header
              t.cell(0,0).text = 'LS'
              t.cell(0,1).text = 'x\u0304 (\u0394x)'
              t.cell(0,2).text = 'y\u0304 (\u0394y)'
              t.cell(0,3).text = 'm\u0304 (\u0394m)'
              t.cell(0,4).text = 'b\u0304 (\u0394b)'
              t.cell(0,5).text = 'r'

            j += 1
            t.cell(j,0).text = marks_vs_LS[i, 6]
            xdx = mean_uncert(float(marks_vs_LS[i, 7]), float(marks_vs_LS[i, 8]))
            t.cell(j,1).text = xdx
            ydy = mean_uncert(marks_vs_LS[i, 9], marks_vs_LS[i, 10])
            t.cell(j,2).text = ydy
            mdm = mean_uncert(marks_vs_LS[i, 11], marks_vs_LS[i, 12])
            t.cell(j,3).text = mdm
            bdb = mean_uncert(marks_vs_LS[i, 13], marks_vs_LS[i, 14])
            t.cell(j,4).text = bdb
            t.cell(j,5).text = round_until_two_non_9(marks_vs_LS[i,15])

         document.add_paragraph()
         if ( float(marks_vs_LS[iref, 3]) <= Lsave ):
          if(float(marks_vs_LS[iref, 3]) == 1):
            document.add_paragraph('Table ' + str(ntab) + ' shows the results for marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' for the whole group, which is formed by the following student: ' + marks_vs_LS[iref, 4] + '.')
          else:
            document.add_paragraph('Table ' + str(ntab) + ' shows the results for marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' for the whole group, which is formed by the following ' + str(marks_vs_LS[iref, 3]) + ' students: ' + marks_vs_LS[iref, 4] + '.')
         else:
          document.add_paragraph('Table ' + str(ntab) + ' shows the results for marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' for the whole group, which is formed by the following ' + str(marks_vs_LS[iref, 3]) + ' students, being the first ' + str(Lsave) + ' ones: ' + marks_vs_LS[iref, 4] + '.')

         printt('-------------------------------------------')
         printt(' Table ' + str(ntab) + ' with the marks and the LS for the whole group done!')
         printt('-------------------------------------------')

#        Clustering-analysis results rendered by K-means algorithm
         if( K_value > 1):
          if( K_value == 2):
           document.add_paragraph('Tables ' + str(ntab + 1) + ' and ' + str(ntab + K_value) + ' include the same results as those presented in Table ' + str(ntab_mean) + ' for the clusters constructed using K-means algorithm.')
          else:
           document.add_paragraph('Tables ' + str(ntab + 1) + ' to ' + str(ntab + K_value) + ' include the same results as those presented in Table ' + str(ntab_mean) + ' for the clusters constructed using K-means algorithm.')

          for k in range(0, K_value):
           print('k : ', k)
           j = 0

           ntab += 1
           title = 0
           print('ntab : ', ntab)
           printt('-------------------------------------------')
           printt(' Table ' + str(ntab) + ' with the marks and the LS for the cluster k = ' + str(k) + '...')
           printt('-------------------------------------------')

           for i in range(0, Lmarks_vs_LS):
#           Average results for the whole group
            if ( (marks_vs_LS[i, 0] == Excel_doc) and (marks_vs_LS[i, 1] == Excel_sheet) and (marks_vs_LS[i, 2] == Excel_column) and (int(marks_vs_LS[i, 5]) == k) ):

              if(title == 0):
                title = 1
                iref = i
#
#               Table title
                if(float(marks_vs_LS[i, 3]) == 1):
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean) + ' for the cluster k=' + str(k) + ' which is formed by one student. The linear regression fails because only student is included.')
                else:
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean) + ' for the cluster k=' + str(k) + ' which is formed by ' + str(marks_vs_LS[i, 3]) + ' students.')
                table_title.alignment = 1  # Center alignment
                title_run = table_title.runs[0]
                title_run.bold = True
#
                table_title.alignment = 1  # Center alignment
                title_run = table_title.runs[0]
                title_run.bold = True

                t = document.add_table(5, 6)
#               Table header
                t.cell(0,0).text = 'LS'
                t.cell(0,1).text = 'x\u0304 (\u0394x)'
                t.cell(0,2).text = 'y\u0304 (\u0394y)'
                t.cell(0,3).text = 'm\u0304 (\u0394m)'
                t.cell(0,4).text = 'b\u0304 (\u0394b)'
                t.cell(0,5).text = 'r'

              j += 1
              t.cell(j,0).text = marks_vs_LS[i, 6]
              xdx = mean_uncert(float(marks_vs_LS[i, 7]), float(marks_vs_LS[i, 8]))
              t.cell(j,1).text = xdx
              ydy = mean_uncert(marks_vs_LS[i, 9], marks_vs_LS[i, 10])
              t.cell(j,2).text = ydy
              mdm = mean_uncert(marks_vs_LS[i, 11], marks_vs_LS[i, 12])
              t.cell(j,3).text = mdm
              bdb = mean_uncert(marks_vs_LS[i, 13], marks_vs_LS[i, 14])
              t.cell(j,4).text = bdb
              print('marks_vs_LS[i,15] : ', marks_vs_LS[i,15])
              t.cell(j,5).text = round_until_two_non_9(marks_vs_LS[i,15])

           document.add_paragraph()
           if ( float(marks_vs_LS[iref, 3]) <= Lsave ):
            if(float(marks_vs_LS[iref, 3]) == 1):
              document.add_paragraph('Table ' + str(ntab) + ' shows the results for the cluster k=' + str(k) + ', which is formed by the following student: ' + marks_vs_LS[iref, 4] + '.')
            else:
              document.add_paragraph('Table ' + str(ntab) + ' shows the results for the cluster k=' + str(k) + ', which is formed by the following ' +  str(marks_vs_LS[iref, 3]) + ' students: ' + marks_vs_LS[iref, 4] + '.')
           else:
            document.add_paragraph('Table ' + str(ntab) +' shows the results for the cluster k=' + str(k) + ', which is formed by ' +  str(marks_vs_LS[iref, 3]) + ' students, being the first ' + str(Lsave) + ' students: ' + marks_vs_LS[iref, 4]+ '.')

          printt(' Table ' + str(ntab) + ' with the marks and the LS for the cluster k = ' +str(k) + ' done!')







  document.add_page_break()
  nsubsec += 1
  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Marks as a function of the principal components', level=2)
  printt(str(nsec) + '.' + str(nsubsec) + ' Marks as a function of the principal components')
  document.add_paragraph('In this section, we report on the relationship between the marks and the principal components.')

  ntab_PC = 0

  for Excel_doc in Excel_documents_marks:
    print('Excel_doc  : ', Excel_doc)

    document.add_heading('Marks of the document ' + Excel_doc, level=3)
    printt(str(nsec) + '.' + str(nsubsec) + ' Marks of the document ' + Excel_doc)

#   Subsubsections (Excel sheets)
    Excel_sheets = [str(row[1]) for row in data_marks if row[0] == Excel_doc]
    print('Excel sheets : ', Excel_sheets)
    Lsheets = len(Excel_sheets)

    for nExcel_sheet in range(0,Lsheets):
      Excel_sheet = Excel_sheets[nExcel_sheet]
      print('Excel sheet : ', Excel_sheet)

#     Tables (each Excel column corresponds to the marks for an evaluating activity
      Excel_columns = [str(row[2]) for row in data_marks if row[0] == Excel_doc and row[1] == Excel_sheet ]
      print('Excel_columns : ', Excel_columns)
      Lcolumns = len(Excel_columns)

      for nExcel_column in range(0,Lcolumns):

        Excel_column = Excel_columns[nExcel_column]
        print('Excel_column : ', Excel_column)
        print('Marks of the column ' + Excel_column + ' of the Excel sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc) # + ' of the Excel document ' + Excel_doc)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Marks of the column {Excel_column} of the sheet {Excel_sheet}")
        run.bold = True

#       Figure
        filename = output_participants_pc_marks + '/Fig_'+ Excel_doc + '_' + str(Excel_sheet) + '_' + Excel_column + '_K' + str(K_value) + '.png'

        if not os.path.exists(filename):
         print(f"The file '{filename}' does NOT exist.")
         document.add_paragraph('No clustering analysis based on K-means algorithm is performed due to the low number of students for the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '. Consequently, no figures or tables relating the marks and the CHAEA punctuations for the different learning styles are included here.')
        else:
         nfig += 1
         printt('-------------------------------------------')
         printt(' Figure ' + str(nfig) + ' with the marks vs PC...')
         printt(' nfig_mean : ' + str(nfig_mean))
         printt('-------------------------------------------')
         if( K_value == 1):
          document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the principal components. The results are presented in Table ' + str(ntab+1) + ', and correspond to the whole group.')
         else:
          if( K_value == 2):
            document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the principal components. The results in gray are presented in Table ' + str(ntab+1) + ', and correspond to the whole group. The colored results are listed in the Tables ' + str(ntab + 2) + ' and ' + str(ntab + 1 + K_value) + ', and are associated with the different clusters formed by the K-mean algorithm.')
          else:
            document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the principal components. The results in gray are presented in Table ' + str(ntab+1) + ', and correspond to the whole group. The colored results are listed in the Tables ' + str(ntab + 2) + ' to ' + str(ntab + 1 + K_value) + ', and are associated with the different clusters formed by the K-mean algorithm.')


         document.add_picture(filename, width=Inches(5))

#        Figure title
         figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Same as Fig. ' + str(nfig_mean) + ' for the  principal components and the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '.')
         figure_title.alignment = 1  # Center alignment
         figure_run = figure_title.runs[0]
         figure_run.bold = True

         printt('-------------------------------------------')
         printt(' Figure ' + str(nfig) + ' with the marks vs PC done!')
         printt('-------------------------------------------')

         ntab += 1
         ntab_PC += 0
         printt('-------------------------------------------')
         printt(' Table ' + str(ntab) + ' with the marks and the PC for the whole group...')
         printt('-------------------------------------------')

         j = 0
         title = 0
         for i in range(0, Lmarks_vs_PC):
          if ( (marks_vs_PC[i, 0] == Excel_doc) and (marks_vs_PC[i, 1] == Excel_sheet) and (marks_vs_PC[i, 2] == Excel_column) and (int(marks_vs_PC[i, 5]) == -1) ):
            if(title == 0):
              title = 1
              iref  = i

#             Table title
              table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean + ntab_PC) + ' for the principal components (the marks are contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ').')
              table_title.alignment = 1  # Center alignment
              title_run = table_title.runs[0]
              title_run.bold = True
#
#             Average results for the whole group
              t = document.add_table(5, 6)
#             Table header
              t.cell(0,0).text = 'PC'
              t.cell(0,1).text = 'x\u0304 (\u0394x)'
              t.cell(0,2).text = 'y\u0304 (\u0394y)'
              t.cell(0,3).text = 'm\u0304 (\u0394m)'
              t.cell(0,4).text = 'b\u0304 (\u0394b)'
              t.cell(0,5).text = 'r'

            j += 1
            t.cell(j,0).text = marks_vs_PC[i, 6]
            xdx = mean_uncert(float(marks_vs_PC[i, 7]), float(marks_vs_PC[i, 8]))
            t.cell(j,1).text = xdx
            ydy = mean_uncert(marks_vs_PC[i, 9], marks_vs_PC[i, 10])
            t.cell(j,2).text = ydy
            mdm = mean_uncert(marks_vs_PC[i, 11], marks_vs_PC[i, 12])
            t.cell(j,3).text = mdm
            bdb = mean_uncert(marks_vs_PC[i, 13], marks_vs_PC[i, 14])
            t.cell(j,4).text = bdb
            t.cell(j,5).text = round_until_two_non_9(marks_vs_PC[i,15])

         document.add_paragraph()
         if(float(marks_vs_PC[iref, 3]) == 1):
          document.add_paragraph('Table ' + str(ntab) + ' shows the results for marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' for the whole group, which is formed by the same student as that considered in Table ' + str(ntab_mean + ntab_PC) + '.')
         else:
          document.add_paragraph('Table ' + str(ntab) + ' shows the results for marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' for the whole group, which is formed by the same ' + str(marks_vs_PC[iref, 3]) + ' students as those considered in Table ' + str(ntab_mean + ntab_PC) + '.')

         ntab_PC += 1

         printt('-------------------------------------------')
         printt(' Table ' + str(ntab) + ' with the marks and the PC for the whole group done!')
         printt('-------------------------------------------')

#        Clustering-analysis results rendered by K-means algorithm
         if( K_value > 1):
          if( K_value == 2):
           document.add_paragraph('Tables ' + str(ntab + 1) + ' and ' + str(ntab + K_value) + ' include the same results as those presented in Tables ' + str(ntab_mean + ntab_PC ) + ' and ' + str(ntab_mean + ntab_PC + K_value - 1) + ' for the clusters constructed using K-means algorithm and for the principal components (instead of the learning styles).')
          else:
           document.add_paragraph('Tables ' + str(ntab + 1) + ' to ' + str(ntab + K_value) + ' include the same results as those presented in Tables ' + str(ntab_mean + ntab_PC ) + ' to ' + str(ntab_mean + ntab_PC + K_value - 1) + ' for the clusters constructed using K-means algorithm and for the principal components (instead of the learning styles).')

          for k in range(0, K_value):
           print('k : ', k)
           j = 0
           ntab += 1
           ntab_PC += 1
           title = 0
           print('ntab : ', ntab)
           printt('-------------------------------------------')
           printt(' Table ' + str(ntab) + ' with the marks and the PC for the cluster k = ' + str(k) + '...')
           printt('-------------------------------------------')

           for i in range(0, Lmarks_vs_PC):
#           Average results for the whole group
            if ( (marks_vs_PC[i, 0] == Excel_doc) and (marks_vs_PC[i, 1] == Excel_sheet) and (marks_vs_PC[i, 2] == Excel_column) and (int(marks_vs_PC[i, 5]) == k) ):

              if(title == 0):
                title = 1
                iref = i

#               Table title
                if(float(marks_vs_PC[i, 3]) == 1):
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean + ntab_PC - 1) + ' for the cluster k=' + str(k) + ' which is formed by one student, which makes the linear regression ill-suited.')
                else:
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean + ntab_PC - 1) + ' for the cluster k=' + str(k) + ' which is formed by ' + str(marks_vs_PC[i, 3]) + ' students.')
                table_title.alignment = 1  # Center alignment
                title_run = table_title.runs[0]
                title_run.bold = True
                t = document.add_table(5, 6)
#               Table header
                t.cell(0,0).text = 'PC'
                t.cell(0,1).text = 'x\u0304 (\u0394x)'
                t.cell(0,2).text = 'y\u0304 (\u0394y)'
                t.cell(0,3).text = 'm\u0304 (\u0394m)'
                t.cell(0,4).text = 'b\u0304 (\u0394b)'
                t.cell(0,5).text = 'r'

              j += 1
              t.cell(j,0).text = marks_vs_PC[i, 6]
              xdx = mean_uncert(float(marks_vs_PC[i, 7]), float(marks_vs_PC[i, 8]))
              t.cell(j,1).text = xdx
              ydy = mean_uncert(marks_vs_PC[i, 9], marks_vs_PC[i, 10])
              t.cell(j,2).text = ydy
              mdm = mean_uncert(marks_vs_PC[i, 11], marks_vs_PC[i, 12])
              t.cell(j,3).text = mdm
              bdb = mean_uncert(marks_vs_PC[i, 13], marks_vs_PC[i, 14])
              t.cell(j,4).text = bdb
              t.cell(j,5).text = round_until_two_non_9(marks_vs_PC[i,15])

           document.add_paragraph()
           if(float(marks_vs_PC[iref, 3]) == 1):
            document.add_paragraph('Table ' + str(ntab) + ' shows the results for the cluster k=' + str(k) + ', which is formed by the same student as that considered in the Table ' + str(ntab_mean + ntab_PC - 1) + '. The linear regression fails because only student is included.')
           else:
            document.add_paragraph('Table ' + str(ntab) + ' shows the results for the cluster k=' + str(k) + ', which is formed by the same ' +  str(marks_vs_PC[iref, 3]) + ' students as those considered in the Table ' + str(ntab_mean + ntab_PC - 1) + '.')

           printt(' Table ' + str(ntab) + ' with the marks and the PC for the cluster k = ' +str(k) + ' done!')

document.add_page_break()

#===================================================
# SECTION 4 OR 5
#===================================================
nsec   += 1
nsubsec = 1
document.add_heading(str(nsec) + '. Students characteristics and recommendations', level=1)
printt('******************************************')
printt(str(nsec) + '. Students characteristics and recommendations')
printt('******************************************')
document.add_paragraph('This section summarizes some of the characteristics associated with the activist, theorist, reflector, and pragmatist learning styles, and provides some recommendation to increase the tendencies towards them, as detailed by Alonso, Gallego & Honey (2007), and Cancino (2014).')

alonso_reference_paragraph = document.add_paragraph()
alonso_reference_paragraph.add_run('Alonso, C. M., Gallego, D. J., & Honey, P. (2007). Los estilos de aprendizaje. Procedimiento de diagnóstico y mejora (7th ed.). Ediciones Mensajero, S. A. U.').italic = True

cancino_reference_paragraph = document.add_paragraph()
cancino_reference_paragraph.add_run('Cancino, O. (2024, March). Blog de la prof de Français: espacio para intercambiar con los estudiantes. add_hyperlink(p, "https://oticar.wordpress.com/estilos-de-aprendizaje/", "https://oticar.wordpress.com/estilos-de-aprendizaje/"').italic = True
#
#==================================================
# Learning-styles characteristics
#==================================================
#
# Activist
activist_main_characteristics = ['Cheerful', 'Improviser', 'Explorer', 'Risk-taker', 'Spontaneous']
activist_other_characteristics = 'creative, innovative, adventurous, revitalizing, inventor, energetic, experience-seeker, idea-generator, bold, protagonist, striking, groundbreaking, talkative, leader, determined, fun, participative, competitive, eager to learn, problem-solver, adaptable.'
#
# Reflector
reflector_main_characteristics = ['Thoughtful', 'Thorough', 'Receptive', 'Analytical', 'Exhaustive']
reflector_other_characteristics = 'observer, collector, patient, careful, detail-oriented, argument-developer, alternative-planner, behavior-studier, data-recorder, researcher, assimilator, report and/or statement writer, slow, distant, cautious, inquisitive, probe-seeker.'
#
# Theorist
theorist_main_characteristics = ['Methodical', 'Logical', 'Objective', 'Critical', 'Structured']
theorist_other_characteristics = 'disciplined, planned, systematic, organized, synthetic, reasoner, thinker, connector, perfectionist, generalizer, hypothesis-seeker, theory-seeker, model-seeker, question-asker, assumption-explorer, concept-seeker, clear-purpose seeker, rationality-seeker, "why"-asker, value-system and criteria-seeker, procedure-inventor, explorer.'
#document.add_paragraph()
# Pragmatist
pragmatist_main_characteristics = ['Experimenter', 'Practical', 'Direct', 'Efficient', 'Realistic']
pragmatist_other_characteristics = 'technical, useful, fast, decisive, planner, positive, concrete, objective, clear, self-confident, organizer, contemporary, problem-solver, applier of learning, action planner.'
#
# Main characteristics
main_characteristics = [activist_main_characteristics, reflector_main_characteristics, theorist_main_characteristics, pragmatist_main_characteristics]
#
# Other characteristics
other_characteristics = [activist_other_characteristics, reflector_other_characteristics, theorist_other_characteristics, pragmatist_other_characteristics]
#
#==================================================
# Affinity and dominant learning style
#==================================================
# Maximum value of the affinity
max_affinity = np.max(affinity)
max_affinity_str = str(round(max_affinity,1))

# Position (index) of the maximum value
max_index_affinity = np.where(affinity == max_affinity)[0]

if(len(max_index_affinity) == 1):
# One dominant learning style
  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Affinity and dominant learning style', level=2)
  i = max_index_affinity[0]
  document.add_paragraph('According to the affinity, the dominant learning style is the ' + List_LS[i] + '. The value of its affinity is equal to ' + max_affinity_str +' %.' )
  document.add_paragraph('The main characteristics of the ' + List_LS[i] + ' learning style are:')
  for point in main_characteristics[i]:
    document.add_paragraph(point, style='ListBullet')
  document.add_paragraph('Furthermore, other characteristics usually associated with the ' + List_LS[i] + ' learning style are being: ' + other_characteristics[i])
else:
# Several dominant learning styles
  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Affinity and dominant learning styles', level=2)
  list_mainLS = List_LS[max_index_affinity[0]]
  for i in range(1, len(max_index_affinity)-1):
    list_mainLS = list_mainLS + ', ' + List_LS[max_index_affinity[i]]
  list_mainLS = list_mainLS + ', and ' + List_LS[max_index_affinity[len(max_index_affinity)-1]] + '.'

# Several learning styles with the same affinity
  document.add_paragraph('The maximum value of the affinity equals ' + max_affinity_str +' %. The dominant learning styles are: ' + list_mainLS)
  for i in max_index_affinity:
    document.add_paragraph('The main characteristics of the ' + List_LS[i] + ' learning style are:')
    for point in main_characteristics[i]:
      document.add_paragraph(point, style='ListBullet')
    document.add_paragraph('Other characteristics usually associated with the ' + List_LS[i] + ' learning style are being: ' + other_characteristics[i])

# Add a comment if the other learning styles are larger than 85%
high_affinity_indices = np.where((affinity >= 85) & (affinity < max_affinity))[0]
print('high_affinity_indices : ', high_affinity_indices)
if( high_affinity_indices.size > 0 ):
  unique_high_affinity_indices = np.setdiff1d(high_affinity_indices, max_index_affinity)
  print('unique_high_affinity_indices : ', unique_high_affinity_indices)
  list_highLS = List_LS[unique_high_affinity_indices[0]]
  for i in range(1, len(unique_high_affinity_indices)-1):
    list_highLS = list_highLS + ', ' + List_LS[unique_high_affinity_indices[i]]
  list_highLS = list_highLS + ', and ' + List_LS[unique_high_affinity_indices[len(unique_high_affinity_indices)-1]] + '.'

  if(len(unique_high_affinity_indices)==1):
    document.add_paragraph('Other learning style with a high affinity (larger than 85%) is: ' + List_LS[unique_high_affinity_indices[0]] + '.')
  else:
    document.add_paragraph('Other learning styles with high affinities (larger than 85%) are: ' + list_highLS)

  for i in unique_high_affinity_indices:
    document.add_paragraph('The main characteristics of the ' + List_LS[i] + ' learning style are:')
    for point in main_characteristics[i]:
      document.add_paragraph(point, style='ListBullet')
    document.add_paragraph('Other characteristics usually associated with the ' + List_LS[i] + ' learning style are being: ' + other_characteristics[i])

# Structure of the eigenvector associated with the main principal component (PC0)
v0 = eigenVectors_percentage[:,0]
#
# LS with the largest component
max_eigen = np.max(v0)
max_index_eigen = np.where(v0 == max_eigen)[0]
LLmax = len(max_index_eigen)
if(LLmax == 1):
  document.add_paragraph("According to the principal components analysis, the differences of the students' learning styles mostly correspond to the " + List_LS[max_index_eigen[0]] + " learning style, as " + str(round(max_eigen,1)) + ' %' + " of the eigenvector associated with the main principal component (PC0) is projected onto it.")
else:
  list_LS = ''
  for i in range(0, LLmax-1):
    list_LS = list_LS + List_LS[max_index_eigen[i]] + ', '
  list_LS = list_LS + 'and ' + List_LS[max_index_eigen[LLmax-1]]
  document.add_paragraph("According to the principal components analysis, the differences of the students' learning styles mostly correspond to the " + list_LS + " learning styles, as " + str(round(max_eigen,1)) + ' %' + " of the eigenvector associated with the main principal component (PC0) is projected onto it.")

# Add a comment on this learning style if it has not been included before
# Combine the arrays to exclude
excluded = np.union1d(max_index_affinity, high_affinity_indices)

# Get elements in max_index_eigen that are not in excluded
different_elements = np.setdiff1d(max_index_eigen, excluded)
print('different_elements : ', different_elements)

for i in different_elements:
  print('i usado en doc : ', i)
  document.add_paragraph('The main characteristics of the ' + List_LS[i] + ' learning style are:')
  for point in main_characteristics[i]:
    document.add_paragraph(point, style='ListBullet')
  document.add_paragraph('Furthermore, other characteristics usually associated with the ' + List_LS[i] + ' learning style are being: ' + other_characteristics[i])
#
# Add a comment if another component of the eigenvector is larger than 40%
high_v0_index = np.where((v0 >= 40) & (v0 < max_eigen))[0]
print('high_v0_index : ', high_v0_index)
if( high_v0_index.size > 0 ):
  unique_high_v0_index = np.setdiff1d(high_v0_index, max_index_affinity)
  print('unique_high_v0_index : ', unique_high_v0_index)
  document.add_paragraph('Other learning style with a high projection on the PC0 (larger than 40%) is: ' + List_LS[high_v0_index[0]] + '.')

  if(unique_high_v0_index == True ):
    document.add_paragraph('The main characteristics of the ' + List_LS[unique_high_v0_index] + ' learning style are:')
    for point in main_characteristics[unique_high_v0_index]:
      document.add_paragraph(point, style='ListBullet')
    document.add_paragraph('Other characteristics usually associated with the ' + List_LS[unique_high_v0_index] + ' learning style are being: ' + other_characteristics[unique_high_v0_index])

nsubsec += 1

#-------------
# Activist
#-------------
i = 0
# if ( xmean[i] >= 13 or xmean[i] <= 9):
document.add_heading('Activist learning style', level=3)
document.add_paragraph("The activist learning style has an average and an uncertainty (in parenthesis) of " + str(mean_uncert(xmean[i],dxmean[i])) + " points." )
#  if ( xmean[i] >= 13 ):
document.add_paragraph("People with scores higher than 13 on the test find more comfortable in learning situations where they can: ")

# Bullet points for high score situations
high_score_points = [
            "Attempt new things, new experiences, new opportunities.",
            "Compete in a team.",
            "Generate ideas without formal or structural limitations.",
            "Solve problems.",
            "Change and vary things.",
            "Address multiple tasks.",
            "Engage in dramatization.",
            "Play roles.",
            "Experience situations of interest or crisis.",
            "Capture attention."]

for point in high_score_points:
  document.add_paragraph(point, style='ListBullet')

# Uncomfortable situations
document.add_paragraph("\nSimilarly, they'll feel uncomfortable in situations that demand:")

# Bullet points for uncomfortable situations
uncomfortable_points = [
            "Presenting topics with a heavy theoretical load: Explaining causes, backgrounds...",
            "Assimilating, analyzing, and interpreting many unclear data.",
            "Paying attention to details.",
            "Working alone, reading, writing, or thinking alone.",
            "Evaluating in advance what is going to be learned.",
            "Assessing what has already been achieved or learned.",
            "Repeating the same activity."]

for point in uncomfortable_points:
  document.add_paragraph(point, style='ListBullet')

document.add_paragraph("\nIf someone has scored below 9, it's advisable to develop certain aspects that characterize this learning style further. Suggestions for improvement include:")

# Suggestions for improvement
improvement_suggestions = [
              "Trying something new, something never done before, at least once a week... Bringing something attention-grabbing to the study or workplace. Reading a newspaper with opposing opinions. Rearranging furniture at home or work.",
              "Practicing initiating conversations with strangers. In large gatherings, parties, conferences, forcing oneself to initiate and sustain conversations with everyone present, if possible. During leisure time, trying to engage in dialogue with strangers or convince them of our ideas.",
              "Deliberately fragmenting the day by cutting and changing activities every half an hour..."]

for point in improvement_suggestions:
  document.add_paragraph(point, style='ListBullet')

# Blocks that inhibit development
document.add_paragraph("\nBlocks that inhibit the development of the activist style:")

blocks = [    "Fear of failure or making mistakes.",
              "Fear of ridicule.",
              "Anxiety towards new or unfamiliar things.",
              "Strong desire to thoroughly think things through in advance.",
              "Self-doubt, lack of self-confidence.",
              "Taking life very seriously, very conscientiously."]

for point in blocks:
  document.add_paragraph(point, style='ListBullet')
#
#-------------
# Reflector
#-------------
i+=1

document.add_heading('Reflector learning style', level=3)
document.add_paragraph("The reflector learning style has an average and an uncertainty (in parenthesis) of " + str(mean_uncert(xmean[i],dxmean[i])) + " points." )
#  if ( xmean[i] >= 18 ):
document.add_paragraph("People with scores higher than 18 on the test learn better in learning situations where they can:")

# Bullet points: preferred reflector learning situations
reflector_strengths = [
        "Observe.",
        "Reflect on activities.",
        "Exchange opinions with others by prior agreement.",
        "Reach decisions at their own pace.",
        "Work without pressures or mandatory deadlines.",
        "Review what's been learned, what has happened.",
        "Investigate thoroughly.",
        "Gather information.",
        "Probe to get to the heart of the matter.",
        "Think before acting."
]

for point in reflector_strengths:
  document.add_paragraph(point, style='List Bullet')

# Discomfort situations
document.add_paragraph("Similarly, they'll feel uncomfortable in situations that demand:")

reflector_discomforts = [
        "Taking the forefront.",
        "Acting as a leader.",
        "Chairing meetings or debates.",
        "Performing in front of observing individuals.",
        "Playing a role.",
        "Participating in situations that require action without planning.",
        "Doing something without prior notice. Presenting an idea spontaneously."]

for point in reflector_discomforts:
  document.add_paragraph(point, style='List Bullet')

document.add_paragraph("If someone has scored below 14, they might consider the following suggestions to enhance the reflector style:")

reflector_suggestions = [
          "Practice observation. Study people's behavior. Note who speaks the most, who interrupts, how often the teacher summarizes... Study non-verbal behavior, when people look at the clock, cross their arms, bite their pencil...",
          "Keep a personal diary. Reflect on the day's events and see if any conclusions can be drawn from them.",
          "Practice review after a meeting or event. Review the sequence of events, what went well, what could be improved. Record a dialogue or conversation on a tape recorder and play it back at least twice. Make a list of lessons learned in this way."]

for point in reflector_suggestions:
  document.add_paragraph(point, style='List Bullet')

# Blocks that hinder reflector style
document.add_paragraph("Blocks that hinder the development of the reflector style:")

reflector_blocks = [
          "Not having enough time to plan and think.",
          "Preferring or enjoying quickly changing from one activity to another.",
          "Being impatient to start action.",
          "Having resistance to listening carefully and analytically.",
          "Resistance to presenting things in writing."]

for point in reflector_blocks:
  document.add_paragraph(point, style='List Bullet')
#
#-------------
# Theorist
#-------------
i+=1

document.add_heading('Theorist learning style', level=3)
document.add_paragraph("The theorist learning style has an average and an uncertainty (in parenthesis) of " + str(mean_uncert(xmean[i],dxmean[i])) + " points." )
#  if ( xmean[i] >= 14 ):
document.add_paragraph("People with scores higher than 14 on the test prefer learning opportunities where they can:")

# Add bullet points for preferred learning situations
preferred_learning = [
        "Feel in structured situations with a clear purpose.",
        "Fit all data into a system, model, concept, or theory.",
        "Have time to methodically explore associations and relationships between ideas, events, and situations.",
        "Have the chance to question.",
        "Participate in a question-and-answer session.",
        "Test methods and logic that form the basis of something."]

for item in preferred_learning:
  document.add_paragraph(item, style='List Bullet')

# Add paragraph about avoided situations
document.add_paragraph("At the same time, they'll avoid learning situations where they have to:")

# Add bullet points for avoided situations
avoided_learning = [
        "Be forced to do something without a clear context or purpose.",
        "Participate in situations where emotions and feelings predominate.",
        "Engage in unstructured activities with uncertain or ambiguous purposes.",
        "Take part in open-ended problems.",
        "Act or decide without a foundation of principles, concepts, policies, or structure."]

for item in avoided_learning:
  document.add_paragraph(item, style='List Bullet')

document.add_paragraph("If someone has scored below 10, they might want to consider the following suggestions to enhance the theorist style:")

# Suggestions to enhance the theorist style
enhancement_suggestions = [
          "Read something dense that provokes thinking for 30 minutes each day (Logic, Linguistics, Theories...). Then try to summarize what they have read using their own words.",
          "Practice detecting inconsistencies or weaknesses in other people's arguments, in reports, in newspaper articles... Take two newspapers with different ideologies and regularly perform a comparative analysis of the differences in their perspectives.",
          "Take a complex situation and analyze it to pinpoint why it developed that way, what could have been done differently, and at what point. Historical or everyday life situations. Analyze how they've used their own time..."]

for item in enhancement_suggestions:
  document.add_paragraph(item, style='List Bullet')

# Add final paragraph about blocks
document.add_paragraph("Blocks that hinder the development of the theorist style:")

# Blocks list
theorist_blocks = [
          "Being swayed by first impressions.",
          "Preferring intuition and subjectivity.",
          "Disliking structured and organized approaches.",
          "Preferring spontaneity and risk-taking."]

for item in theorist_blocks:
  document.add_paragraph(item, style='List Bullet')
#
#-------------
# Pragmatist
#-------------
i+=1
## In Cantino, it is considered xmean[i] >= 16 instead of xmean[i] >= 14.
## We set it equal to 14, which corresponds to a high tendency towards the
## pragmatic learning style, as in the other ones.
document.add_heading('Pragmatist learning style', level=3)
document.add_paragraph("The pragmatist learning style has an average and an uncertainty (in parenthesis) of " + str(mean_uncert(xmean[i],dxmean[i])) + " points." )
#  if ( xmean[i] >= 14 ):
document.add_paragraph("People with scores higher than 16 on the test find more comfortable in situations that allow them to:")

# Preferred pragmatist learning situations
pragmatist_strengths = [
        "Learn techniques for doing things with evident practical advantages.",
        "Be exposed to a model they can emulate.",
        "Acquire immediately applicable techniques in your work.",
        "Have the immediate possibility to apply what they've learned, to experiment.",
        "Develop action plans with a clear result.",
        "Give directions, suggest shortcuts."]

for point in pragmatist_strengths:
  document.add_paragraph(point, style='List Bullet')

# Discomfort situations
document.add_paragraph("Likewise, they'll find difficulty in situations that demand:")

pragmatist_discomforts = [
        "Realizing that learning isn't related to an immediate need that they recognize or cannot see.",
        "Perceiving that the learning doesn't have immediate importance or practical benefit.",
        "Learning something that is distant from reality.",
        "Learning theories and general principles.",
        "Working without clear instructions on how to do it."]

for point in pragmatist_discomforts:
  document.add_paragraph(point, style='List Bullet')

document.add_paragraph("If someone has scored below 10, the following suggestions will be useful for enhancing the pragmatist style:")

pragmatist_suggestions = [
          "Gather techniques, practical ways of doing things. They can cover anything that might be useful: analytical techniques, interpersonal skills, assertiveness, time-saving presentation techniques, statistics, memory improvement techniques...",
          "Seek help from individuals with demonstrated experience. They can observe our technique and advise on how to improve it.",
          "Focus on developing action plans in meetings and discussions of all kinds. These action plans should be specific and have deadlines. Make it a rule never to leave a meeting, debate, or class without a list of actions for yourself."]

for point in pragmatist_suggestions:
  document.add_paragraph(point, style='List Bullet')

# Blocks that hinder development
document.add_paragraph("Blocks that hinder the development of the pragmatist style:")

pragmatist_blocks = [
          "Interest in the perfect solution rather than practicality.",
          "Considering useful techniques as exaggerated simplifications.",
          "Always leaving topics open and not committing to specific actions.",
          "Believing that others' ideas won't work when applied to their situation.",
          "Enjoying marginal topics or getting lost in them."]

for point in pragmatist_blocks:
  document.add_paragraph(point, style='List Bullet')
#
#==================================================
# FINAL CONTACT PAGE
#==================================================
document.add_page_break()
document.add_paragraph(' ')
closing_paragraph = document.add_paragraph()
closing_paragraph.add_run(
    "If you have any comments or suggestions, feel free to reach out by sending an email with your feedback to fabio.revuelta@upm.es.").italic = True
document.add_paragraph(' ')
#
#==================================================
# End report
#==================================================
# Get the current date and time
current_datetime = datetime.datetime.now()
#
# Format the date and time as a string
formatted_datetime = current_datetime.strftime("%Y-%m-%d %H:%M:%S")
printt(' ')
printt('Program continues on: ' + formatted_datetime)
printt(' ')
#
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Saving summary report...')
printt('-------------------------------------------')
# Docx
printt('   Docx...')
report_name='CHAEA_learning_styles_summary_report'
report_docx=report_name+'.docx'
report_docx_file = output_gen+'/'+report_docx
report_pdf_file  = output_gen+'/'+report_name+'.pdf'
printt('     Saving ' + report_docx + '...')
document.save(report_docx_file)
printt('   Docx saved!')
printt(' ')
#
# PDF
def convert_docx_to_pdf(operating_system, docx_file):
  print('OS : ', operating_system)
  if(operating_system == 'linux'):
    try:
      subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", "output", docx_file])
      printt("     Report successfully saved in PDF format, too!")
    except Exception as e:
      printt(f"Conversion failed: {e}")
  elif(operating_system == 'win'):
    try:
      subprocess.run(["docx2pdf", report_docx_file, report_pdf_file])
      printt("     Report successfully saved in PDF format, too!")
    except Exception as e:
      printt(f"Conversion failed: {e}")
  elif(operating_system == 'mac'):
    try:
      subprocess.run(["docx2pdf", report_docx_file, report_pdf_file])
      printt("     Report successfully saved in PDF format, too!")
    except Exception as e:
      printt(f"Conversion failed: {e}")
#
#
printt('   PDF...')
convert_docx_to_pdf(operating_system, report_docx_file)
printt('   PDF saved!')
#
printt('-------------------------------------------')
printt(' Summary report saved!')
printt('-------------------------------------------')
#---------------------------------------------------
#
#===================================================
#===================================================
printt('===========================================')
printt('SUMMARY REPORT DONE!')
printt('===========================================')
printt('  ')
printt('  ')
printt('  ')
printt('  ')
printt('===========================================')
printt('E-MAIL CONTACT...')
printt('===========================================')
def send_feedback():
    email = "fabio.revuelta@upm.es"
    subject = "Feedback on CHAEA3S"
    body = "Hello Fabio,\n\nI have a suggestion for CHAEA3S software package:\n\n"

    # Create a mailto link
    mailto_link = f"mailto:{email}?subject={subject}&body={body.replace(' ', '%20')}"

    # Ask user if they want to send feedback
    user_input = input("Would you like to send feedback with any comments or suggestions? If so, type 'y' or 'yes' (in uppercase or lowercase): ").strip()

    if user_input in ["y", "Y", "yes", "YES", "Yes", "yEs", "yeS", "YEs", "YeS", "yES"]:
        print(f"Opening email client... The email will be sent to: {email}")
        webbrowser.open(mailto_link)
    else:
        print("Feedback request canceled.")

# Run the function
if __name__ == "__main__":
    send_feedback()
printt('===========================================')
printt('E-MAIL CONTACT DONE!')
printt('===========================================')
printt(' ')
printt(' ')
#
# Get the current date and time
current_datetime = datetime.datetime.now()
#
# Format the date and time as a string
formatted_datetime = current_datetime.strftime("%Y-%m-%d %H:%M:%S")
#===================================================
#===================================================
#===================================================
#===================================================
printt('===========================================')
printt('===========================================')
printt('BINGO PITINGO!')
printt('===========================================')
printt('===========================================')
printt(' ')
printt('===========================================')
printt('===========================================')
printt('Program finishes on: ' + formatted_datetime)
printt('===========================================')
printt('===========================================')
printt(' ')
#==================================================================
#==================================================================
#==================================================================
#==================================================================#!/user/bin/env python
#===================================================
#===================================================
# CHAEA3S:
# CHAEA AUTOMATIC ANALYSIS OF LEARNING STYLES
#===================================================
#===================================================
# VERSION 7.g
#
# This program implements a principal component 
# analysis of the learning styles of a group of
# students as prescribed by CHAEA:
# activist, theorist, pragmatist, and reflector.
# The program is based on the program 
# learning_styles_v7e.py, by Fabio Revuelta.
#
# Further information can be found in the reference:
# J. Ablanque, V. Gabaldon, P. Almendros, J. C. Losada,
# R. M. Benito,  and F. Revuelta
# “CHAEA3S: A software package for comprehensive analysis
# of learning styles and academic performance”
# Computers and Education (2025).
#
# The software provided here is distributed on an “as-is” basis,
# without any warranties or guarantees of any kind.
# While we have made every effort to ensure its accuracy
# and reliability, we cannot be held responsible for any
# unintended consequences, errors, or issues that may arise from its use.
# Users are encouraged to thoroughly test the software,
# review the source code, and exercise due diligence before deploying it 
# in any critical environment.
#
# Version 7.g also performs two aditional optional calculations,
# if the corresponding input data is provided.
# On the one hand, if the learning styles of the
# teacher(s)/instructor(s) are given as input(s),
# the learning styles of the students are plotted as
# a function of them on the plain of the two main principal
# components using a certain color.
# On the other hand, if the marks of the students are provided,
# a study of their relationship with the learning styles is also
# conducted as a funciton of the original learning-styles classification,
# and as a function of the principal components.
# Furthermore, a clustering analysis based on K-means algorithm
# is also performed.
#
# In version 7.g, the default web browser is opened after execution
# in case the used wants to contact the software authors.
#
# By using this software, you agree to hold the developers harmless 
# from any liability, damages, or losses resulting from its use.
# Additionally, we kindly request that you cite the reference mentioned
# above when sharing or distributing it, and 
# acknowledge Fabio Revuelta as the author of this software.
# 
# Feel free to adapt this disclaimer further to match your specific context.
# And remember to give credit to us :)
#
# Copyright 2025 Fabio Revuelta
# 
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# 
#     http://www.apache.org/licenses/LICENSE-2.0
# 
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
# Author: Fabio Revuelta
#         Grupo de Sistemas Complejos
#         Universidad Politecnica de Madrid
#         Madrid, August 2025
#
#===================================================
# PRELIMINARIES
#===================================================
from mpl_toolkits.mplot3d import Axes3D
from mpl_toolkits.mplot3d import proj3d
from matplotlib.patches import FancyArrowPatch
from matplotlib import cm
#
import pandas as pd 
import matplotlib
import matplotlib.pyplot as plt
import numpy as np
import os, fnmatch
from matplotlib.ticker import  MultipleLocator, FormatStrFormatter
from decimal import Decimal
from numpy.linalg import eig
#
import statistics
import scipy.stats as st
from mpl_toolkits.mplot3d.art3d import Poly3DCollection
import random
#
from random import seed
from random import gauss
from scipy.stats import linregress
from scipy.optimize import curve_fit
#
from sklearn.cluster import KMeans
#
import itertools
import datetime
#
# Library to contact via e-mail
import webbrowser
#
# Necessary libraries for the summary report
import subprocess
#
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
#
#
# Necessary for the log file
import os
import re
import sys
import datetime
#
## For Windows
#import codecs
##sys.stdout = codecs.getwriter('utf-8')(sys.stdout.detach())   
#
#
#===================================================
# CHAEA3S.LOG FILE
#===================================================
# Current folder
current_folder = os.getcwd()
print('Current folder : ' + current_folder)
print(' ')
#
# Output folders
output_gen = current_folder + '/output'
os.makedirs(output_gen, exist_ok=True) 
print('Output general folder         : ' + output_gen)
print(' ')
#
# Define the file name for logging
log_file = "output/chaea3s.log"
#
if os.path.exists(log_file):
  os.remove(log_file)
  print(f"File '{log_file}' removed successfully.")
else:
  print(f"File '{log_file}' did not previously exist.")
#
print(f"File '{log_file}' created.")
print()
#
# Create a new log_file
with open(log_file, "w") as f:
  f.write('==============================\n')
  f.write('CHAEA3S.LOG\n')
  f.write('==============================\n')
  f.write('This file contains all the information on the execution of CHAEA3S package\n')
  f.write('\n')
  pass
#
# Define a custom print function
def printt(*args):
  print(*args, end=' ')
#   
  with open(log_file, "a", encoding='utf-8') as f:
    f.write(str(*args)+'\n')
#
# Redirect stdout to log file
#sys.stdout = open(log_file, "a")
#
# Get the current date and time
current_datetime = datetime.datetime.now()
#
# Format the date and time as a string
formatted_datetime = current_datetime.strftime("%Y-%m-%d %H:%M:%S")
printt('Program starts on: ' + formatted_datetime)
printt(' ')
#
# Check operating system
# (necessary to transform the docx report in pdf)
if(sys.platform == 'linux' or sys.platform == 'linux2'):
  operating_system = 'linux'
  printt('The operating system is Linu.x')
elif(sys.platform == 'darwin'):
  operating_system = 'mac'
  printt('The operating system is macOS (Darwin).')
elif(sys.platform == 'win32' or sys.platform == 'win64'):
  operating_system = 'win'
  printt('The operating system is Windows.')
#
printt(' ')
#
#===================================================
# PLOT PARAMETERS
#===================================================
#----------------------------------------------------
def write_number_with_decimals(x, n):
  return "{:.{}f}".format(x, n)
#----------------------------------------------------
def find_number_position(string, number):
  try:
    position = string.index(str(number))
    return position
  except ValueError:
    return -1  # Number not found in the string
#----------------------------------------------------
def mean_uncert(x, dx):
# This subroutine returns a string with the average mean
# and the corresponding uncertainty with the correct
# number of decimals
# Convert numbers to strings if necessary
  if isinstance(x, str):
    x_str  = x
    xin    = float(x)
  else:
    x_str  = str(x)
    xin    = x

  if isinstance(dx, str):
    dx_str = dx
    dxin   = float(dx)
  else:
    dx_str = str(dx)

# Check if dx==nan
  if ( dx_str == 'nan' or dx_str == '0.0'):
     x_out = write_number_with_decimals(xin, 1) + '(nan)'
  else:
#   Get the three most significant digits of dx
#   (in case the point is included)
    for digit in dx_str:
      if(digit != '0' and digit != '.'):
        dx0 = int(digit)
        break

    pos0 = find_number_position(dx_str, dx0)
    dec  = find_number_position(dx_str, '.')
    if(dec == -1):
      dx_str = dx_str + '.0'
      dec  = find_number_position(dx_str, '.')

    if(dx0 != 1):
#     dx0 > 1
      if( pos0 < dec ):
        dx_out = dx_str[pos0]
        for j in range(pos0, dec-1):
          dx_out = dx_out + '0'
      else:
        dx_out = dx_str[0:pos0+1]
    else:
#     dx0 = 1
      if( pos0 > dec ):
        if(pos0+1 == len(dx_str)):
          dx_out = dx_str + '0'
        else:
          dx_out = dx_str[0:pos0+2]
      else:
        if(pos0+1 == dec):
          if(pos0+2 == len(dx_str)):
            dx_out = dx_str + '0'
          else:
            dx_out = dx_str[0:pos0+3]
        else:
          dx_out = dx_str[pos0] + dx_str[pos0+1]
          for j in range(pos0+1, dec-1):
            dx_out = dx_out + '0'

#   Count the number of decimals
    num_decimals = len(dx_out.split('.')[1]) if '.' in dx_out else 0

#   Get the same number of decimals from x
    x_out = write_number_with_decimals(xin, num_decimals) + '(' + dx_out + ')'

  return x_out
#----------------------------------------------------
def round_until_two_non_9(xin):

    if isinstance(xin, str):
      x = float(xin)
    else:
      x = xin

    if (-1 < x < 1):
      s = f"{x:.20f}"  # Convert to string with many decimals
      sign = "-" if x < 0 else ""
      decimals = s.split(".")[1]

      result = "0."
      non9_count = 0
      for d in decimals:
        result += d
        if d != '9':
          non9_count += 1
        if non9_count == 2:
           break
      output = str(sign + result)
    else:
      if( x == -1 ):
        output = '-1.0'
      else:
        if( x == 1 ):
          output = '1.0'
        else:
          if(np.isnan(x)):
            output = 'nan'
          else:
            raise ValueError("x must be between -1 and 1")

    return output
#----------------------------------------------------
#----------------------------------------------------
printt('===========================================')
printt('PLOTTING PARAMETERS...')
printt('===========================================')
printt(' ')
#
# Labels for the learning styles (LSs)
List_LS        = ['activist', 'reflector', 'theorist', 'pragmatist']
Label_LS       = ['Activist', 'Reflector', 'Theorist', 'Pragmatist']
Label_LS_print = ['Activist  ', 'Reflector ', 'Theorist  ', 'Pragmatist']
#
# Labels for the principal components (PCs)
Label_PC = [0, 1, 2, 3]
#
Label_PC   = list(itertools.chain(range(0, len(Label_LS))))
Label_PCPC = list(itertools.chain(range(0, len(Label_LS))))

for i in range(0, len(Label_LS)):
  Label_PC[i] = str(Label_PC[i])
  Label_PCPC[i] = 'PC' + Label_PC[i]
#
# Tendency labels
Label_tendencies       = ['Very low', 'Low', 'Moderate', 'High', 'Very high']
Label_tendencies_print = ['Very low ', 'Low      ', 'Moderate ', 'High     ', 'Very high']
#
#---------------------------------------------------
# Size, Label_LS and ticks of the figures
w_fig = 28
h_fig = w_fig / 1.61
#
# Original: labels = 30, ticks = 0.9*labels
labelsize = 50
ticksize  = 0.90 * labelsize
#
# Quartiles
labelsize_quartile = 50
ticksize_quartile  = 0.9 * labelsize_quartile
#
# W(PR)
labelsize_PR = 50
ticksize_PR  = 0.9 * labelsize_PR
#
# W(PC)
labelsize_PC = 50
ticksize_PC  = 0.9 * labelsize_PC
#
# 3D plots
labelsize_3D = 50
ticksize_3D  = 0.9 * labelsize_3D
labelpad_3D  = labelsize_3D
#
width_PC_vectors_3D = [9, 7, 5, 4]
width_PC_vectors_2D = [9, 7, 5, 5]
#
colors_proj     = 'tab:blue'
edgecolors_proj = 'tab:blue'
#
# Colors of the projections on the Cartesian planes
# if each PC corresponds to a different LS
colors_proj_PC     = 'tab:blue'
edgecolors_proj_PC = 'tab:blue'
#
# Colors of the projections on the Cartesian planes
# if more than one PC is associated with the same LS
colors_proj_PC2     = 'tab:blue'
edgecolors_proj_PC2 = 'tab:blue'
#
# Mean average potting parameters
size_mean_ref = 30000
mean_symbol   = '*'
mean_color    = 'tab:blue'
mean_color2   = 'gray'
#
# Uncertainty of the average mean
mean_uncert_color = 'pink'
mean_alpha        = 0.5
#
# 2D and 3D plots for the PCs
twenty_symbol = ['s', 's', 's', 's', 's', 'D']
twenty_color  = ['r', 'green', 'purple', 'orange', 'k', 'brown']
twenty_line   = ['--', '--', '--', '--', '--', 'dotted']
twenty_alpha  = 0.5
#
twenty_width_3D = [2, 2, 1, 1, 1, 1, 2]
twenty_width_2D = [2, 2, 1, 1, 2, 2, 2]
twenty_width_W  = [6, 6, 6, 6, 3, 4, 2]
#
tendency_color = ['pink', 'brown', 'b']
#
origin_symbol = '^'
origin_color  = 'k'
origin_alpha  = 0.5


# Clustering parameters
cluster_symbol = ['o', 's', '^', 'v', '*', 'D']
cluster_color  = ['dodgerblue', 'darkred', 'forestgreen', 'navy', 'goldenrod', 'saddlebrown' ]
cluster_alpha  = 0.5
#---------------------------------------------------
# Number of plots
Nplots = 1
printt('  Nplots : ' + str(Nplots))
# If Nplots = 1, then only one 2D and 3D
# representions as a function of the LSs are
# performed. Otherwise, all (2D and 3D)
# possible combinations are consided
#
#---------------------------------------------------
# Plotting limits for the plots as a function 
# of the LSs
xmin =  0
xmax = 20.4
#
ymin = xmin
ymax = xmax
#
zmin = xmin
zmax = xmax
#
xmin3D = xmin
xmax3D = xmax
#
ymin3D = xmin3D
ymax3D = xmax3D
#
zmin3D = xmin3D
zmax3D = xmax3D
#
printt( ' xmin : ' + str(xmin))
printt( ' xmax : ' + str(xmax))
printt(' ')
printt( ' ymin : ' + str(ymin))
printt( ' ymax : ' + str(ymax))
printt(' ')
printt( ' zmin : ' + str(zmin))
printt( ' zmax : ' + str(zmax))
printt(' ')
#---------------------------------------------------
# Parameters for the LS statistical analysis
LS_mean_line  = 'dotted'
LS_mean_width = 2.5
LS_mean_alpha = 0.05
#
LS_color = ['r', 'green', 'purple', 'orange', 'pink']
LS_line  = ['-',     '--', 'dashdot',    'dotted']
LS_areas = ['b', 'orange', 'green', 'pink']
LS_alpha = 0.25
LS_alpha2 = 0.60
LS_alpha3 = 1.0 #0.60
#
affinity_alpha = 1.0 #0.60
#
LS_hist_alpha  = 0.25
LS_hist_alpha2 = 0.40
#
LS_W_width      = 4
#
# Parameters of the participation ratios (PRs)
Label_PR = ['PR (LS)', 'PR (PC)']
#
PR_color = ['r', 'b']
PR_line  = ['-',     '--', 'dashdot',    'dotted']
PR_areas = ['brown', 'orange', 'green', 'pink']
PR_alpha = 0.40
#
PR_mean_line  = 'dotted'
PR_mean_width = 2.5
PR_mean_alpha = 0.05
#
PR_W_width      = 2.5
PC_W_width      = 4
#
# Parameters of the quartiles of the PRs
quartile_width = 2
quartile_color = 'k'
quartile_line  = 'dotted'
#
PR_hist_quartile_color = ['r', 'b', 'green', 'orange', 'purple', 'pink']
PR_hist_quartile_alpha = 1.0
#---------------------------------------------------
# Parameters of the eigenvectors and the projections
# on them (principal components, PCs)
PC_color = ['r', 'green', 'purple', 'orange', 'pink', 'brown']
PC_line  = ['-', '--', 'dashdot', 'dotted']
PC_line_bis = ['-', '-', '-', '-']
PC_areas = ['b', 'orange', 'green', 'pink']
PC_width = 2.5
#
PC_mean_line  = 'dotted'
PC_mean_width = 2.5
PC_mean_alpha = 0.05
#
PC_alpha = 0.25
PC_alpha2 = 0.60
#
PC_hist_alpha  = 0.25
PC_hist_alpha2 = 0.40
#
Weibull_width = 2
#
max_iter_Weib = 5000
#---------------------------------------------------
# Parameters the PCs
colors_proj_lines = ['r', 'purple']
style_proj_lines  = ['-', '--']
colors_proj_areas = ['b', 'orange', 'green', 'pink']
#---------------------------------------------------
# Possible combinations of the learning styles in pairs
learning_pairs = [[0,1], [0,2], [0,3], [1,2], [1,3], [2, 3]]
#
# Possible combinations of the learning styles in trios
learning_trios = [[0,1,2], [0,1,3], [0,2,3], [1,2,3]]
#----------------------------------------------------------
printt('-------------------------------------------')
printt('  Definition of tendency/plotting functions...')
printt('-------------------------------------------')
printt(' ')
# Tendency depending on the reference LS
# (activist -a-, theorist-t-, pragmatist -p-, reflector -r-).
#
def scatter_tendency(ls, value):
  # Activist
  if(ls == 'Activist'):
    if value<=6:
      tendency='vl'
    else:
      if value<=8:
        tendency='l'
      else:
        if value<=12:
          tendency='m'
        else:
          if value<=14:
            tendency='h'
          else:
            if value<=20:
              tendency='vh'
            else:
              printt('Wrong input for the ls :' + ls)
  else:
    # Reflector
    if(ls == 'Reflector'):
      if value<=10:
        tendency='vl'
      else:
        if value<=13:
          tendency='l'
        else:
          if value<=17:
            tendency='m'
          else:
            if value<=19:
              tendency='h'
            else:
              if value<=20:
                tendency='vh'
              else:
                printt('Wrong input for the ls :' + ls)
    else:
      # Theorist
      if(ls == 'Theorist'):
        if value<=6:
          tendency='vl'
        else:
          if value<=9:
            tendency='l'
          else:
            if value<=13:
              tendency='m'
            else:
              if value<=15:
                tendency='h'
              else:
                if value<=20:
                  tendency='vh'
                else:
                  printt('Wrong input for the ls :', ls)
      else:
        # Pragmatist
        if(ls == 'Pragmatist'):
          if value<=8:
            tendency='vl'
          else:
            if value<=10:
              tendency='l'
            else:
              if value<=13:
                tendency='m'
              else:
                if value<=15:
                  tendency='h'
                else:
                  if value<=20:
                    tendency='vh'
                  else:
                    printt('Wrong input for the ls :', ls)
  return tendency
#===================================================
def tendency_long_name(tendency):
  if(tendency == 'vl'):
    return 'Very low'
  else:
    if(tendency == 'vll'):
      return 'Very low/ Low'
    else:
      if(tendency == 'l'):
        return 'Low'
      else:
        if(tendency == 'lm'):
          return 'Low/ Moderate'
        else:
          if(tendency == 'm'):
            return 'Moderate'
          else:
            if(tendency == 'mh'):
              return 'Moderate/ High'
            else:
              if(tendency == 'h'):
                return 'High'
              else:
                if(tendency == 'hvh'):
                  return 'High/ Very high'
                else:
                  if(tendency == 'vh'):
                    return 'Very high'
                  else:
                    return 'Error'
#===================================================
def tendency_intermediate(ls, value):
  # Here, we also take into account the transition values
  # for the group of students, i.e., 6.5 points for active
  # is not very low (it is larger than 6) but smaller than 
  # low (it is smaller than 7).
  # Activist
  if(ls == 'Activist'):
    if value<=6:
      tendency='vl'
    else:
      if value<7:
        tendency='vll'
      else:
        if value<=8:
          tendency='l'
        else:
          if value<9:
            tendency='lm'
          else:
            if value<=12:
              tendency='m'
            else:
              if value<13:
                tendency='mh'
              else:
                if value<=14:
                  tendency='h'
                else:
                  if value<15:
                    tendency='hvh'
                  else:
                    if value<=20:
                      tendency='vh'
                    else:
                      printt('Wrong input for the ls :', ls)
  else:
    # Reflector
    if(ls == 'Reflector'):
      if value<=10:
        tendency='vl'
      else:
        if value<11:
          tendency='vll'
        else:
          if value<=13:
            tendency='l'
          else:
            if value<14:
              tendency='lm'
            else:
              if value<=17:
                tendency='m'
              else:
                if value<18:
                  tendency='mh'
                else:
                  if value<=19:
                    tendency='h'
                  else:
                    if value<20:
                      tendency='hvh'
                    else:
                      if value<=20:
                        tendency='vh'
                      else:
                        printt('Wrong input for the ls :', ls)
    else:
      # Theorist
      if(ls == 'Theorist'):
        if value<=6:
          tendency='vl'
        else:
          if value<7:
            tendency='vll'
          else:
            if value<=9:
              tendency='l'
            else:
              if value<10:
                tendency='lm'
              else:
                if value<=13:
                  tendency='m'
                else:
                  if value<14:
                    tendency='mh'
                  else:
                    if value<=15:
                      tendency='h'
                    else:
                      if value<16:
                        tendency='hvh'
                      else:
                        if value<=20:
                          tendency='vh'
                        else:
                          printt('Wrong input for the ls :', ls)
      else:
      # Pragmatist
        if(ls == 'Pragmatist'):
          if value<=8:
            tendency='vl'
          else:
            if value<9:
              tendency='vll'
            else:
              if value<=10:
                tendency='l'
              else:
                if value<11:
                  tendency='lm'
                else:
                  if value<=13:
                    tendency='m'
                  else:
                    if value<14:
                      tendency='mh'
                    else:
                      if value<=15:
                        tendency='h'
                      else:
                        if value<16:
                          tendency='hvh'
                        else:
                          if value<=20:
                            tendency='vh'
                          else:
                            printt('Wrong input for the ls :', ls)
  return tendency
#
#===================================================
def scatter_properties(tendency, value):
  # It returns a color and a number, depending on whether
  # the tendency towards a certain learning style is
  # very low/low (blueish), moderate (greenish) or high/very high (redish).
  #
  # Scatter color and symbol
  if(tendency=='vl' or tendency=='l'):
    scatter_color='b'
    scatter_symbol='o'
  else:
    if(tendency=='m'):
      scatter_color='b'
      scatter_symbol='o'
    else:
      if(tendency=='h' or tendency=='vh'):
        scatter_color='b'
        scatter_symbol='o'
  #
  # Scatter filling
  alphamin=0.1
  alphamax=0.8
  scatter_alpha=alphamin+(alphamax-alphamin)*value/20
  #
  # Scatter size
  scattersizemax = 150*100/L # The more points, the smaller the scatter size
  scattersizemin =  20*100/L
  #
  scatter_size=scattersizemin+(scattersizemax-scattersizemin)*value/20
  if ( scatter_size < 1):
    scatter_size = 1
  else:
    if ( scatter_size > 500):
      scatter_size = 500
  #               
  return [tendency, scatter_color, scatter_size, scatter_alpha, scatter_symbol]
printt('-------------------------------------------')
printt('  Definition of tendency/plotting functions done!')
printt('-------------------------------------------')
#
#
#---------------------------------------------------
# Fitting functions
#---------------------------------------------------
#
printt('-------------------------------------------')
printt('    Fitting functions...')
printt('-------------------------------------------')
printt(' ')
# We fit the projections of the LSs of each student and the
# corresponding probabilities using Weibull distributions.
# The fitting is performed considering the cumulative 
# distribution functions 
# (regular Weibull distribution for the probabilities
# and the translated Weibull distribution for the
# projections)
#
# Cumulative function for Weibull distribution
def Wweibull(x, alpha_Weibull, k_Weibull):
  return 1 - np.exp(-(x/alpha_Weibull)**k_Weibull)
#
# Probability density function for Weibull distribution
def Pweibull(x, alpha_Weibull, k_Weibull):
  return (k_Weibull/alpha_Weibull) * (x/alpha_Weibull)**(k_Weibull-1) * np.exp(-(x/alpha_Weibull)**k_Weibull)
#  
# Cumulative function for the translated Weibull distribution
# Translated Weibull distribution for the projections on the PCs
def Wweibull_translated(x, alpha_Weibull, k_Weibull, theta_Weibull):
  return 1 - np.exp(-((x-theta_Weibull)/alpha_Weibull)**k_Weibull)
# 
# Probability density function for the translated Weibull distribution
def Pweibull_translated(x, alpha_Weibull, k_Weibull, theta_Weibull):
  return (k_Weibull/alpha_Weibull) * ((x-theta_Weibull)/alpha_Weibull)**(k_Weibull-1) * np.exp(-((x-theta_Weibull)/alpha_Weibull)**k_Weibull)
printt(' ')
#
printt('-------------------------------------------')
printt('    Fitting functions done!')
printt('-------------------------------------------')
printt(' ')
#
#===================================================
#===================================================
#
#===================================================
printt('===========================================')
printt('PLOTTING PARAMETERS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# FOLDERS
#===================================================
printt('===========================================')
printt('FOLDERS...')
printt('===========================================')
printt(' ')
#
# Current folder
current_folder = os.getcwd()
printt('Current folder : ' + current_folder)
printt(' ')
#
# Input folder
input_folder = current_folder + '/input/'
printt('Input folder   : ' + input_folder)
printt(' ')
#
# Output folders
output_gen = current_folder + '/output'
os.makedirs(output_gen, exist_ok=True) 
printt('Output general folder         : ' + output_gen)
printt(' ')
#
output = output_gen + '/figs/'
os.makedirs(output, exist_ok=True) 
printt('Output figures folder         : ' + output)
printt(' ')
#       
output_participants = output + '/participants'
os.makedirs(output_participants, exist_ok=True) 
printt('Output-participants folder : ' + output_participants)
printt(' ')
#
output_statistics = output + '/statistics'
os.makedirs(output_statistics, exist_ok=True) 
printt('Output-statistics folder : ' + output_statistics)
printt(' ')
#
output_statistics_pr = output_statistics + '/pr'
os.makedirs(output_statistics_pr, exist_ok=True) 
printt('Output-PR folder : ' + output_statistics_pr)
printt(' ')
#
output_statistics_ls = output_statistics + '/ls'
os.makedirs(output_statistics_ls, exist_ok=True) 
printt('Output-statistics-LS folder : ' + output_statistics_ls)
printt(' ')
#
# Marks folders
output_participants_ls_marks = output_participants + '/marks'
os.makedirs(output_participants_ls_marks, exist_ok=True)
printt('Output-LS-marks folder : ' + output_participants_ls_marks)
#
output_participants_pc_marks = output_participants + '/marks'
printt('Output-PC-marks folder : ' + output_participants_pc_marks)
printt('  ')
printt('===========================================')
printt('FOLDERS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#
#
#
#===================================================
# INPUT READING
#===================================================
printt('===========================================')
printt('READING INPUT FILES...')
printt('===========================================')
printt(' ')
nan = 0
#
studentsin = [] # List with the names of all the students
students   = [] # List with the names of the students that have correct input data
data       = [] # Input data

data_marks = ['Excel', 'Sheet', 'Student'] # List with the names of the students that have correct input data and their marks or grades (if given)
datain_error_students = []
#
# Iteration over all input files (.xls and .xlsx)
for files in os.listdir(input_folder):
# 
# xls and xlsx files
  if ((files.endswith('.xls')) or (files.endswith('.xlsx'))):
    studentsin.append(files)
#
studentsin = sorted(studentsin)
#
for filei in studentsin:
#
    data_in = []
    opt = 0
    try:
      data_in = pd.read_excel(input_folder + filei, header=None, sheet_name='CÁLCULO')
    except FileNotFoundError:
      print(f"File '{filei}' not found.")
    except ValueError:
      print("Worksheet 'CÁLCULO' not found. Trying alternative sheet names...")
      try:
        data_in = pd.read_excel(input_folder + filei, header=None, sheet_name='CALCULO')
      except ValueError:
        try:
          data_in = pd.read_excel(input_folder + filei, header=None, sheet_name='CALCULATION')
          opt = 1
        except ValueError:
          print("Neither 'CALCULATIONS', 'CALCULO' nor 'CÁLCULO' sheet found.")

    if not data_in.empty:
      data_LS = data_in.loc[29].values          # Extract the LS data (I)
      if( opt != 1):
        line    = np.delete(data_LS, [0,1,6,7,8]) # Extract the LS data (II)
      else:
        line    = np.delete(data_LS, [0,1,2,7,8]) # Extract the LS data (II)
#
    file_line = filei+'      '+str(line[0])+' '+str(line[1])+' '+str(line[2])+' '+str(line[3])
#
    try:   
      n0 = float(line[0])
      n1 = float(line[1])
      n2 = float(line[2])
      n3 = float(line[3])
      error = 'False'      
    except ValueError:
      error = 'True'
#
    if ( (error == 'True') or
       ( np.isnan(n0) or np.isnan(n1) or np.isnan(n2) or np.isnan(n3) 
         or (n0 < 0) or (n0 > 20) or (n1 < 0) or (n1 > 20)
         or (n2 < 0) or (n2 > 20) or (n3 < 0) or (n3 > 20) ) ):

      printt('Invalid input file')
      printt(file_line)
      if(len(datain_error_students) == 0):
        datain_error_students.append(filei)
      else:
        datain_error_students.append(', ' + filei)

    else:
      printt(file_line)
      students.append(filei)                        # Students' names added (I)
      data_marks.append(os.path.splitext(filei)[0]) # Students' names added (II)
      data.append([n0, n1, n2, n3])                 # LS values added to data matrix

datain_error_students.append('.')
data = np.vstack(data) # Stack the list
#
shape = np.shape(data)
#
L = shape[0]
K = shape[1]
#
xa = data[:, 0]
xr = data[:, 1]
xt = data[:, 2]
xp = data[:, 3]
#
printt(' xa : ' + str(xa))
printt(' ')
printt(' xr : ' + str(xr))
printt(' xt : ' + str(xt))
printt(' ')
printt(' xp : ' + str(xp))
printt(' ')
#
printt(' ')
#
#-------------------------------------------------------
# OPTIONAL INPUT FILE(S) WITH EDUCATORS' LEARNING STYLES
#-------------------------------------------------------
# The corresponding Excel files must be
# all contained in a folder placed in the input folder
# named educator, professor, teacher, or instructor.
# If more than one folder is included,
# the program only considers the data included first
# in the folder educator. If professor folder does
# not exist, then the data in professor is considered.
# If none of them appear, then the data in the
# teacher folder is used. If none of the previous
# three folders are included, then the data in the
# instructor folder is analyzed.
# Finally, if no folder named educator, professor,
# teacher, or instructor is inluded in the input
# folder, no data for the educator is used.

# Input file with the LSs of the educator(s)/professor(s)/teacher(s)/instuctor(s)
vprofessor = ['educator', 'professor', 'teacher', 'instructor']
vProfessor = ['Educator', 'Professor', 'Teacher', 'Instructor']

nfiles_prof = 0
Lprof = 0

prof_ref = vprofessor[0]
Prof_ref = vProfessor[0]
# Check if one or more of the folders exists
for professor in vprofessor:
  input_folder_prof = input_folder + professor + '/'
  if os.path.exists(input_folder_prof) and os.path.isdir(input_folder_prof):
    printt(f"The folder '{input_folder_prof}' exists.")
    prof_ref = professor
    Prof_ref = vProfessor[nfiles_prof]
    nfiles_prof += 1
    input_folder_prof_ref = input_folder_prof
  else:
    printt(f"The folder '{input_folder_prof}' does not exist.")

printt('nfiles_prof :' + str(nfiles_prof))

if ( nfiles_prof == 0):
  printt(' The learning styles of the educator(s) have not been provided')
else:
  datain_error_profs = []
  if ( nfiles_prof >= 2):
    printt(' Error: Conflict detected. Multiple folders associated with the learning styles of the educator(s). Please, change so that all the input files are contained in a single folder with one of these names:')
    for i in range(0, len(vprofessor)):
      printt(vprofessor[i])
  else:
    if ( nfiles_prof == 1):
      printt(' The learning styles of the educator(s) have been provided')

      input_folder_prof = input_folder_prof_ref

      profs       = []
      data_prof   = []
      profs_files = []

      for files in os.listdir(input_folder_prof):
#
#       xls and xlsx files
        if ((files.endswith('.xls')) or (files.endswith('.xlsx'))):
          profs_files.append(files)
#
      profs_files = sorted(profs_files)
#
      for filei in profs_files:
#
          profi = filei.rsplit(".", 1)[0]
          data_prof_in = []
          opt = 0
          try:
            data_prof_in = pd.read_excel(input_folder_prof + filei, header=None, sheet_name='CÁLCULO')
          except FileNotFoundError:
            printt(f"File '{filei}' not found.")
          except ValueError:
            printt("Worksheet 'CÁLCULO' not found. Trying alternative sheet names...")
            try:
              data_prof_in = pd.read_excel(input_folder_prof + filei, header=None, sheet_name='CALCULO')
            except ValueError:
              try:
                data_prof_in = pd.read_excel(input_folder_prof + filei, header=None, sheet_name='CALCULATION')
                opt = 1
              except ValueError:
                printt("Neither 'CALCULATIONS' nor 'CALCULO' sheet found.")

          if not data_in.empty:
            data_prof_LS = data_prof_in.loc[29].values          # Extract the LS data (I)
            if( opt != 1):
              line    = np.delete(data_prof_LS, [0,1,6,7,8]) # Extract the LS data (II)
            else:
              line    = np.delete(data_prof_LS, [0,1,2,7,8]) # Extract the LS data (II)
#
          file_line = filei+'      '+str(line[0])+' '+str(line[1])+' '+str(line[2])+' '+str(line[3])
#
          try:
            n0 = float(line[0])
            n1 = float(line[1])
            n2 = float(line[2])
            n3 = float(line[3])
            error = 'False'
          except ValueError:
            error = 'True'
#
          if ( (error == 'True') or
             ( np.isnan(n0) or np.isnan(n1) or np.isnan(n2) or np.isnan(n3)
               or (n0 < 0) or (n0 > 20) or (n1 < 0) or (n1 > 20)
               or (n2 < 0) or (n2 > 20) or (n3 < 0) or (n3 > 20) ) ):
#
            printt('error in input file : ' + str(file_line))
            if(len(datain_error_profs) == 0):
              datain_error_profs.append(filei)
            else:
              datain_error_profs.append(', ' + filei)
#
          else:
            printt(file_line)
            profs.append(profi)                  # Educator/Professor/Teacher/Instructor added
            data_prof.append([n0, n1, n2, n3])   # LS values added to data matrix
#
      datain_error_profs.append('.')

      data_prof = np.vstack(data_prof) # Stack the list data_prof
      Lprof = len(profs)
      printt(' ')
      printt('Educators : ')
#
#     Folder to save the output figures with the educators
      output_average_profile_prof = output_statistics_ls + '/' + prof_ref
      os.makedirs(output_average_profile_prof, exist_ok=True)
      printt('Output-average profile-prof folder : ' + output_average_profile_prof)

      output_participants_pc_prof = output_participants + '/' + prof_ref
      os.makedirs(output_participants_pc_prof, exist_ok=True)
      printt('Output-participants-PC_prof folder : ' + output_participants_pc_prof)
      printt('  ')

      for i in range(0, Lprof):
        printt(profs[i])
        printt(str(data_prof[i]))
#===================================================
#===================================================
#
#
#
#------------------------------------------------
# OPTIONAL INPUT FILE(S) WITH THE STUDENTS' MARKS
#------------------------------------------------
# The corresponding Excel files must be
# all contained in a folder placed in the input
# folder named marks, grades or qualitications.
# The corresponding data are named as a function of
# the names of the (i) Excel document, (ii) Excel sheet,
# and (iii) column in the Excel sheet.
# Each column is related to the marks obtained by the
# students in a different activity (laboratory, exams, etc.)
# and/or academic year.
# Thus, the results for different academic years and/or
# subjects can be saved all of them in a single Excel sheet,
# or in different documents.
# If the input folder contains folders named marks, grades and
# qualitications, only the data in the folder marks is considered;
# If the input folder contains folders named grades and
# qualitications, only the data in the folder grades is considered;
# Finally, if no folder named marks, grades or qualitications
# is inluded in the input folder,
# no calculation with the students' marks is conducted.

# Input file(s) with the students' marks/grades/qualifications
# (all of them must be in the same folder)

K_value = 0
vmarks  = ['marks', 'grades', 'qualifications']

pattern_K = re.compile(r"^marks_K(\d+)$")  # Pattern to match "marks_K<number>"

nfiles_marks = 0

# Check if one or more of the folders exists

# Check for marks_K* folders inside input_folder
for marks in vmarks:
#     printt('marks  :  ' + marks)
#     Check if the folder without the number of clusters (K) exists
      input_folder_mark = input_folder + marks + '/'
      if os.path.exists(input_folder_mark) and os.path.isdir(input_folder_mark):
        printt(f"The folder '{input_folder_mark}' exists.")
        nfiles_marks += 1
        input_folder_marks_ref = marks
        input_folder_marks     = input_folder_mark
#       The default K value is 2
        K_value = 2
      else:
        printt(f"The folder '{input_folder_mark}' does not exist.")

#     Check if the folder with the number of clusters exists, e.g., marks_K3 for 3 clusters
      for folder in os.listdir(input_folder):
       folder_path = os.path.join(input_folder, folder)
#      printt('folder  :  ' + folder)
       match = re.match(rf"^{marks}_K(\d+)$", folder)
       if match:
        K_value = int(match.group(1))
        print(f"Found file '{folder}' with K value: {K_value}")
        nfiles_marks += 1
        input_folder_marks_ref = folder
        input_folder_marks     = folder_path + '/'

printt('nfiles_marks :' + str(nfiles_marks))
printt( )

if ( nfiles_marks == 0):
  printt(' No file with the marks is provided, and then no clustering analysis is performed')
  K_value = 0
else:
  if ( nfiles_marks >= 2):
    printt(' Error: Conflict detected. Multiple folders associated with the marks of the educator(s). Please, change so that all the input files are contained in a single folder with one of these names:')
    for i in range(0, len(vmarks)):
      printt(vmarks[i])
    printt(' (If desired, change the number of the files, e.g., to marks_K3 to perform a K-means clustering with K=3 )')
    K_value = 0
  else:
    if ( nfiles_marks == 1):
      printt(' The marks of the students have been provided in the file')
      printt(' ' + input_folder_marks_ref)
      printt(' have been provided')
      printt(' ' + input_folder_marks)
#
# Marks folders (if K_value > 0)
if ( K_value == 0 ):
#
# K = 0
# No marks are provided, and the no clustering analysis is conducted.
#
# Check if the mark folders exist and, if not empty, remove them
  if os.path.exists(output_participants_ls_marks) and os.path.isdir(output_participants_ls_marks):
    if not os.listdir(output_participants_ls_marks):
      os.rmdir(output_participants_ls_marks)

  if os.path.exists(output_participants_pc_marks) and os.path.isdir(output_participants_pc_marks):
    if not os.listdir(output_participants_pc_marks):
      os.rmdir(output_participants_pc_marks)

else:
#
# K > 0
# Marks are provided.
# For K = 1 the whole group is studied; for K > 1, a K-means clustering is performed.

  if os.path.exists(output_participants_ls_marks) and os.path.isdir(output_participants_ls_marks):
    if not os.listdir(output_participants_ls_marks):
      os.rmdir(output_participants_ls_marks)

  output_participants_ls_marks = output_participants + '/' + input_folder_marks_ref
  os.makedirs(output_participants_ls_marks, exist_ok=True)
  printt('Output-LS-marks folder : ' + output_participants_ls_marks)
  printt('  ')

  if os.path.exists(output_participants_pc_marks) and os.path.isdir(output_participants_pc_marks):
    if not os.listdir(output_participants_pc_marks):
      os.rmdir(output_participants_pc_marks)

  output_participants_pc_marks = output_participants + '/' + input_folder_marks_ref
  os.makedirs(output_participants_pc_marks, exist_ok=True)
  printt('Output-PC-marks folder : ' + output_participants_pc_marks)
  printt('  ')

  excel_marks = []
  jref        = 0

# Run over all the Excel documents (and their corresponding sheets)
# to read the marks of the students
  for files in os.listdir(input_folder_marks):
#
#   xls and xlsx files
    if ((files.endswith('.xls')) or (files.endswith('.xlsx'))):
      excel_marks.append(files)
#
  excel_marks = sorted(excel_marks)
  printt("Excel files with the marks of the students : " + str(excel_marks))
  printt("\n")
#
  for filei in excel_marks:
    print()
    print('===============================')
    print('Excel document : ', filei)
    print('===============================')

    try:
      all_sheets = pd.read_excel(input_folder_marks + filei, sheet_name=None)  # Read all sheets into a dictionary
      for sheet_name, data_marks_in in all_sheets.items():
        print()
        print('-------------------------------')
        print(f"Sheet name     : {sheet_name}")
        print('-------------------------------')
        print("\n")
        print(data_marks_in)
        print('size : ', data_marks_in.shape)
        (Lstudents_tot, Lmarks) = data_marks_in.shape
        print('size : ', Lstudents_tot, Lmarks)
        print("\n")

        add_marks = 0

        if( data_marks_in.columns[0].strip().lower() == 'student' or data_marks_in.columns[0].strip().lower() == 'students'):
          print("The first column is named 'Student(s)' (case-insensitive).")
        else:
          print(f"The first column is named '{data_marks_in.columns[0]}', not 'Student'.")
          break

#       Save only the marks of those students whose LS questionnaires are provided
#       'student(s)' column (first column)
#
#       Check if any of the students whose LS marks are provided is also
#       contained in the list of students whose LS have been previouly given
        for j in range(0, L):
          studentj = os.path.splitext(students[j])[0]
          print('studentj : ', studentj)

          for i in range(Lstudents_tot):
            studenti = data_marks_in.iloc[i, 0]
            #print('  studenti :' + str(studenti))

            if(studenti == studentj):
              print('  studenti = studentj = ' + str(studenti))  # Print all elements in the column

#             Add new rows for the marks (if required)
              if(add_marks == 0):
                print('Enlarge data_marks')
                print('np.shape(data_marks) : ', np.shape(data_marks))
                print('data_marks : ', data_marks)
                data_marks = np.vstack([data_marks, np.full((Lmarks - 1, L + 3), np.nan)])
                print('np.shape(data_marks) : ', np.shape(data_marks))
                add_marks = 1

              for k in range(1, Lmarks):
                print('jref, k : ', jref, k)
#               Add to the first line of data_marks the name of the Excel document
                data_marks[k+jref, 0] = filei
#               Add to the second line of data_marks the name of the Excel sheet
                data_marks[k+jref, 1] = sheet_name
#               Add to the third line of data_marks the name of the column (mark)
                data_marks[k+jref, 2] = data_marks_in.columns[k]

              for k in range(1, Lmarks):
#               Add the marks in the corresponding positions
                markij = data_marks_in.iloc[i, k]
                print('    markij              : ', markij)
                if isinstance(markij, str):  # Check if it's a string
                  print(f"data_marks_in[{i}, {j}] is an (empty?) string.")
                  markij =  np.nan
                elif pd.isna(markij):  # Check for NaN or None (emptiness)
                  print(f"data_marks_in[{i}, {j}] is empty (NaN or None).")
                  markij =  np.nan

                data_marks[k + jref, j + 3] = markij
              break
        print(f"Sheet name     : {sheet_name}")
        if ( add_marks == 1 ):
          jref += Lmarks - 1
#        jref = jref + Lmarks - 1

    except ValueError:
        print('')

  for i in range(0, len(data_marks)):
    print('    student ' + str(data_marks[i]))
  print('end students')

  data_marks = np.vstack(data_marks)
  shape = np.shape(data_marks)
  printt('data_marks : ' + str(data_marks))
  printt('shape : ' + str(shape))

  if ( len(shape) > 0 and shape[1] > 1 ):
    Lmarks_activities = shape[0]
    Lmarks_students   = shape[1]
  else:
    Lmarks_activities = 0
    Lmarks_students   = 0
    Lmarks = 0
  printt(' Lmarks_activities : ' + str(Lmarks_activities))
  printt(' Lmarks_students   : ' + str(Lmarks_students))
  printt(' ')
#
#
#===================================================
#===================================================
#
#
#
#
#===================================================
# AVERAGE MEAN AND UNCERTAINTIES
#===================================================
printt('===========================================')
printt('AVERAGE MEAN AND UNCERTAINTIES...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
# Average mean
printt('-------------------------------------------')
printt(' Average mean...')
printt('-------------------------------------------')
shape = np.shape(data)
#
L = shape[0]
K = shape[1]
printt('  Input data size')
printt('    L :' + str(L))
printt('    K :' + str(K))
printt(' ')
#
# Size of the average value 0n the plots
size_mean = size_mean_ref / L
if ( size_mean < 1):
  size_mean = 4
else:
  if ( size_mean > 500):
    size_mean = 800
#
xamean = np.mean(xa)
xrmean = np.mean(xr)
xtmean = np.mean(xt)
xpmean = np.mean(xp)
#
xall     = np.concatenate( data )
xmeanall = np.mean(xall)
#
xmean = [xamean, xrmean, xtmean, xpmean]
#
printt('  Average mean')
printt('    mean(xa),    # : ' + str(xamean))
printt('    mean(xr),    # : ' + str(xrmean))
printt('    mean(xt),    # : ' + str(xtmean))
printt('    mean(xp),    # : ' + str(xpmean))
printt(' ')
printt('    mean(all),   # : ' + str(xmeanall))
printt('------------------------------------------')
printt(' Average mean done!')
printt('------------------------------------------')
printt(' ')
#
# Length of the eigenvectors when plotted
# (in 2D and 3D) as a function of the LSs
distmeanmax = 30. - np.min( xmean )
#
#---------------------------------------------------
# Uncertainties
printt('------------------------------------------')
printt('  Uncertainties...')
printt('------------------------------------------')
# Create 95% confidence interval for population mean weight.
#
confidence_interval_ACT=st.t.interval(confidence=0.95, df=len(data)-1, loc=np.mean(xa), scale=st.sem(xa))
uncert_abs_ACT=0.5*(confidence_interval_ACT[1]-confidence_interval_ACT[0])
#
confidence_interval_REF=st.t.interval(confidence=0.95, df=len(data)-1, loc=np.mean(xr), scale=st.sem(xr)) 
uncert_abs_REF=0.5*(confidence_interval_REF[1]-confidence_interval_REF[0])
#
confidence_interval_THEO=st.t.interval(confidence=0.95, df=len(data)-1, loc=np.mean(xt), scale=st.sem(xt)) 
uncert_abs_THEO=0.5*(confidence_interval_THEO[1]-confidence_interval_THEO[0])
#
confidence_interval_PRA=st.t.interval(confidence=0.95, df=len(data)-1, loc=np.mean(xp), scale=st.sem(xp)) 
uncert_abs_PRA=0.5*(confidence_interval_PRA[1]-confidence_interval_PRA[0])
#
dxmean = [uncert_abs_ACT, uncert_abs_REF, uncert_abs_THEO, uncert_abs_PRA]
#
confidence_interval_ALL=st.t.interval(confidence=0.95, df=len(xall)-1, loc=np.mean(xall), scale=st.sem(xall)) 
uncert_abs_ALL=0.5*(confidence_interval_ALL[1]-confidence_interval_ALL[0])
#
printt('  Uncertainties')
printt('    Uncert(xa),  # : ' + str(uncert_abs_ACT))
printt('    Uncert(xr),  # : ' + str(uncert_abs_REF))
printt('    Uncert(xt),  # : ' + str(uncert_abs_THEO))
printt('    Uncert(xp),  # : ' + str(uncert_abs_PRA))
printt(' ')
printt('    Uncert(all), # : ' + str(uncert_abs_ALL))
printt('------------------------------------------')
printt(' Uncertainties done!')
printt('------------------------------------------')
printt(' ')
#---------------------------------------------------
# Average profile
printt('------------------------------------------')
printt(' Average profile...')
printt('------------------------------------------')
# Length of axes and spacing between ticks.
xmin, xmax, ymin, ymax = -20, 20, -20, 20
ticks_frequency = 5
#
# Plot
fig, ax = plt.subplots(figsize = ( w_fig, h_fig )) 
#
# Coordenates (x, y) for the average mean
xs = [xrmean, 0, -xpmean,0]
ys = [0, xamean, 0, -xtmean]
ax.scatter(xs, ys)
#
# Connect the previous points
vertices_x=xs.copy()
vertices_x.append(xrmean)
vertices_y=ys.copy()
vertices_y.append(0)
#
# Average values
printt(' vertices_x : ' + str(vertices_x))
printt(' vertices_y : ' + str(vertices_y))
plt.plot(vertices_x, vertices_y, ls='--')
#
# Plot error barrs
printt(str(-xpmean-uncert_abs_PRA))
printt(str(xamean+uncert_abs_ACT))
printt(str(xpmean+uncert_abs_PRA))
#
#---------------------------------------------------
# Function that returns the value at x
# of a line that crosses points (x1, y1) and (x2, y2)
printt('-------------------------------------------')
printt('    Interpolating functions...')
printt('-------------------------------------------')
def line(x1, y1, x2, y2, x):
  m = (y2-y1)/(x2-x1)
  return y1 + m * ( x - x1 )  
#
ysupneg = line(-xpmean-uncert_abs_PRA, 0,
               0, xamean+uncert_abs_ACT,
               -xpmean+uncert_abs_PRA)
#
yinfneg = line(-xpmean-uncert_abs_PRA, 0,
               0, -xtmean-uncert_abs_THEO,
               -xpmean+uncert_abs_PRA)
#
ysuppos = line(xrmean+uncert_abs_REF, 0,
               0, xamean+uncert_abs_ACT,
               xrmean-uncert_abs_REF)
#
yinfpos = line(xrmean+uncert_abs_REF, 0,
               0, -xtmean-uncert_abs_THEO,
               xrmean-uncert_abs_REF)
#
printt('-------------------------------------------')
printt('    Interpolating functions done!')
printt('-------------------------------------------')
#---------------------------------------------------
printt(' ')
#
plt.fill_between([-xpmean-uncert_abs_PRA, -xpmean+uncert_abs_PRA, 0,
                   xrmean-uncert_abs_REF,  xrmean+uncert_abs_REF],
                 [0,  0, xamean-uncert_abs_ACT,  0, 0],
                 [0,  ysupneg, xamean+uncert_abs_ACT, ysuppos, 0],
                 facecolor='pink', alpha=0.5)
#
plt.fill_between([-xpmean-uncert_abs_PRA, -xpmean+uncert_abs_PRA, 0,
                   xrmean-uncert_abs_REF,  xrmean+uncert_abs_REF],
                 [0, yinfneg, -xtmean-uncert_abs_THEO, yinfpos, 0],
                 [0,  0, -xtmean+uncert_abs_THEO, 0, 0],
                 facecolor='pink', alpha=0.5)
#
x_error=[uncert_abs_REF,0,uncert_abs_PRA,0]
y_error=[0,uncert_abs_ACT,0,uncert_abs_THEO]
#
plt.errorbar(xs, ys, xerr = x_error, yerr=y_error, fmt='o', ecolor = 'red',color='blue',elinewidth = 5, capsize=10)
#
# Same scale for all axes
ax.set(xlim=(xmin-1, xmax+1), ylim=(ymin-1, ymax+1), aspect='equal')
#
# Set bottom and left spines as x and y axes of coordinate system
ax.spines['bottom'].set_position('zero')
ax.spines['left'].set_position('zero')
#
# Remove top and right spines
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
#
# Create 'x' and 'y' labels placed at the end of the axes
plt.text(0,  22, Label_LS[0], fontsize=labelsize, horizontalalignment='center')
plt.text( 22, 0, Label_LS[1], fontsize=labelsize, verticalalignment  ='center')
plt.text(0, -24, Label_LS[2], fontsize=labelsize, horizontalalignment='center')
plt.text(-34, 0, Label_LS[3], fontsize=labelsize, verticalalignment  ='center')
#
# Create custom major ticks to determine position of tick labels
x_ticks = np.arange(xmin, xmax+1, ticks_frequency)
y_ticks = np.arange(ymin, ymax+1, ticks_frequency)
ax.set_xticks(x_ticks[x_ticks != 0])
ax.set_yticks(y_ticks[y_ticks != 0])
#
# Rename negative parts of the axes
labels = [item.get_text() for item in ax.get_xticklabels()]
printt('labels : ' + str(labels))
labels[0] = 20
labels[1] = 15
labels[2] = 10
labels[3] = 5
ax.set_xticklabels(labels)
#
labels = [item.get_text() for item in ax.get_yticklabels()]
printt('labels : ' + str(labels))
labels[0] = 20
labels[1] = 15
labels[2] = 10
labels[3] = 5
ax.set_yticklabels(labels)
#
ax.set_xticks(np.arange(xmin, xmax+1), minor=True)
ax.set_yticks(np.arange(ymin, ymax+1), minor=True)
#
ax.tick_params(axis='both', which='both', labelsize = ticksize)
#
#plt.show()
filename_average_profile=output_statistics_ls + '/Fig_averageprofile.png'
printt('Saving ' + filename_average_profile)
plt.savefig(filename_average_profile)
printt(' ')
printt('------------------------------------------')
printt(' Average profile done!')
printt('------------------------------------------')
printt(' ')

# Average profile
printt('------------------------------------------')
printt(' Average profile '+ professor + '...')
printt('------------------------------------------')
#
filename_average_profile_prof = [ ]
for i in range(0, Lprof):
# Plot
  fig, ax = plt.subplots(figsize = ( w_fig, h_fig ))
#
# Coordenates (x, y) for the average mean
  ax.scatter(xs, ys)
  plt.plot(vertices_x, vertices_y, ls='--')

# Coordinates (x, y) for the professors
  xs_prof = [data_prof[i, 1], 0, -data_prof[i, 2],0]
  ys_prof = [0, data_prof[i,0], 0, -data_prof[i, 3]]
  ax.scatter(xs_prof, ys_prof, c='g', s=2)
#
# Connect the previous points
  vertices_x_prof=xs_prof.copy()
  vertices_x_prof.append(data_prof[i, 1])
  vertices_y_prof=ys_prof.copy()
  vertices_y_prof.append(0)
#
# Average values
  printt(' vertices_x_prof : ' + str(vertices_x_prof))
  printt(' vertices_y_prof : ' + str(vertices_y_prof))
  plt.plot(vertices_x_prof, vertices_y_prof, ls='-', c = 'g', lw=2)
#
# Plot error barrs
  printt(str(-xpmean-uncert_abs_PRA))
  printt(str(xamean+uncert_abs_ACT))
  printt(str(xpmean+uncert_abs_PRA))
#
#---------------------------------------------------
  printt(' ')
#
  plt.fill_between([-xpmean-uncert_abs_PRA, -xpmean+uncert_abs_PRA, 0,
                   xrmean-uncert_abs_REF,  xrmean+uncert_abs_REF],
                   [0,  0, xamean-uncert_abs_ACT,  0, 0],
                   [0,  ysupneg, xamean+uncert_abs_ACT, ysuppos, 0],
                   facecolor='pink', alpha=0.5)
#
  plt.fill_between([-xpmean-uncert_abs_PRA, -xpmean+uncert_abs_PRA, 0,
                     xrmean-uncert_abs_REF,  xrmean+uncert_abs_REF],
                   [0, yinfneg, -xtmean-uncert_abs_THEO, yinfpos, 0],
                   [0,  0, -xtmean+uncert_abs_THEO, 0, 0],
                   facecolor='pink', alpha=0.5)
#
  x_error=[uncert_abs_REF,0,uncert_abs_PRA,0]
  y_error=[0,uncert_abs_ACT,0,uncert_abs_THEO]
#
  plt.errorbar(xs, ys, xerr = x_error, yerr=y_error, fmt='o', ecolor = 'red',color='blue',elinewidth = 5, capsize=10)
#
# Same scale for all axes
  ax.set(xlim=(xmin-1, xmax+1), ylim=(ymin-1, ymax+1), aspect='equal')
#
# Set bottom and left spines as x and y axes of coordinate system
  ax.spines['bottom'].set_position('zero')
  ax.spines['left'].set_position('zero')
#
# Remove top and right spines
  ax.spines['top'].set_visible(False)
  ax.spines['right'].set_visible(False)
#
# Create 'x' and 'y' labels placed at the end of the axes
  plt.text(0,  22, Label_LS[0], fontsize=labelsize, horizontalalignment='center')
  plt.text( 22, 0, Label_LS[1], fontsize=labelsize, verticalalignment  ='center')
  plt.text(0, -24, Label_LS[2], fontsize=labelsize, horizontalalignment='center')
  plt.text(-34, 0, Label_LS[3], fontsize=labelsize, verticalalignment  ='center')
#
# Create custom major ticks to determine position of tick labels
  x_ticks = np.arange(xmin, xmax+1, ticks_frequency)
  y_ticks = np.arange(ymin, ymax+1, ticks_frequency)
  ax.set_xticks(x_ticks[x_ticks != 0])
  ax.set_yticks(y_ticks[y_ticks != 0])
#
# Rename negative parts of the axes
  labels = [item.get_text() for item in ax.get_xticklabels()]
  printt('labels : ' + str(labels))
  labels[0] = 20
  labels[1] = 15
  labels[2] = 10
  labels[3] = 5
  ax.set_xticklabels(labels)
#
  labels = [item.get_text() for item in ax.get_yticklabels()]
  printt('labels : ' + str(labels))
  labels[0] = 20
  labels[1] = 15
  labels[2] = 10
  labels[3] = 5
  ax.set_yticklabels(labels)
#
  ax.set_xticks(np.arange(xmin, xmax+1), minor=True)
  ax.set_yticks(np.arange(ymin, ymax+1), minor=True)
#
  ax.tick_params(axis='both', which='both', labelsize = ticksize)
#
#plt.show()
  filename_average_profile_prof.append( output_average_profile_prof + '/Fig_averageprofile_prof_' + profs[i] + '.png' )
  printt('Saving ' + filename_average_profile_prof[i])
  plt.savefig(filename_average_profile_prof[i])
  printt(' ')
  printt('------------------------------------------')
  printt(' Average profile with '+ professor + profs[i] + ' done!')
  printt('------------------------------------------')
printt(' ')
printt('===========================================')
printt('AVERAGE MEAN AND UNCERTAINTIES DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# TENDENCIES TO THE DIFFERENT LSs
#===================================================
printt('===========================================')
printt('TENDENCIES...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
# Average mean
printt('-------------------------------------------')
printt(' Tendencies in %...')
printt('-------------------------------------------')
# Matrix with the tendencies (very low, low,
# moderate, high and very high) towards the
# different LSs.
# Each line is associated with a different LS.
# Each column corresponds to a different
# tendency.
tendency_matrix     = np.zeros((K, 5))
tendency_matrix_all = [['' for _ in range(K)] for _ in range(L)]
tendency_vector  = ['', '', '', '', '']
#
for i in range(0,L):
  printt(str(i))
  tendency_vector[0] = scatter_tendency('Activist',   xa[i])
  tendency_vector[1] = scatter_tendency('Reflector',  xr[i])
  tendency_vector[2] = scatter_tendency('Theorist',   xt[i])
  tendency_vector[3] = scatter_tendency('Pragmatist', xp[i])
#
  tendency_matrix_all[i][0] = tendency_vector[0]
  tendency_matrix_all[i][1] = tendency_vector[1]
  tendency_matrix_all[i][2] = tendency_vector[2]
  tendency_matrix_all[i][3] = tendency_vector[3]
#
  for j in range(0,K):
    if ( tendency_vector[j] == 'vl' ) :
      tendency_matrix[j,0] = tendency_matrix[j,0] + 1
    else:
      if ( tendency_vector[j] == 'l' ) :
        tendency_matrix[j,1] = tendency_matrix[j,1] + 1
      else:
        if ( tendency_vector[j] == 'm' ) :
          tendency_matrix[j,2] = tendency_matrix[j,2] + 1
        else:
          if ( tendency_vector[j] == 'h' ) :
            tendency_matrix[j,3] = tendency_matrix[j,3] + 1
          else:
            if ( tendency_vector[j] == 'vh' ) :
              tendency_matrix[j,4] = tendency_matrix[j,4] + 1
#
tendency_matrix_percentage = tendency_matrix * 100 / L
#
for i in range(0,K):
  printt('  ' + Label_LS[i] + ' learning style')
  printt('    Tendency      No.     %')
  for j in range(0,len(Label_tendencies)):
    printt('   ' + Label_tendencies_print[j] + '   ' + str(int(tendency_matrix[i, j])) + '   ' + str(round(tendency_matrix_percentage[i, j], 1)))
  printt(' ')
#
printt('Learning Style ' + Label_tendencies_print[0] + Label_tendencies_print[1] +
                         Label_tendencies_print[2]   +
                         Label_tendencies_print[3]   + Label_tendencies_print[4])
printt('                No.  %    No.  %    No.  %    No.  %    No.  %')
for i in range(0,K):
  printt('  ' + Label_LS_print[i] + ' ' + str(int(tendency_matrix[i, 0])) + str(round(tendency_matrix_percentage[i, 0], 1))  + '   ' + str(int(tendency_matrix[i, 1])) + str(round(tendency_matrix_percentage[i, 1], 1)) + '   ' + str(int(tendency_matrix[i, 2])) + str(round(tendency_matrix_percentage[i, 2], 1)) + '   ' + str(int(tendency_matrix[i, 3])) + str(round(tendency_matrix_percentage[i, 3], 1)) + '   ' + str(int(tendency_matrix[i, 4])) + str(round(tendency_matrix_percentage[i, 4], 1)) )
printt(' ')
#
Ntot_tendency    = np.zeros(len(Label_tendencies))
percent_tendency = np.zeros(len(Label_tendencies))
#
for j in range(0,len(Label_tendencies)):
  for i in range(0,K):
    Ntot_tendency[j]    = Ntot_tendency[j] + tendency_matrix[i, j]
    percent_tendency[j] = percent_tendency[j] + tendency_matrix[i, j] * tendency_matrix_percentage[i, j]  
#
  percent_tendency[j] = percent_tendency[j] / Ntot_tendency[j]
#
printt('   Average    ' + str(int(Ntot_tendency[0])) + str(round(percent_tendency[0], 1)) + '   ' + str(int(Ntot_tendency[1])) + str(round(percent_tendency[1], 1)) + '   ' + str(int(Ntot_tendency[2])) + str(round(percent_tendency[2], 1)) + '   ' + str(int(Ntot_tendency[3])) + str(round(percent_tendency[3], 1)) + '   ' + str(int(Ntot_tendency[4])) + str(round(percent_tendency[4], 1)))
#
fig, ax = plt.subplots(figsize = ( w_fig, h_fig )) #, layout='constrained')
#
ymax = np.max(tendency_matrix_percentage)
#
# Fill areas below the curve
for j in range(0,K):
  for i in range(0,len(Label_tendencies)):
#
    percentage = tendency_matrix_percentage[j,i]
    xmin_percentage = -2+i+(j-2)*0.2
    xmax_percentage = xmin_percentage + 0.2
    plt.fill_between([xmin_percentage, xmax_percentage], [percentage, percentage], facecolor = LS_color[j], alpha = LS_alpha3)  
#
# Tick parameters
  ax.set_xticks([-2, -1, 0, 1, 2], Label_tendencies)
  ax.tick_params(axis='both', which='major', labelsize = ticksize_quartile)
#
# Labels
  ax.set_xlabel('Tendency',   fontsize = labelsize_quartile)
  ax.set_ylabel('% Students', fontsize = labelsize_quartile)
#
ymax = ymax*1.05
ax.set_xlim( [ -2.5, 2.5] )
ax.set_ylim( [ 0.05, ymax] )
#
filename_tendencies=output_statistics_ls + '/Fig_tendencies.png'
printt('Saving ' + filename_tendencies)
plt.savefig(filename_tendencies)
printt('------------------------------------------')
printt(' Tendencies done!')
printt('------------------------------------------')
printt(' ')
#
printt('===========================================')
printt('TENDENCIES DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# AFINITY
#===================================================
printt('===========================================')
printt('AFFINITY...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Affinity in %...')
printt('-------------------------------------------')
#
affinity  = np.zeros(K)
#
for i in range(0,K):
  for j in range(2,len(Label_tendencies)):
    affinity[i] = affinity[i] + tendency_matrix_percentage[i,j] 
#
fig, ax = plt.subplots(figsize = ( w_fig, h_fig ))
#
ymax = np.max(affinity)
#
# Fill areas below the curve
printt( '   Learning style     %')
for j in range(0,K):
  printt( '  ' + Label_LS_print[j] + '    ' + str(round(affinity[j], 1)))
  plt.fill_between([j-0.4, j+0.4], [affinity[j], affinity[j]], facecolor = LS_color[j], alpha = affinity_alpha)  
#
# Tick parameters
  x = np.arange(len(Label_LS))
  ax.set_xticks(x, Label_LS)
  ax.tick_params(axis='both', which='major', labelsize = ticksize_quartile)
#
# Labels
  ax.set_xlabel('LS',         fontsize = labelsize_quartile)
  ax.set_ylabel('% Students', fontsize = labelsize_quartile)
#
printt( '   Average        ' + str(round(percent_tendency[2]+percent_tendency[3]+percent_tendency[4], 1)))
#
ymax = ymax*1.05
ax.set_xlim( [ -0.5, 3.5] )
ax.set_ylim( [ 0.05, ymax] )
#
printt(' ')
filename_affinity=output_statistics_ls + '/Fig_affinity.png'
printt('Saving ' + filename_affinity)
plt.savefig(filename_affinity)
printt('------------------------------------------')
printt(' Tendencies done!')
printt('------------------------------------------')
printt(' ') 
#
printt('===========================================')
printt('AFFINITY DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# PROBABILITIES IN THE ORIGINAL LS BASIS SET
#===================================================
#
printt('-------------------------------------------')
printt(' Probability of each student on the LSs...')
printt('-------------------------------------------')
printt(' ')
# Calculate the probability to have a
# certain LS for each student
probLS = np.zeros((L,K))
#
for i in range(0,L):
  normLSj2 = 0
  for j in range(0,K):
    normLSj2 = normLSj2 + data[i,j]*data[i,j]
#
# Compute the probability
# (it is set to zero if the vector is zero)
  for j in range(0,K):
    if ( normLSj2 > 0):
      probLS[i, j] = 100 * data[i,j] * data[i,j] / normLSj2
#
printt('   data   : ' + str(data))
printt(' ')
printt('   probLS : ' + str(probLS))
printt(' ')
#
# Mean values and uncertainties of the probabilities
probLSmean = np.array([np.mean(probLS[:,0]), np.mean(probLS[:,1]), np.mean(probLS[:,2]), np.mean(probLS[:,3])])
#
uncert_abs_probLS = [0, 0, 0, 0]
for j in range(0,K):
  probj = probLS[:,j]
  confidence_interval_probLS=st.t.interval(confidence=0.95, df=len(probj)-1, loc=probLSmean[j], scale=st.sem(probj)) 
  uncert_abs_probLS[j]=0.5*(confidence_interval_probLS[1]-confidence_interval_probLS[0])
#
printt('  The mean values of the corresponding probabilities')
printt('  do not nullify (as all of them are positive or zero)')
printt('    probLSmean     : ' + str(probLSmean))
printt('    Uncert(probLS) : ' + str(uncert_abs_probLS))
printt(' ')
#
printt(' Probability of each student on the LSs done!')
printt('-------------------------------------------')
printt(' ')
printt(' ')
printt('===========================================')
printt('PROBABILITIES IN ORIGINAL BASIS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# STATISTICAL ANALYSIS IN THE LS ORIGINAL BASIS
#===================================================
printt('===========================================')
printt('STATISTICAL ANALYSIS LS ORIGINAL BASIS...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Analysis of the cumulative probability')
printt(' distribution (W(prob)) of the probabilities')
printt(' of the different LSs (plot and fitting')
printt(' using a Weibull distribution)...')
printt('-------------------------------------------')
printt(' ')
printt('-------------------------------------------')
printt(' W(LS)...')  
printt('-------------------------------------------')
xWeib   = np.arange(0,20.1, 0.001)
#
ymin  = 0
ymax  = 1.01
ymaxx = 10000000000
#
parameters_Weibull_LS20 = np.zeros((K,2))
#
for ils in range(0, K):
  printt('   LS 0-20 :' + Label_LS[ils])
  x = data[:, ils]
  x = np.sort(x)
#
  y = np.zeros(L)
  y[0] = 0
  dy = 1 / float(L-1)
  for i in range(1,L):
    y[i] = y[i-1] + dy
#
# Fitting using Weibull distribution
  param = curve_fit(Wweibull, x, y, [xmean[ils], 1], maxfev = max_iter_Weib)
  printt('    ' + str(param))
  [alpha_Weibull, k_Weibull] = param[0]
  parameters_Weibull_LS20[ils] = [alpha_Weibull, k_Weibull]
  printt('    alpha_Weibull : ' + str(alpha_Weibull))
  printt('    k_Weibull     : ' + str(k_Weibull))  
  printt(' ')
#
# Create a figure with 4 panels,
# each one related to a different LS
fig, axs = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), sharex = True, sharey = True, layout='constrained')  
#
panels = [[0, 0], [0, 1], [1, 0], [1, 1]]
#
for ils in range(0, K):
  printt('   LS 0-20 :' + Label_LS[ils])
  x = data[:, ils]
  x = np.sort(x)
#
  y = np.zeros(L)
  y[0] = 0
  dy = 1 / float(L-1)
  for i in range(1,L):
    y[i] = y[i-1] + dy
#
# Staircase
  x2 = np.zeros(2*L+2)
  y2 = np.zeros(2*L+2)
  x2[0]  =   0
  y2[0]  =   0
  x2[2*L+1] = x[L-1]
  y2[2*L+1] = 1
#
# Fitting using Weibull distribution
  [alpha_Weibull, k_Weibull] = parameters_Weibull_LS20[ils]
  printt('    alpha_Weibull : ' + str(alpha_Weibull))
  printt('    k_Weibull     : ' + str(k_Weibull))  
  printt(' ')
#
  for i in range(0,L):
    x2[2*i+1] = x[i]
    x2[2*i+2] = x[i]

    y2[2*i+1] = y[i]
    y2[2*i+2] = y[i] + dy
  y2[2*L] = 1    
#
# Find the points close to 5, 10, and 15
# in order to correcly color the
# LS that lie in the ranges [0, 5), 
# (10, 15), and (15,20]
  ic   = [-1,  -1, -1]
  xc   = [ 0,  5, 10]
  xref = [ 5, 10, 15]
#
  for i in range(0,len(ic)):
    for j in range(0, L):
      if( x[j] > xc[i] and x[j] < xref[i] ):
        ic[i] = j
        xc[i] = x[j]
#
  axs[panels[ils][0], panels[ils][1]].set_xlim( [ 0, 20] )
  axs[panels[ils][0], panels[ils][1]].set_ylim( [ ymin, ymax] )
#
  axs[panels[ils][0], panels[ils][1]].tick_params(axis='both', which='major', labelsize = ticksize_PC)
#
  axs[panels[ils][0], panels[ils][1]].plot(x2, y2,                                 lw = PC_width,   color = PC_color[ils], ls = PC_line[ils])
  axs[panels[ils][0], panels[ils][1]].plot(xWeib,Wweibull(xWeib, alpha_Weibull, k_Weibull) , lw = PC_W_width, color = PC_color[ils], ls = PC_line_bis[ils])  
#
# Plot the average value as a vertical line
# we do not plot the uncertainties as they can be very large
  axs[panels[ils][0], panels[ils][1]].plot([xmean[ils] , xmean[ils] ], [ymin, ymaxx], lw = LS_mean_width, color = LS_color[ils], ls = LS_mean_line)
  axs[panels[ils][0], panels[ils][1]].fill_between([xmean[ils]-dxmean[ils], xmean[ils]+dxmean[ils]], [ymin, ymaxx], color = LS_color[ils], alpha = LS_mean_alpha )
#      
# Text label
  axs[panels[ils][0], panels[ils][1]].text(20*0.05, 0.9*ymax, Label_LS[ils],
        fontsize = labelsize,
        horizontalalignment='left',
        verticalalignment='top')
#
axs[1, 0].tick_params(axis='both', which='major', labelsize=ticksize)
#
axs[1, 0].set_xlabel('Contribution', fontsize=labelsize)
axs[1, 0].set_ylabel('W [Contribution]', fontsize=labelsize)
#
filename_statistics_ls_w = output_statistics_ls + '/Fig_LS20_W2.png'
printt('Saving ' + filename_statistics_ls_w)
plt.savefig(filename_statistics_ls_w)
#
printt('-------------------------------------------')
printt(' W(LS) done!')  
printt('-------------------------------------------')
printt(' ')
#
printt('-------------------------------------------')
printt(' W(probLS)...')  
printt('-------------------------------------------')
xWeib   = np.arange(0, 1.1, 0.001)
parameters_Weibull_LS = np.zeros((K,2))
#
#
for ils in range(0, K):
  printt('   LS :' + Label_LS[ils])
  x = probLS[:, ils] / 100
  x = np.sort(x)
#
  y = np.zeros(L)
  y[0] = 0
  dy = 1 / float(L-1)
  for i in range(1,L):
    y[i] = y[i-1] + dy
#
# Fitting using Weibull distribution
  param = curve_fit(Wweibull, x, y, [probLSmean[ils], 1], maxfev = max_iter_Weib)
  printt('    ' + str(param))
  [alpha_Weibull, k_Weibull] = param[0]
  parameters_Weibull_LS[ils] = [alpha_Weibull, k_Weibull]
  printt('    alpha_Weibull : ' + str(alpha_Weibull))
  printt('    k_Weibull : ' + str(k_Weibull))
  printt(' ')
#
printt('-------------------------------------------')
printt(' W(probLS) done!')  
printt('-------------------------------------------')
printt(' ')
#
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Probabilities of the LSs...')
printt('-------------------------------------------')
#
xWeib   = np.arange(0, 120, 0.005)
ymaxtot = 0
ymin  = 0
ymax  = 1.01
ymaxx   = 10000000000
#
# Create a figure with 4 panels,
# each one related to a different LS
fig, axs = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), sharex = True, sharey = True, layout='constrained')  
#
panels = [[0, 0], [0, 1], [1, 0], [1, 1]]
#
for ils in range(0, K):
  printt('   LS 0-20 :' + Label_LS[ils])
#
  x = data[:, ils]
  x = np.sort(x)
#
  [alpha_Weibull, k_Weibull] = parameters_Weibull_LS20[ils]
  printt('    alpha_Weibull : ' + str(alpha_Weibull))
  printt('    k_Weibull : ' + str(k_Weibull))
  printt(' ')
#
# Plot the average value as a vertical line
# we do not plot the uncertainties as they can be very large
  axs[panels[ils][0], panels[ils][1]].plot([xmean[ils], xmean[ils]], [ymin, ymaxx], lw = LS_mean_width, color = LS_color[ils], ls = LS_mean_line)
  axs[panels[ils][0], panels[ils][1]].fill_between([xmean[ils]-dxmean[ils], xmean[ils]+dxmean[ils]],  [ymin, ymaxx], color = LS_color[ils], alpha = LS_mean_alpha )
#
  y, x, _= axs[panels[ils][0], panels[ils][1]].hist(x, bins=20, color = LS_color[ils], alpha = LS_hist_alpha)
  ymax = y.max()
#
  ymax = y.max()
  if ymax > ymaxtot:
    ymaxtot = ymax
#
# Probability distribution for Weibull function (scaled)
  factor = ymax / Pweibull(alpha_Weibull, alpha_Weibull, k_Weibull)
  axs[panels[ils][0], panels[ils][1]].plot(xWeib,factor*Pweibull(xWeib, alpha_Weibull, k_Weibull) , lw = LS_W_width,   color = LS_color[ils], ls = LS_line[0])  
#
  axs[panels[ils][0], panels[ils][1]].tick_params(axis='both', which='major', labelsize=ticksize)
#
  axs[panels[ils][0], panels[ils][1]].set_xlim( [ 0, 20] )
  axs[panels[ils][0], panels[ils][1]].set_ylim( [ 0, ymaxtot] )   
#
# Text label
  axs[panels[ils][0], panels[ils][1]].text(20*0.05, 0.9*ymaxtot, Label_LS[ils],
        fontsize = labelsize,
        horizontalalignment='left',
        verticalalignment='top')
#
axs[1, 0].tick_params(axis='both', which='major', labelsize=ticksize)
#
axs[1, 0].set_xlabel('Contribution', fontsize=labelsize)
axs[1, 0].set_ylabel('P [Contribution]', fontsize=labelsize)
#
filename_statistics_ls = output_statistics_ls + '/Fig_LS20_P2.png'
printt('Saving ' + filename_statistics_ls)
plt.savefig(filename_statistics_ls)
printt('-------------------------------------------')
printt(' Probabilities of the LSs done!')
printt('-------------------------------------------')
printt(' ')
printt('===========================================')
printt('STATISTICAL ANALYSIS LS ORIGINAL BASIS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# COVARIANCE MATRIX
#===================================================
printt('===========================================')
printt('COVARIANCE MATRIX...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Differences between input data and mean...')
printt('-------------------------------------------')
printt(' ')
# Difference between the input data and the mean
ddata = np.zeros((L,K))
for i in range(0,L):
  for j in range(0,K):
    ddata[i,j] = data[i,j] - xmean[j]
#
printt('------------------------------------------')
printt(' Differences done!')
printt('------------------------------------------')
printt(' ')
#---------------------------------------------------
# Covariance matrix
printt('------------------------------------------')
printt(' Construction of the covariance matrix...')
printt('------------------------------------------')
#
covX = np.cov([xa, xr, xt, xp])
#
printt('  covX = ' + str(covX))
printt(' ')
#
trace_covX = 0
for i in range(0,K):
  trace_covX = trace_covX + covX[i,i]
#
printt('  tr(covX) = ' + str(trace_covX))
printt(' ')
printt('------------------------------------------')
printt(' Covariance matrix constructed!')
printt('------------------------------------------')
printt(' ')
#---------------------------------------------------
printt('------------------------------------------')
printt(' Eigensystem...')
printt('------------------------------------------')
eigenValues, eigenVectors = eig(covX)
printt('------------------------------------------')
printt(' done!')
printt('------------------------------------------')
printt(' ')
printt('------------------------------------------')
printt(' Sorting the eigenvalues...')
printt('------------------------------------------')
idx = eigenValues.argsort()[::-1]  
printt('  Eigenvalues ordering')
printt('   idx    : ' + str(idx))
printt(' ')
eigenValues = eigenValues[idx]
eigenVectors = eigenVectors[:,idx]
#
# The eigenvectores are saved in the columns 
# of the matrix of the eigenVectors
printt('  EigenValues  : ' + str(eigenValues))
printt(' ')
disp0 = 100/trace_covX *  eigenValues[0]
disp1 = 100/trace_covX *  eigenValues[1]
disp2 = 100/trace_covX *  eigenValues[2]
disp3 = 100/trace_covX *  eigenValues[3]
printt('  % Dispersion (PC0) : ' + str(disp0))
printt('  % Dispersion (PC1) : ' + str(disp1) + ' % Dispersion (PC0+PC1)         : ' + str(disp0 + disp1))
printt('  % Dispersion (PC2) : ' + str(disp2) + ' % Dispersion (PC0+PC1+PC2)     : ' + str(disp0 + disp1 + disp2))
printt('  % Dispersion (PC3) : ' + str(disp3) + ' % Dispersion (PC0+PC1+PC2+PC3) : ' + str(disp0 + disp1 + disp2 + disp3))
printt(' ')
printt('  EigenVectors : ' + str(eigenVectors))
printt(' ')
printt('------------------------------------------')
printt(' Sorting the eigenvalues done!')
printt('------------------------------------------')
printt(' ')
printt('------------------------------------------')
printt(' EigenVectors normalization (norm, max)...')
printt('------------------------------------------')
eigenVectors_max1       = np.zeros((K,K))
eigenVectors_percentage = np.zeros((K,K))
#
for j in range(0,K):
  norm=0.0
  vj = eigenVectors[:,j]
#
  vjmax = np.max(np.abs(vj))
  vjmin = np.min(vj)  
  if vjmax == -vjmin:
    vjmax = vjmin
  eigenVectors_max1[:,j] = vj/vjmax
#
  for l in range(0,K):
    norm = norm + vj[l] * vj[l]
  norm = np.sqrt(norm)    
#
  eigenVectors[:,j] = vj/norm
  eigenVectors_percentage[:,j] = eigenVectors[:,j]*eigenVectors[:,j]*100
#
printt('  (Renormalized) Sorted eigenVectors (norm=1): ' + str(eigenVectors))
printt(' ')
printt('  Sorted eigenVectors (maximum = 1): ' + str(eigenVectors_max1))
printt(' ')
printt('  Sorted eigenVectors (%)          : ' + str(eigenVectors_percentage))
printt(' ')
printt('------------------------------------------')
printt(' EigenVectors normalization done!')
printt('------------------------------------------')

printt(' ')
printt('------------------------------------------')
printt(' Correspondence between eigenvectors-LS...')
printt('------------------------------------------')
# We only plot the two/three principal directions.
# Thus, we redifine the learning_pairs and
# learning_trios vectors so that their first
# component gives the two LSs of interest
#
lsref = [0, 0, 0]
for j in range(0,3):
  Vector = np.abs(eigenVectors[:,j])
  Vectormax = Vector[0]
  for i in range(1,K):
    if (Vector[i] > Vectormax):
      lsref[j]  = i
      Vectormax = Vector[i]
#
# If two of the terms in lsref are equal
# we simply take lsref equal to the
# first element in learning_trios
if( lsref[0] == lsref[1] or lsref[0] == lsref[2] or lsref[1] == lsref[2] ):
  lsref = learning_trios[0]
# The points in the plots are shown as diamonds instead as circles
#  
printt(' lsref : ' + str(lsref))
printt(' ')
#
for i in range(0, len(learning_pairs)):
  if(([lsref[0], lsref[1]] == learning_pairs[i]) or 
      [lsref[1], lsref[0]] == learning_pairs[i]):
    ipairs = i
#    
for i in range(0, len(learning_trios)):
  if(([lsref[0], lsref[1], lsref[2]] == learning_trios[i]) or 
     ([lsref[0], lsref[2], lsref[1]] == learning_trios[i]) or
     ([lsref[1], lsref[0], lsref[2]] == learning_trios[i]) or 
     ([lsref[1], lsref[2], lsref[0]] == learning_trios[i]) or
     ([lsref[2], lsref[0], lsref[1]] == learning_trios[i]) or 
     ([lsref[2], lsref[1], lsref[0]] == learning_trios[i])):
    itrios = i
#
printt('  Pairs of learning styles')
printt('    ipairs                 : ' + str(ipairs))
printt('    learning_pairs[ipairs] : ' + str(learning_pairs[ipairs]))
printt(' ')
printt('  Trios of learning styles')
printt('    itrios                 : ' + str(itrios))
printt('    learning_trios[itrios] : ' + str(learning_trios[itrios]))
printt(' ')
printt('------------------------------------------')
printt(' Correspondence eigenvectors-LS done!')
printt('------------------------------------------')
printt(' ')
printt('===========================================')
printt('COVARIANCE MATRIX DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#

#===================================================
# PRINCIPAL COMPONENT ANALYSIS (PCA).
# PROJECTIONS ON THE EIGENVECTORS OF THE 
# COVARIANCE MATRIX, AND STATISTICAL ANALYSIS
#===================================================
printt('===========================================')
printt('PRINCIPAL COMPONENT ANALYSIS...')
printt('===========================================')
printt(' ')
printt('===========================================')
printt('PROJECTION ON THE EIGENVECTORS...')
printt('===========================================')
#---------------------------------------------------
# Projections on the principal directions of the 
# maximal "pure" LSs given by vectors with one 
# component equal to 20, and the rest to 0
# (original LS in lines 0 to 3 of the
# matrix LS; the projections are given in lines 0 to 3
# of the matrix projectLS), the origin (line 4) and the
# maximal possible state (20, 20, 20, 20) (line 5).
# The j-th of the matrix projectLS gives the projection 
# on the j-th eigenvector
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Projections on the principal directions')
printt(' of some reference states...')
printt('-------------------------------------------')
#
LS        = np.zeros((K+2,K))
projectLS = np.zeros((K+2,K))
#
# Maximal "pure" LS
for i in range(0,K):
  LS[i,i] = 20
#
# Maximal possible state (20, 20, 20, 20)
for j in range(0,K):
  LS[5,j] = 20
#
printt('  LS associated with the pure states,')
printt('    the origin and the maximal state')
printt('                    (20, 20, 20, 20)')
printt(' ')
printt('    LS        : '  + str(LS))
printt(' ')
#
# Difference with the mean value
for i in range(0,K+2):
  for j in range(0,K):
    LS[i,j] = LS[i,j] - xmean[j]
#
for i  in range(0,K+2):  
  for j  in range(0,K):
    vj = eigenVectors[:,j]
#
    for l  in range(0,K):
      projectLS[i,j]    = projectLS[i,j] + vj[l]*LS[i,l]  
printt('    LS-xmean  : ' + str(LS))
printt(' ')
printt('    projectLS : ' + str(projectLS))
printt(' ')
printt('-------------------------------------------')
printt(' Projections of some reference states done!')
printt('-------------------------------------------')
printt(' ')
#---------------------------------------------------
# Projection of the (ith) data on the jth eigenvector
printt('-------------------------------------------')
printt(' Projections of data on eigenvectors...')
printt('-------------------------------------------')
printt(' ')
proj = np.zeros((L,K))
prob = np.zeros((L,K))
#
for i  in range(0,L):
  norm_ddata = 0.0
  for j  in range(0,K):
    norm_ddata = norm_ddata + ddata[i,j]*ddata[i,j]
  norm_ddata = np.sqrt( norm_ddata )
#
  for j  in range(0,K):
    vj = eigenVectors[:,j]
#
    for l  in range(0,K):
      proj[i,j] = proj[i,j] + vj[l]*ddata[i,l]  
#
    cos = proj[i,j]/norm_ddata
    prob[i,j]  = cos * cos * 100
printt('-------------------------------------------')
printt(' Projections of data on eigenvectors done!')
printt('-------------------------------------------')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Mean values for the projections')
printt(' and probabilities...')
printt('-------------------------------------------')
printt(' ')
# Mean values of the projections (they must zero)
projmean = np.array([np.mean(proj[:,0]), np.mean(proj[:,1]), np.mean(proj[:,2]), np.mean(proj[:,3])])
printt('  The mean values of the projections must nullify')
printt('    projmean : ' + str(projmean))
printt(' ')
printt('  Uncertainty associated with the mean values of the projections')
uncert_abs_proj = [0, 0, 0, 0]
for j in range(0,K):
  projj = proj[:,j]
  confidence_interval_proj=st.t.interval(confidence=0.95, df=len(projj)-1, loc=projmean[j], scale=st.sem(projj))  
  uncert_abs_proj[j]=0.5*(confidence_interval_proj[1]-confidence_interval_proj[0])
#
printt('    Uncert(proj) : ' + str(uncert_abs_proj))
printt(' ')
#---------------------------------------------------
# Mean values and uncertainties of the probabilities
probmean = np.array([np.mean(prob[:,0]), np.mean(prob[:,1]), np.mean(prob[:,2]), np.mean(prob[:,3])])
#
uncert_abs_prob = [0, 0, 0, 0]
for j in range(0,K):
  probj = prob[:,j]
  confidence_interval_prob=st.t.interval(confidence=0.95, df=len(probj)-1, loc=probmean[j], scale=st.sem(probj)) 
  uncert_abs_prob[j]=0.5*(confidence_interval_prob[1]-confidence_interval_prob[0])
#
printt('  The mean values of the corresponding probabilities')
printt('  do not nullify (as all of them are positive or zero)')
printt('    probmean     : ' + str(probmean))
printt('    Uncert(prob) : ' + str(uncert_abs_prob))
printt(' ')
#---------------------------------------------------
# Cumulative values of the probabilities obtained by adding all of them
# from 0 to 0, 0 to 1, 0 to 2, and 0 to 3)
probmeantot = np.array([np.mean(prob[:,0]), np.mean(prob[:,0])+np.mean(prob[:,1]), np.mean(prob[:,0])+np.mean(prob[:,1])+np.mean(prob[:,2]), np.mean(prob[:,0])+np.mean(prob[:,1])+np.mean(prob[:,2])+np.mean(prob[:,3])])
printt('  Cumulative values of the probabilities obtained by adding all of them')
printt('    probmeantot     : ' + str(probmeantot))
printt(' ')
printt('-------------------------------------------')
printt(' Project. and prob. mean values done!')
printt('-------------------------------------------')
printt(' ')
printt('===========================================')
printt('PROJECTION ON THE EIGENVECTORS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#
#

#===================================================
# PARTICIPATION RATIOS OF THE DATA IN THE
# LEARNING-STYLES'S ORIGINAL BASIS SET
# (ACTIVIST, REFLECTOR, THEORIST, AND PRAGMATIST)
# AND IN THE BASIS SET FORMED BY THE EIGENVECTORS
# OF THE COVARIANCE MATRIX
#===================================================
printt('===========================================')
printt('PARTICIPATION RATIOS (PRs)...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Computation of the PRs...')
printt('-------------------------------------------')
printt(' ')
# Participation ratios in the original LS basis set
# (column 0) and in the basis set formed by the
# covariance eigenvectors (column 1)
pr   = np.zeros((L,2))
#
# Notice that the PRs for a state with coefficientes C_i
# are defined as
#
# PR = (\sum Ci^2)^2/\sum C_i^4
#
# Instead of as
#
# PR = \sum Ci^2 /\sum C_i^4
#
# In this way, in a basis set formed by N elements
#
#  1 <= PR <= N.
#
# Notice that PR may equal 0 when all coefficients
# nullify. Then, we impose a value equal to 1.
#
for i  in range(0,L):
# LSs basis set
# (activist, reflector, theorist, pragmatist)
  sum2 = 0.0
  sum4 = 0.0
  for j  in range(0,K):
    p2   = data[i, j] * data[i, j]
    sum2 = sum2 + p2
    sum4 = sum4 + p2 * p2
#
  if( p2 == 0):
#   We impose that the PR equals 1
    pr[i,0] = 1 # 0
  else:
    pr[i,0] = sum2 * sum2 / sum4
#
# Basis set formed by the eigenfunctions
  sum2 = 0.0
  sum4 = 0.0
  for j  in range(0,K):
    p2   = proj[i, j] * proj[i, j]
    sum2 = sum2 + p2
    sum4 = sum4 + p2 * p2
#
  if( p2 == 0):
#   We impose that the PR equals 1
    pr[i,1] = 1
  else:
    pr[i,1] = sum2 * sum2 / sum4
#
printt('-------------------------------------------')
printt(' Computation of the PRs done!')
printt('-------------------------------------------')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Mean value and uncertainties of the PRs...')
printt('-------------------------------------------')
printt(' ')
prmean = [ np.mean(pr[:,0]), np.mean(pr[:,1])]
#
confidence_interval_pr0=st.t.interval(confidence=0.95, df=len(pr[:,0])-1, loc=np.mean(pr[:,0]), scale=st.sem(pr[:,0])) 
uncert_abs_pr0    =0.5*(confidence_interval_pr0[1]-confidence_interval_pr0[0])
#
confidence_interval_pr1=st.t.interval(confidence=0.95, df=len(pr[:,1])-1, loc=np.mean(pr[:,1]), scale=st.sem(pr[:,1])) 
uncert_abs_pr1    =0.5*(confidence_interval_pr1[1]-confidence_interval_pr1[0])
#
uncert_abs_pr = [ uncert_abs_pr0, uncert_abs_pr1 ]
printt('  prmean     : ' + str(prmean))
printt('  Uncert(PR) : ' + str(uncert_abs_pr))
printt(' ')
printt('-------------------------------------------')
printt(' Mean value and uncertainties of PRs done!')
printt('-------------------------------------------')
printt(' ')
#
#---------------------------------------------------
# Plot of the participation ratios...
printt('-------------------------------------------')
printt(' Analysis of the cumulative probability')
printt(' distribution (W(PR)) of the PRs')
printt(' (plot and fitting using a Weibull)')
printt(' distribution)...')
printt('-------------------------------------------')
#
fig, ax = plt.subplots(figsize = ( w_fig, h_fig )) #, layout='constrained')
#
ymaxx   = 10000000000
#
parameters_Weibull_PR = np.zeros((K,2))
xWeib   = np.arange(1, 4.1, 0.001)
#
for ipr in range(0, 2):
  if(ipr == 0):
    printt('  W(PR) in the LS original basis set')
  else:
    printt('  W(PR) in the eigenvectors basis set')
  printt(' ')
#
  x = pr[:, ipr]
  x = np.sort(x)
#
  y = np.zeros(L)
  y[0] = 0
  dy = 1 / float(L-1)
  for i in range(1,L):
    y[i] = y[i-1] + dy
#
# Staircase
  x2 = np.zeros(2*L+2)
  y2 = np.zeros(2*L+2)
  x2[0]  =   0
  y2[0]  =   0
  x2[2*L+1] = x[L-1]
  y2[2*L+1] = 1
#
# Find the points close to 1, 2, and 3
# in order to correcly color the
# PR that lie in the ranges [0, 1), 
# (1, 2), (2,3), and (3,4]
  ic   = [-1, -1]
  xc   = [ 1,  2]
  xref = [ 2,  3]
#
  for i in range(0,len(ic)):
    for j in range(0, L):
      if( x[j] > xc[i] and x[j] < xref[i] ):
        ic[i] = j
        xc[i] = x[j]
#
# Fitting using Weibull distribution
  param = curve_fit(Wweibull, x, y, [prmean[ipr], 1], maxfev = max_iter_Weib)

  [alpha_Weibull, k_Weibull] = param[0]
  pcov = param[1]
  parameters_Weibull_PR[ipr] = [alpha_Weibull, k_Weibull]
  printt('  alpha_Weibull : ' + str(alpha_Weibull))
  printt('  k_Weibull : ' + str(k_Weibull))
  printt(' ')
  printt(' ')
#
  for i in range(0,L):
    x2[2*i+1] = x[i]
    x2[2*i+2] = x[i]

    y2[2*i+1] = y[i]
    y2[2*i+2] = y[i] + dy
  y2[2*L] = 1    
#
# Plotting limits
  ymin_W = 0.00
  ymax_W = 1.01 
  ax.set_xlim( [ 1, 4] )
  ax.set_ylim( [ ymin_W, ymax_W] )
#
# Tick parameters
  ax.tick_params(axis='both', which='major', labelsize = ticksize_PR)
#
# Labels
  ax.set_xlabel('PR', fontsize = labelsize_PR)
  ax.set_ylabel('W [PR]',      fontsize = labelsize_PR)
#
# Plot cumulative distribution and Weibull fitting function
  ax.plot(x2, y2, lw = PR_W_width, color = PR_color[ipr], ls = PR_line[ipr])
  ax.plot(xWeib,Wweibull(xWeib, alpha_Weibull, k_Weibull) , lw = PR_W_width, color = PR_color[ipr], ls = PR_line[ipr])  
#
# Plot the average value as a vertical line
  ax.plot([prmean[ipr], prmean[ipr]], [ymin, ymax], lw = PR_mean_width, color = PR_color[ipr], ls = PR_mean_line)
  plt.fill_between([prmean[ipr]-uncert_abs_pr[ipr], prmean[ipr]+uncert_abs_pr[ipr]], [ymin, ymaxx], color = PR_color[ipr], alpha = PR_mean_alpha )  
#
filename_PR_W = output_statistics_pr + '/Fig_PR_W.png'
printt('Saving ' + filename_PR_W)
plt.savefig(filename_PR_W)
printt(' ')
printt('-------------------------------------------')
printt(' W(PR) done!')
printt('-------------------------------------------')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Probability distribution (P(PR)) of the PRs')
printt(' (plot and fitting using Weibull function...)')
printt('-------------------------------------------')
#
fig, ax = plt.subplots(figsize = ( w_fig, h_fig )) #, layout='constrained')
#
xWeib   = np.arange(0.8, 4.2, 0.001)
xmaxtot = 0
ymaxtot = 0
ymaxx   = 10000000000
#
for ipr in range(0, 2):
  if(ipr == 0):
    printt('  W(PR) in the LS original basis set')
  else:
    printt('  W(PR) in the eigenvectors basis set')
  printt(' ')
  x = pr[:, ipr]
#
# Tick parameters
  ax.tick_params(axis='both', which='major', labelsize = ticksize_PR)
#
# Labels
  ax.set_xlabel('PR', fontsize = labelsize_PR)
  ax.set_ylabel('P [PR]',      fontsize = labelsize_PR)
#
  y, x, _= ax.hist(x, bins=15, color = PR_color[ipr], alpha = PR_alpha) # We create a histogram with 15 blocks instead of just bins=10
  xmax = x.max()
  if xmax > xmaxtot:
    xmaxtot = xmax
#
  ymax = y.max()
  if ymax > ymaxtot:
    ymaxtot = ymax
#
# Probability distribution for Weibull function, scaled
# in such a way that its maximum coincides with that of the histogram
  [alpha_Weibull, k_Weibull] = parameters_Weibull_PR[ipr]
  factor = ymax / Pweibull(alpha_Weibull, alpha_Weibull, k_Weibull)
  ax.plot(xWeib,factor*Pweibull(xWeib, alpha_Weibull, k_Weibull) , lw = PR_W_width,   color = PR_color[ipr], ls = PR_line[ipr])  
#
# Plot the average value as a vertical line
  plt.fill_between([prmean[ipr]-uncert_abs_pr[ipr], prmean[ipr]+uncert_abs_pr[ipr]], [ymin, ymaxx], color = PR_color[ipr], alpha = PR_mean_alpha )
  ax.plot([prmean[ipr], prmean[ipr]], [ymin, ymaxx], lw = PR_mean_width, color = PR_color[ipr], ls = PR_mean_line)
#
# Plotting limits 
ymax = ymaxtot * 1.05
ax.set_xlim( [ 0.99, 4.01] )
ax.set_xlim( [ 1, xmaxtot] )
ax.set_ylim( [ 0, ymax] )

filename_PR_P = output_statistics_pr + '/Fig_PR_P.png'
printt('Saving ' + filename_PR_P)
plt.savefig(filename_PR_P)
printt(' ')
printt('-------------------------------------------')
printt(' P(PR) done!')
printt('-------------------------------------------')
printt(' ')
printt(' ')
printt('===========================================')
printt('PARTICIPATION RATIOS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#
#
#===================================================
# 4D PLOT OF THE LS
# (THE FOURTH LS IS INCLUDED IN THE CHARACTERISTICS
# OF THE SYMBOLS)
#===================================================
printt('===========================================')
printt('4D PLOT OF THE LSs...')
printt('===========================================')
printt(' ')
#----------------------------------------------------------
Llt = len(learning_trios)
printt('-------------------------------------------')
printt(' Llt : ' + str(Llt))
printt('-------------------------------------------')
printt(' ')

for ilt in range(itrios,itrios+1):

  fig = plt.subplots(figsize = ( w_fig, h_fig )) #, layout='constrained')
  ax  = plt.axes(projection ="3d")

  ax.set_xlim([ xmin3D, xmax3D] )
  ax.set_ylim([ ymin3D, ymax3D] )
  ax.set_zlim([ zmin3D, zmax3D] )

  ax.set_xticks([5, 10, 15])
  ax.set_yticks([5, 10, 15])
  ax.set_zticks([5, 10, 15])
  ax.tick_params(axis='both', which='major', labelsize = ticksize_3D)
  
#----------------------------------------------------------
# Scatter as a function of PCi, PCj and PCk
# with i, j, k = 0, 1, 2, or 3.
  printt('-------------------------------------------')
  printt('   ilt : ' + str(ilt))

  index0 = learning_trios[ilt][0]
  index1 = learning_trios[ilt][1]
  index2 = learning_trios[ilt][2]
  
  for i in range(0,K):
    if( i != index0 and i!= index1  and i!= index2 ):
      index3 = i

# Label_LS
  Label_LS0 = Label_LS[index0]
  Label_LS1 = Label_LS[index1]
  Label_LS2 = Label_LS[index2]

  printt('Labels LS   :     ' + str(Label_LS0) + ' vs ' + str(Label_LS1) + ' vs ' + str(Label_LS2))
  printt(' ')        
  printt('     ' + str(Label_LS0) + ' vs ' + str(Label_LS1) + ' vs ' + str(Label_LS2))
  printt('      Mean :' + str(xmean[index0]) + str(xmean[index1]) + str(xmean[index2]))
  printt(' ')
  printt('      index0 : ' + str(index0))
  printt('      index1 : ' + str(index1))     
  printt('      index2 : ' + str(index2))
  printt('     (index3 : ' + str(index3) + ')')
  
  ax.set_xlabel(Label_LS0, fontsize = labelsize_3D, labelpad = labelpad_3D)
  ax.set_ylabel(Label_LS1, fontsize = labelsize_3D, labelpad = labelpad_3D)
  ax.set_zlabel(Label_LS2, fontsize = labelsize_3D, labelpad = labelpad_3D)
                   
  vx0 = data[:,index0]
  vx1 = data[:,index1]
  vx2 = data[:,index2]
  vx3 = data[:,index3]

# Eigenvectors
  for j in range(0,K):
    vj = eigenVectors_max1[:,j]*distmeanmax

    v0 = xmean[index0]+vj[index0]
    v1 = xmean[index1]+vj[index1]
    v2 = xmean[index2]+vj[index2]              
    v3 = xmean[index3]+vj[index3]

    ax.plot([xmean[index0], v0], [xmean[index1], v1], [xmean[index2], v2], lw=width_PC_vectors_3D[j], color=PC_color[j], ls=PC_line_bis[j])

# Shaded are in the PC1-PC2 plane
  vj0 = eigenVectors[:,0]*distmeanmax
  v00 = xmean[index0]+vj0[index0]
  v01 = xmean[index1]+vj0[index1]
  v02 = xmean[index2]+vj0[index2]        
  v03 = xmean[index3]+vj0[index3]
    
  vj1 = eigenVectors[:,1]*distmeanmax
  v10 = xmean[index0]+vj1[index0]
  v11 = xmean[index1]+vj1[index1]
  v12 = xmean[index2]+vj1[index2]        
  v13 = xmean[index3]+vj1[index3]
  
  xx = [xmean[index0], v00, v10]
  yy = [xmean[index1], v01, v11]
  zz = [xmean[index2], v02, v12]

# Points
  for i in range(0,L):
    x0 = vx0[i]
    x1 = vx1[i]
    x2 = vx2[i]
    x3 = vx3[i]

    tendency3 = scatter_tendency(Label_LS[index3], x3)
    [tendency3, scatter_color3, scatter_size3, scatter_alpha3, scatter_symbol3]=scatter_properties(tendency3, x3)

    ax.scatter3D(x0, x1, x2, s=scatter_size3, marker=scatter_symbol3, color=scatter_color3,      alpha=scatter_alpha3)
    
#   Projections on the Cartesian planes
    colorsx     = colors_proj
    edgecolorsx = edgecolors_proj
      
    colorsy     = colors_proj
    edgecolorsy = edgecolors_proj

    colorsz     = colors_proj
    edgecolorsz = edgecolors_proj
    
    ax.scatter(x1, x2, s=scatter_size3, marker=scatter_symbol3, color=colorsx,  alpha=scatter_alpha3, zdir='x')
    ax.scatter(x0, x2, s=scatter_size3, marker=scatter_symbol3, color=colorsy,  alpha=scatter_alpha3, zdir='y')
    ax.scatter(x0, x1, s=scatter_size3, marker=scatter_symbol3, color=colorsz,  alpha=scatter_alpha3, zdir='z')
            
# Average mean
  ax.scatter3D(xmean[index0], xmean[index1], xmean[index2], s=size_mean, marker=mean_symbol, color='k')
  
# Projections of the average mean on the Cartesian planes
  ax.scatter(xmean[index1], xmean[index2], s=size_mean, marker=mean_symbol, color=colorsx, zdir='x')
  
  ax.scatter(xmean[index0], xmean[index2], s=size_mean, marker=mean_symbol, color=colorsy, zdir='y')
    
  ax.scatter(xmean[index0], xmean[index1], s=size_mean, marker=mean_symbol, color=colorsz, zdir='z')
  
# Fill uncertainty areas for the average mean on the projections
  xx = [xmin3D, xmin3D, xmin3D, xmin3D]
  yy = [xmean[index1]-dxmean[index1],
        xmean[index1]+dxmean[index1],
        xmean[index1]+dxmean[index1],
        xmean[index1]-dxmean[index1]]
  zz = [xmean[index2]-dxmean[index2],
        xmean[index2]-dxmean[index2],
        xmean[index2]+dxmean[index2],
        xmean[index2]+dxmean[index2]]
  verts = [(xx[i],yy[i],zz[i]) for i in range(len(xx))]
  ax.add_collection3d(Poly3DCollection([verts],facecolor = mean_uncert_color, alpha = mean_alpha))

  xx = [xmean[index0]-dxmean[index0],
        xmean[index0]+dxmean[index0],
        xmean[index0]+dxmean[index0],
        xmean[index0]-dxmean[index0]]
  zz = [xmean[index2]-dxmean[index2],
        xmean[index2]-dxmean[index2],
        xmean[index2]+dxmean[index2],
        xmean[index2]+dxmean[index2]]
  yy = [ymin3D, ymin3D, ymin3D, ymin3D]
  verts = [(xx[i],yy[i],zz[i]) for i in range(len(xx))]
  ax.add_collection3d(Poly3DCollection([verts],facecolor = mean_uncert_color, alpha = mean_alpha))
  
  xx = [xmean[index0]-dxmean[index0],
        xmean[index0]+dxmean[index0],
        xmean[index0]+dxmean[index0],
        xmean[index0]-dxmean[index0]]
  yy = [xmean[index1]-dxmean[index1],
        xmean[index1]-dxmean[index1],
        xmean[index1]+dxmean[index1],
        xmean[index1]+dxmean[index1]]
  zz = [zmin3D, zmin3D, zmin3D, zmin3D]
  verts = [(xx[i],yy[i],zz[i]) for i in range(len(xx))]
  ax.add_collection3d(Poly3DCollection([verts],facecolor = mean_uncert_color, alpha = mean_alpha))

# Origin
  orig0 = LS[4, index0] + xmean[index0]
  orig1 = LS[4, index1] + xmean[index1]
  orig2 = LS[4, index2] + xmean[index2]
  ax.scatter3D(orig0, orig1, orig2, s = size_mean, marker = origin_symbol, color = origin_color, alpha = origin_alpha)  

# Pure maximal learning styles
# (only those whith the LS that appear in the axes)
  for i in [index0, index1, index2]:  
    x0 = LS[i, index0] + xmean[index0]
    x1 = LS[i, index1] + xmean[index1]
    x2 = LS[i, index2] + xmean[index2]
    ax.scatter3D(x0, x1, x2,  s = size_mean, marker = twenty_symbol[i], color = PC_color[i], alpha = twenty_alpha)  
    ax.plot([orig0, x0], [orig1, x1], [orig2, x2], color = twenty_color[i], ls = twenty_line[i], lw = twenty_width_3D[i])

# Maximal state (20, 20, 20, 20)
  ax.plot(xx,yy,zz,label="line plot")

  ax.view_init(25,40)

  filename_ls_3d = output_participants + '/Fig_'+Label_LS0+'_'+Label_LS1+'_'+Label_LS2+'.png'
  printt('Saving ' + filename_ls_3d)
  plt.savefig(filename_ls_3d)

  printt(' Plots for the LSs done!')
  printt('-------------------------------------------')
  printt(' ')
printt('===========================================')
printt('4D PLOT OF THE LSs DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#
#===================================================
#===================================================
#
#
#
#

printt('===========================================')
printt(' PROJECTIONS ON PCs AND STATISTICS...')
printt('===========================================')
printt(' ')
#---------------------------------------------------
printt('-------------------------------------------')
printt(' W(s) for the projections on the PCs...')
printt('-------------------------------------------')
parameters_Weibull_PC20 = np.zeros((K,3))
  
xmin = proj.min()
xmax = proj.max()
  
xmin_proj_tot = xmin
xmax_proj_tot = xmax

xWeib_proj = np.arange(xmin*2.5, xmax*2.5, (xmax-xmin)/100000)

for ipc in range(0, K):
  printt('   PC -20 to 20 : ' + str(ipc))
  x = proj[:, ipc]
  x = np.sort(x)

  xmin_proj = x.min()
 
  y = np.zeros(L)
  y[0] = 0
  dy = 1 / float(L-1)
  for i in range(1,L):
    y[i] = y[i-1] + dy

# Fitting using Weibull distribution
  param = curve_fit(Wweibull, x-xmin_proj, y, [projmean[ipc]-xmin_proj, 1], maxfev = max_iter_Weib)

  [alpha_Weibull, k_Weibull] = param[0]
  theta_Weibull = xmin_proj
  parameters_Weibull_PC20[ipc] = [alpha_Weibull, k_Weibull, theta_Weibull]
  printt('   alpha_Weibull     : ' + str(alpha_Weibull))
  printt('   k_Weibull         : ' + str(k_Weibull))
  printt('   theta_Weibull     : ' + str(theta_Weibull))
  printt(' ')
  
for ilp in range(0, 1):

  fig, ((ax_hor, ax_no), (ax, ax_ver)) = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), layout='constrained')

# Limits
  ax_no.axis('off')
  ax.set_xlim([ xmin, xmax] )
  ax.set_ylim([ ymin, ymax] )

# Ticks
  ax.tick_params(axis='both', which='major', labelsize = ticksize)

#----------------------------------------------------------
# Scatter as a function of PC0, PC1, PC2, and PC3
  printt('-------------------------------------------')
  printt(' Plotting scatter 2D on the principal axes')
  printt(' and probabilities on the PCs...')

  printt(' PC : ' + str(ilp))
  printt(' ')
  index0 = learning_pairs[ilp][0]
  index1 = learning_pairs[ilp][1]
  index2 = -1
 
  for i in range(0,K):
    if( i != index0 and i!= index1 ):
      if ( index2 == -1):
        index2 = i
      else:
        index3 = i
#
# MAIN PANEL 
# Label PCs
  Label_PC0 = Label_PC[index0]
  Label_PC1 = Label_PC[index1]
  
  Label_PCPC0 = Label_PCPC[index0]
  Label_PCPC1 = Label_PCPC[index1]

  printt(' ( ' + str(Label_PCPC0) + ' vs ' + str(Label_PCPC1) + ' )')
  printt(' ')

  ax.set_xlabel(Label_PCPC0, fontsize=labelsize)
  ax.set_ylabel(Label_PCPC1, fontsize=labelsize)
    
# LS values
  vx0 = data[:,index0]
  vx1 = data[:,index1]
  vx2 = data[:,index2]
  vx3 = data[:,index3]

# Plot the projection of the coefficients PC0 and PC1
  vy0 = proj[:,index0]
  vy1 = proj[:,index1]
  vy2 = proj[:,index2]
  vy3 = proj[:,index3]

# Plot limits
  xmin = np.min( [np.min(vy0), np.min(projectLS[:,index0])])
  xmax = np.max( [np.max(vy0), np.max(projectLS[:,index0])])
  
  ymin = np.min( [np.min(vy1), np.min(projectLS[:,index1])])
  ymax = np.max( [np.max(vy1), np.max(projectLS[:,index1])])
  
  if ( xmin < 0 ):
   xmin = xmin * 1.05
  else:
   xmin = xmin * 0.95
   
  if ( xmax < 0 ):
   xmax = xmax * 0.95
  else:
   xmax = xmax * 1.05
   
  if ( ymin < 0 ):
   ymin = ymin * 1.05
  else:
   ymin = ymin * 0.95
   
  if ( ymax < 0 ):
   ymax = ymax * 0.95
  else:
   ymax = ymax * 1.05
   
  printt('   xmin : ' + str(xmin))
  printt('   xmax : ' + str(xmax))
  printt('   ymin : ' + str(ymin))
  printt('   ymax : ' + str(ymax))
  printt(' ')
  
# Points
  for i in range(0,L):
#  printt(i)
    x0 = vx0[i]
    x1 = vx1[i]
    x2 = vx2[i]
    x3 = vx3[i]
    
    y0 = vy0[i]
    y1 = vy1[i]
    y2 = vy2[i]
    y3 = vy3[i]

    tendency2 = scatter_tendency(Label_LS[index2], x2)
    tendency3 = scatter_tendency(Label_LS[index3], x3)
    [tendency2, scatter_color2, scatter_size2, scatter_alpha2, scatter_symbol2]=scatter_properties(tendency2, x2)
    [tendency3, scatter_color3, scatter_size3, scatter_alpha3, scatter_symbol3]=scatter_properties(tendency3, x3)

#   Color depends on whether we are making a plot with PC' vs PC1 or not
#   (see projections in the 3D plots as a funciton of the original LSs)
    if (ilp == 0):
      ax.scatter(y0, y1, s=scatter_size2, marker=scatter_symbol2,      color = colors_proj_PC, alpha=scatter_alpha3)
      ax.scatter(y0, y1, s=scatter_size2, marker=scatter_symbol2, edgecolors = colors_proj_PC, facecolors='none')
    else:
      ax.scatter(y0, y1, s=scatter_size2, marker=scatter_symbol2,      color = colors_proj, alpha=scatter_alpha3)
      ax.scatter(y0, y1, s=scatter_size2, marker=scatter_symbol2, edgecolors = colors_proj, facecolors='none')
      
# Origin
  orig0 = projectLS[4, index0]
  orig1 = projectLS[4, index1]
  ax.scatter(orig0, orig1, s = size_mean, marker = origin_symbol, color = origin_color, alpha = origin_alpha)  

# Pure maximal learning styles
  for i in range(0,K):  
    y0 = projectLS[i, index0]
    y1 = projectLS[i, index1]
    ax.scatter(y0, y1,   s = size_mean, marker = twenty_symbol[i], color = PC_color[i], alpha = twenty_alpha)  
    ax.plot([orig0, y0], [orig1, y1], color = twenty_color[i], ls = twenty_line[i], lw = twenty_width_2D[i])

  i = K + 1
  y0 = projectLS[i, index0]
  y1 = projectLS[i, index1]
  ax.scatter(y0, y1,   s = size_mean, marker = twenty_symbol[i], color = twenty_color[i], alpha = twenty_alpha)  
  ax.plot([orig0, y0], [orig1, y1], color    = twenty_color[i], ls = twenty_line[i], lw = twenty_width_2D[i])
  
# Average mean
  ax.scatter(0, 0, s = size_mean, marker=mean_symbol, color = mean_color2)
  
  ax.set_xlim([ xmin, xmax] )
  ax.set_ylim([ ymin, ymax] )
#
#**********************************************************************
# SECONDARY PANELS (PROBABILITY DENSITY FUNCTIONS)
#
  xWeib   = np.arange(xmin, xmax, (xmax-xmin)/10000)
  yWeib   = np.arange(ymin, ymax, (ymax-ymin)/10000)
#
#-------------------------------------------------------
# Projection on the PC associated with the horizontal axis
  ax_hor.set_ylabel('PDF', fontsize=labelsize)
  ax_hor.set_xlim([ xmin, xmax] )
  
  ipc = index0
  x = proj[:, ipc]
  x = np.sort(x)
     
  [alpha_Weibull, k_Weibull, theta_Weibull] = parameters_Weibull_PC20[ipc]
  printt('   alpha_Weibull : ' + str(alpha_Weibull))
  printt('   k_Weibull     : ' + str(k_Weibull))
  printt('   theta_Weibull : ' + str(theta_Weibull))
  printt(' ')
  
  ax_hor.tick_params(axis='both', which='major', labelsize=ticksize, labelbottom = False)
 
  y, x, _= ax_hor.hist(x, bins=20, color = PC_color[ipc], alpha=0.5)
  ymaxx = y.max()

# Probability distribution for Weibull function, scaled
  factor = ymaxx / Pweibull(alpha_Weibull, alpha_Weibull,  k_Weibull)
  ax_hor.plot(xWeib_proj,factor*Pweibull_translated(xWeib_proj, alpha_Weibull, k_Weibull, theta_Weibull), lw = PC_W_width,   color = PC_color[ipc], ls = PC_line[0]) #, ls = PC_line[ipc])
  
  ymaxx = ymaxx * 1.205
  ax_hor.set_ylim( [ 0, ymaxx] )   
 
# Plot as a vertical line the origin
  ax_hor.plot([projectLS[4,ipc], projectLS[4,ipc]], [ymin, ymaxx], color = twenty_color[4], ls = twenty_line[4], lw = twenty_width_W[6])
  
# Mark with lines the origin and the maximal (pure) learning styles  
  for i in range(0,K):  
    ax_hor.plot([projectLS[i, ipc], projectLS[i, ipc]], [ymin, ymaxx], color = twenty_color[i], ls = twenty_line[i], lw = twenty_width_W[6])

# Mark as lines the maximal point (20, 20, 20, 20)
  ax_hor.plot([projectLS[5,ipc], projectLS[5,ipc]], [ymin, ymaxx], color=twenty_color[5], ls = twenty_line[5], lw = twenty_width_W[6])
#
#-------------------------------------------------------
# Projection on the PC associated with the vertical axis
  ax_ver.set_xlabel('PDF', fontsize=labelsize)
  ax_ver.set_ylim([ ymin, ymax] )
  
  ipc = index1
  x = proj[:, ipc]
  x = np.sort(x)
     
  [alpha_Weibull, k_Weibull, theta_Weibull] = parameters_Weibull_PC20[ipc]
  printt('   alpha_Weibull : ' + str(alpha_Weibull))
  printt('   k_Weibull     : ' + str(k_Weibull))
  printt('   theta_Weibull : ' + str(theta_Weibull))
  printt(' ')
  
  ax_ver.tick_params(axis='both', which='major', labelsize=ticksize, labelleft = False)
 
  y, x, _= ax_ver.hist(x, bins=20, color = PC_color[ipc], alpha=0.5, orientation='horizontal')
  ymaxx = y.max()

# Probability distribution for Weibull function, scaled
  factor = ymaxx / Pweibull(alpha_Weibull, alpha_Weibull,  k_Weibull)
  ax_ver.plot(factor*Pweibull_translated(yWeib, alpha_Weibull, k_Weibull, theta_Weibull), yWeib, lw = PC_W_width,   color = PC_color[ipc], ls = PC_line[0]) #, ls = PC_line[ipc])
  
  ymaxx = ymaxx * 1.20
  ax_ver.set_xlim( [ 0, ymaxx ] )    # Plot as a vertical line the origin
  ax_ver.plot([xmin, ymaxx], [projectLS[4,ipc], projectLS[4,ipc]], color = twenty_color[4], ls = twenty_line[4], lw = twenty_width_W[6])
  
# Mark with lines the origin and the maximal (pure) learning styles  
  for i in range(0,K):  
    ax_ver.plot([xmin, ymaxx], [projectLS[i, ipc], projectLS[i, ipc]], color = twenty_color[i], ls = twenty_line[i], lw = twenty_width_W[6])

# Mark as lines the maximal point (20, 20, 20, 20)
  ax_ver.plot([xmin, ymaxx], [projectLS[5,ipc], projectLS[5,ipc]], color=twenty_color[5], ls = twenty_line[5], lw = twenty_width_W[6])

  filename_PC0PC1 = output + '/Fig_'+Label_PCPC0+'_'+Label_PCPC1+'_P.png'
  plt.savefig(filename_PC0PC1)
  printt('   Saving ' + filename_PC0PC1)
  plt.savefig(filename_PC0PC1)
  printt(' ')
  printt(' Plots for the PCs done!')
  printt('-------------------------------------------')
  printt(' ')
printt('===========================================')
printt(' PROJECTIONS ON PCs AND STATISTICS DONE!')
printt('===========================================')
printt(' ')
printt(' ')
printt(' ')
printt(' ')
#********************************************

printt('===========================================')
printt(' PROJECTIONS ON PCs AND PROFESSORS LS...')
printt('===========================================')
printt(' ')

printt('-------------------------------------------')
printt(' Projections of professors data on eigenvectors...')
printt('-------------------------------------------')
printt(' ')
# Difference between the input data and the mean
if(Lprof!=0):
 ddata_prof = np.zeros((Lprof,K))
 for i in range(0,Lprof):
  for j in range(0,K):
    ddata[i,j] = data_prof[i,j] - xmean[j]

 proj_prof = np.zeros((L,K))
#
 for i  in range(0,Lprof):
  norm_ddata = 0.0
  for j  in range(0,K):
    norm_ddata = norm_ddata + ddata_prof[i,j]*ddata_prof[i,j]
  norm_ddata = np.sqrt( norm_ddata )
#
  for j  in range(0,K):
    vj = eigenVectors[:,j]
#
    for l  in range(0,K):
      proj_prof[i,j] = proj_prof[i,j] + vj[l]*ddata[i,l]
#
 printt('-------------------------------------------')
 printt(' Projections of professors data on eigenvectors done!')
 printt('-------------------------------------------')
 printt(' ')

pr_prof   = np.zeros((Lprof,2))

for i  in range(0,Lprof):
  if (i == 0):
    printt('-------------------------------------------')
    printt(' Participation ratios of the professors... ')
    printt('-------------------------------------------')

# LSs basis set
# (activist, reflector, theorist, pragmatist)
  sum2 = 0.0
  sum4 = 0.0
  for j  in range(0,K):
    p2   = data_prof[i, j] * data_prof[i, j]
    sum2 = sum2 + p2
    sum4 = sum4 + p2 * p2
#
  if( p2 == 0):
#   We impose that the PR equals 1
    pr_prof[i,0] = 1 # 0
  else:
    pr_prof[i,0] = sum2 * sum2 / sum4
#
# Basis set formed by the eigenfunctions
  sum2 = 0.0
  sum4 = 0.0
  for j  in range(0,K):
    p2   = proj_prof[i, j] * proj_prof[i, j]
    sum2 = sum2 + p2
    sum4 = sum4 + p2 * p2
#
  if( p2 == 0):
#   We impose that the PR equals 1
    pr_prof[i,1] = 1 # 0
  else:
    pr_prof[i,1] = sum2 * sum2 / sum4

  printt('-------------------------------------------')
  printt(' Participation ratios of the professors done!')
  printt('-------------------------------------------')
  printt(' ')



for i_prof in range(0, Lprof):
  profi = profs[i_prof]
  printt('Professor : ' + profi )
  fig, axs = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), sharex = True, sharey = True, layout='constrained')  
  
  for ils in range(0,K):

    axs[panels[ils][0],panels[ils][1]].tick_params(axis='both', which='major', labelsize=ticksize)  
#----------------------------------------------------------
#   Scatter as a function of PC0, PC1, PC2, and PC3
    printt('-------------------------------------------')
    printt(' Plotting scatter 2D on PCs (professor clusters)...')

    printt(' ls : ' + str(ilp))
    printt('  ')
  
    tendencyref = scatter_tendency(Label_LS[ils], data_prof[i_prof, ils])
    [tendencyprof, scatter_colorref, scatter_sizeref, scatter_alpharef, scatter_symbolprof]=scatter_properties(tendencyref,  data_prof[i_prof, ils])
    
    index0 = learning_pairs[ilp][0]
    index1 = learning_pairs[ilp][1]
    index2 = -1
 
    for i in range(0,K):
      if( i != index0 and i!= index1 ):
        if ( index2 == -1):
          index2 = i
        else:
          index3 = i
    
#   Label PCs
    Label_PC0 = Label_PC[index0]
    Label_PC1 = Label_PC[index1]
  
    Label_PCPC0 = Label_PCPC[index0]
    Label_PCPC1 = Label_PCPC[index1]

    printt(' ( ' + Label_PCPC0 + ' vs ' + Label_PCPC1 + ' )')
    printt('  ')

    ax.set_xlabel(Label_PCPC0, fontsize=labelsize)
    ax.set_ylabel(Label_PCPC1, fontsize=labelsize)
    
#   LS values
    vx0 = data[:,index0]
    vx1 = data[:,index1]
    vx2 = data[:,index2]
    vx3 = data[:,index3]

#   Considered LS for clustering
#   vxref = data[:,ils]
  
#   Plot the projection of the coefficients PC0 and PC1
    vy0 = proj[:,index0]
    vy1 = proj[:,index1]
    vy2 = proj[:,index2]
    vy3 = proj[:,index3]

#   Plot limits
    xmin = np.min( [np.min(vy0), np.min(projectLS[:,index0])])
    xmax = np.max( [np.max(vy0), np.max(projectLS[:,index0])])
  
    ymin = np.min( [np.min(vy1), np.min(projectLS[:,index1])])
    ymax = np.max( [np.max(vy1), np.max(projectLS[:,index1])])
  
    if ( xmin < 0 ):
      xmin = xmin * 1.05
    else:
     xmin = xmin * 0.95
   
    if ( xmax < 0 ):
      xmax = xmax * 0.95
    else:
      xmax = xmax * 1.05
   
    if ( ymin < 0 ):
      ymin = ymin * 1.05
    else:
      ymin = ymin * 0.95
   
    if ( ymax < 0 ):
      ymax = ymax * 0.95
    else:
      ymax = ymax * 1.05
   
    printt( ' xmin : ' + str(xmin))
    printt( ' xmax : ' + str(xmax))
    printt( ' ymin : ' + str(ymin))
    printt( ' ymax : ' + str(ymax))
    printt('  ')
  
#   Points
    for i in range(0,L):
      x0 = vx0[i]
      x1 = vx1[i]

      vxrefi = data[i,ils]
        
      y0 = vy0[i]
      y1 = vy1[i]

      tendencyref = scatter_tendency(Label_LS[ils], vxrefi)
      [tendencyref, scatter_colorref, scatter_sizeref, scatter_alpharef, scatter_symbolref]=scatter_properties(tendencyref, vxrefi)

      if(tendencyref == tendencyprof):
        scatter_sizeref   = 2*scatter_sizeref
        scatter_symbolref = 'X'

#     Color depends on the tendency
      if( tendencyref == 'vl' or tendencyref == 'l' ):
        colorref = tendency_color[0]
      else:
        if( tendencyref == 'm' ):
          colorref = tendency_color[1]
        else:
          colorref = tendency_color[2]
    
      axs[panels[ils][0],panels[ils][1]].scatter(y0, y1, s=scatter_sizeref, marker=scatter_symbolref,      color = colorref, alpha=scatter_alpharef)
      axs[panels[ils][0],panels[ils][1]].scatter(y0, y1, s=scatter_sizeref, marker=scatter_symbolref, edgecolors = colorref, facecolors='none')

#   Points professors
    for i in range(0, Lprof):

#     Plot the projection of the coefficients PC0 and PC1
      tendencyref = scatter_tendency(Label_LS[ils], data_prof[i, ils])
      [tendencyref, scatter_colorref, scatter_sizeref, scatter_alpharef, scatter_symbolref]=scatter_properties(tendencyref, data_prof[i, ils])

      if( tendencyref == 'vl' or tendencyref == 'l' ):
        colorref = tendency_color[0]
      else:
        if( tendencyref == 'm' ):
          colorref = tendency_color[1]
        else:
          colorref = tendency_color[2]

      y0 = proj_prof[i, index0]
      y1 = proj_prof[i, index1]
      if(i == i_prof ):
        axs[panels[ils][0],panels[ils][1]].scatter(y0, y1, s=2*scatter_sizeref, marker='P', color = colorref, alpha=scatter_alpharef)
        axs[panels[ils][0],panels[ils][1]].scatter(y0, y1, s=2*scatter_sizeref, marker='P', edgecolors = colorref, facecolors='none')
      else:
        axs[panels[ils][0],panels[ils][1]].scatter(y0, y1, s=scatter_sizeref, marker='P', edgecolors = colorref, facecolors='none')

#   Origin
    orig0 = projectLS[4, index0]
    orig1 = projectLS[4, index1]
    axs[panels[ils][0],panels[ils][1]].scatter(orig0, orig1, s = size_mean, marker = origin_symbol, color = origin_color, alpha = origin_alpha)  

#   Pure maximal learning styles
    for i in range(0,K):  
      y0 = projectLS[i, index0]
      y1 = projectLS[i, index1]
      axs[panels[ils][0],panels[ils][1]].scatter(y0, y1,   s = size_mean, marker = twenty_symbol[i], color = PC_color[i], alpha = twenty_alpha)  
      axs[panels[ils][0],panels[ils][1]].plot([orig0, y0], [orig1, y1], color = twenty_color[i], ls = twenty_line[i], lw = twenty_width_2D[i])

    i = K + 1
    y0 = projectLS[i, index0]
    y1 = projectLS[i, index1]
    axs[panels[ils][0],panels[ils][1]].scatter(y0, y1,   s = size_mean, marker = twenty_symbol[i], color = twenty_color[i], alpha = twenty_alpha)  
    axs[panels[ils][0],panels[ils][1]].plot([orig0, y0], [orig1, y1], color    = twenty_color[i], ls = twenty_line[i], lw = twenty_width_2D[i])

#   Average mean
    axs[panels[ils][0],panels[ils][1]].scatter(0, 0, s = size_mean, marker=mean_symbol, color = mean_color2)

#   Educators
    axs[panels[ils][0],panels[ils][1]].scatter(proj_prof[i_prof, 0], proj_prof[i_prof, 1], s = size_mean, marker = 'P', color = 'k', alpha = origin_alpha)

# Text label
    axs[panels[ils][0], panels[ils][1]].text(xmin + 0.1*(xmax-xmin), ymax - 0.1*(ymax-ymin), Label_LS[ils],
        fontsize = labelsize,
        horizontalalignment='left',
        verticalalignment='top')
  
  axs[panels[ils][0],panels[ils][1]].set_xlim([ xmin, xmax] )
  axs[panels[ils][0],panels[ils][1]].set_ylim([ ymin, ymax] )

  axs[1, 0].set_xlabel(Label_PCPC0, fontsize=labelsize)
  axs[1, 0].set_ylabel(Label_PCPC1, fontsize=labelsize)

  filename = output_participants_pc_prof + '/Fig_'+Label_PCPC0+'_'+Label_PCPC1+'_prof_' + profi + '.png'
  plt.savefig(filename)
  printt('   Saving ' + filename)
  plt.savefig(filename)

printt('  ')
printt(' Plots for the professor clusters done!')
printt('-------------------------------------------')
printt('  ')

printt('===========================================')
printt('2D PLOT OF THE PCs WITH THE SCATTER')
printt('SIZE AND COLORS A FUNCTION OF PROFESSORS LS')
printt('DONE!')
printt('===========================================')
printt('  ')
printt('  ')
printt('  ')
printt('  ')
#
#===================================================
#===================================================
# OPTIONAL ANALYSIS OF STUDENTS' MARKS
# (peformed only if their marks are provided)
#===================================================
#===================================================
print('K_value : ', K_value)

if ( K_value > 0 and Lmarks > 0 ):
 printt('===========================================')
 printt('ANALYSIS OF THE MARKS OF THE STUDENTS...')
 printt('===========================================')

 printt('-------------------------------------------')
 printt(' Plots of students marks vs original LS...')
 printt('-------------------------------------------')

#Clustering figure for the LS clustering
 fig_clustering, ax_clustering = plt.subplots(figsize=(w_fig, h_fig), layout='constrained')

 panels = [[0, 0], [0, 1], [1, 0], [1, 1]]

 printt('np.shape(data_marks) : ' + str(np.shape(data_marks)))

 printt('data_marks    : ' + str(data_marks))
 marks_vs_LS = []

 printt(' K_value          : ' + str(K_value) )

 if(K_value > len(cluster_color)):
   printt('Please, include more colors in the cluster_color array.')
   sys.exit()

 for i in range(1, Lmarks_activities):

   printt(' Evaluating activity : ' + data_marks[i, 2] + ' contained in the sheet ' + data_marks[i, 1] + ' of ' + data_marks[i, 0])
   printt('  ')

   fig, axs = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), sharex = True, sharey = True, layout='constrained')

   students = data_marks[0, 3:]
   marks    = np.array(data_marks[i, 3:], dtype=float)

#  Create two vectors with the marks and the LS
#  and a string with the names of the students
#  (if marks are correctly defined)
   vls       = []
   vstudents = []
   vmarks    = []
   for j in range(0, L):
    if(~np.isnan(marks[j])):
      vstudents.append(students[j])
      vls.append([data[j,0], data[j,1], data[j,2], data[j,3]])
      vmarks.append(marks[j])

   vstudents = np.stack(vstudents)
   vls    = np.stack(vls)
   vmarks = np.array(vmarks)
   vmarks = vmarks.astype(float)

#  Plot limits
   xmin = np.min( np.min(data) )
   xmax = np.max( np.max(data) )

   ymin = -0.5/1.05
   ymax = 10
   if( np.max(vmarks) > ymax):
    ymax = np.max(vmarks)

   if ( xmin < 0 ):
    xmin = xmin * 1.05
   else:
    xmin = xmin * 0.95

   if ( xmax < 0 ):
    xmax = xmax * 0.95
   else:
    xmax = xmax * 1.05

   if ( ymin < 0 ):
    ymin = ymin * 1.05
   else:
    ymin = ymin * 0.95

   if ( ymax < 0 ):
    ymax = ymax * 0.95
   else:
    ymax = ymax * 1.05

   printt( ' xmin : ' + str(xmin))
   printt( ' xmax : ' + str(xmax))
   printt( ' ymin : ' + str(ymin))
   printt( ' ymax : ' + str(ymax))
   printt('  ')

   if( K_value <= len(vmarks) ):
    for ils in range(0,K):
     print('panels : ', panels[ils][0], panels[ils][1])
     axs[panels[ils][0], panels[ils][1]].tick_params(axis='both', which='major', labelsize=ticksize)
#----------------------------------------------------------
#    Scatter of the marks as a function of the LS  points
     printt(' LS : ' + str(ils))
     printt('  ')

     Label_LS0 = Label_LS[ils]
     printt(' ( marks vs ' + Label_LS0 + ' )')
     printt('  ')

     vx0 = vls[:,ils]
     vx1 = vmarks

     printt('   Clustering for ' + Label_LS0 + '...')

     array=vx1.reshape(-1, 1)

#    Create and fit the K-Means model with the optimal number of clusters
     kmeans = KMeans(n_clusters=K_value, init='k-means++', random_state=42)
     kmeans.fit(array)

#    Get the cluster labels for each data point
     k_labels = kmeans.labels_

#    Average means and uncertainties for the clusters
     clusters_means  = [ ]
     clusters_uncert = [ ]
     clusters_index  = [ ]
     list_students = ''
     for k in range(0, K_value):
      values = []

      for kk in range(0,len(k_labels)):
        if(k == k_labels[kk]):
          if(list_students == ''):
            list_students = vstudents[kk]
          else:
            list_students = list_students + ', ' + vstudents[kk]
          values.append( [ vx0[kk], vx1[kk] ] )
      shape = np.shape(values)
      clusters_index.append(shape[0])
      values = np.stack(values)

#     Average means
      x0_mean = np.mean( values[:,0] )
      x1_mean = np.mean( values[:,1] )

#     Uncertainties
      confidence_interval=st.t.interval(confidence=0.95, df=len(values[:,0])-1, loc=x0_mean, scale=st.sem(values[:,0]))
      dx0 = 0.5*(confidence_interval[1]-confidence_interval[0])
      confidence_interval=st.t.interval(confidence=0.95, df=len(values[:,1])-1, loc=x0_mean, scale=st.sem(values[:,1]))
      dx1 = 0.5*(confidence_interval[1]-confidence_interval[0])
      clusters_means.append([ x0_mean, x1_mean ])
      clusters_uncert.append( [dx0, dx1] )

     clusters_means  = np.stack(clusters_means)
     clusters_uncert = np.stack(clusters_uncert)
     printt('   Clustering for ' + Label_LS0 + ' done!')

#    Linear regression
     if np.all(vx0 == vx0[0]):
       print("Error: All vx0 values are identical. Linear regression cannot be performed.")
       m = b = delta_m = delta_b = rvalue = pvalue = np.nan
     else:
       m, b, rvalue, pvalue, delta_m = linregress(vx0, vx1)
#      Compute uncertainty in intercept (y0)
       delta_b  = delta_m * np.sqrt( np.sum(vx0 * vx0) / len(vx0) )

#    Average means
     x0_mean = np.mean( vx0 )
     x1_mean = np.mean( vx1 )
#    Uncertainties
     confidence_interval=st.t.interval(confidence=0.95, df=len(vx0)-1, loc=x0_mean, scale=st.sem(vx0))
     dx0 = 0.5*(confidence_interval[1]-confidence_interval[0])
     confidence_interval=st.t.interval(confidence=0.95, df=len(vx1)-1, loc=x1_mean, scale=st.sem(vx1))
     dx1 = 0.5*(confidence_interval[1]-confidence_interval[0])

     marks_vs_LS.append([ data_marks[i, 0], data_marks[i, 1], data_marks[i, 2], len(vx0), list_students, -1, Label_LS0, x0_mean, dx0, x1_mean, dx1, m,  delta_m, b, delta_b, rvalue ] )

     def lin_regr2(m, b, x):
      return m*x+b

#    The confidence interval is well defined within the lines
#    ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
#    since x is always positive
     x_linregr  = [ xmin, xmax ]
     y_linregr  = [ lin_regr2(m, b, x_linregr[0]), lin_regr2(m, b, x_linregr[1]) ]
     y_linregr1 = [ lin_regr2(m-delta_m, b-delta_b, x_linregr[0]), lin_regr2(m-delta_m, b-delta_b, x_linregr[1]) ]
     y_linregr2 = [ lin_regr2(m+delta_m, b+delta_b, x_linregr[0]), lin_regr2(m+delta_m, b+delta_b, x_linregr[1]) ]

     axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr,  l, c='k')
     axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr1, l, ls='--', c='gray') # LS_color[ils])
     axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr2, l, ls='--', c='gray') # LS_color[ils])
     axs[panels[ils][0], panels[ils][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1[0], y_linregr1[1]], [y_linregr2[0], y_linregr2[1]], facecolor = 'gray', alpha = LS_alpha*0.25) # LS_color[ils]

#    Scatter points
     for k in range(0, len(vx0)):
      axs[panels[ils][0], panels[ils][1]].scatter(vx0[k], vx1[k], s = scatter_size2, marker=cluster_symbol[k_labels[k]], color = cluster_color[k_labels[k]], alpha = cluster_alpha)

     for k in range(0,K_value):
      vx0_K = []
      vx1_K = []
      list_students = ''
      for kk in range(0, len(vx0)):
        if(k == k_labels[kk]):
          if( list_students == ''):
            list_students = vstudents[kk]
          else:
            list_students = list_students + ', ' + vstudents[kk]
          vx0_K.append(vx0[kk])
          vx1_K.append(vx1[kk])

      vx0_K =  np.array( vx0_K )
      vx1_K =  np.array( vx1_K )
      printt('k          : ' + str(k))
      printt('list_students : ' + str(list_students))
      printt('vx0_K      : ' + str(vx0_K))
      printt('vx1_K      : ' + str(vx1_K))

#     Check if all vx0_K values are identical
      if np.all(vx0_K == vx0_K[0]):
        print("Error: All vx0_K values are identical. Linear regression cannot be performed.")
        m_K = b_K = delta_m_K = delta_b_K = rvalue_K = np.nan
      else:
        m_K, b_K, rvalue_K, pvalue_K, delta_m_K = linregress(vx0_K, vx1_K)
        delta_b_K  = delta_m_K * np.sqrt( np.sum(vx0_K * vx0_K) / len(vx0_K) )

      printt('    m_K            : ' + str(m_K))
      printt('    delta_m_K      : ' + str(delta_m_K))
      printt('    b_K            : ' + str(b_K))
      printt('    delta_b_K      : ' + str(delta_b_K))
      printt('    rvalue_K            : ' + str(rvalue_K))
      printt('              ')

      marks_vs_LS.append([ data_marks[i, 0], data_marks[i, 1], data_marks[i, 2], clusters_index[k], list_students, k, Label_LS0, clusters_means[k, 0], clusters_uncert[k, 0], clusters_means[k, 1], clusters_uncert[k, 1], m_K, delta_m_K, b_K, delta_b_K, rvalue_K ] )

#       The confidence interval is well defined within the lines
#       ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
#        since x is always positive
      y_linregr_K  = [ lin_regr2(m_K, b_K, x_linregr[0]), lin_regr2(m_K, b_K, x_linregr[1]) ]
      y_linregr1_K = [ lin_regr2(m_K-delta_m_K, b_K-delta_b_K, x_linregr[0]), lin_regr2(m_K-delta_m_K, b_K-delta_b_K,  x_linregr[1]) ]
      y_linregr2_K = [ lin_regr2(m_K+delta_m_K, b_K+delta_b_K, x_linregr[0]), lin_regr2(m_K+delta_m_K, b_K+delta_b_K,  x_linregr[1]) ]

      axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr_K,  l, c=cluster_color[k])
      axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr1_K, l, ls='--', c=cluster_color[k])
      axs[panels[ils][0], panels[ils][1]].plot(x_linregr, y_linregr2_K, l, ls='--', c=cluster_color[k])
      axs[panels[ils][0], panels[ils][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1_K[0], y_linregr1_K[1]],  [y_linregr2_K[0], y_linregr2_K[1]], facecolor = cluster_color[k], alpha = LS_alpha*0.25) # LS_color[ils]

#    Average mean
     axs[panels[ils][0], panels[ils][1]].scatter(x0_mean, x1_mean, s = size_mean, marker = mean_symbol, color = mean_color2)

     for k in range(0,K_value):
#     Average means for the clusters
      axs[panels[ils][0], panels[ils][1]].scatter(clusters_means[k, 0], clusters_means[k, 1], s = size_mean, marker=mean_symbol, color = cluster_color[k])

#    Text label
     axs[panels[ils][0], panels[ils][1]].text(xmin + 0.1*(xmax-xmin), ymax - 0.1*(ymax-ymin), Label_LS0,
        fontsize = labelsize,
        horizontalalignment='left',
        verticalalignment='top')

    axs[1, 0].set_xlim([ xmin, xmax] )
    axs[1, 0].set_ylim([ ymin, ymax] )

    axs[1, 0].set_xlabel('Points', fontsize=labelsize)
    axs[1, 0].set_ylabel('Marks', fontsize=labelsize)

    filename = output_participants_ls_marks + '/Fig_'+ data_marks[i, 0] + '_' + data_marks[i, 1] + '_' + data_marks[i, 2] + '_LS_K' +str(K_value) + '.png'
    plt.savefig(filename)
    printt('   Saving ' + filename)
    plt.savefig(filename)

 printt('   LS clustering done!')
 marks_vs_LS = np.stack(marks_vs_LS)
 print('marks_vs_LS : ', marks_vs_LS)

 printt('  ')
 printt('-------------------------------------------')
 printt(' Plots of students marks vs original LS done!')
 printt('-------------------------------------------')


 printt('-------------------------------------------')
 printt(' Plots of students marks vs PCs...')
 printt('-------------------------------------------')

 marks_vs_PC = []

 for i in range(1, Lmarks_activities):

   printt(' Evaluating activity : ' + data_marks[i, 2] + ' contained in the sheet ' + data_marks[i, 1] + ' of ' + data_marks[i, 0])
   printt('  ')

   fig, axs = plt.subplots(2, 2, figsize = ( w_fig, h_fig ), sharex = True, sharey = True, layout='constrained')

   students = data_marks[0, 3:]
   marks    = np.array(data_marks[i, 3:], dtype=float)

#  Create two vectors with the marks and the LS
#  (if marks are correctly defined)
   vpc   = []
   vmarks = []
   for j in range(0, L):
     if(~np.isnan(marks[j])):
       vpc.append([proj[j,0], proj[j,1], proj[j,2], proj[j,3]])
       vmarks.append(marks[j])

   vpc    = np.stack(vpc)
   vmarks = np.array(vmarks)
   vmarks = vmarks.astype(float)

#  Plot limits
   xmin = np.min( np.min(proj) )
   xmax = np.max( np.max(proj) )

   ymin = -0.5/1.05
   ymax = 10
   if( np.max(vmarks) > ymax):
    ymax = np.max(vmarks)

   if ( xmin < 0 ):
    xmin = xmin * 1.05
   else:
    xmin = xmin * 0.95

   if ( xmax < 0 ):
    xmax = xmax * 0.95
   else:
    xmax = xmax * 1.05

   if ( ymin < 0 ):
    ymin = ymin * 1.05
   else:
    ymin = ymin * 0.95

   if ( ymax < 0 ):
    ymax = ymax * 0.95
   else:
    ymax = ymax * 1.05

   printt( ' xmin : ' + str(xmin))
   printt( ' xmax : ' + str(xmax))
   printt( ' ymin : ' + str(ymin))
   printt( ' ymax : ' + str(ymax))
   printt('  ')

   if( K_value <= len(vmarks) ):
    for ipc in range(0,K):
     print('panels : ', panels[ipc][0], panels[ipc][1])
     axs[panels[ipc][0], panels[ipc][1]].tick_params(axis='both', which='major', labelsize=ticksize)
#----------------------------------------------------------
#    Scatter of the marks as a function of the LS  points
     printt(' PC : ' + str(ipc))
     printt('  ')

     Label_PC0 = Label_PCPC[ipc]
     printt(' ( marks vs ' + Label_PC0 + ' )')
     printt('  ')

     vx0 = vpc[:,ipc]
     vx1 = vmarks

     printt('   Clustering for ' + Label_PC0 + '...')

     array=vx1.reshape(-1, 1)

#    Create and fit the K-Means model with the optimal number of clusters
     kmeans = KMeans(n_clusters=K_value, init='k-means++', random_state=42)
     kmeans.fit(array)

#    Get the cluster labels for each data point
     k_labels = kmeans.labels_

#    Average means and uncertainties for the clusters
     clusters_means  = [ ]
     clusters_uncert = [ ]
     clusters_index  = [ ]
     for k in range(0, K_value):
      values = []

      for kk in range(0,len(k_labels)):
        if(k == k_labels[kk]):
          values.append( [ vx0[kk], vx1[kk] ] )
      shape = np.shape(values)
      clusters_index.append(shape[0])
      values = np.stack(values)

#     Average means
      x0_mean = np.mean( values[:,0] )
      x1_mean = np.mean( values[:,1] )

#     Uncertainties
      confidence_interval=st.t.interval(confidence=0.95, df=len(values[:,0])-1, loc=x0_mean, scale=st.sem(values[:,0]))
      dx0 = 0.5*(confidence_interval[1]-confidence_interval[0])
      confidence_interval=st.t.interval(confidence=0.95, df=len(values[:,1])-1, loc=x0_mean, scale=st.sem(values[:,1]))
      dx1 = 0.5*(confidence_interval[1]-confidence_interval[0])
      clusters_means.append([ x0_mean, x1_mean ])
      clusters_uncert.append( [dx0, dx1] )

     clusters_means  = np.stack(clusters_means)
     clusters_uncert = np.stack(clusters_uncert)

#    Linear regression
     if np.all(vx0 == vx0[0]):
      print("Error: All vx0 values are identical. Linear regression cannot be performed.")
      m = b = delta_m = delta_b = rvalue = pvalue = np.nan
     else:
      m, b, rvalselected_kue, pvalue, delta_m = linregress(vx0, vx1)
#     Compute uncertainty in intercept (y0)
      delta_b  = delta_m * np.sqrt( np.sum(vx0 * vx0) / len(vx0) )

#    Average means
     x0_mean = np.mean( vx0 )
     x1_mean = np.mean( vx1 )

#    Uncertainties
     confidence_interval=st.t.interval(confidence=0.95, df=len(vx0)-1, loc=x0_mean, scale=st.sem(vx0))
     dx0 = 0.5*(confidence_interval[1]-confidence_interval[0])
     confidence_interval=st.t.interval(confidence=0.95, df=len(vx1)-1, loc=x1_mean, scale=st.sem(vx1))
     dx1 = 0.5*(confidence_interval[1]-confidence_interval[0])

     marks_vs_PC.append([ data_marks[i, 0], data_marks[i, 1], data_marks[i, 2], len(vx0), '', -1, Label_PC0,  x0_mean, dx0, x1_mean, dx1, m, delta_m, b, delta_b, rvalue ] )

     def lin_regr2(m, b, x):
      return m*x+b

#   The confidence interval for x > 0
#   ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
     x_linregr  = [ 0, xmax ]
     y_linregr  = [ lin_regr2(m, b, x_linregr[0]), lin_regr2(m, b, x_linregr[1]) ]
     y_linregr1 = [ lin_regr2(m-delta_m, b-delta_b, x_linregr[0]), lin_regr2(m-delta_m, b-delta_b, x_linregr[1]) ]
     y_linregr2 = [ lin_regr2(m+delta_m, b+delta_b, x_linregr[0]), lin_regr2(m+delta_m, b+delta_b, x_linregr[1]) ]

     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr,  l, c='k')
     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr1, l, ls='--', c='gray')
     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr2, l, ls='--', c='gray')
     axs[panels[ipc][0], panels[ipc][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1[0], y_linregr1[1]], [y_linregr2[0], y_linregr2[1]], facecolor = 'gray', alpha = LS_alpha*0.25)

#   The confidence interval for x < 0
#   ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
     x_linregr  = [ xmin, 0 ]
     y_linregr  = [ lin_regr2(m, b, x_linregr[0]), lin_regr2(m, b, x_linregr[1]) ]
     y_linregr1 = [ lin_regr2(m-delta_m, b+delta_b, x_linregr[0]), lin_regr2(m-delta_m, b+delta_b, x_linregr[1]) ]
     y_linregr2 = [ lin_regr2(m+delta_m, b-delta_b, x_linregr[0]), lin_regr2(m+delta_m, b-delta_b, x_linregr[1]) ]

     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr,  l, c='k')
     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr1, l, ls='--', c='gray')
     axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr2, l, ls='--', c='gray')
     axs[panels[ipc][0], panels[ipc][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1[0], y_linregr1[1]], [y_linregr2[0], y_linregr2[1]], facecolor = 'gray', alpha = LS_alpha*0.25)

#    Scatter points
     for k in range(0, len(vx0)):
      axs[panels[ipc][0], panels[ipc][1]].scatter(vx0[k], vx1[k], s = scatter_size2, marker=cluster_symbol[k_labels[k]], color = cluster_color[k_labels[k]], alpha = cluster_alpha)

     for k in range(0,K_value):
      vx0_K = []
      vx1_K = []
      for kk in range(0, len(vx0)):
        if(k == k_labels[kk]):
          vx0_K.append(vx0[kk])
          vx1_K.append(vx1[kk])

      vx0_K =  np.array( vx0_K )
      vx1_K =  np.array( vx1_K )

      printt('k     : ' + str(k))
      printt('vx0_K : ' + str(vx0_K))
      printt('vx1_K : ' + str(vx1_K))

      if np.all(vx0_K == vx0_K[0]):
        print("Error: All vx0_K values are identical. Linear regression cannot be performed.")
        m_L = b_K = delta_m_K = delta_b_K = delta_b_K = rvalue_K = np.nan
      else:
        m_K, b_K, rvalue_K, pvalue_K, delta_m_K = linregress(vx0_K, vx1_K)
        delta_b_K  = delta_m_K * np.sqrt( np.sum(vx0_K * vx0_K) / len(vx0_K) )

      printt('    m_K            : ' + str(m_K))
      printt('    delta_m_K      : ' + str(delta_m_K))
      printt('    b_K            : ' + str(b_K))
      printt('    delta_b_K      : ' + str(delta_b_K))
      printt('    rvalue_K            : ' + str(rvalue_K))
      printt('              ')

      marks_vs_PC.append([ data_marks[i, 0], data_marks[i, 1], data_marks[i, 2], clusters_index[k], '', k, Label_PC0, clusters_means[k, 0], clusters_uncert[k, 0], clusters_means[k, 1], clusters_uncert[k, 1], m_K, delta_m_K, b_K, delta_b_K, rvalue_K ] )

#     The confidence interval for x > 0
#     ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
      x_linregr  = [ 0, xmax ]
      y_linregr_K  = [ lin_regr2(m_K, b_K, x_linregr[0]), lin_regr2(m_K, b_K, x_linregr[1]) ]
      y_linregr1_K = [ lin_regr2(m_K-delta_m_K, b_K-delta_b_K, x_linregr[0]), lin_regr2(m_K-delta_m_K, b_K-delta_b_K, x_linregr[1]) ]
      y_linregr2_K = [ lin_regr2(m_K+delta_m_K, b_K+delta_b_K, x_linregr[0]), lin_regr2(m_K+delta_m_K, b_K+delta_b_K, x_linregr[1]) ]

      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr_K,  l, c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr1_K, l, ls='--', c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr2_K, l, ls='--', c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1_K[0], y_linregr1_K[1]], [y_linregr2_K[0], y_linregr2_K[1]], facecolor = cluster_color[k], alpha = LS_alpha*0.25)

#     The confidence interval for x < 0
#     ( m_K - delta_m_K )*x+ ( b_K - delta_b_K) and ( m_K + delta_m_K )*x+ ( b_K + delta_b_K)
      x_linregr  = [ xmin, 0 ]
      y_linregr_K  = [ lin_regr2(m_K, b_K, x_linregr[0]), lin_regr2(m_K, b_K, x_linregr[1]) ]
      y_linregr1_K = [ lin_regr2(m_K-delta_m_K, b_K-delta_b_K, x_linregr[0]), lin_regr2(m_K-delta_m_K, b_K+delta_b_K, x_linregr[1]) ]
      y_linregr2_K = [ lin_regr2(m_K+delta_m_K, b_K+delta_b_K, x_linregr[0]), lin_regr2(m_K+delta_m_K, b_K-delta_b_K, x_linregr[1]) ]

      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr_K,  l, c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr1_K, l, ls='--', c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].plot(x_linregr, y_linregr2_K, l, ls='--', c=cluster_color[k])
      axs[panels[ipc][0], panels[ipc][1]].fill_between([x_linregr[0], x_linregr[1]], [y_linregr1_K[0], y_linregr1_K[1]], [y_linregr2_K[0], y_linregr2_K[1]], facecolor = cluster_color[k], alpha = LS_alpha*0.25)

#    Average mean
     axs[panels[ipc][0], panels[ipc][1]].scatter(x0_mean, x1_mean, s = size_mean, marker = mean_symbol, color = mean_color2)

     for k in range(0,K_value):
#     Average means for the clusters
      axs[panels[ipc][0], panels[ipc][1]].scatter(clusters_means[k, 0], clusters_means[k, 1], s = size_mean, marker=mean_symbol, color = cluster_color[k])

#    Text label
     axs[panels[ipc][0], panels[ipc][1]].text(xmin + 0.1*(xmax-xmin), ymax - 0.1*(ymax-ymin), Label_PC0,
        fontsize = labelsize,
        horizontalalignment='left',
        verticalalignment='top')

    axs[1, 0].set_xlim([ xmin, xmax] )
    axs[1, 0].set_ylim([ ymin, ymax] )

    axs[1, 0].set_xlabel('Points', fontsize=labelsize)
    axs[1, 0].set_ylabel('Marks', fontsize=labelsize)

    filename = output_participants_pc_marks + '/Fig_'+ data_marks[i, 0] + '_' + data_marks[i, 1] + '_' + data_marks[i, 2] + '_PC_K' +str(K_value) + '.png'
    plt.savefig(filename)
    printt('   Saving ' + filename)
    plt.savefig(filename)
    printt('  ')

 marks_vs_PC = np.stack(marks_vs_PC)

 printt('-------------------------------------------')
 printt(' Plots of students marks vs PCs done!')
 printt('-------------------------------------------')
 printt('  ')
 printt('===========================================')
 printt('ANALYSIS OF THE MARKS OF THE STUDENTS DONE!')
 printt('===========================================')
 printt('  ')
 printt('  ')
 printt('  ')
 printt('  ')
#
#
# Get the current date and time
current_datetime = datetime.datetime.now()
#
# Format the date and time as a string
formatted_datetime = current_datetime.strftime("%Y-%m-%d %H:%M:%S")
printt('Program continues on: ' + formatted_datetime)
printt(' ')
#
#===================================================
#===================================================
#===================================================
#===================================================
# SUMMARY REPORT
#===================================================
#===================================================
# Most of the results of the analysis are
# included in the summary report, which is saved in
# docx and pdf formats.
# Further information can be found in the
# log file (output/chaea3s.log).
#
# If the number of students is larger than
# Lmax = 500, then only the first 150 are saved.
#
if ( L > 150):
  Lsave = 150
else:
  Lsave = L

#===================================================
# File location and title
printt('===========================================')
printt('SUMMARY REPORT...')
printt('===========================================')
printt('output : ' + output)
folder_name = os.path.basename(output)
document    = Document()
#
#===================================================
# TITLE AND INTRO
#===================================================
# Title
title = 'Summary report of CHAEA learning styles by CHAEA'
title += u'\u00B3'  # Adding superscript 3
title += 'S package'

document.add_heading(title, level=0)

document.add_paragraph('This report contains the most important results of the analysis that is conducted to unveil the learning styles that are present in the group of students under study. The analysis is based on the learning styles considered by CHAEA: activist, reflector, theorist, and pragmatist. Unless otherwise stated, the uncertainties throughout the document (in parenthesis) have been obtained using a t-Student distribution with a confidence interval of 95%.')

if ( K_value == 0 ):
# No marks provided, i.e., no clustering conducted
  document.add_paragraph('The report is structured as follows. First, Section 1 is devoted to the individual and global statistical analysis. Here, the importance of the different learning styles for each individual student can be found (both quantitatively as well as qualitatively). Then, average means and confidence intervals are presented, along with the affinities and the Probability Density Functions. Second, Section 2 discusses the principal component analysis. The eigenvalues and eigenvectors of the covariance matrix are first introduced. Subsequently, the learning styles of the students in the principal components basis set is presented. Next, a reduced dimensional representation of the data is conducted. Third, the participation ratios are presented in Section 3. The values for the original (activist, reflector, theorist, and pragmatist) and in the principal components basis sets are listed. Here, a statistical analysis of the distribution of the participation ratios is performed. The report concludes in Section 4 with a brief summary of the main characteristics of the students, along with some recommendations.')
else:
  if ( K_value == 1 ):
# Marks provided, but no clustering conducted as K=1
    document.add_paragraph('The report is structured as follows. First, Section 1 is devoted to the individual and global statistical analysis. Here, the importance of the different learning styles for each individual student can be found (both quantitatively as well as qualitatively). Then, average means and confidence intervals are presented, along with the affinities and the Probability Density Functions. Second, Section 2 discusses the principal component analysis. The eigenvalues and eigenvectors of the covariance matrix are first introduced. Subsequently, the learning styles of the students in the principal components basis set is presented. Next, a reduced dimensional representation of the data is conducted. Third, the participation ratios are presented in Section 3. The values for the original (activist, reflector, theorist, and pragmatist) and in the principal components basis sets are listed. Here, a statistical analysis of the distribution of the participation ratios is performed. Section 4 is devoted to the analysis of marks of the whole group of students, and their relationship to the CHAEA learning styles and the principal components. The report concludes in Section 5 with a brief summary of the main characteristics of the students, along with some recommendations.')
  else:
#   Marks provided, and no clustering conducted as K>1
    document.add_paragraph('The report is structured as follows. First, Section 1 is devoted to the individual and global statistical analysis. Here, the importance of the different learning styles for each individual student can be found (both quantitatively as well as qualitatively). Then, average means and confidence intervals are presented, along with the affinities and the Probability Density Functions. Second, Section 2 discusses the principal component analysis. The eigenvalues and eigenvectors of the covariance matrix are first introduced. Subsequently, the learning styles of the students in the principal components basis set is presented. Next, a reduced dimensional representation of the data is conducted. Third, the participation ratios are presented in Section 3. The values for the original (activist, reflector, theorist, and pragmatist) and in the principal components basis sets are listed. Here, a statistical analysis of the distribution of the participation ratios is performed. Section 4 presents the clustering analysis of the data based on K-means algorithm with K=' + str(K_value)+ ' clusters. In this section, the relationship between the students marks, and the CHAEA learning styles and the principal components The report concludes in Section 5 with a brief summary of the main characteristics of the students, along with some recommendations.')

# Further information in our reference
our_reference = document.add_paragraph()

our_reference.add_run("Further information at: ").italic = True

document.add_paragraph(' ')
our_reference.add_run("J. Ablanque, V. Gabaldon, P. Almendros, J. C. Losada, R. M. Benito, and F. Revuelta. ").italic = True
document.add_paragraph(' ')
#
#
our_reference.add_run("CHAEA3S: A software package for comprehensive analysis of learning styles and academic performance. ").italic = True
document.add_paragraph()

our_reference.add_run("Computers and Education (2025).").italic = True

#our_reference.add_run("If you have any comments or suggestions, feel free to reach out by sending an email with your feedback to fabio.revuelta@upm.es.").italic = True

document.add_page_break()
#===================================================
# SECTION 1
#===================================================
nsec    = 1
nsubsec = 1
#==================================================
document.add_heading(str(nsec) + '. Individual and global statistical analysis', level=1)
printt('******************************************')
printt(str(nsec) + '. Individual and global statistical analysis')
printt('******************************************')
#==================================================
document.add_paragraph('In this section we present the individual and global statistical analysis of the learning styles as originally defined in CHAEA (activist, reflector, theorist, and pragmatist). This section is divided in two parts. First, the (quantitative and qualitative) importance of the different learning styles for each individual student can be found. Second, a global analysis is performed, where average means, confidence intervals, affinities, and the probability density functions can be found.')
#
#---------------------------------------------------
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Individual analysis', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Individual analysis')
ntab = 1
ntab = 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the number of points...')
printt('-------------------------------------------')
document.add_heading('Quantitative description of the learning styles for each individual student', level=3)

if ( L <= Lsave ):
  document.add_paragraph('Table ' + str(ntab) + ' shows the number of points (from 0 to 20) that each student gets in CHAEA for the different learning styles.')
else:
  document.add_paragraph('Table ' + str(ntab) + ' shows the number of points (from 0 to 20) that each student gets in CHAEA for the different learning styles (only the first ' + str(Lsave) + ' students are included).')

# Table 1 title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Points related the learning styles for each of the students.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

# Table 1
t = document.add_table(Lsave+1, K+1)

# Table header
t.cell(0,0).text = 'Student'
for j in range(0,K):
  t.cell(0,j+1).text = Label_LS_print[j]

for i in range(0,Lsave):
  t.cell(i+1,0).text = students[i]
  for j in range(0,K):
    t.cell(i+1,j+1).text = str(data[i,j])

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the number of points done!')
printt('-------------------------------------------')
printt(' ')

#Lerror = len(datain_error_students)
if ( datain_error_students != ['.'] ):
  document.add_paragraph()
  document.add_paragraph('The following students were not considered in the analysis because of errors in the CHAEA files:')
  document.add_paragraph(datain_error_students)

document.add_page_break()

ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the individual tendencies...')
printt('-------------------------------------------')
document.add_heading('Qualitative description of the learning styles for each individual student', level=3)

if ( L <= Lsave ):
  document.add_paragraph('Table ' + str(ntab) + ' shows the individual qualitative tendency (very low, low, moderate, high, or very high) that each student has towards each of the learning styles.')
else:
  document.add_paragraph('Table ' + str(ntab) + ' shows the individual qualitative tendency (very low, low, moderate, high, or very high) that each student has towards each of the learning styles (only the first ' + str(Lsave) + ' students are included).')



# Table 2 title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Tendency towards the learning styles for each of the students.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(Lsave+1, K+1)

# Table header
t.cell(0,0).text = 'Student'
for j in range(0,K):
  t.cell(0,j+1).text = Label_LS_print[j]

for i in range(0,Lsave):
  t.cell(i+1,0).text = students[i]
  for j in range(0,K):
    t.cell(i+1,j+1).text = tendency_long_name(tendency_matrix_all[i][j])

document.add_page_break()
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the individual tendencies done!')
printt('-------------------------------------------')
printt(' ')







nsubsec += 1
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Global analysis', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Global analysis')
ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the tendencies...')
printt('-------------------------------------------')
document.add_heading('Global tendencies towards the learning styles of the students (I)', level=3)

document.add_paragraph('Table ' + str(ntab) + ' presents the number and percentage of students with the same tendency (very low, low, moderate, high, or very high) towards each of the learning styles. The bottom line shows the average tendency. These results are also shown as a barr graphic in Fig. 1.')

# Table 4 title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Number and percentage of students with the same tendency towards each of the learning styles.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+3, 6)

t.cell(0,0).text = 'Tendency'
for j in range(0,5):
  t.cell(0,j+1).text = Label_tendencies_print[j]
  t.cell(1,j+1).text = 'No.  %'

for i in range(0,K):
  t.cell(i+2,0).text = Label_LS_print[i]
  for j in range(0,5):
    t.cell(i+2,j+1).text = str(int(tendency_matrix[i, j])) + '   ' + str( round(tendency_matrix_percentage[i, j], 1))

t.cell(K+2,0).text = 'Total'
for j in range(0,5):
  t.cell(K+2,j+1).text = str(int(Ntot_tendency[j])) + '   ' + str(round(percent_tendency[j], 1))

document.add_page_break()
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the tendencies done!')
printt('-------------------------------------------')
printt(' ')

nfig = 1
# Figure 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the tendencies...')
printt('-------------------------------------------')
document.add_picture(filename_tendencies, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Bar graphic showing the percentage of students with a very low (leftmost), low (left), moderate (middle), high (right) and very high (rightmost) tendency towards the activist (red), reflector (green), theorist (purple) and pragmatist (orange) learning styles. The shown results correspond to the values of Table 3.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

document.add_page_break()

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the tendencies done!')
printt('-------------------------------------------')



ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the global tendencies...')
printt('-------------------------------------------')
document.add_heading('Global tendencies towards the learning styles of the students (II)', level=3)

document.add_paragraph('Table ' + str(ntab) + ' shows the average mean and the uncertainties for the different learning styles, along with the corresponding qualitative tendency, and the corresponding affinity. For the shake of clarity, the average profile of the students given by the average means and the corresponding uncertainties is shown in Fig. 2, while the affinities are represented as a barr graphic in Fig. 3.')
#
#The bottom line shows the average tendency.
#
# Table 4 title
table_title = document.add_paragraph('Table 4. Numerical average values of each learning style with the uncertainties (in parenthesis) along with their qualitative tendency, and the corresponding affinity.')
# The last line provides the average results for the whole data set (no tendency is provided in this case).')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+1, 4)

# Table header
t.cell(0,0).text = 'Learning style'
t.cell(0,1).text = 'Average mean (Uncertainty)' #'μ (Δμ)'
t.cell(0,2).text = 'Tendency'
t.cell(0,3).text = 'Affinity (%)'

for j in range(0,K):
  t.cell(j+1,0).text = Label_LS_print[j]
  xdx = mean_uncert(xmean[j], dxmean[j])
  t.cell(j+1,1).text = xdx
  tendency = tendency_long_name(tendency_intermediate(Label_LS[j], xmean[j]))
  t.cell(j+1,2).text = tendency
  t.cell(j+1,3).text = str(round(affinity[j], 1))

document.add_page_break()
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the global tendencies done!')
printt('-------------------------------------------')
printt(' ')

# Figure 2
nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the average profile...')
printt('-------------------------------------------')
document.add_picture(filename_average_profile, width=Inches(5))
x
# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Average profile of the learning styles of the students.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the average profile done!')
printt('-------------------------------------------')


# Figure 3
nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the affinities...')
printt('-------------------------------------------')

document.add_picture(filename_affinity, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure 3. Affinity of the learning styles. The bar graphic gives the percentage of students that have a noticeable tendency towards the activist (red), reflector (green), theorist (purple) and pragmatist (orange) learning styles.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the affinities done!')
printt('-------------------------------------------')





nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the PDFs of the LSs...')
printt('-------------------------------------------')

document.add_paragraph('Figure ' + str(nfig) + ' shows histograms with the probability distributions of the results of Table 1. The continuous lines show fittings with Weibull distributions with the parameters shown in Table 5. These fittings have been performed on the corresponding cumulative distributions given by the staircases shown in Figure 5.')

document.add_picture(filename_statistics_ls, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure 4. Histograms showing the probability distributions for the activist (red), reflector (green), theorist (purple), and pragmatist (yellow) learning styles. The continuous lines show fittings provided by the Weibull distributions with the parameters contained in Table 5. The vertical dotted lines mark the average values, and the shaded areas around them the corresponding confidence intervals.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the PDFs of the LSs done!')
printt('-------------------------------------------')

nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with W(s) of the LSs...')
printt('-------------------------------------------')

document.add_picture(filename_statistics_ls_w, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure 5. Same as Fig. 4 for the cumulative distributions (staircases).')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

document.add_page_break()
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with W(s) of the LSs done!')
printt('-------------------------------------------')


ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the Weibull parameters for the PDFs of the LSs...')
printt('-------------------------------------------')

# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Parameters of the Weibull distributions that fit the probability distributions (histograms) shown in Figs. 4 and 5. The location parameter is set equal to θ = 0.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+1, 3)

# Table header
t.cell(0,0).text = 'Learning style'
t.cell(0,1).text = 'α'
t.cell(0,2).text = 'k'

for j in range(0,K):
  t.cell(j+1,0).text = Label_LS_print[j]
  [alpha_Weibull, k_Weibull] = parameters_Weibull_LS20[j]
  t.cell(j+1,1).text = str(round(alpha_Weibull,2))
  t.cell(j+1,2).text = str(round(k_Weibull,2))

document.add_page_break()
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the Weibull parameters for the PDFs of the LSs done!')
printt('-------------------------------------------')





#---------------------------------------------------
# OPTIONAL SUBSECTION WITH THE PROFESSOR'S
# LEARNING STYLES

if(Lprof == 0):
  printt('No results for the ' + prof_ref + 's are provided.')
else:
  nsubsec += 1
  ntab += 1
  if(Lprof == 1):
    nsubsec += 1
    document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Summary of the learning styles of the ' + prof_ref, level=2)
    printt(str(nsec) + '.' + str(nsubsec) + ' Summary of the learning styles of the ' + prof_ref)
    document.add_paragraph('In this section we present the learning styles of the ' + prof_ref + ' as originally defined in CHAEA (activist, reflector, theorist, and pragmatist), and compare them to those of the students.')
  else:
    document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Summary of the learning styles of the ' + prof_ref + 's', level=2)
    printt(str(nsec) + '.' + str(nsubsec) + ' Summary of the learning styles of the ' + prof_ref + 's')
    document.add_paragraph('In this section we present the learning styles of the ' + prof_ref + 's as originally defined in CHAEA (activist, reflector, theorist, and pragmatist), and compare them to those of the students.')

  printt('-------------------------------------------')
  printt(' Table ' + str(ntab) + ' with the number of points of the ' + prof_ref + '(s)...')
  printt('-------------------------------------------')
  if(Lprof == 1):
    document.add_heading('Quantitative description of the learning styles for the ' + prof_ref, level=3)
    document.add_paragraph('Table ' + str(ntab) + ' shows the number of points (from 0 to 20) that the ' + prof_ref + ' gets in CHAEA for the different learning styles.')
  else:
    document.add_heading('Quantitative description of the learning styles for each individual ' + prof_ref, level=3)
    document.add_paragraph('Table ' + str(ntab) + ' shows the number of points (from 0 to 20) that each ' + prof_ref + ' gets in CHAEA for the different learning styles.')

  # Table 6 title
  if(Lprof == 1):
    table_title = document.add_paragraph('Table ' + str(ntab) + '. Points related the learning styles for the ' + prof_ref + '.')
  else:
    table_title = document.add_paragraph('Table ' + str(ntab) + '. Points related the learning styles for each of the ' + prof_ref + 's.')
  table_title.alignment = 1  # Center alignment
  title_run = table_title.runs[0]
  title_run.bold = True

  # Table 6
  t = document.add_table(Lprof+1, K+1)

  # Table header
  t.cell(0,0).text = Prof_ref
  for j in range(0,K):
    t.cell(0,j+1).text = Label_LS_print[j]

  for i in range(0,Lprof):
    t.cell(i+1,0).text = profs[i]
    for j in range(0,K):
      t.cell(i+1,j+1).text = str(data_prof[i,j])

  if ( datain_error_profs != ['.'] ):
    document.add_paragraph()
    document.add_paragraph('The following ' + prof_ref + '(s) were not considered in the analysis because of errors in the CHAEA files:')
    document.add_paragraph(datain_error_profs)
  printt('-------------------------------------------')
  printt(' Table ' + str(ntab) + ' with the number of points of the ' + prof_ref + '(s) done!')
  printt('-------------------------------------------')
  printt(' ')

  ntab += 1
  printt('-------------------------------------------')
  printt(' Table ' + str(ntab) + ' with the individual tendencies of the ' + prof_ref + '(s)...')
  printt('-------------------------------------------')

  if(Lprof == 1):
    document.add_heading('Qualitative description of the learning styles for the ' + prof_ref, level=3)
    document.add_paragraph('Table ' + str(ntab) + ' shows the qualitative tendency (very low, low, moderate, high, or very high) that the ' + prof_ref + ' has towards each of the learning styles.')
  else:
    document.add_heading('Qualitative description of the learning styles for each individual ' + prof_ref, level=3)
    document.add_paragraph('Table ' + str(ntab) + ' shows the individual qualitative tendency (very low, low, moderate, high, or very high) that each ' + professor + ' has towards each of the learning styles.')

  # Table 7 title
  table_title = document.add_paragraph('Table ' + str(ntab) + '. Tendency towards the learning styles for each of the ' + prof_ref + 's.')
  table_title.alignment = 1  # Center alignment
  title_run = table_title.runs[0]
  title_run.bold = True

  t = document.add_table(Lprof+1, K+1)
  tendency_prof = ['', '', '', '', '']

  # Table header
  t.cell(0,0).text = Prof_ref
  for j in range(0,K):
    t.cell(0,j+1).text = Label_LS_print[j]

  for i in range(0,Lprof):
    t.cell(i+1,0).text = profs[i]

    tendency_prof[0] = scatter_tendency('Activist',   data_prof[i, 0])
    tendency_prof[1] = scatter_tendency('Reflector',  data_prof[i, 1])
    tendency_prof[2] = scatter_tendency('Theorist',   data_prof[i, 2])
    tendency_prof[3] = scatter_tendency('Pragmatist', data_prof[i, 3])

    for j in range(0,K):
      t.cell(i+1,j+1).text = tendency_long_name(tendency_prof[j])

  document.add_page_break()
  printt('-------------------------------------------')
  printt(' Table ' + str(ntab) + ' with the individual tendencies of the ' + prof_ref + '(s) done!')
  printt('-------------------------------------------')
  printt(' ')


  printt('-------------------------------------------')
  printt(' Figures with the average profile of the students + the profile of the ' + prof_ref + 's ...')
  printt('-------------------------------------------')
  nfig_prof0 = nfig + 1
  if(Lprof == 1):
    document.add_heading('Comparison of the learning-styles profile of the students and of the ' + prof_ref, level=3)
    document.add_paragraph('Figure ' + str(nfig+1) + ' shows the learning style profiles of the ' + prof_ref + ' superimposed with the average profile of the students shown in Fig. 2.')
  else:
    document.add_heading('Comparison of the learning-styles profile of the students and of each individual ' + prof_ref, level=3)
    document.add_paragraph('Figures ' + str(nfig+1) + ' to ' + str(nfig + Lprof) +' show the learning style profiles of the ' + prof_ref + 's superimposed with the average profile of the students shown in Fig. 2.')

  for i_prof in range(0,Lprof):
    profi = profs[i_prof]
    nfig  = nfig + 1
    filename_average_profile_prof=output_average_profile_prof + '/Fig_averageprofile_prof_' + profi +'.png'
    document.add_picture(filename_average_profile_prof, width=Inches(5))

#   Figure title
    if(i_prof == 0):
      figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Profile of the learning styles of the ' + prof_ref + ' "' + profi + '" (green continuous line) superimposed with the average profile of the students shown in Fig. 2.')
    else:
      figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Same as Fig. ' + str(nfig_prof0) + ' for the ' + prof_ref + ' of ' + profi + '.')
    figure_title.alignment = 1  # Center alignment
    figure_run = figure_title.runs[0]
    figure_run.bold = True

  printt('-------------------------------------------')
  printt(' Figure with the average profile of the students + the profile of the ' + professor + 's done!')
  printt('-------------------------------------------')










document.add_page_break()
#===================================================
# SECTION 2
#==================================================
nsec   += 1
nsubsec = 1
document.add_heading(str(nsec) + '. Principal component analysis', level=1)
printt('******************************************')
printt(str(nsec) + '. Principal component analysis')
printt('******************************************')
#==================================================

document.add_paragraph('This section presents the principal component analysis of the learning styles. It is organized as follows. First, the eigenvalues and eigenvectors of the covariance matrix are presented. Second, the description of the learning styles of the students in the principal components basis set is discussed. Finally, a reduced dimensional representation of the data is conducted.')

document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Structure of the principal components', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Structure of the principal components')

document.add_heading('Eigenvalues and dispersion of the covariance matrix', level=3)

ntab += 1
document.add_paragraph('In this section, we discuss the essentials of the principal components. For this purpose, Table ' + str(ntab) + ' shows the four eigenvalues of the covariance matrix. Here, not only their values are listed but also their contribution to the total dispersion (given by the trace of the covariance matrix tr(K)='+str(round(trace_covX,2))+', which equals the sum of all the eigenvalues) as a percentage. Likewise, the dispersion accounted by solely the principal component with the largest eigenvalue Σ_0(%)=λ_0*100/tr(K), and by combining the principal components with the two, three, and four largest eigenvalues Σ_1(%)=(λ_0+λ_1)*100/tr(K), Σ_2(%)=(λ_0+λ_1+λ_2)*100/tr(K), and by four Σ_3(%)=(λ_0+λ_1+λ_2+λ_3)*100/tr(K)=100%, respectively, are given.')

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the eigenvalues of the PCs...')
printt('-------------------------------------------')

# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Eigenvalues of the covariance matrix given by their corresponding values λ_i and as a percentage of the total dispersion, and percentage of total dispersion Σ_i accounted by combination of the principal components with the eigenvalues λ_j, being j≤i.')

table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+1, 4)

# Table header
t.cell(0,0).text = 'Principal component (i)'
t.cell(0,1).text = 'λ_i'
t.cell(0,2).text = 'λ_i(%)'
t.cell(0,3).text = 'Σ_i(%)'

sigma_percent = 0.0
for j in range(0,K):
  t.cell(j+1,0).text = Label_PC[j]
  lj            = eigenValues[j]
  lj_percent    = 100/trace_covX * lj
  sigma_percent = sigma_percent + lj_percent
  t.cell(j+1,1).text = str(round(lj,2))
  t.cell(j+1,2).text = str(round(lj_percent,1))
  t.cell(j+1,3).text = str(round(sigma_percent,1))
  if(j==2):
    sigma01 = str(round(sigma_percent,1))

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the eigenvalues of the PCs done!')
printt('-------------------------------------------')

document.add_page_break()
document.add_heading('Eigenvectors of the covariance matrix', level=3)
ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the structure of the eigenvectors...')
printt('-------------------------------------------')

document.add_paragraph('Table ' + str(ntab) + ' shows the structure of the eigenvectors of the covariance matrix in the basis set of CHAEA learning styles. Each cell contains the percentage of the eigenvector of the principal component 0, 1, 2, and 3 that is projected on the corresponding learning style (activist, reflector, theorist, and pragmatist).')

# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Structure (as percentages) of the eigenvectors of the covariance matrix in the basis set of CHAEA learning styles.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+1, 5)

# Table header
t.cell(0,0).text = 'Principal component'

for j in range(0,K):
  t.cell(0,j+1).text = Label_LS[j]

sigma_percent = 0.0
for i in range(0,K):
  t.cell(i+1,0).text = Label_PC[i]
  for j in range(0,K):
    t.cell(i+1,j+1).text = str(round(eigenVectors_percentage[j,i],1))

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the structure of the eigenvectors done!')
printt('-------------------------------------------')







document.add_page_break()
nsubsec += 1
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Individual analysis of the principal components', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Individual analysis of the principal components')
ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the individual analysis of PCs...')
printt('-------------------------------------------')
document.add_heading('Quantitative description of the students in the basis set of the principal components', level=3)

document.add_paragraph('Table ' + str(ntab) + ' shows (as percentages) the structure of the learning styles in the basis set formed by the principal components. The table also includes the percentage of the learning styles of each student that is described by combining the two (sum of the percentages for the principal components 0 and 1) or three (sum of the percentages for the principal components 0, 1, and 2) principal components with the largest eigenvalues. Recall that when the four principal components are considered, 100% of the learning style of the student is reproduced.')

# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Structure of the learning styles of each of the students in the basis set of principal components.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

# Table 1
t = document.add_table(Lsave+1, 7)

# Table header
t.cell(0,0).text = 'Student'
for j in range(0,K):
  t.cell(0,j+1).text = Label_PC[j]

#Σ
t.cell(0,5).text = '0+1'
t.cell(0,6).text = '0+1+2'

for i in range(0,Lsave):
  t.cell(i+1,0).text = students[i]
  for j in range(0,K):
    t.cell(i+1,j+1).text = str(round(prob[i,j], 2))

  t.cell(i+1,5).text = str(round(prob[i,0]+prob[i,1], 2))
  t.cell(i+1,6).text = str(round(prob[i,0]+prob[i,1]+prob[i,2], 2))

document.add_page_break()
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the individual analysis of PCs done!')
printt('-------------------------------------------')
printt(' ')


































nsubsec += 1
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Global analysis of the principal components', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Global analysis of the principal components')
nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the 4D plot of the LSs...')
printt('-------------------------------------------')
nfig_4d = nfig

document.add_paragraph('Figure ' + str(nfig) + ' shows a four-dimensional representation of the learning styles of the students as a function of the contributions to three of the CHAEA learning styles. The tendency towards the remaining learning style is implicitly shown in the size and color shades of the points (bigger and darker colors imply a larger tendency). The principal directions shown as continuous lines (PC0 , red; PC1, green; PC2, purple; PC3, orange) emerge from centroid, which is given by the average mean of the data set and is shown as a black star. The projection of the points and the average mean is shown in gray. The projection confidence interval of the average mean is also shown. The origin, which is defined as (0, 0, 0, 0) in the 4D space of the learning styles is shown as a black triangle. The colored points correspond to the maximal pure learning styles (e.g., (20, 0, 0, 0) for activist (red), or (0, 20, 0, 0) for reflector (green) are also marked.')
document.add_picture(filename_ls_3d, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Four-dimensional representation of the learning styles of the students as a function of the CHAEA learning styles.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

document.add_page_break()
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the 4D plot of the LSs done!')
printt('-------------------------------------------')



nfig += 1
nfig_PC_statistics = nfig
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the projection on PC0-PC1...')
printt('-------------------------------------------')
document.add_paragraph('The central panel in Fig. ' + str(nfig) + ' shows the projection of the learning styles of the student represented in Fig. ' + str(nfig_4d) + ' on the two main principal components (PC0 and PC1), along with their histograms. The data surround the average mean (gray star). The points (20, 0, 0, 0), (0, 20, 0, 0), (0, 0, 20, 0), and (0, 0, 0, 20), which correspond, respectively, to the maximal pure activist (red), reflector (green), theorist (purple), and pragmatist (orange) learning styles in the original basis set are also marked, and joined with a dashed line to the corresponding origin (0, 0, 0, 0) (black triangle). The point (20, 20, 20, 20) is also shown (brown diamond).')

document.add_paragraph('The reduced dimensional representation of the data given by Fig. ' + str(nfig_4d) + ' is usually meaningful if Σ_1(%)≥70.0%; in the case under study we have that Σ_1(%)='+sigma01+'. The top and left panels of Fig. ' + str(nfig) + ' show the histrograms of the data as a function of PC0 (red) and PC1 (green), which are fitted using Weibull distributions with the parameters listed in Table ' + str(ntab + 1) + '.')
document.add_picture(filename_PC0PC1, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. The projection of the learning styles of the students represented in Fig. ' + str(nfig_4d) + ' on the two main principal components (PC0 and PC1), along with their histograms as a function of PC0 (red) and PC1 (green).')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

document.add_page_break()
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the projection on PC0-PC1 done!')
printt('-------------------------------------------')

ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the Weibull parameters for the PDFs of the PCs...')
printt('-------------------------------------------')
# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Parameters of the Weibull distributions that fit the probability distributions (histograms) of the projection of the learning styles on the principal-components basis set (see histograms in Fig. ' + str(nfig) + '). The location parameter θ is set equal to the smallest projection for each component.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(K+1, 4)

# Table header
t.cell(0,0).text = 'Principal component'
t.cell(0,1).text = 'α'
t.cell(0,2).text = 'k'
t.cell(0,3).text = 'θ'

for j in range(0,K):
  t.cell(j+1,0).text = Label_PC[j]
  [alpha_Weibull, k_Weibull, theta_Weibull] = parameters_Weibull_PC20[j]
  t.cell(j+1,1).text = str(round(alpha_Weibull,2))
  t.cell(j+1,2).text = str(round(k_Weibull,2))
  t.cell(j+1,3).text = str(round(theta_Weibull,2))

document.add_page_break()

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the Weibull parameters for the PDFs of the PCs done!')
printt('-------------------------------------------')












if ( Lprof > 0 ):
 nsubsec += 1
 nfig_prof0 = nfig + 1
 if ( Lprof == 1 ):
  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Relation between the learning styles of the students and of the '+ prof_ref +' in the principal-component space', level=2)
  printt(str(nsec) + '.' + str(nsubsec) + ' Relation between the learning styles of the students and of the '+ prof_ref +' in the principal-component space')
  document.add_paragraph('Figure ' + str(nfig_prof0) + ' shows the projection of the learning styles of the students on the two main principal components (PC0 and PC1), as represented in the central panel of Fig. ' + str(nfig_PC_statistics) + ', along with the results for the ' + prof_ref + ' (gray cross). The students with the same qualitative tendency as the ' + prof_ref + ' have been represented with colored crosses.')
 else:
  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Relation between the learning styles of the students and of the '+ prof_ref +'s in the principal-component space', level=2)
  printt(str(nsec) + '.' + str(nsubsec) + ' Relation between the learning styles of the students and of the '+ prof_ref +'s in the principal-component space')
  document.add_paragraph('Figures ' + str(nfig_prof0) + ' to '  + str(nfig_prof0 + Lprof) + ' show the projection of the learning styles of the students on the two main principal components (PC0 and PC1), as represented in the central panel of Fig. ' + str(nfig_PC_statistics) + ', along with the results for the ' + prof_ref + 's (gray crosses). The students with the same qualitative tendency as the ' + prof_ref + 's have been represented with colored crosses.')

 printt('-------------------------------------------')
 printt(' Figure(s) with the LSs of the students and of the '+ prof_ref +'(s) on PC0-PC1...')
 printt('-------------------------------------------')
 for i_prof in range(0, Lprof):
  profi = profs[i_prof]
  nfig  = nfig + 1

  printt('output_participants_pc_prof:' + output_participants_pc_prof)
  filename = output_participants_pc_prof + '/Fig_'+Label_PCPC0+'_'+Label_PCPC1+'_prof_' + profi + '.png'
  document.add_picture(filename, width=Inches(5))

# Figure title
  printt('i_prof : ' + str(i_prof))
  printt('prof_ref : ' + profi)
  if(i_prof == 0):
    figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Projection of the learning styles of the students on the two main principal components (PC0 and PC1), as represented in the central panel of Fig. ' + str(nfig_PC_statistics) + ' along with the results of the ' + prof_ref + ' of ' + profi + ' (gray cross). The colored crosses mark the students with the same qualitative tendency towards the corresponding learning style (light shade for very low and low tendencies, medium shade for moderate tendency, and dark shade for high and very high tendencies).')
  else:
    figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Same as Fig. ' + str(nfig_prof0) + ' for the ' + prof_ref + ' of ' + profi + '.')

  figure_title.alignment = 1  # Center alignment
  figure_run = figure_title.runs[0]
  figure_run.bold = True

  printt('-------------------------------------------')
  printt(' Figure ' + str(nfig) + ' with the projection of the LS of the students and professor ' + profi +' on PC0-PC1 done!')
  printt('-------------------------------------------')


document.add_page_break()
#===================================================
# SECTION 3
#===================================================
#==================================================
nsec   += 1
nsubsec = 1
document.add_heading(str(nsec) + '. Participation ratios', level=1)
printt('******************************************')
printt(str(nsec) + '. Participation ratios')
printt('******************************************')
#==================================================

document.add_paragraph('Third, the participation ratios are presented in this section. The values for the CHAEA (activist, reflector, theorist, and pragmatist), and for the principal-components basis sets are listed. A statistical analysis of the distribution of the participation ratios is also included.')

nsubsec = 1
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Individual analysis of the participation ratios', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Individual analysis of the participation ratios')

ntab += 1
printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the PRs...')
printt('-------------------------------------------')

document.add_paragraph('Table ' + str(ntab) + ' shows the participation ratios of each individual student for CHAEA (activist, reflector, theorist, and pragmatist) and for the principal-components basis sets. For the case under study, this parameter lies between 1 and 4. The smaller the participation ratio, the better.')

# Table 1 title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Participation ratios of each individual student for CHAEA learning styles (LS) and for the principal-components (PC) basis sets.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

# Table 1
t = document.add_table(Lsave+1, 3)

# Table header
t.cell(0,0).text = 'Student'
t.cell(0,1).text = 'LS'
t.cell(0,2).text = 'PC'

for i in range(0,Lsave):
  t.cell(i+1,0).text = students[i]
  t.cell(i+1,1).text = str(round(pr[i,0],2))
  t.cell(i+1,2).text = str(round(pr[i,1],2))

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the PRs done!')
printt('-------------------------------------------')
printt(' ')


document.add_page_break()
nsubsec += 1
document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Global analysis of the participation ratios', level=2)
printt(str(nsec) + '.' + str(nsubsec) + ' Global analysis of the participation ratios')

nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the PDF of the PRs...')
printt('-------------------------------------------')
ntab += 1
document.add_paragraph('Figure ' + str(nfig) + ' shows the probability distribution functions of the participation ratios for the learning-styles basis set (red) and for the principal components (blue). The vertical dashed lines mark the average means along with their corresponding confidence intervals (shaded areas), whose values can be found in Table ' + str(ntab) + '. The corresponding cumulative distributions used in the fitting are shown in Fig. ' + str(nfig+1) + '.')

document.add_picture(filename_PR_P, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Probability distribution functions of the participation ratios associated with the learning-styles basis set (red) and with the principal components (blue). The vertical dashed lines mark the average means along with their corresponding confidence intervals (shaded areas), whose values can be found in Table ' + str(ntab) + '. The continuous lines show fitting functions given by Weibull distributions with parameters shown in Table ' + str(ntab) + '.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with the PDF of the PRs done!')
printt('-------------------------------------------')


nfig += 1
printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with W(s) of the PRs...')
printt('-------------------------------------------')

document.add_paragraph('')
document.add_picture(filename_PR_W, width=Inches(5))

# Figure title
figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Cumulative distributions for the results of Fig. ' + str(nfig-1) + '.')
figure_title.alignment = 1  # Center alignment
figure_run = figure_title.runs[0]
figure_run.bold = True

printt('-------------------------------------------')
printt(' Figure ' + str(nfig) + ' with W(s) of the PRs done!')
printt('-------------------------------------------')




printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the average parameters for the PRs...')
printt('-------------------------------------------')

# Table title
table_title = document.add_paragraph('Table ' + str(ntab) + '. Average mean and corresponding uncertainty (in parenthesis) of the participation ratios for the learning-styles basis set and for the principal components. α and k are, respectively, the shape and scale parameters of the Weibull distributions the fit the probability distributions (histograms) of Fig. ' + str(nfig) + '. The location parameter is set equal to θ = 0.')
table_title.alignment = 1  # Center alignment
title_run = table_title.runs[0]
title_run.bold = True

t = document.add_table(3, 4)

# Table header
t.cell(0,0).text = 'Basis set'
t.cell(0,1).text = 'Average mean (Uncertainty)'
t.cell(0,2).text = 'α'
t.cell(0,3).text = 'k'

t.cell(1,0).text = 'Learning styles'
t.cell(2,0).text = 'Principal components'

for j in range(0,2):
  xdx = mean_uncert(prmean[j], uncert_abs_pr[j])
  t.cell(j+1,1).text = xdx
  [alpha_Weibull, k_Weibull] = parameters_Weibull_PR[j]
  t.cell(j+1,2).text = str(round(alpha_Weibull,2))
  t.cell(j+1,3).text = str(round(k_Weibull,2))

printt('-------------------------------------------')
printt(' Table ' + str(ntab) + ' with the average parameters for the PRs done!')
printt('-------------------------------------------')



print('ntab : ', ntab)
ntab0_mark = ntab + 1
document.add_page_break()
#===================================================
# SECTION 4 (IF MARKS ARE PROVIDED)
#===================================================
if ( K_value > 0 ):
 nsec   += 1
 nsubsec = 1
 document.add_heading(str(nsec) + '. Analysis of the students marks', level=1)
 printt('******************************************')
 printt(str(nsec) + '. Analysis of the students marks')
 printt('******************************************')
 if ( Lmarks == 0 ):
  document.add_paragraph('The marks provided in the input folder ' + input_folder_marks_ref + ' have not been analyzed as the students included are not the same ones as those that fullfilled CHAEA.')
 else:
  document.add_paragraph('This section is devoted to the analysis of the students marks contained in the input folder ' + input_folder_marks_ref + '.')

  document.add_paragraph('The tables present the averages and uncertainties (in parenthesis) of the learning styles And the principal components (x) and the marks (y) of the whole group, along with the parameters of the linear regressions that fit the marks as a function of the number of points in the corresponding learning style or principal components, being m the slope, b the intercept, and r the regression parameter. Recall that the closest the regression parameter to -1 or 1, the better the fitting. Conversely, a regression parameter close to zero indicates the lack of correlations between the marks and the corresponding variable.')

  document.add_paragraph('The results marked as nan indicate failures in the linear regressions because, e.g., having too few points.')

  df = pd.DataFrame(marks_vs_LS)
  if df.isna().values.any():
    document.add_paragraph('The cells with nan indicate the failure of the linear regression.')

  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Marks as a function of CHAEA learning styles', level=2)
  printt(str(nsec) + '.' + str(nsubsec) + ' Marks as a function of CHAEA learning styles')
  document.add_paragraph('In this section, we report on the relationship between the marks and CHAEA learning styles.')

  column0_strings = [str(row[0]) for row in data_marks]
  Excel_documents_marks = column0_strings
  del Excel_documents_marks[0]
  Excel_documents_marks = list(dict.fromkeys(Excel_documents_marks))
  print('Excel_documents_marks : ', Excel_documents_marks)

  Lmarks_vs_LS = np.shape(marks_vs_LS)[0]
  Lmarks_vs_PC = np.shape(marks_vs_PC)[0]

  ntab_mean = ntab + 1
  nfig_mean = nfig + 1

  for Excel_doc in Excel_documents_marks:
    print('Excel_doc  : ', Excel_doc)

    document.add_heading('Marks of the document ' + Excel_doc, level=3)
    printt(str(nsec) + '.' + str(nsubsec) + ' Marks of the document ' + Excel_doc)

#   Subsubsections (Excel sheets)
    Excel_sheets = [str(row[1]) for row in data_marks if row[0] == Excel_doc]
    print('Excel sheets : ', Excel_sheets)
    Lsheets = len(Excel_sheets)

    for nExcel_sheet in range(0,Lsheets):
#      if(Lsheets == 1):marks_vs_LS
#        Excel_sheet = Excel_sheets
#      else:
      Excel_sheet = Excel_sheets[nExcel_sheet]
      print('Excel sheet : ', Excel_sheet)

#     Tables (each Excel column corresponds to the marks for an evaluating activity
      Excel_columns = [str(row[2]) for row in data_marks if row[0] == Excel_doc and row[1] == Excel_sheet ]
      print('Excel_columns : ', Excel_columns)
      Lcolumns = len(Excel_columns)

      for nExcel_column in range(0,Lcolumns):

        Excel_column = Excel_columns[nExcel_column]
        print('Excel_column : ', Excel_column)
        print('Marks of the column ' + Excel_column + ' of the Excel sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc) # + ' of the Excel document ' + Excel_doc)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Marks of the column {Excel_column} of the sheet {Excel_sheet}")
        run.bold = True

        nfig += 1
        # Figure
        printt('-------------------------------------------')
        printt(' Figure ' + str(nfig) + ' with the marks vs LS...')
        printt('-------------------------------------------')

        filename = output_participants_ls_marks + '/Fig_'+ Excel_doc + '_' + str(Excel_sheet) + '_' + Excel_column + '_K' + str(K_value) + '.png'
        if not os.path.exists(filename):
         nfig -= 1
         print(f"The file '{filename}' does NOT exist.")
         document.add_paragraph('No clustering analysis based on K-means algorithm is performed due to the low number of students for the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '. Consequently, no figures or tables relating the marks and the CHAEA punctuations for the different learning styles are included here.')
        else:
         print(f"The file '{filename}' exists.")
         if( K_value == 1):
          document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the CHAEA punctuations for the different learning styles. The results are presented in Table ' + str(ntab+1) + ', and correspond to the whole group.')
         else:
          if( K_value == 2):
            document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the CHAEA punctuations for the different learning styles. The results in gray are presented in Table ' + str(ntab+1) + ', and correspond to the whole group. The colored results are listed in the Tables ' + str(ntab + 2) + ' and ' +str(ntab + 1 + K_value) + ', and are associated with the different clusters formed by the K-mean algorithm.')
          else:
            document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the CHAEA punctuations for the different learning styles. The results in gray are presented in Table ' + str(ntab+1) + ', and correspond to the whole group. The colored results are listed in the Tables ' + str(ntab + 2) + ' to ' +str(ntab + 1 + K_value) + ', and are associated with the different clusters formed by the K-mean algorithm.')

         document.add_picture(filename, width=Inches(5))

#        Figure title
         if(nfig == nfig_mean):
          if(K_value == 1):
           figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the CHAEA punctuations for the different learning styles. Students are depicted with circles, while cluster centroids are marked with stars. Linear fits are shown as continuous lines, and their confidence intervals are illustrated in shaded areas.')
          else:
           figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the CHAEA punctuations for the different learning styles. Each color represents a distinct cluster determined using the K-means algorithm with K=' + str(K_value)+'. Students are depicted with circles, while cluster centroids are marked with stars. Linear fits are shown as continuous lines, and their confidence intervals are illustrated in shaded areas. The results for the whole group are indicated in gray.')
         else:
           figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Same as Fig. ' + str(nfig_mean) + ' for the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '.')

         figure_title.alignment = 1  # Center alignment
         figure_run = figure_title.runs[0]
         figure_run.bold = True

         printt('-------------------------------------------')
         printt(' Figure ' + str(nfig) + ' with the marks vs LS done!')
         printt('-------------------------------------------')

         ntab += 1
         printt('-------------------------------------------')
         printt(' Table ' + str(ntab) + ' with the marks and the LS for the whole group...')
         printt('-------------------------------------------')

         j = 0
         title = 0
         for i in range(0, Lmarks_vs_LS):
          if ( (marks_vs_LS[i, 0] == Excel_doc) and (marks_vs_LS[i, 1] == Excel_sheet) and (marks_vs_LS[i, 2] == Excel_column) and (int(marks_vs_LS[i, 5]) == -1) ):
            if(title == 0):
              title = 1
              iref  = i
#             Table title
              if(ntab == ntab_mean):
                if(float(marks_vs_LS[i, 3]) == 1):
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Results for the whole group, which is formed by only ' + str(marks_vs_LS[i, 3]) + ' student. The table presents the average values and uncertainties (in parenthesis) for the learning styles (x), for the marks (y) contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ', along with the parameters of the linear regression, which cannot fit the mark as a function of the number of points in the corresponding learning style because only one point is provided (m: slope, b: intercept, r: regression parameter).')
                else:
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Results for the whole group, which is formed by ' + str(marks_vs_LS[i, 3]) + ' students. The table presents the average values and uncertainties (in parenthesis) for the learning styles (x), for the marks (y) contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ', along with the parameters of the linear regressions that fit the marks as a function of the number of points in the corresponding learning style (m: slope, b: intercept, r: regression parameter).')
              else:
                if(float(marks_vs_LS[i, 3]) == 1):
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean) + ' for the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '. The linear regression fails because only student is included.')
                else:
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean) + ' for the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '.')

              table_title.alignment = 1  # Center alignment
              title_run = table_title.runs[0]
              title_run.bold = True
#
#             Average results for the whole group
              t = document.add_table(5, 6)
#             Table header
              t.cell(0,0).text = 'LS'
              t.cell(0,1).text = 'x\u0304 (\u0394x)'
              t.cell(0,2).text = 'y\u0304 (\u0394y)'
              t.cell(0,3).text = 'm\u0304 (\u0394m)'
              t.cell(0,4).text = 'b\u0304 (\u0394b)'
              t.cell(0,5).text = 'r'

            j += 1
            t.cell(j,0).text = marks_vs_LS[i, 6]
            xdx = mean_uncert(float(marks_vs_LS[i, 7]), float(marks_vs_LS[i, 8]))
            t.cell(j,1).text = xdx
            ydy = mean_uncert(marks_vs_LS[i, 9], marks_vs_LS[i, 10])
            t.cell(j,2).text = ydy
            mdm = mean_uncert(marks_vs_LS[i, 11], marks_vs_LS[i, 12])
            t.cell(j,3).text = mdm
            bdb = mean_uncert(marks_vs_LS[i, 13], marks_vs_LS[i, 14])
            t.cell(j,4).text = bdb
            t.cell(j,5).text = round_until_two_non_9(marks_vs_LS[i,15])

         document.add_paragraph()
         if ( float(marks_vs_LS[iref, 3]) <= Lsave ):
          if(float(marks_vs_LS[iref, 3]) == 1):
            document.add_paragraph('Table ' + str(ntab) + ' shows the results for marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' for the whole group, which is formed by the following student: ' + marks_vs_LS[iref, 4] + '.')
          else:
            document.add_paragraph('Table ' + str(ntab) + ' shows the results for marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' for the whole group, which is formed by the following ' + str(marks_vs_LS[iref, 3]) + ' students: ' + marks_vs_LS[iref, 4] + '.')
         else:
          document.add_paragraph('Table ' + str(ntab) + ' shows the results for marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' for the whole group, which is formed by the following ' + str(marks_vs_LS[iref, 3]) + ' students, being the first ' + str(Lsave) + ' ones: ' + marks_vs_LS[iref, 4] + '.')

         printt('-------------------------------------------')
         printt(' Table ' + str(ntab) + ' with the marks and the LS for the whole group done!')
         printt('-------------------------------------------')

#        Clustering-analysis results rendered by K-means algorithm
         if( K_value > 1):
          if( K_value == 2):
           document.add_paragraph('Tables ' + str(ntab + 1) + ' and ' + str(ntab + K_value) + ' include the same results as those presented in Table ' + str(ntab_mean) + ' for the clusters constructed using K-means algorithm.')
          else:
           document.add_paragraph('Tables ' + str(ntab + 1) + ' to ' + str(ntab + K_value) + ' include the same results as those presented in Table ' + str(ntab_mean) + ' for the clusters constructed using K-means algorithm.')

          for k in range(0, K_value):
           print('k : ', k)
           j = 0

           ntab += 1
           title = 0
           print('ntab : ', ntab)
           printt('-------------------------------------------')
           printt(' Table ' + str(ntab) + ' with the marks and the LS for the cluster k = ' + str(k) + '...')
           printt('-------------------------------------------')

           for i in range(0, Lmarks_vs_LS):
#           Average results for the whole group
            if ( (marks_vs_LS[i, 0] == Excel_doc) and (marks_vs_LS[i, 1] == Excel_sheet) and (marks_vs_LS[i, 2] == Excel_column) and (int(marks_vs_LS[i, 5]) == k) ):

              if(title == 0):
                title = 1
                iref = i
#
#               Table title
                if(float(marks_vs_LS[i, 3]) == 1):
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean) + ' for the cluster k=' + str(k) + ' which is formed by one student. The linear regression fails because only student is included.')
                else:
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean) + ' for the cluster k=' + str(k) + ' which is formed by ' + str(marks_vs_LS[i, 3]) + ' students.')
                table_title.alignment = 1  # Center alignment
                title_run = table_title.runs[0]
                title_run.bold = True
#
                table_title.alignment = 1  # Center alignment
                title_run = table_title.runs[0]
                title_run.bold = True

                t = document.add_table(5, 6)
#               Table header
                t.cell(0,0).text = 'LS'
                t.cell(0,1).text = 'x\u0304 (\u0394x)'
                t.cell(0,2).text = 'y\u0304 (\u0394y)'
                t.cell(0,3).text = 'm\u0304 (\u0394m)'
                t.cell(0,4).text = 'b\u0304 (\u0394b)'
                t.cell(0,5).text = 'r'

              j += 1
              t.cell(j,0).text = marks_vs_LS[i, 6]
              xdx = mean_uncert(float(marks_vs_LS[i, 7]), float(marks_vs_LS[i, 8]))
              t.cell(j,1).text = xdx
              ydy = mean_uncert(marks_vs_LS[i, 9], marks_vs_LS[i, 10])
              t.cell(j,2).text = ydy
              mdm = mean_uncert(marks_vs_LS[i, 11], marks_vs_LS[i, 12])
              t.cell(j,3).text = mdm
              bdb = mean_uncert(marks_vs_LS[i, 13], marks_vs_LS[i, 14])
              t.cell(j,4).text = bdb
              print('marks_vs_LS[i,15] : ', marks_vs_LS[i,15])
              t.cell(j,5).text = round_until_two_non_9(marks_vs_LS[i,15])

           document.add_paragraph()
           if ( float(marks_vs_LS[iref, 3]) <= Lsave ):
            if(float(marks_vs_LS[iref, 3]) == 1):
              document.add_paragraph('Table ' + str(ntab) + ' shows the results for the cluster k=' + str(k) + ', which is formed by the following student: ' + marks_vs_LS[iref, 4] + '.')
            else:
              document.add_paragraph('Table ' + str(ntab) + ' shows the results for the cluster k=' + str(k) + ', which is formed by the following ' +  str(marks_vs_LS[iref, 3]) + ' students: ' + marks_vs_LS[iref, 4] + '.')
           else:
            document.add_paragraph('Table ' + str(ntab) +' shows the results for the cluster k=' + str(k) + ', which is formed by ' +  str(marks_vs_LS[iref, 3]) + ' students, being the first ' + str(Lsave) + ' students: ' + marks_vs_LS[iref, 4]+ '.')

          printt(' Table ' + str(ntab) + ' with the marks and the LS for the cluster k = ' +str(k) + ' done!')







  document.add_page_break()
  nsubsec += 1
  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Marks as a function of the principal components', level=2)
  printt(str(nsec) + '.' + str(nsubsec) + ' Marks as a function of the principal components')
  document.add_paragraph('In this section, we report on the relationship between the marks and the principal components.')

  ntab_PC = 0

  for Excel_doc in Excel_documents_marks:
    print('Excel_doc  : ', Excel_doc)

    document.add_heading('Marks of the document ' + Excel_doc, level=3)
    printt(str(nsec) + '.' + str(nsubsec) + ' Marks of the document ' + Excel_doc)

#   Subsubsections (Excel sheets)
    Excel_sheets = [str(row[1]) for row in data_marks if row[0] == Excel_doc]
    print('Excel sheets : ', Excel_sheets)
    Lsheets = len(Excel_sheets)

    for nExcel_sheet in range(0,Lsheets):
      Excel_sheet = Excel_sheets[nExcel_sheet]
      print('Excel sheet : ', Excel_sheet)

#     Tables (each Excel column corresponds to the marks for an evaluating activity
      Excel_columns = [str(row[2]) for row in data_marks if row[0] == Excel_doc and row[1] == Excel_sheet ]
      print('Excel_columns : ', Excel_columns)
      Lcolumns = len(Excel_columns)

      for nExcel_column in range(0,Lcolumns):

        Excel_column = Excel_columns[nExcel_column]
        print('Excel_column : ', Excel_column)
        print('Marks of the column ' + Excel_column + ' of the Excel sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc) # + ' of the Excel document ' + Excel_doc)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"Marks of the column {Excel_column} of the sheet {Excel_sheet}")
        run.bold = True

#       Figure
        filename = output_participants_pc_marks + '/Fig_'+ Excel_doc + '_' + str(Excel_sheet) + '_' + Excel_column + '_K' + str(K_value) + '.png'

        if not os.path.exists(filename):
         print(f"The file '{filename}' does NOT exist.")
         document.add_paragraph('No clustering analysis based on K-means algorithm is performed due to the low number of students for the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '. Consequently, no figures or tables relating the marks and the CHAEA punctuations for the different learning styles are included here.')
        else:
         nfig += 1
         printt('-------------------------------------------')
         printt(' Figure ' + str(nfig) + ' with the marks vs PC...')
         printt(' nfig_mean : ' + str(nfig_mean))
         printt('-------------------------------------------')
         if( K_value == 1):
          document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the principal components. The results are presented in Table ' + str(ntab+1) + ', and correspond to the whole group.')
         else:
          if( K_value == 2):
            document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the principal components. The results in gray are presented in Table ' + str(ntab+1) + ', and correspond to the whole group. The colored results are listed in the Tables ' + str(ntab + 2) + ' and ' + str(ntab + 1 + K_value) + ', and are associated with the different clusters formed by the K-mean algorithm.')
          else:
            document.add_paragraph('Figure ' + str(nfig) + ' shows the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' as a function of the principal components. The results in gray are presented in Table ' + str(ntab+1) + ', and correspond to the whole group. The colored results are listed in the Tables ' + str(ntab + 2) + ' to ' + str(ntab + 1 + K_value) + ', and are associated with the different clusters formed by the K-mean algorithm.')


         document.add_picture(filename, width=Inches(5))

#        Figure title
         figure_title = document.add_paragraph('Figure ' + str(nfig) + '. Same as Fig. ' + str(nfig_mean) + ' for the  principal components and the marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + '.')
         figure_title.alignment = 1  # Center alignment
         figure_run = figure_title.runs[0]
         figure_run.bold = True

         printt('-------------------------------------------')
         printt(' Figure ' + str(nfig) + ' with the marks vs PC done!')
         printt('-------------------------------------------')

         ntab += 1
         ntab_PC += 0
         printt('-------------------------------------------')
         printt(' Table ' + str(ntab) + ' with the marks and the PC for the whole group...')
         printt('-------------------------------------------')

         j = 0
         title = 0
         for i in range(0, Lmarks_vs_PC):
          if ( (marks_vs_PC[i, 0] == Excel_doc) and (marks_vs_PC[i, 1] == Excel_sheet) and (marks_vs_PC[i, 2] == Excel_column) and (int(marks_vs_PC[i, 5]) == -1) ):
            if(title == 0):
              title = 1
              iref  = i

#             Table title
              table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean + ntab_PC) + ' for the principal components (the marks are contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ').')
              table_title.alignment = 1  # Center alignment
              title_run = table_title.runs[0]
              title_run.bold = True
#
#             Average results for the whole group
              t = document.add_table(5, 6)
#             Table header
              t.cell(0,0).text = 'PC'
              t.cell(0,1).text = 'x\u0304 (\u0394x)'
              t.cell(0,2).text = 'y\u0304 (\u0394y)'
              t.cell(0,3).text = 'm\u0304 (\u0394m)'
              t.cell(0,4).text = 'b\u0304 (\u0394b)'
              t.cell(0,5).text = 'r'

            j += 1
            t.cell(j,0).text = marks_vs_PC[i, 6]
            xdx = mean_uncert(float(marks_vs_PC[i, 7]), float(marks_vs_PC[i, 8]))
            t.cell(j,1).text = xdx
            ydy = mean_uncert(marks_vs_PC[i, 9], marks_vs_PC[i, 10])
            t.cell(j,2).text = ydy
            mdm = mean_uncert(marks_vs_PC[i, 11], marks_vs_PC[i, 12])
            t.cell(j,3).text = mdm
            bdb = mean_uncert(marks_vs_PC[i, 13], marks_vs_PC[i, 14])
            t.cell(j,4).text = bdb
            t.cell(j,5).text = round_until_two_non_9(marks_vs_PC[i,15])

         document.add_paragraph()
         if(float(marks_vs_PC[iref, 3]) == 1):
          document.add_paragraph('Table ' + str(ntab) + ' shows the results for marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' for the whole group, which is formed by the same student as that considered in Table ' + str(ntab_mean + ntab_PC) + '.')
         else:
          document.add_paragraph('Table ' + str(ntab) + ' shows the results for marks contained in the column ' + Excel_column + ' of the sheet ' + str(Excel_sheet) + ' of the document ' + Excel_doc + ' for the whole group, which is formed by the same ' + str(marks_vs_PC[iref, 3]) + ' students as those considered in Table ' + str(ntab_mean + ntab_PC) + '.')

         ntab_PC += 1

         printt('-------------------------------------------')
         printt(' Table ' + str(ntab) + ' with the marks and the PC for the whole group done!')
         printt('-------------------------------------------')

#        Clustering-analysis results rendered by K-means algorithm
         if( K_value > 1):
          if( K_value == 2):
           document.add_paragraph('Tables ' + str(ntab + 1) + ' and ' + str(ntab + K_value) + ' include the same results as those presented in Tables ' + str(ntab_mean + ntab_PC ) + ' and ' + str(ntab_mean + ntab_PC + K_value - 1) + ' for the clusters constructed using K-means algorithm and for the principal components (instead of the learning styles).')
          else:
           document.add_paragraph('Tables ' + str(ntab + 1) + ' to ' + str(ntab + K_value) + ' include the same results as those presented in Tables ' + str(ntab_mean + ntab_PC ) + ' to ' + str(ntab_mean + ntab_PC + K_value - 1) + ' for the clusters constructed using K-means algorithm and for the principal components (instead of the learning styles).')

          for k in range(0, K_value):
           print('k : ', k)
           j = 0
           ntab += 1
           ntab_PC += 1
           title = 0
           print('ntab : ', ntab)
           printt('-------------------------------------------')
           printt(' Table ' + str(ntab) + ' with the marks and the PC for the cluster k = ' + str(k) + '...')
           printt('-------------------------------------------')

           for i in range(0, Lmarks_vs_PC):
#           Average results for the whole group
            if ( (marks_vs_PC[i, 0] == Excel_doc) and (marks_vs_PC[i, 1] == Excel_sheet) and (marks_vs_PC[i, 2] == Excel_column) and (int(marks_vs_PC[i, 5]) == k) ):

              if(title == 0):
                title = 1
                iref = i

#               Table title
                if(float(marks_vs_PC[i, 3]) == 1):
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean + ntab_PC - 1) + ' for the cluster k=' + str(k) + ' which is formed by one student, which makes the linear regression ill-suited.')
                else:
                  table_title = document.add_paragraph('Table ' + str(ntab) + '. Same as Table ' + str(ntab_mean + ntab_PC - 1) + ' for the cluster k=' + str(k) + ' which is formed by ' + str(marks_vs_PC[i, 3]) + ' students.')
                table_title.alignment = 1  # Center alignment
                title_run = table_title.runs[0]
                title_run.bold = True
                t = document.add_table(5, 6)
#               Table header
                t.cell(0,0).text = 'PC'
                t.cell(0,1).text = 'x\u0304 (\u0394x)'
                t.cell(0,2).text = 'y\u0304 (\u0394y)'
                t.cell(0,3).text = 'm\u0304 (\u0394m)'
                t.cell(0,4).text = 'b\u0304 (\u0394b)'
                t.cell(0,5).text = 'r'

              j += 1
              t.cell(j,0).text = marks_vs_PC[i, 6]
              xdx = mean_uncert(float(marks_vs_PC[i, 7]), float(marks_vs_PC[i, 8]))
              t.cell(j,1).text = xdx
              ydy = mean_uncert(marks_vs_PC[i, 9], marks_vs_PC[i, 10])
              t.cell(j,2).text = ydy
              mdm = mean_uncert(marks_vs_PC[i, 11], marks_vs_PC[i, 12])
              t.cell(j,3).text = mdm
              bdb = mean_uncert(marks_vs_PC[i, 13], marks_vs_PC[i, 14])
              t.cell(j,4).text = bdb
              t.cell(j,5).text = round_until_two_non_9(marks_vs_PC[i,15])

           document.add_paragraph()
           if(float(marks_vs_PC[iref, 3]) == 1):
            document.add_paragraph('Table ' + str(ntab) + ' shows the results for the cluster k=' + str(k) + ', which is formed by the same student as that considered in the Table ' + str(ntab_mean + ntab_PC - 1) + '. The linear regression fails because only student is included.')
           else:
            document.add_paragraph('Table ' + str(ntab) + ' shows the results for the cluster k=' + str(k) + ', which is formed by the same ' +  str(marks_vs_PC[iref, 3]) + ' students as those considered in the Table ' + str(ntab_mean + ntab_PC - 1) + '.')

           printt(' Table ' + str(ntab) + ' with the marks and the PC for the cluster k = ' +str(k) + ' done!')

document.add_page_break()

#===================================================
# SECTION 4 OR 5
#===================================================
nsec   += 1
nsubsec = 1
document.add_heading(str(nsec) + '. Students characteristics and recommendations', level=1)
printt('******************************************')
printt(str(nsec) + '. Students characteristics and recommendations')
printt('******************************************')
document.add_paragraph('This section summarizes some of the characteristics associated with the activist, theorist, reflector, and pragmatist learning styles, and provides some recommendation to increase the tendencies towards them, as detailed by Alonso, Gallego & Honey (2007), and Cancino (2014).')

alonso_reference_paragraph = document.add_paragraph()
alonso_reference_paragraph.add_run('Alonso, C. M., Gallego, D. J., & Honey, P. (2007). Los estilos de aprendizaje. Procedimiento de diagnóstico y mejora (7th ed.). Ediciones Mensajero, S. A. U.').italic = True

cancino_reference_paragraph = document.add_paragraph()
cancino_reference_paragraph.add_run('Cancino, O. (2024, March). Blog de la prof de Français: espacio para intercambiar con los estudiantes. add_hyperlink(p, "https://oticar.wordpress.com/estilos-de-aprendizaje/", "https://oticar.wordpress.com/estilos-de-aprendizaje/"').italic = True
#
#==================================================
# Learning-styles characteristics
#==================================================
#
# Activist
activist_main_characteristics = ['Cheerful', 'Improviser', 'Explorer', 'Risk-taker', 'Spontaneous']
activist_other_characteristics = 'creative, innovative, adventurous, revitalizing, inventor, energetic, experience-seeker, idea-generator, bold, protagonist, striking, groundbreaking, talkative, leader, determined, fun, participative, competitive, eager to learn, problem-solver, adaptable.'
#
# Reflector
reflector_main_characteristics = ['Thoughtful', 'Thorough', 'Receptive', 'Analytical', 'Exhaustive']
reflector_other_characteristics = 'observer, collector, patient, careful, detail-oriented, argument-developer, alternative-planner, behavior-studier, data-recorder, researcher, assimilator, report and/or statement writer, slow, distant, cautious, inquisitive, probe-seeker.'
#
# Theorist
theorist_main_characteristics = ['Methodical', 'Logical', 'Objective', 'Critical', 'Structured']
theorist_other_characteristics = 'disciplined, planned, systematic, organized, synthetic, reasoner, thinker, connector, perfectionist, generalizer, hypothesis-seeker, theory-seeker, model-seeker, question-asker, assumption-explorer, concept-seeker, clear-purpose seeker, rationality-seeker, "why"-asker, value-system and criteria-seeker, procedure-inventor, explorer.'
#document.add_paragraph()
# Pragmatist
pragmatist_main_characteristics = ['Experimenter', 'Practical', 'Direct', 'Efficient', 'Realistic']
pragmatist_other_characteristics = 'technical, useful, fast, decisive, planner, positive, concrete, objective, clear, self-confident, organizer, contemporary, problem-solver, applier of learning, action planner.'
#
# Main characteristics
main_characteristics = [activist_main_characteristics, reflector_main_characteristics, theorist_main_characteristics, pragmatist_main_characteristics]
#
# Other characteristics
other_characteristics = [activist_other_characteristics, reflector_other_characteristics, theorist_other_characteristics, pragmatist_other_characteristics]
#
#==================================================
# Affinity and dominant learning style
#==================================================
# Maximum value of the affinity
max_affinity = np.max(affinity)
max_affinity_str = str(round(max_affinity,1))

# Position (index) of the maximum value
max_index_affinity = np.where(affinity == max_affinity)[0]

if(len(max_index_affinity) == 1):
# One dominant learning style
  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Affinity and dominant learning style', level=2)
  i = max_index_affinity[0]
  document.add_paragraph('According to the affinity, the dominant learning style is the ' + List_LS[i] + '. The value of its affinity is equal to ' + max_affinity_str +' %.' )
  document.add_paragraph('The main characteristics of the ' + List_LS[i] + ' learning style are:')
  for point in main_characteristics[i]:
    document.add_paragraph(point, style='ListBullet')
  document.add_paragraph('Furthermore, other characteristics usually associated with the ' + List_LS[i] + ' learning style are being: ' + other_characteristics[i])
else:
# Several dominant learning styles
  document.add_heading(str(nsec) + '.' + str(nsubsec) + ' Affinity and dominant learning styles', level=2)
  list_mainLS = List_LS[max_index_affinity[0]]
  for i in range(1, len(max_index_affinity)-1):
    list_mainLS = list_mainLS + ', ' + List_LS[max_index_affinity[i]]
  list_mainLS = list_mainLS + ', and ' + List_LS[max_index_affinity[len(max_index_affinity)-1]] + '.'

# Several learning styles with the same affinity
  document.add_paragraph('The maximum value of the affinity equals ' + max_affinity_str +' %. The dominant learning styles are: ' + list_mainLS)
  for i in max_index_affinity:
    document.add_paragraph('The main characteristics of the ' + List_LS[i] + ' learning style are:')
    for point in main_characteristics[i]:
      document.add_paragraph(point, style='ListBullet')
    document.add_paragraph('Other characteristics usually associated with the ' + List_LS[i] + ' learning style are being: ' + other_characteristics[i])

# Add a comment if the other learning styles are larger than 85%
high_affinity_indices = np.where((affinity >= 85) & (affinity < max_affinity))[0]
print('high_affinity_indices : ', high_affinity_indices)
if( high_affinity_indices.size > 0 ):
  unique_high_affinity_indices = np.setdiff1d(high_affinity_indices, max_index_affinity)
  print('unique_high_affinity_indices : ', unique_high_affinity_indices)
  list_highLS = List_LS[unique_high_affinity_indices[0]]
  for i in range(1, len(unique_high_affinity_indices)-1):
    list_highLS = list_highLS + ', ' + List_LS[unique_high_affinity_indices[i]]
  list_highLS = list_highLS + ', and ' + List_LS[unique_high_affinity_indices[len(unique_high_affinity_indices)-1]] + '.'

  if(len(unique_high_affinity_indices)==1):
    document.add_paragraph('Other learning style with a high affinity (larger than 85%) is: ' + List_LS[unique_high_affinity_indices[0]] + '.')
  else:
    document.add_paragraph('Other learning styles with high affinities (larger than 85%) are: ' + list_highLS)

  for i in unique_high_affinity_indices:
    document.add_paragraph('The main characteristics of the ' + List_LS[i] + ' learning style are:')
    for point in main_characteristics[i]:
      document.add_paragraph(point, style='ListBullet')
    document.add_paragraph('Other characteristics usually associated with the ' + List_LS[i] + ' learning style are being: ' + other_characteristics[i])

# Structure of the eigenvector associated with the main principal component (PC0)
v0 = eigenVectors_percentage[:,0]
#
# LS with the largest component
max_eigen = np.max(v0)
max_index_eigen = np.where(v0 == max_eigen)[0]
LLmax = len(max_index_eigen)
if(LLmax == 1):
  document.add_paragraph("According to the principal components analysis, the differences of the students' learning styles mostly correspond to the " + List_LS[max_index_eigen[0]] + " learning style, as " + str(round(max_eigen,1)) + ' %' + " of the eigenvector associated with the main principal component (PC0) is projected onto it.")
else:
  list_LS = ''
  for i in range(0, LLmax-1):
    list_LS = list_LS + List_LS[max_index_eigen[i]] + ', '
  list_LS = list_LS + 'and ' + List_LS[max_index_eigen[LLmax-1]]
  document.add_paragraph("According to the principal components analysis, the differences of the students' learning styles mostly correspond to the " + list_LS + " learning styles, as " + str(round(max_eigen,1)) + ' %' + " of the eigenvector associated with the main principal component (PC0) is projected onto it.")

# Add a comment on this learning style if it has not been included before
# Combine the arrays to exclude
excluded = np.union1d(max_index_affinity, high_affinity_indices)

# Get elements in max_index_eigen that are not in excluded
different_elements = np.setdiff1d(max_index_eigen, excluded)
print('different_elements : ', different_elements)

for i in different_elements:
  print('i usado en doc : ', i)
  document.add_paragraph('The main characteristics of the ' + List_LS[i] + ' learning style are:')
  for point in main_characteristics[i]:
    document.add_paragraph(point, style='ListBullet')
  document.add_paragraph('Furthermore, other characteristics usually associated with the ' + List_LS[i] + ' learning style are being: ' + other_characteristics[i])
#
# Add a comment if another component of the eigenvector is larger than 40%
high_v0_index = np.where((v0 >= 40) & (v0 < max_eigen))[0]
print('high_v0_index : ', high_v0_index)
if( high_v0_index.size > 0 ):
  unique_high_v0_index = np.setdiff1d(high_v0_index, max_index_affinity)
  print('unique_high_v0_index : ', unique_high_v0_index)
  document.add_paragraph('Other learning style with a high projection on the PC0 (larger than 40%) is: ' + List_LS[high_v0_index[0]] + '.')

  if(unique_high_v0_index == True ):
    document.add_paragraph('The main characteristics of the ' + List_LS[unique_high_v0_index] + ' learning style are:')
    for point in main_characteristics[unique_high_v0_index]:
      document.add_paragraph(point, style='ListBullet')
    document.add_paragraph('Other characteristics usually associated with the ' + List_LS[unique_high_v0_index] + ' learning style are being: ' + other_characteristics[unique_high_v0_index])

nsubsec += 1

#-------------
# Activist
#-------------
i = 0
# if ( xmean[i] >= 13 or xmean[i] <= 9):
document.add_heading('Activist learning style', level=3)
document.add_paragraph("The activist learning style has an average and an uncertainty (in parenthesis) of " + str(mean_uncert(xmean[i],dxmean[i])) + " points." )
#  if ( xmean[i] >= 13 ):
document.add_paragraph("People with scores higher than 13 on the test find more comfortable in learning situations where they can: ")

# Bullet points for high score situations
high_score_points = [
            "Attempt new things, new experiences, new opportunities.",
            "Compete in a team.",
            "Generate ideas without formal or structural limitations.",
            "Solve problems.",
            "Change and vary things.",
            "Address multiple tasks.",
            "Engage in dramatization.",
            "Play roles.",
            "Experience situations of interest or crisis.",
            "Capture attention."]

for point in high_score_points:
  document.add_paragraph(point, style='ListBullet')

# Uncomfortable situations
document.add_paragraph("\nSimilarly, they'll feel uncomfortable in situations that demand:")

# Bullet points for uncomfortable situations
uncomfortable_points = [
            "Presenting topics with a heavy theoretical load: Explaining causes, backgrounds...",
            "Assimilating, analyzing, and interpreting many unclear data.",
            "Paying attention to details.",
            "Working alone, reading, writing, or thinking alone.",
            "Evaluating in advance what is going to be learned.",
            "Assessing what has already been achieved or learned.",
            "Repeating the same activity."]

for point in uncomfortable_points:
  document.add_paragraph(point, style='ListBullet')

document.add_paragraph("\nIf someone has scored below 9, it's advisable to develop certain aspects that characterize this learning style further. Suggestions for improvement include:")

# Suggestions for improvement
improvement_suggestions = [
              "Trying something new, something never done before, at least once a week... Bringing something attention-grabbing to the study or workplace. Reading a newspaper with opposing opinions. Rearranging furniture at home or work.",
              "Practicing initiating conversations with strangers. In large gatherings, parties, conferences, forcing oneself to initiate and sustain conversations with everyone present, if possible. During leisure time, trying to engage in dialogue with strangers or convince them of our ideas.",
              "Deliberately fragmenting the day by cutting and changing activities every half an hour..."]

for point in improvement_suggestions:
  document.add_paragraph(point, style='ListBullet')

# Blocks that inhibit development
document.add_paragraph("\nBlocks that inhibit the development of the activist style:")

blocks = [    "Fear of failure or making mistakes.",
              "Fear of ridicule.",
              "Anxiety towards new or unfamiliar things.",
              "Strong desire to thoroughly think things through in advance.",
              "Self-doubt, lack of self-confidence.",
              "Taking life very seriously, very conscientiously."]

for point in blocks:
  document.add_paragraph(point, style='ListBullet')
#
#-------------
# Reflector
#-------------
i+=1

document.add_heading('Reflector learning style', level=3)
document.add_paragraph("The reflector learning style has an average and an uncertainty (in parenthesis) of " + str(mean_uncert(xmean[i],dxmean[i])) + " points." )
#  if ( xmean[i] >= 18 ):
document.add_paragraph("People with scores higher than 18 on the test learn better in learning situations where they can:")

# Bullet points: preferred reflector learning situations
reflector_strengths = [
        "Observe.",
        "Reflect on activities.",
        "Exchange opinions with others by prior agreement.",
        "Reach decisions at their own pace.",
        "Work without pressures or mandatory deadlines.",
        "Review what's been learned, what has happened.",
        "Investigate thoroughly.",
        "Gather information.",
        "Probe to get to the heart of the matter.",
        "Think before acting."
]

for point in reflector_strengths:
  document.add_paragraph(point, style='List Bullet')

# Discomfort situations
document.add_paragraph("Similarly, they'll feel uncomfortable in situations that demand:")

reflector_discomforts = [
        "Taking the forefront.",
        "Acting as a leader.",
        "Chairing meetings or debates.",
        "Performing in front of observing individuals.",
        "Playing a role.",
        "Participating in situations that require action without planning.",
        "Doing something without prior notice. Presenting an idea spontaneously."]

for point in reflector_discomforts:
  document.add_paragraph(point, style='List Bullet')

document.add_paragraph("If someone has scored below 14, they might consider the following suggestions to enhance the reflector style:")

reflector_suggestions = [
          "Practice observation. Study people's behavior. Note who speaks the most, who interrupts, how often the teacher summarizes... Study non-verbal behavior, when people look at the clock, cross their arms, bite their pencil...",
          "Keep a personal diary. Reflect on the day's events and see if any conclusions can be drawn from them.",
          "Practice review after a meeting or event. Review the sequence of events, what went well, what could be improved. Record a dialogue or conversation on a tape recorder and play it back at least twice. Make a list of lessons learned in this way."]

for point in reflector_suggestions:
  document.add_paragraph(point, style='List Bullet')

# Blocks that hinder reflector style
document.add_paragraph("Blocks that hinder the development of the reflector style:")

reflector_blocks = [
          "Not having enough time to plan and think.",
          "Preferring or enjoying quickly changing from one activity to another.",
          "Being impatient to start action.",
          "Having resistance to listening carefully and analytically.",
          "Resistance to presenting things in writing."]

for point in reflector_blocks:
  document.add_paragraph(point, style='List Bullet')
#
#-------------
# Theorist
#-------------
i+=1

document.add_heading('Theorist learning style', level=3)
document.add_paragraph("The theorist learning style has an average and an uncertainty (in parenthesis) of " + str(mean_uncert(xmean[i],dxmean[i])) + " points." )
#  if ( xmean[i] >= 14 ):
document.add_paragraph("People with scores higher than 14 on the test prefer learning opportunities where they can:")

# Add bullet points for preferred learning situations
preferred_learning = [
        "Feel in structured situations with a clear purpose.",
        "Fit all data into a system, model, concept, or theory.",
        "Have time to methodically explore associations and relationships between ideas, events, and situations.",
        "Have the chance to question.",
        "Participate in a question-and-answer session.",
        "Test methods and logic that form the basis of something."]

for item in preferred_learning:
  document.add_paragraph(item, style='List Bullet')

# Add paragraph about avoided situations
document.add_paragraph("At the same time, they'll avoid learning situations where they have to:")

# Add bullet points for avoided situations
avoided_learning = [
        "Be forced to do something without a clear context or purpose.",
        "Participate in situations where emotions and feelings predominate.",
        "Engage in unstructured activities with uncertain or ambiguous purposes.",
        "Take part in open-ended problems.",
        "Act or decide without a foundation of principles, concepts, policies, or structure."]

for item in avoided_learning:
  document.add_paragraph(item, style='List Bullet')

document.add_paragraph("If someone has scored below 10, they might want to consider the following suggestions to enhance the theorist style:")

# Suggestions to enhance the theorist style
enhancement_suggestions = [
          "Read something dense that provokes thinking for 30 minutes each day (Logic, Linguistics, Theories...). Then try to summarize what they have read using their own words.",
          "Practice detecting inconsistencies or weaknesses in other people's arguments, in reports, in newspaper articles... Take two newspapers with different ideologies and regularly perform a comparative analysis of the differences in their perspectives.",
          "Take a complex situation and analyze it to pinpoint why it developed that way, what could have been done differently, and at what point. Historical or everyday life situations. Analyze how they've used their own time..."]

for item in enhancement_suggestions:
  document.add_paragraph(item, style='List Bullet')

# Add final paragraph about blocks
document.add_paragraph("Blocks that hinder the development of the theorist style:")

# Blocks list
theorist_blocks = [
          "Being swayed by first impressions.",
          "Preferring intuition and subjectivity.",
          "Disliking structured and organized approaches.",
          "Preferring spontaneity and risk-taking."]

for item in theorist_blocks:
  document.add_paragraph(item, style='List Bullet')
#
#-------------
# Pragmatist
#-------------
i+=1
## In Cantino, it is considered xmean[i] >= 16 instead of xmean[i] >= 14.
## We set it equal to 14, which corresponds to a high tendency towards the
## pragmatic learning style, as in the other ones.
document.add_heading('Pragmatist learning style', level=3)
document.add_paragraph("The pragmatist learning style has an average and an uncertainty (in parenthesis) of " + str(mean_uncert(xmean[i],dxmean[i])) + " points." )
#  if ( xmean[i] >= 14 ):
document.add_paragraph("People with scores higher than 16 on the test find more comfortable in situations that allow them to:")

# Preferred pragmatist learning situations
pragmatist_strengths = [
        "Learn techniques for doing things with evident practical advantages.",
        "Be exposed to a model they can emulate.",
        "Acquire immediately applicable techniques in your work.",
        "Have the immediate possibility to apply what they've learned, to experiment.",
        "Develop action plans with a clear result.",
        "Give directions, suggest shortcuts."]

for point in pragmatist_strengths:
  document.add_paragraph(point, style='List Bullet')

# Discomfort situations
document.add_paragraph("Likewise, they'll find difficulty in situations that demand:")

pragmatist_discomforts = [
        "Realizing that learning isn't related to an immediate need that they recognize or cannot see.",
        "Perceiving that the learning doesn't have immediate importance or practical benefit.",
        "Learning something that is distant from reality.",
        "Learning theories and general principles.",
        "Working without clear instructions on how to do it."]

for point in pragmatist_discomforts:
  document.add_paragraph(point, style='List Bullet')

document.add_paragraph("If someone has scored below 10, the following suggestions will be useful for enhancing the pragmatist style:")

pragmatist_suggestions = [
          "Gather techniques, practical ways of doing things. They can cover anything that might be useful: analytical techniques, interpersonal skills, assertiveness, time-saving presentation techniques, statistics, memory improvement techniques...",
          "Seek help from individuals with demonstrated experience. They can observe our technique and advise on how to improve it.",
          "Focus on developing action plans in meetings and discussions of all kinds. These action plans should be specific and have deadlines. Make it a rule never to leave a meeting, debate, or class without a list of actions for yourself."]

for point in pragmatist_suggestions:
  document.add_paragraph(point, style='List Bullet')

# Blocks that hinder development
document.add_paragraph("Blocks that hinder the development of the pragmatist style:")

pragmatist_blocks = [
          "Interest in the perfect solution rather than practicality.",
          "Considering useful techniques as exaggerated simplifications.",
          "Always leaving topics open and not committing to specific actions.",
          "Believing that others' ideas won't work when applied to their situation.",
          "Enjoying marginal topics or getting lost in them."]

for point in pragmatist_blocks:
  document.add_paragraph(point, style='List Bullet')
#
#==================================================
# FINAL CONTACT PAGE
#==================================================
document.add_page_break()
document.add_paragraph(' ')
closing_paragraph = document.add_paragraph()
closing_paragraph.add_run(
    "If you have any comments or suggestions, feel free to reach out by sending an email with your feedback to fabio.revuelta@upm.es.").italic = True
document.add_paragraph(' ')
#
#==================================================
# End report
#==================================================
# Get the current date and time
current_datetime = datetime.datetime.now()
#
# Format the date and time as a string
formatted_datetime = current_datetime.strftime("%Y-%m-%d %H:%M:%S")
printt(' ')
printt('Program continues on: ' + formatted_datetime)
printt(' ')
#
#---------------------------------------------------
printt('-------------------------------------------')
printt(' Saving summary report...')
printt('-------------------------------------------')
# Docx
printt('   Docx...')
report_name='CHAEA_learning_styles_summary_report'
report_docx=report_name+'.docx'
report_docx_file = output_gen+'/'+report_docx
report_pdf_file  = output_gen+'/'+report_name+'.pdf'
printt('     Saving ' + report_docx + '...')
document.save(report_docx_file)
printt('   Docx saved!')
printt(' ')
#
# PDF
def convert_docx_to_pdf(operating_system, docx_file):
  print('OS : ', operating_system)
  if(operating_system == 'linux'):
    try:
      subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", "output", docx_file])
      printt("     Report successfully saved in PDF format, too!")
    except Exception as e:
      printt(f"Conversion failed: {e}")
  elif(operating_system == 'win'):
    try:
      subprocess.run(["docx2pdf", report_docx_file, report_pdf_file])
      printt("     Report successfully saved in PDF format, too!")
    except Exception as e:
      printt(f"Conversion failed: {e}")
  elif(operating_system == 'mac'):
    try:
      subprocess.run(["docx2pdf", report_docx_file, report_pdf_file])
      printt("     Report successfully saved in PDF format, too!")
    except Exception as e:
      printt(f"Conversion failed: {e}")
#
#
printt('   PDF...')
convert_docx_to_pdf(operating_system, report_docx_file)
printt('   PDF saved!')
#
printt('-------------------------------------------')
printt(' Summary report saved!')
printt('-------------------------------------------')
#---------------------------------------------------
#
#===================================================
#===================================================
printt('===========================================')
printt('SUMMARY REPORT DONE!')
printt('===========================================')
printt('  ')
printt('  ')
printt('  ')
printt('  ')
printt('===========================================')
printt('E-MAIL CONTACT...')
printt('===========================================')
def send_feedback():
    email = "fabio.revuelta@upm.es"
    subject = "Feedback on CHAEA3S"
    body = "Hello Fabio,\n\nI have a suggestion for CHAEA3S software package:\n\n"

    # Create a mailto link
    mailto_link = f"mailto:{email}?subject={subject}&body={body.replace(' ', '%20')}"

    # Ask user if they want to send feedback
    user_input = input("Would you like to send feedback with any comments or suggestions? If so, type 'y' or 'yes' (in uppercase or lowercase): ").strip()

    if user_input in ["y", "Y", "yes", "YES", "Yes", "yEs", "yeS", "YEs", "YeS", "yES"]:
        print(f"Opening email client... The email will be sent to: {email}")
        webbrowser.open(mailto_link)
    else:
        print("Feedback request canceled.")

# Run the function
if __name__ == "__main__":
    send_feedback()
printt('===========================================')
printt('E-MAIL CONTACT DONE!')
printt('===========================================')
printt(' ')
printt(' ')
#
# Get the current date and time
current_datetime = datetime.datetime.now()
#
# Format the date and time as a string
formatted_datetime = current_datetime.strftime("%Y-%m-%d %H:%M:%S")
#===================================================
#===================================================
#===================================================
#===================================================
printt('===========================================')
printt('===========================================')
printt('BINGO PITINGO!')
printt('===========================================')
printt('===========================================')
printt(' ')
printt('===========================================')
printt('===========================================')
printt('Program finishes on: ' + formatted_datetime)
printt('===========================================')
printt('===========================================')
printt(' ')
#==================================================================
#==================================================================
#==================================================================
#==================================================================
